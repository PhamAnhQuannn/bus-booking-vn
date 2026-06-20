"""Generate Vietnamese legal dossier report as .doc — detailed version."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

DARK_BLUE = RGBColor(0, 51, 102)
DARK_RED = RGBColor(128, 0, 0)
DARK_GREEN = RGBColor(0, 100, 0)


# ── helpers ──────────────────────────────────────────────────
def add_center(text, size=18, bold=True, color=DARK_BLUE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = 'Times New Roman'

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = DARK_BLUE

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 80, 130)

def h3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.name = 'Times New Roman'

def para(text, bold=False, italic=False, color=None, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    if color:
        r.font.color.rgb = color
    return p

def bullet(text, bold_prefix='', size=11):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = 'Times New Roman'
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.name = 'Times New Roman'
    else:
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.name = 'Times New Roman'
        if not p.runs:
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.font.name = 'Times New Roman'

def note(text):
    para(text, italic=True, color=DARK_RED, size=11)

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Light Grid Accent 1'
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Times New Roman'
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            c = row.cells[i]
            c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(11)
            r.font.name = 'Times New Roman'
    doc.add_paragraph()
    return t


# ═════════════════════════════════════════════════════════════
# TITLE
# ═════════════════════════════════════════════════════════════
add_center('BÁO CÁO CHI TIẾT', size=20)
add_center('HỒ SƠ PHÁP LÝ CẦN HOÀN THÀNH', size=16)
add_center('Dự án: Sàn đặt vé xe khách trực tuyến (Bus Booking)', size=13, bold=False)
doc.add_paragraph()

note('⚠ Lưu ý: Công ty đã có Giấy chứng nhận đăng ký doanh nghiệp (ERC) và Mã số thuế (MST) '
     'từ hoạt động kinh doanh vận tải bê tông cho công trình xây dựng. '
     'Báo cáo này tập trung vào các hồ sơ pháp lý BỔ SUNG cần thiết để vận hành sàn TMĐT đặt vé xe khách.')

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════
# SECTION 1 — NARROWED: already have company
# ═════════════════════════════════════════════════════════════
h1('1. Bổ sung ngành nghề kinh doanh (VSIC)')

para('VSIC = Vietnam Standard Industrial Classification — Hệ thống ngành kinh tế Việt Nam 2025',
     italic=True, size=11, color=DARK_GREEN)

h2('1.1 Là gì?')
para('VSIC là hệ thống mã ngành kinh tế do Tổng cục Thống kê ban hành. '
     'Mỗi doanh nghiệp phải đăng ký mã ngành phù hợp với hoạt động kinh doanh thực tế. '
     'Công ty hiện tại có mã ngành vận tải — cần BỔ SUNG thêm mã ngành CNTT/TMĐT để hợp pháp hóa hoạt động sàn đặt vé.')

h2('1.2 Tại sao cần?')
bullet('Hoạt động sàn TMĐT đặt vé xe khách KHÔNG thuộc mã ngành vận tải hiện tại')
bullet('Nếu không bổ sung → vi phạm Luật Doanh nghiệp (kinh doanh ngoài ngành đăng ký)')
bullet('Cần mã ngành CNTT/TMĐT để đăng ký sàn TMĐT tại online.gov.vn (Bước 2)')

h2('1.3 Mã ngành cần bổ sung')
table(
    ['Mã VSIC', 'Mô tả', 'Lý do'],
    [
        ['6201', 'Lập trình máy tính (Computer programming)', 'Phát triển phần mềm sàn đặt vé'],
        ['6209', 'Hoạt động dịch vụ CNTT khác', 'Vận hành nền tảng công nghệ'],
        ['4791', 'Bán lẻ qua TMĐT (E-commerce retail)', 'Bán vé xe khách trực tuyến'],
        ['6311', 'Xử lý dữ liệu, hosting (Data processing)', 'Lưu trữ và xử lý dữ liệu khách hàng'],
    ]
)

note('⚠ QUAN TRỌNG: KHÔNG chọn mã 4921 (Vận tải hành khách đường bộ). '
     'Sàn đặt vé là trung gian TMĐT, KHÔNG phải đơn vị vận tải. '
     'Chọn sai mã → bị yêu cầu Giấy phép kinh doanh vận tải + giới hạn vốn nước ngoài.')

h2('1.4 Cách thực hiện')
bullet('Bước 1: ', 'Soạn hồ sơ thay đổi nội dung đăng ký doanh nghiệp (Mẫu theo Nghị định 01/2021/NĐ-CP)')
bullet('Bước 2: ', 'Nộp tại Sở Kế hoạch và Đầu tư (Sở KH&ĐT) nơi công ty đặt trụ sở')
bullet('Bước 3: ', 'Nhận ERC cập nhật với mã ngành mới')

h2('1.5 Hồ sơ cần nộp')
bullet('Thông báo thay đổi nội dung đăng ký doanh nghiệp (theo mẫu)')
bullet('Quyết định của chủ sở hữu/hội đồng thành viên về việc bổ sung ngành nghề')
bullet('Bản sao ERC hiện tại')
bullet('CMND/CCCD người đại diện pháp luật')

h2('1.6 Thời gian + Chi phí')
table(
    ['Hạng mục', 'Thời gian', 'Chi phí'],
    [
        ['Xử lý tại Sở KH&ĐT', '3-5 ngày làm việc', 'Miễn phí (lệ phí nhà nước)'],
        ['Thuê dịch vụ pháp lý hỗ trợ', '1-2 ngày chuẩn bị', '1-3 triệu VND'],
    ]
)

# ═════════════════════════════════════════════════════════════
# SECTION 2 — TMĐT
# ═════════════════════════════════════════════════════════════
h1('2. Đăng ký sàn TMĐT — Thương mại điện tử')

para('TMĐT = Thương mại điện tử (E-Commerce)\n'
     'MOIT = Ministry of Industry and Trade — Bộ Công Thương',
     italic=True, size=11, color=DARK_GREEN)

h2('2.1 Là gì?')
para('Đăng ký sàn giao dịch TMĐT là thủ tục bắt buộc với mọi website/ứng dụng '
     'cho phép bên thứ ba (nhà xe) bán hàng hóa/dịch vụ (vé xe) thông qua nền tảng. '
     'Sàn đặt vé xe khách = sàn giao dịch TMĐT theo định nghĩa Nghị định 85/2021.')

h2('2.2 Tại sao cần?')
bullet('Bắt buộc theo Nghị định 85/2021/NĐ-CP (hướng dẫn Luật TMĐT)')
bullet('Nghị định 117/2025 (Luật TMĐT 2025) — yêu cầu mới có hiệu lực')
bullet('Mức phạt nếu không đăng ký: ', '40-60 triệu VND (Nghị định 98/2020)')
bullet('Không đăng ký → không thể hợp pháp thu phí/hoa hồng từ nhà xe')

h2('2.3 Cách thực hiện')
bullet('Bước 1: ', 'Truy cập Cổng thông tin quản lý hoạt động TMĐT: online.gov.vn')
bullet('Bước 2: ', 'Tạo tài khoản doanh nghiệp trên cổng')
bullet('Bước 3: ', 'Chọn loại hình "Sàn giao dịch TMĐT" (KHÔNG phải "Website TMĐT bán hàng")')
bullet('Bước 4: ', 'Điền thông tin và upload hồ sơ theo yêu cầu')
bullet('Bước 5: ', 'Chờ Bộ Công Thương thẩm định và phê duyệt')

h2('2.4 Nội dung / Hồ sơ cần nộp')
table(
    ['Tài liệu', 'Mô tả chi tiết'],
    [
        ['ERC (bản sao)', 'Giấy chứng nhận ĐKDN — đã có, cần bản cập nhật với mã VSIC TMĐT'],
        ['Quy chế hoạt động sàn', 'Quy định quyền/nghĩa vụ của sàn, nhà xe, hành khách;\ncơ chế xử lý tranh chấp;\nquy trình đăng ký nhà xe trên sàn'],
        ['Chính sách bảo mật\n(Privacy Policy)', 'Tuân thủ PDPL 2025 — mô tả dữ liệu thu thập,\nmục đích, thời hạn lưu, quyền của chủ thể dữ liệu'],
        ['Điều khoản sử dụng\n(ToS = Terms of Service)', 'Quyền/nghĩa vụ người dùng, giới hạn trách nhiệm,\nđiều kiện hủy vé, chính sách hoàn tiền'],
        ['Cơ chế giải quyết tranh chấp', 'Quy trình khiếu nại, thời gian xử lý,\ntrách nhiệm giữa sàn — nhà xe — hành khách'],
        ['Quy trình xác minh nhà xe', 'Mô tả KYB: kiểm tra ERC, giấy phép vận tải,\nbảo hiểm, giấy chứng thực tuyến đường'],
    ]
)

h2('2.5 Thời gian + Chi phí')
table(
    ['Hạng mục', 'Thời gian', 'Chi phí'],
    [
        ['Bộ Công Thương thẩm định', '2-4 tuần', 'Miễn phí'],
        ['Soạn quy chế + chính sách (có luật sư)', '1-2 tuần', '10-30 triệu VND'],
        ['Tự soạn (tham khảo mẫu)', '2-4 tuần', 'Nội bộ'],
    ]
)

note('⚠ ĐÂY LÀ CHẶN CỨNG — không thể vận hành sàn hợp pháp cho đến khi được phê duyệt.')

# ═════════════════════════════════════════════════════════════
# SECTION 3 — DPO + DPIA + CDTIA (expanded)
# ═════════════════════════════════════════════════════════════
h1('3. Bảo vệ dữ liệu cá nhân')

para('Cơ sở pháp lý: PDPL 2025 (Luật Bảo vệ dữ liệu cá nhân số 91/2025/QH15) '
     'và Nghị định 356/2025/NĐ-CP\nNộp cho: Cục A05 / BCA (Bộ Công an)',
     italic=True, size=11, color=DARK_GREEN)

# ── 3A: DPO ──
h2('3.1 DPO — Data Protection Officer (Cán bộ bảo vệ dữ liệu cá nhân)')

h3('Là gì?')
para('DPO là cá nhân hoặc tổ chức được bổ nhiệm chịu trách nhiệm giám sát '
     'toàn bộ hoạt động xử lý dữ liệu cá nhân của doanh nghiệp. '
     'DPO là đầu mối liên lạc giữa doanh nghiệp với Bộ Công an (A05) '
     'và với chủ thể dữ liệu (khách hàng, nhà xe).')

h3('Tại sao cần?')
bullet('Bắt buộc theo PDPL 2025 — mọi tổ chức xử lý dữ liệu cá nhân')
bullet('Startup có miễn trừ 5 năm (đến 2031) NHƯNG: ', 'sàn đặt vé xử lý dữ liệu thanh toán '
       '= dữ liệu "nhạy cảm" theo Nghị định 356/2025 → KHÔNG ĐƯỢC MIỄN TRỪ')
bullet('Không bổ nhiệm DPO → vi phạm PDPL, phạt đến 5% doanh thu năm trước')

h3('Ai có thể làm DPO?')
table(
    ['Phương án', 'Yêu cầu', 'Ưu/Nhược điểm'],
    [
        ['Nhân sự nội bộ kiêm nhiệm', 'Tối thiểu 2 năm kinh nghiệm\nbảo vệ dữ liệu cá nhân',
         'Ưu: tiết kiệm chi phí\nNhược: khó tìm người đủ chuyên môn'],
        ['Tuyển nhân sự chuyên trách', 'Chuyên gia privacy/compliance\nBáo cáo trực tiếp ban lãnh đạo',
         'Ưu: chuyên nghiệp\nNhược: chi phí lương cao'],
        ['Thuê ngoài (Outsourced DPO)', 'Công ty tư vấn có giấy phép\nKinh nghiệm PDPL 2025',
         'Ưu: nhanh nhất cho startup\nNhược: chi phí 5-10 triệu VND/tháng'],
    ]
)

h3('DPO làm những gì?')
bullet('Chuẩn bị và nộp DPIA + CDTIA cho Cục A05')
bullet('Điều phối xử lý sự cố vi phạm dữ liệu (thông báo A05 trong 72 giờ)')
bullet('Quản lý đồng ý (consent) — thu thập, rút lại, lưu trữ bằng chứng')
bullet('Xử lý yêu cầu quyền dữ liệu từ khách hàng (truy cập/sửa/xóa) — SLA 10-20 ngày')
bullet('Đào tạo nhân viên về bảo vệ dữ liệu hàng năm')
bullet('Độc lập với bộ phận kinh doanh — không bị chi phối bởi lợi ích thương mại')

h3('Cách thực hiện')
bullet('Bước 1: ', 'Chọn phương án bổ nhiệm (nội bộ hoặc thuê ngoài)')
bullet('Bước 2: ', 'Soạn Quyết định bổ nhiệm DPO — gồm: họ tên, chứng chỉ, '
       'cam kết độc lập, quyền báo cáo trực tiếp ban lãnh đạo')
bullet('Bước 3: ', 'Nộp thông báo bổ nhiệm DPO cho Cục A05 / BCA (thường nộp chung với DPIA)')
bullet('Hạn chót: ', 'TRƯỚC khi thu thập dữ liệu người dùng đầu tiên (= trước soft launch)')

h3('Thời gian + Chi phí')
table(
    ['Hạng mục', 'Thời gian', 'Chi phí'],
    [
        ['Bổ nhiệm nội bộ', '1 tuần', 'Không phát sinh thêm (kiêm nhiệm)'],
        ['Thuê DPO ngoài', '1-2 tuần tìm + ký HĐ', '5-10 triệu VND/tháng'],
        ['Nộp thông báo A05', 'Cùng lúc với DPIA', 'Miễn phí'],
    ]
)

# ── 3B: DPIA ──
h2('3.2 DPIA — Data Protection Impact Assessment (Đánh giá tác động xử lý dữ liệu cá nhân)')

h3('Là gì?')
para('DPIA là văn bản đánh giá toàn diện các rủi ro phát sinh từ việc thu thập, '
     'lưu trữ, xử lý dữ liệu cá nhân của người dùng. Đây là hồ sơ bắt buộc nộp cho '
     'Bộ Công an (Cục A05) theo mẫu quy định của Nghị định 356/2025.')

h3('Tại sao cần?')
bullet('Bắt buộc theo PDPL 2025 — KHÔNG có ngoại lệ cho bất kỳ tổ chức nào')
bullet('Hạn chót: nộp trong vòng ', '60 ngày kể từ ngày bắt đầu thu thập dữ liệu cá nhân')
bullet('Không nộp → phạt đến ', '3 tỷ VND hoặc 5% doanh thu năm trước')

h3('Nội dung DPIA gồm những gì?')
para('DPIA bao gồm các phần sau (theo mẫu Nghị định 356/2025):', bold=True, size=11)
doc.add_paragraph()

table(
    ['Phần', 'Nội dung', 'Chi tiết cho sàn đặt vé'],
    [
        ['A. Kiểm kê dữ liệu', 'Liệt kê MỌI loại dữ liệu cá nhân\nthu thập',
         'Tên, SĐT, email, lịch sử đặt vé,\ntoken thanh toán, địa chỉ IP,\nOTP attempts, thông tin nhà xe'],
        ['B. Cơ sở pháp lý', 'Lý do pháp lý cho từng\nhoạt động xử lý',
         'Đồng ý (đăng ký), hợp đồng (đặt vé),\nnghĩa vụ pháp lý (thuế/hóa đơn),\nlợi ích hợp pháp (bảo mật OTP)'],
        ['C. Phân loại nhạy cảm', 'Xác định dữ liệu nào\nlà "nhạy cảm"',
         'Dữ liệu thanh toán = TÀI CHÍNH → nhạy cảm\nĐịa chỉ IP = VỊ TRÍ → nhạy cảm\nCần đồng ý RIÊNG BIỆT cho từng loại'],
        ['D. Thời hạn lưu trữ', 'Bao lâu giữ từng loại dữ liệu,\ncơ chế xóa tự động',
         'Booking: 5 năm (Luật Kế toán)\nThanh toán/HĐĐT: 10 năm (NĐ 123)\nOTP logs: 90 ngày\nSession: hết hạn JWT (15 phút)'],
        ['E. Biện pháp bảo mật', 'Mã hóa, kiểm soát truy cập,\nche giấu dữ liệu trong log',
         'Mã hóa database, HTTPS/TLS,\nche SĐT trong log (đã triển khai),\nJWT ký HS256, CSRF double-submit'],
        ['F. Đánh giá rủi ro', 'Liệt kê từng rủi ro +\nmức độ + biện pháp giảm thiểu',
         'Lộ dữ liệu thanh toán (Cao)\nChuyển dữ liệu ra nước ngoài (Cao)\nLộ hành trình di chuyển (TB)\nSĐT làm ID chính (TB)'],
        ['G. Quy trình vi phạm', 'Thông báo BCA trong 72 giờ\nThông báo người dùng bị ảnh hưởng',
         'Phát hiện → DPO + CEO (0-2 giờ)\nBáo A05 (72 giờ)\nBáo NHNN nếu liên quan thanh toán\nĐiều tra + báo cáo (30 ngày)'],
        ['H. Cơ chế đồng ý', 'Ảnh chụp giao diện consent,\nmô tả cơ chế rút đồng ý',
         'Toggle riêng từng mục đích\nKHÔNG tick sẵn (pre-ticked prohibited)\nKHÔNG gộp marketing + dịch vụ\nRút đồng ý 1 click'],
    ]
)

h3('Cách thực hiện — Từng bước')
bullet('Bước 1 (Tuần 1-2): ', 'Kiểm kê toàn bộ dữ liệu — liệt kê mọi loại dữ liệu, nơi lưu, ai truy cập')
bullet('Bước 2 (Tuần 2-3): ', 'Ghi nhận cơ sở pháp lý cho từng hoạt động xử lý')
bullet('Bước 3 (Tuần 3-4): ', 'Đánh giá rủi ro — xác định mức độ nghiêm trọng + biện pháp giảm thiểu')
bullet('Bước 4 (Tuần 4-6): ', 'Thu thập tài liệu đính kèm: Privacy Policy, Consent UI screenshots, '
       'DPA với VNPay/MoMo/eSMS/MISA/hosting, Breach runbook')
bullet('Bước 5 (Tuần 6-8): ', 'Điền biểu mẫu DPIA theo mẫu Nghị định 356/2025 '
       '(KHUYẾN NGHỊ MẠNH: thuê luật sư rà soát)')
bullet('Bước 6 (Tuần 8-10): ', 'Nộp DPIA cho Cục A05 qua cổng BCA')
bullet('Bước 7: ', 'Chờ thẩm định 15 ngày — nếu thiếu sót có 30 ngày bổ sung')

h3('Sau khi nộp — Nghĩa vụ liên tục')
bullet('Cập nhật DPIA mỗi ', '6 tháng')
bullet('Cập nhật trong ', '10 ngày nếu có thay đổi trọng yếu (thêm loại dữ liệu mới, đổi vendor, etc.)')
bullet('Lưu trữ tất cả phiên bản DPIA — ', 'vô thời hạn (audit trail)')

h3('Thời gian + Chi phí')
table(
    ['Hạng mục', 'Thời gian', 'Chi phí'],
    [
        ['Chuẩn bị DPIA (có luật sư)', '6-10 tuần', '10-30 triệu VND'],
        ['Chuẩn bị DPIA (tự làm)', '12-16 tuần', 'Chi phí nhân sự nội bộ'],
        ['BCA thẩm định', '15 ngày', 'Miễn phí'],
        ['Bổ sung nếu thiếu sót', '30 ngày', 'Tùy mức độ'],
    ]
)

# ── 3C: CDTIA ──
h2('3.3 CDTIA — Cross-border Data Transfer Impact Assessment '
   '(Đánh giá tác động chuyển dữ liệu xuyên biên giới)')

h3('Là gì?')
para('CDTIA (còn gọi là TIA — Transfer Impact Assessment) là văn bản đánh giá rủi ro '
     'khi dữ liệu cá nhân của người dùng Việt Nam được chuyển ra nước ngoài. '
     'Bắt buộc khi sử dụng dịch vụ cloud/hosting đặt ngoài Việt Nam.')

h3('Tại sao sàn đặt vé cần CDTIA?')
bullet('Vercel Singapore = app hosting ở nước ngoài → dữ liệu transit qua Singapore')
bullet('Bắt buộc theo PDPL 2025 + Nghị định 356/2025 + Nghị định 53/2022')
bullet('Hạn chót: nộp trong vòng ', '60 ngày kể từ lần chuyển dữ liệu đầu tiên ra nước ngoài')
bullet('Không nộp → phạt đến ', '5% doanh thu năm trước + BCA có thể yêu cầu DỪNG chuyển dữ liệu')

h3('Nội dung CDTIA')
table(
    ['Phần', 'Nội dung cần có'],
    [
        ['Danh sách luồng chuyển DL', 'Liệt kê: dữ liệu nào → nhà cung cấp nào → quốc gia nào\n'
         'VD: Booking data → Vercel → Singapore\n'
         'Payment tokens → VNPay → Vietnam (KHÔNG chuyển ra ngoài)'],
        ['Giải trình sự cần thiết', 'Tại sao phải chuyển ra nước ngoài (tốc độ, CDN, etc.);\n'
         'không có giải pháp tương đương trong nước'],
        ['Biện pháp bảo vệ', 'Hợp đồng ràng buộc (DPA) với nhà cung cấp;\n'
         'mã hóa dữ liệu khi truyền (TLS) và khi lưu;\n'
         'kiểm soát truy cập'],
        ['Đánh giá rủi ro quốc gia đích', 'Singapore: đánh giá mức độ bảo vệ DL cá nhân;\n'
         'khung pháp lý PDPA (Singapore) tương đương'],
        ['Bằng chứng đồng ý', 'Người dùng đã được thông báo và đồng ý\n'
         'việc dữ liệu chuyển ra nước ngoài'],
    ]
)

h3('Phương án giảm thiểu rủi ro')
table(
    ['Phương án', 'Mô tả', 'Đánh giá'],
    [
        ['A. DB đặt tại Việt Nam', 'PostgreSQL trên Viettel IDC / VNPT / FPT Cloud;\n'
         'Vercel Singapore chỉ làm compute/CDN',
         'Tốt nhất cho compliance\nThêm 5-15ms latency (chấp nhận được)'],
        ['B. Nộp CDTIA + ký DPA', 'Giữ nguyên kiến trúc hiện tại;\nnộp CDTIA + DPA với hosting provider',
         'Đơn giản hơn\nNhưng BCA có thể yêu cầu chuyển về VN sau'],
    ]
)

h3('Thời gian + Chi phí')
table(
    ['Hạng mục', 'Thời gian', 'Chi phí'],
    [
        ['Chuẩn bị CDTIA (có luật sư)', '1-2 tuần', '5-15 triệu VND'],
        ['Nộp cho A05 (chung với DPIA)', 'Cùng lúc', 'Miễn phí'],
        ['Chuyển DB về VN (nếu chọn PA A)', '2-4 tuần kỹ thuật', '~2-5 triệu VND/tháng hosting'],
    ]
)

# ═════════════════════════════════════════════════════════════
# SECTION 4 — TAX & E-INVOICE
# ═════════════════════════════════════════════════════════════
h1('4. Thuế & Hóa đơn điện tử (HĐĐT)')

para('Cơ sở pháp lý: Nghị định 123/2020, Nghị định 70/2025, Thông tư 32/2025\n'
     'Cơ quan: GDT = General Department of Taxation — Tổng cục Thuế',
     italic=True, size=11, color=DARK_GREEN)

note('⚠ Công ty ĐÃ CÓ MST + đăng ký VAT từ hoạt động vận tải. '
     'Phần này tập trung vào hồ sơ BỔ SUNG cho hoạt động sàn TMĐT.')

h2('4.1 Phán quyết thuế về vai trò xuất HĐĐT')

h3('Là gì?')
para('Phán quyết thuế (tax ruling) là văn bản xác nhận vai trò của sàn trong việc xuất hóa đơn điện tử. '
     'Sàn TMĐT có thể: (a) ủy quyền nhà xe tự xuất HĐĐT, hoặc (b) sàn xuất HĐĐT thay nhà xe. '
     'Cần xác nhận rõ với cơ quan thuế để tránh tranh chấp sau này.')

h3('Cách thực hiện')
bullet('Bước 1: ', 'Gửi công văn hỏi cơ quan thuế quản lý trực tiếp về vai trò xuất HĐĐT')
bullet('Bước 2: ', 'Nhận phản hồi bằng văn bản (thường 15-30 ngày)')
bullet('Bước 3: ', 'Dựa trên phản hồi → ký hợp đồng ủy quyền phù hợp với từng nhà xe')

h2('4.2 Hợp đồng ủy quyền xuất HĐĐT')

h3('Là gì?')
para('Hợp đồng ủy quyền xuất HĐĐT = Authorization Agreement for E-Invoice Issuance. '
     'Theo Điều 17 Nghị định 123/2020 (sửa đổi bởi NĐ 70/2025), bên bán (nhà xe) '
     'có thể ủy quyền cho bên thứ ba (sàn) xuất HĐĐT thay mình cho hành khách.')

h3('Nội dung hợp đồng')
bullet('Thông tin bên ủy quyền (nhà xe): tên, MST, địa chỉ')
bullet('Thông tin bên được ủy quyền (sàn): tên, MST, địa chỉ')
bullet('Phạm vi ủy quyền: loại hóa đơn, mặt hàng/dịch vụ, thời hạn')
bullet('Trách nhiệm của mỗi bên: kê khai thuế, lưu trữ hóa đơn')
bullet('Thông báo cho cơ quan thuế: gửi bản sao hợp đồng cho Chi cục Thuế quản lý nhà xe')

h3('Cách thực hiện')
bullet('Ký với TỪNG nhà xe đối tác (không ký chung)')
bullet('Gửi bản sao hợp đồng cho Chi cục Thuế quản lý nhà xe')
bullet('Lưu bản gốc tại sàn — thời hạn lưu 10 năm')

h2('4.3 Đăng ký nhà cung cấp HĐĐT + Chữ ký số')

h3('Là gì?')
para('Để xuất HĐĐT hợp lệ, sàn cần: (1) đăng ký sử dụng phần mềm HĐĐT từ nhà cung cấp '
     'được GDT chứng nhận (khuyến nghị: MISA), và (2) mua chữ ký số (Digital Signature / CA) '
     'từ nhà cung cấp được GDT chấp thuận.')

h3('Nhà cung cấp chữ ký số (CA = Certificate Authority)')
table(
    ['Nhà cung cấp', 'Ghi chú'],
    [
        ['VNPT-CA', 'Phổ biến nhất, tích hợp tốt với MISA'],
        ['Viettel-CA', 'Giá cạnh tranh'],
        ['FPT-CA', 'Hỗ trợ kỹ thuật tốt'],
    ]
)

h3('Thời gian + Chi phí')
table(
    ['Hạng mục', 'Thời gian', 'Chi phí'],
    [
        ['Đăng ký MISA HĐĐT', '3-5 ngày', '~1-3 triệu VND/năm'],
        ['Mua chữ ký số', '1-2 ngày', '~1-2 triệu VND/năm'],
        ['Tích hợp API MISA', '1-2 tuần dev', 'Nội bộ'],
    ]
)

# ═════════════════════════════════════════════════════════════
# SECTION 5 — SMS
# ═════════════════════════════════════════════════════════════
h1('5. Viễn thông & SMS — ⚠ CHẶN CỨNG')

para('Cơ sở pháp lý: Nghị định 91/2020/NĐ-CP (chống tin nhắn rác)\n'
     'Cơ quan: Từng nhà mạng (Viettel, VNPT/VinaPhone, Mobifone)',
     italic=True, size=11, color=DARK_GREEN)

h2('5.1 Đăng ký SMS Brandname')

h3('Là gì?')
para('SMS Brandname là tin nhắn hiển thị TÊN THƯƠNG HIỆU (VD: "BUSBOOK") thay vì số điện thoại ngẫu nhiên. '
     'Bắt buộc cho OTP xác thực và thông báo đặt vé. Nếu không đăng ký → tin nhắn bị nhà mạng chặn hoặc '
     'hiển thị số lạ → người dùng không tin tưởng.')

h3('Tại sao là CHẶN CỨNG?')
bullet('Không có SMS Brandname → không gửi được OTP → người dùng không đăng ký/đăng nhập được')
bullet('Thời gian phê duyệt: ', '2-4 tuần MỖI nhà mạng — không rút ngắn được')
bullet('Phải nộp cho TỪNG nhà mạng riêng biệt')

h3('Cách thực hiện')
bullet('Bước 1: ', 'Chọn nhà cung cấp SMS aggregator (VD: eSMS, SpeedSMS, Stringee) '
       'HOẶC nộp trực tiếp cho từng nhà mạng')
bullet('Bước 2: ', 'Chuẩn bị hồ sơ đăng ký (xem bên dưới)')
bullet('Bước 3: ', 'Nộp cho từng nhà mạng: Viettel, VNPT/VinaPhone, Mobifone')
bullet('Bước 4: ', 'Chờ phê duyệt 2-4 tuần / nhà mạng')
bullet('Bước 5: ', 'Sau phê duyệt → tích hợp API gửi SMS qua aggregator')

h3('Hồ sơ cần nộp cho mỗi nhà mạng')
bullet('Bản sao ERC (Giấy chứng nhận ĐKDN)')
bullet('Bản sao MST (Mã số thuế)')
bullet('Mẫu tin nhắn (SMS templates) — VD: "Ma OTP cua ban la {code}. Het han sau 5 phut."')
bullet('Tên Brandname muốn đăng ký (VD: "BUSBOOK")')
bullet('Mục đích sử dụng (OTP xác thực, thông báo đặt vé, etc.)')

h2('5.2 Xác minh Zalo OA (Official Account)')

h3('Là gì?')
para('Zalo OA = tài khoản doanh nghiệp chính thức trên Zalo. '
     'Cần xác minh để sử dụng ZNS (Zalo Notification Service) — '
     'kênh thông báo đặt vé thay thế/bổ sung cho SMS, chi phí thấp hơn SMS.')

h3('Cách thực hiện')
bullet('Đăng ký tại oa.zalo.me')
bullet('Upload ERC + CMND/CCCD người đại diện')
bullet('Chờ Zalo xác minh (3-7 ngày)')

h3('Thời gian + Chi phí')
table(
    ['Hạng mục', 'Thời gian', 'Chi phí'],
    [
        ['SMS Brandname (mỗi nhà mạng)', '2-4 tuần', '~500-2.000 VND/SMS (theo gói)'],
        ['SMS qua aggregator (eSMS)', '1-2 tuần setup', '~350-800 VND/SMS'],
        ['Zalo OA xác minh', '3-7 ngày', 'Miễn phí xác minh; ZNS ~200 VND/tin'],
    ]
)

# ═════════════════════════════════════════════════════════════
# SECTION 6 — OPERATOR KYB
# ═════════════════════════════════════════════════════════════
h1('6. Hồ sơ thu thập từ mỗi nhà xe (Operator KYB)')

para('KYB = Know Your Business — Xác minh doanh nghiệp đối tác\n'
     'Cơ sở pháp lý: Nghị định 85/2021, Nghị định 10/2020, Luật TMĐT 2025',
     italic=True, size=11, color=DARK_GREEN)

h2('6.1 Là gì?')
para('KYB là quy trình xác minh tính hợp pháp của nhà xe trước khi cho phép '
     'đăng bán vé trên sàn. Sàn TMĐT có NGHĨA VỤ PHÁP LÝ xác minh đối tác — '
     'nếu nhà xe vi phạm và sàn không xác minh → sàn chịu trách nhiệm liên đới.')

h2('6.2 Danh sách hồ sơ cần thu thập từ mỗi nhà xe')

table(
    ['Hồ sơ', 'Viết tắt / Giải thích', 'Cơ sở pháp lý', 'Chi tiết'],
    [
        ['Giấy chứng nhận ĐKDN\n(ERC)', 'ERC = Enterprise Registration\nCertificate',
         'Nghị định 85/2021', 'Xác nhận nhà xe là doanh nghiệp\nhợp pháp, còn hoạt động'],
        ['Mã số thuế (MST)', 'MST = Tax Identification\nNumber',
         'Nghị định 85/2021', 'Để ký HĐ ủy quyền HĐĐT\nvà đối soát thuế'],
        ['Giấy phép KDVT', 'KDVT = Kinh doanh\nvận tải hành khách',
         'Nghị định 10/2020/NĐ-CP', 'Xác nhận nhà xe được phép\nvận chuyển hành khách'],
        ['Bảo hiểm hành khách', 'Chứng chỉ bảo hiểm\ntrách nhiệm dân sự',
         'NĐ 03/2021,\nsửa đổi NĐ 67/2023', 'Bắt buộc cho mọi xe khách;\nkiểm tra hạn hiệu lực'],
        ['Chứng thực tuyến đường', 'Giấy phép hoạt động\ntuyến cố định',
         'Sở GTVT cấp', 'Xác nhận nhà xe chạy đúng\ntuyến đã đăng ký'],
        ['Tài khoản ngân hàng', 'Xác minh TK NH\ncủa nhà xe',
         '—', 'Cho VNPay split-settlement;\nchuyển tiền vé về nhà xe'],
    ]
)

h2('6.3 Quy trình KYB trên sàn')
bullet('Bước 1: ', 'Nhà xe đăng ký tài khoản operator trên sàn')
bullet('Bước 2: ', 'Upload bản scan/ảnh các giấy tờ trên')
bullet('Bước 3: ', 'Admin sàn xác minh: đối chiếu ERC trên dangkykinhdoanh.gov.vn, '
       'kiểm tra giấy phép KDVT, kiểm tra hạn bảo hiểm')
bullet('Bước 4: ', 'Phê duyệt hoặc yêu cầu bổ sung')
bullet('Bước 5: ', 'Thiết lập cron kiểm tra: cảnh báo trước 60 ngày khi giấy phép/bảo hiểm hết hạn')

# ═════════════════════════════════════════════════════════════
# SECTION 7 — PSP
# ═════════════════════════════════════════════════════════════
h1('7. Thanh toán — Onboarding PSP')

para('PSP = Payment Service Provider — Nhà cung cấp dịch vụ thanh toán\n'
     'Áp dụng cho: VNPay, MoMo',
     italic=True, size=11, color=DARK_GREEN)

h2('7.1 Là gì?')
para('PSP Onboarding = quy trình đăng ký merchant (tài khoản đối tác) với cổng thanh toán '
     'để nhận thanh toán trực tuyến từ khách hàng. Mỗi PSP có quy trình riêng nhưng '
     'yêu cầu hồ sơ tương tự.')

h2('7.2 Hồ sơ cần nộp cho VNPay / MoMo')
table(
    ['Tài liệu', 'Chi tiết'],
    [
        ['ERC (bản sao công chứng)', 'Giấy chứng nhận ĐKDN — đã có'],
        ['MST', 'Mã số thuế — đã có'],
        ['Tài khoản ngân hàng doanh nghiệp', 'Sổ TK hoặc xác nhận từ ngân hàng'],
        ['Website hoạt động', 'Phải có: Privacy Policy, ToS,\nthông tin liên hệ, quy trình hoàn tiền'],
        ['CMND/CCCD người đại diện', 'Bản sao công chứng'],
        ['Mô tả dịch vụ', 'Giải thích mô hình sàn đặt vé xe khách'],
    ]
)

h2('7.3 Cách thực hiện')
bullet('Bước 1: ', 'Đăng ký sandbox (môi trường test) — có thể làm TRƯỚC khi có ERC cập nhật')
bullet('Bước 2: ', 'Tích hợp kỹ thuật API trong sandbox')
bullet('Bước 3: ', 'Nộp hồ sơ production khi website sẵn sàng')
bullet('Bước 4: ', 'PSP kiểm tra website + hồ sơ → phê duyệt merchant')

h2('7.4 Thời gian + Chi phí')
table(
    ['Hạng mục', 'Thời gian', 'Chi phí'],
    [
        ['Sandbox VNPay', '1-3 ngày', 'Miễn phí'],
        ['Production VNPay', '1-2 tuần', 'Phí giao dịch ~1.1-1.5%'],
        ['Sandbox MoMo', '1-3 ngày', 'Miễn phí'],
        ['Production MoMo', '1-2 tuần', 'Phí giao dịch ~1.0-1.5%'],
    ]
)

# ═════════════════════════════════════════════════════════════
# SECTION 8 — STANDARD CONTRACT
# ═════════════════════════════════════════════════════════════
h1('8. Đăng ký hợp đồng mẫu')

para('T&Cs = Terms & Conditions — Điều khoản & Điều kiện\n'
     'Cơ sở pháp lý: Luật 19/2023 (Luật Bảo vệ quyền lợi NTD), Nghị định 55/2024',
     italic=True, size=11, color=DARK_GREEN)

h2('8.1 Là gì?')
para('Hợp đồng mẫu (Standard-form Contract) là hợp đồng soạn sẵn mà người tiêu dùng '
     'chỉ có thể chấp nhận hoặc từ chối, không thương lượng từng điều khoản. '
     'Điều khoản mua vé trực tuyến trên sàn = hợp đồng mẫu → phải đăng ký với Bộ Công Thương.')

h2('8.2 Tại sao cần?')
bullet('Luật 19/2023 yêu cầu đăng ký hợp đồng mẫu với MOIT cho một số ngành')
bullet('Sàn TMĐT bán vé = cung cấp dịch vụ cho NTD (người tiêu dùng) qua hợp đồng mẫu')
bullet('Không đăng ký → hợp đồng có thể bị tòa tuyên vô hiệu nếu tranh chấp')

h2('8.3 Nội dung hợp đồng mẫu')
bullet('Điều kiện mua vé, hủy vé, đổi vé')
bullet('Chính sách hoàn tiền')
bullet('Giới hạn trách nhiệm của sàn vs nhà xe')
bullet('Quyền và nghĩa vụ của hành khách')
bullet('Cơ chế giải quyết tranh chấp')
bullet('Bảo vệ dữ liệu cá nhân (tham chiếu Privacy Policy)')

h2('8.4 Cách thực hiện')
bullet('Bước 1: ', 'Soạn hợp đồng mẫu (có luật sư rà soát)')
bullet('Bước 2: ', 'Nộp đăng ký tại Sở Công Thương hoặc Bộ Công Thương')
bullet('Bước 3: ', 'Chờ thẩm định (15-30 ngày)')
bullet('Bước 4: ', 'Cập nhật trên website sau khi được phê duyệt')

h2('8.5 Thời gian + Chi phí')
table(
    ['Hạng mục', 'Thời gian', 'Chi phí'],
    [
        ['Soạn hợp đồng mẫu (có luật sư)', '1-2 tuần', '5-15 triệu VND'],
        ['Thẩm định MOIT', '15-30 ngày', 'Miễn phí'],
    ]
)

# ═════════════════════════════════════════════════════════════
# SECTION 9 — TRADEMARK
# ═════════════════════════════════════════════════════════════
h1('9. Đăng ký nhãn hiệu (Trademark)')

para('Cơ quan: Cục Sở hữu trí tuệ (NOIP = National Office of Intellectual Property)\n'
     'Hệ thống: First-to-file (ai nộp trước được quyền)',
     italic=True, size=11, color=DARK_GREEN)

h2('9.1 Là gì?')
para('Đăng ký nhãn hiệu = bảo hộ tên thương hiệu + logo của sàn đặt vé. '
     'Việt Nam áp dụng hệ thống first-to-file: ai nộp đơn trước thì có quyền, '
     'bất kể ai sử dụng trước. Nếu không đăng ký → đối thủ hoặc bên thứ ba '
     'có thể đăng ký tên "Bus Booking" và kiện ngược lại.')

h2('9.2 Nhóm nhãn hiệu cần đăng ký')
table(
    ['Nhóm (Nice Classification)', 'Mô tả', 'Lý do'],
    [
        ['Nhóm 39', 'Dịch vụ vận tải', 'Hoạt động đặt vé xe khách'],
        ['Nhóm 35', 'Quản lý kinh doanh, quảng cáo', 'Hoạt động sàn TMĐT, marketing'],
        ['Nhóm 42', 'Dịch vụ CNTT, phần mềm', 'Nền tảng công nghệ, SaaS'],
    ]
)

h2('9.3 Cách thực hiện')
bullet('Bước 1: ', 'Tra cứu nhãn hiệu tại ipvietnam.gov.vn — kiểm tra trùng/tương tự')
bullet('Bước 2: ', 'Nộp đơn đăng ký tại Cục SHTT (trực tiếp hoặc qua đại diện SHTT)')
bullet('Bước 3: ', 'Thẩm định hình thức (1-2 tháng)')
bullet('Bước 4: ', 'Công bố đơn (2 tháng)')
bullet('Bước 5: ', 'Thẩm định nội dung (9-12 tháng)')
bullet('Bước 6: ', 'Cấp văn bằng bảo hộ')

h2('9.4 Thời gian + Chi phí')
table(
    ['Hạng mục', 'Thời gian', 'Chi phí'],
    [
        ['Toàn bộ quy trình', '12-18 tháng', '~1-3 triệu VND lệ phí nhà nước / nhóm'],
        ['Có đại diện SHTT / luật sư', 'Giảm rủi ro từ chối', 'Thêm ~10-20 triệu VND'],
        ['Bảo hộ', '10 năm, gia hạn không giới hạn', '—'],
    ]
)

note('⚠ NỘP SỚM — hệ thống first-to-file, ai nộp trước được quyền. '
     'Rủi ro bị "chiếm" tên thương hiệu rất cao tại Việt Nam.')

# ═════════════════════════════════════════════════════════════
# SECTION 10 — RETENTION
# ═════════════════════════════════════════════════════════════
h1('10. Lưu trữ hồ sơ (Record Retention)')

h2('10.1 Là gì?')
para('Nghĩa vụ lưu trữ hồ sơ kinh doanh theo quy định pháp luật. '
     'Mỗi loại hồ sơ có thời hạn lưu trữ tối thiểu khác nhau. '
     'Vi phạm → phạt hành chính + rủi ro pháp lý khi tranh chấp.')

h2('10.2 Bảng thời hạn lưu trữ')
table(
    ['Loại hồ sơ', 'Thời hạn lưu', 'Cơ sở pháp lý', 'Ghi chú'],
    [
        ['Hồ sơ đặt vé (booking)', '5 năm', 'Luật Kế toán', 'Bao gồm: mã booking,\nthông tin khách, tuyến, ghế'],
        ['Hồ sơ thanh toán', '10 năm', 'Nghị định 123/2020', 'Giao dịch VNPay/MoMo,\ntrạng thái, mã tham chiếu'],
        ['Hóa đơn điện tử (HĐĐT)', '10 năm', 'Nghị định 123/2020\nNghị định 70/2025',
         'Hóa đơn xuất cho hành khách\nvà hóa đơn hoa hồng'],
        ['Hồ sơ vi phạm dữ liệu', '5 năm', 'PDPL 2025\nNĐ 356/2025',
         'Báo cáo sự cố, biên bản\nđiều tra, biện pháp khắc phục'],
        ['Hồ sơ DPIA (tất cả phiên bản)', 'Vô thời hạn', 'PDPL 2025', 'Audit trail — lưu mọi\nphiên bản cập nhật'],
        ['Hồ sơ đồng ý (consent logs)', 'Suốt thời gian đồng ý\n+ 1 năm sau rút', 'PDPL 2025',
         'Timestamp, phiên bản policy,\nuser ID, mục đích'],
        ['Hợp đồng với nhà xe', '10 năm sau kết thúc HĐ', 'Bộ luật Dân sự', 'HĐ ủy quyền HĐĐT,\nHĐ hợp tác'],
    ]
)

h2('10.3 Yêu cầu kỹ thuật')
bullet('Xóa tự động: ', 'phải có cron job tự động xóa dữ liệu khi hết thời hạn lưu')
bullet('KHÔNG xóa thủ công: ', 'trừ khi có yêu cầu quyền dữ liệu từ chủ thể (DSAR)')
bullet('Backup: ', 'phải có bản sao lưu đảm bảo khôi phục được trong thời hạn lưu trữ')

# ═════════════════════════════════════════════════════════════
# SUMMARY TIMELINE
# ═════════════════════════════════════════════════════════════
doc.add_page_break()
h1('Tổng kết: Lộ trình thực hiện')

note('⚠ Giả định: Công ty đã có ERC + MST từ hoạt động vận tải bê tông. '
     'Lộ trình dưới đây chỉ tính các bước BỔ SUNG cho sàn TMĐT.')

table(
    ['Tuần', 'Hạng mục', 'Chi tiết'],
    [
        ['1-2', 'Bổ sung mã VSIC', 'Nộp thay đổi ĐKDN tại Sở KH&ĐT;\nnhận ERC cập nhật trong 3-5 ngày'],
        ['1-2', 'Bổ nhiệm DPO', 'Chọn phương án + soạn quyết định bổ nhiệm'],
        ['1-2', 'Nộp đăng ký nhãn hiệu', 'Tra cứu + nộp đơn tại Cục SHTT (12-18 tháng xử lý)'],
        ['2-4', 'Đăng ký SMS Brandname', 'Nộp cho 3 nhà mạng — CHẶN CỨNG 2-4 tuần'],
        ['2-4', 'Đăng ký Zalo OA', 'Xác minh tại oa.zalo.me (3-7 ngày)'],
        ['2-4', 'Sandbox VNPay + MoMo', 'Đăng ký sandbox, bắt đầu tích hợp kỹ thuật'],
        ['3-6', 'Soạn quy chế sàn + Privacy Policy + ToS', 'Có luật sư rà soát'],
        ['3-6', 'Đăng ký sàn TMĐT', 'Nộp tại online.gov.vn — CHẶN CỨNG 2-4 tuần'],
        ['4-8', 'Chuẩn bị DPIA + CDTIA', 'Kiểm kê dữ liệu, đánh giá rủi ro,\nthu thập DPA từ vendors'],
        ['4-8', 'Đăng ký HĐĐT + chữ ký số', 'MISA + VNPT-CA/Viettel-CA'],
        ['4-8', 'Phán quyết thuế vai trò xuất HĐĐT', 'Gửi công văn hỏi cơ quan thuế'],
        ['6-10', 'Nộp DPIA + CDTIA + DPO', 'Nộp trọn gói cho Cục A05 / BCA'],
        ['6-10', 'Production VNPay + MoMo', 'Nộp hồ sơ merchant production'],
        ['8-12', 'Đăng ký hợp đồng mẫu', 'Nộp tại Sở/Bộ Công Thương'],
        ['8-12', 'Ký HĐ ủy quyền HĐĐT', 'Ký với từng nhà xe đối tác'],
    ]
)

# ═════════════════════════════════════════════════════════════
# PENALTY SUMMARY
# ═════════════════════════════════════════════════════════════
h1('Tổng hợp mức phạt vi phạm')

table(
    ['Vi phạm', 'Mức phạt', 'Cơ sở pháp lý'],
    [
        ['Không đăng ký sàn TMĐT', '40-60 triệu VND', 'NĐ 98/2020'],
        ['Không nộp DPIA', 'Đến 3 tỷ VND hoặc 5% doanh thu', 'PDPL 2025'],
        ['Chuyển DL ra nước ngoài không phép', 'Đến 5% doanh thu', 'PDPL 2025, NĐ 356/2025'],
        ['Thông báo vi phạm DL trễ > 72 giờ', 'Đến 3 tỷ VND', 'PDPL 2025'],
        ['Không bổ nhiệm DPO', 'Phạt hành chính (chưa quy định cụ thể)', 'PDPL 2025'],
        ['Kinh doanh ngoài ngành VSIC', 'Cảnh cáo + phạt 10-15 triệu VND', 'NĐ 50/2016'],
        ['Không đăng ký hợp đồng mẫu', 'Hợp đồng có thể bị tuyên vô hiệu', 'Luật 19/2023'],
    ]
)

# ═════════════════════════════════════════════════════════════
# GLOSSARY (expanded)
# ═════════════════════════════════════════════════════════════
doc.add_page_break()
h1('Phụ lục: Bảng giải thích viết tắt đầy đủ')

table(
    ['Viết tắt', 'Tiếng Anh', 'Tiếng Việt'],
    [
        ['A05', 'Cybersecurity & Hi-tech Crime Prevention Dept', 'Cục An ninh mạng và Phòng chống\ntội phạm sử dụng công nghệ cao (thuộc BCA)'],
        ['BCA', 'Ministry of Public Security', 'Bộ Công an'],
        ['CA', 'Certificate Authority', 'Tổ chức chứng thực chữ ký số'],
        ['CCCD', '—', 'Căn cước công dân'],
        ['CDTIA', 'Cross-border Data Transfer\nImpact Assessment', 'Đánh giá tác động chuyển\ndữ liệu xuyên biên giới'],
        ['CMND', '—', 'Chứng minh nhân dân (cũ, nay là CCCD)'],
        ['CNTT', 'Information Technology (IT)', 'Công nghệ thông tin'],
        ['DPA', 'Data Processing Agreement', 'Thỏa thuận xử lý dữ liệu\n(hợp đồng giữa controller và processor)'],
        ['DPO', 'Data Protection Officer', 'Cán bộ bảo vệ dữ liệu cá nhân'],
        ['DPIA', 'Data Protection Impact Assessment', 'Đánh giá tác động xử lý\ndữ liệu cá nhân'],
        ['DSAR', 'Data Subject Access Request', 'Yêu cầu quyền truy cập\ndữ liệu cá nhân từ chủ thể'],
        ['ĐKDN', '—', 'Đăng ký doanh nghiệp'],
        ['ERC', 'Enterprise Registration Certificate', 'Giấy chứng nhận đăng ký doanh nghiệp'],
        ['GDT', 'General Department of Taxation', 'Tổng cục Thuế'],
        ['GTVT', '—', 'Giao thông Vận tải'],
        ['HĐĐT', 'E-Invoice', 'Hóa đơn điện tử'],
        ['IRC', 'Investment Registration Certificate', 'Giấy chứng nhận đăng ký đầu tư\n(cho doanh nghiệp FDI)'],
        ['KDVT', '—', 'Kinh doanh vận tải'],
        ['KH&ĐT', '—', 'Kế hoạch và Đầu tư'],
        ['KYB', 'Know Your Business', 'Xác minh doanh nghiệp đối tác'],
        ['MISA', '—', 'Nhà cung cấp phần mềm kế toán\nvà hóa đơn điện tử'],
        ['MOIT', 'Ministry of Industry and Trade', 'Bộ Công Thương'],
        ['MPS', 'Ministry of Public Security', 'Bộ Công an'],
        ['MST', 'Tax Identification Number (TIN)', 'Mã số thuế'],
        ['NĐ', '—', 'Nghị định'],
        ['NHNN', 'State Bank of Vietnam (SBV)', 'Ngân hàng Nhà nước Việt Nam'],
        ['NOIP', 'National Office of\nIntellectual Property', 'Cục Sở hữu trí tuệ'],
        ['NTD', '—', 'Người tiêu dùng'],
        ['OA', 'Official Account', 'Tài khoản chính thức (Zalo)'],
        ['OTP', 'One-Time Password', 'Mật khẩu dùng một lần'],
        ['PDPD', 'Personal Data Protection Decree', 'Nghị định Bảo vệ dữ liệu cá nhân\n(Nghị định 13/2023)'],
        ['PDPL', 'Personal Data Protection Law', 'Luật Bảo vệ dữ liệu cá nhân\n(Luật 91/2025/QH15)'],
        ['PSP', 'Payment Service Provider', 'Nhà cung cấp dịch vụ thanh toán'],
        ['SaaS', 'Software as a Service', 'Phần mềm dạng dịch vụ'],
        ['SHTT', 'Intellectual Property (IP)', 'Sở hữu trí tuệ'],
        ['SĐT', '—', 'Số điện thoại'],
        ['TLS', 'Transport Layer Security', 'Giao thức bảo mật truyền dữ liệu\n(HTTPS sử dụng TLS)'],
        ['TMĐT', 'E-Commerce', 'Thương mại điện tử'],
        ['ToS', 'Terms of Service', 'Điều khoản sử dụng'],
        ['T&Cs', 'Terms & Conditions', 'Điều khoản & Điều kiện'],
        ['VAT', 'Value Added Tax', 'Thuế giá trị gia tăng (GTGT)'],
        ['VSIC', 'Vietnam Standard Industrial\nClassification', 'Hệ thống ngành kinh tế Việt Nam'],
        ['ZNS', 'Zalo Notification Service', 'Dịch vụ thông báo Zalo'],
    ]
)

# ═════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════
out_path = os.path.join('D:', os.sep, 'Bus-Booking', 'documentation', 'business', 'HO-SO-PHAP-LY-CHI-TIET.doc')
doc.save(out_path)
print(f'Saved: {out_path}')
