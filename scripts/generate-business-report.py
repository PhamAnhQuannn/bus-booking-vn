#!/usr/bin/env python3
"""Generate comprehensive Vietnamese business documentation report as .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── MLA FORMAT: Times New Roman 12pt, double-spaced, 1-inch margins ────
from docx.shared import Emu
from docx.enum.text import WD_LINE_SPACING

# Page margins: 1 inch = 914400 EMU
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(12)

style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = FONT_SIZE
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
# MLA: first-line indent 0.5 inch for body paragraphs (applied selectively)

for i in range(1, 5):
    h = doc.styles[f'Heading {i}']
    h.font.name = FONT_NAME
    h.font.color.rgb = None
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h.paragraph_format.space_after = Pt(0)
    if i == 1:
        h.font.size = Pt(16)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif i == 2:
        h.font.size = Pt(14)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif i == 3:
        h.font.size = Pt(12)
        h.font.bold = True
        h.font.italic = True
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        h.font.size = Pt(12)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def set_cell_font(cell, font_name=FONT_NAME, font_size=Pt(11)):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size

def add_decision_box(doc, title, options, tradeoffs, choice, reason):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run(f'QUYẾT ĐỊNH: {title}')
    r.bold = True
    r.font.name = FONT_NAME
    r.font.size = FONT_SIZE

    if options:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.5)
        r2 = p2.add_run('Các phương án: ')
        r2.bold = True
        r2.font.name = FONT_NAME
        r2.font.size = FONT_SIZE
        run2 = p2.add_run(options)
        run2.font.name = FONT_NAME
        run2.font.size = FONT_SIZE

    if tradeoffs:
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Inches(0.5)
        r3 = p3.add_run('Đánh đổi: ')
        r3.bold = True
        r3.font.name = FONT_NAME
        r3.font.size = FONT_SIZE
        run3 = p3.add_run(tradeoffs)
        run3.font.name = FONT_NAME
        run3.font.size = FONT_SIZE

    p4 = doc.add_paragraph()
    p4.paragraph_format.left_indent = Inches(0.5)
    r4 = p4.add_run(f'→ Lựa chọn: {choice}')
    r4.bold = True
    r4.font.name = FONT_NAME
    r4.font.size = FONT_SIZE

    if reason:
        p5 = doc.add_paragraph()
        p5.paragraph_format.left_indent = Inches(0.5)
        r5 = p5.add_run('Lý do: ')
        r5.bold = True
        r5.font.name = FONT_NAME
        r5.font.size = FONT_SIZE
        run5 = p5.add_run(reason)
        run5.font.name = FONT_NAME
        run5.font.size = FONT_SIZE

    doc.add_paragraph()

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = FONT_NAME
                run.font.size = Pt(11)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="000000" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_font(cell, font_size=Pt(11))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            set_cell_font(cell)
    doc.add_paragraph()
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = FONT_NAME
        r.font.size = FONT_SIZE
        run2 = p.add_run(text)
        run2.font.name = FONT_NAME
        run2.font.size = FONT_SIZE
    else:
        run1 = p.add_run(text)
        run1.font.name = FONT_NAME
        run1.font.size = FONT_SIZE

def add_bold_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT_NAME
    r.font.size = FONT_SIZE
    return p

# ════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════

for _ in range(8):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BÁO CÁO TỔNG HỢP\nTÀI LIỆU KINH DOANH')
r.bold = True
r.font.name = FONT_NAME
r.font.size = Pt(24)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Nền Tảng Đặt Vé Xe Khách BB')
r2.font.name = FONT_NAME
r2.font.size = Pt(18)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('"Shopify cho Nhà Xe" — Hệ thống đặt vé xe khách trực tuyến\nthị trường Việt Nam')
r3.font.name = FONT_NAME
r3.font.size = FONT_SIZE
r3.italic = True

for _ in range(6):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_info = info.add_run('Phiên bản: 1.0  |  Ngày: 17/06/2026\nNguồn: 40 tài liệu nghiên cứu kinh doanh\nBảo mật: NỘI BỘ')
r_info.font.name = FONT_NAME
r_info.font.size = FONT_SIZE

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  MỤC LỤC
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('MỤC LỤC', level=1)

toc_items = [
    'PHẦN 0: GIỚI THIỆU & BẢNG THUẬT NGỮ VIẾT TẮT',
    'PHẦN 1: THỊ TRƯỜNG VIỆT NAM & CHIẾN LƯỢC GIA NHẬP',
    '    1.1 Quy mô thị trường',
    '    1.2 Top 5 hành động ưu tiên trong 90 ngày',
    '    1.3 Điều nên DỪNG và nên BẮT ĐẦU',
    '    1.4 Chuỗi gia nhập thị trường 4 giai đoạn',
    '    1.5 Điều KHÔNG LÀM trong Năm 1',
    'PHẦN 2: CÁC BÊN LIÊN QUAN (Stakeholder Map)',
    '    2.1 Bên liên quan chính',
    '    2.2 Bên liên quan thứ cấp',
    '    2.3 Bên liên quan bậc ba',
    '    2.4 Top 5 bên quyết định sống còn',
    '    2.5 Bên liên quan đặc thù Việt Nam',
    'PHẦN 3: MA TRẬN RỦI RO',
    '    3.1 Top 15 rủi ro xếp hạng',
    '    3.2 Khoảng trống tính năng & tính năng tương lai',
    'PHẦN 4: CHÂN DUNG NGƯỜI DÙNG (Personas)',
    '    4.1 Nhà xe — 5 phân khúc',
    '    4.2 Hành khách — 6 phân khúc',
    '    4.3 Quản trị viên — 4 vai trò',
    '    4.4 KPI cho Nhà đầu tư',
    'PHẦN 5: ĐỐI THỦ CẠNH TRANH',
    '    5.1 Bản đồ cạnh tranh (Tier 1/2/3)',
    '    5.2 Ma trận so sánh tính năng',
    '    5.3 So sánh hoa hồng & thanh toán',
    '    5.4 Tâm lý nhà xe & nguyên nhân rời bỏ',
    '    5.5 Benchmark CAC/LTV',
    '    5.6 Phủ sóng địa lý',
    '    5.7 Lợi thế cạnh tranh BB',
    'PHẦN 6: MÔ HÌNH KINH DOANH',
    '    6.1 Cơ cấu hoa hồng',
    '    6.2 Kinh tế đơn vị',
    '    6.3 Mô hình giá kép',
    '    6.4 Nguồn doanh thu bổ sung',
    '    6.5 Playbook thu hút nhà xe',
    '    6.6 Kênh thu hút khách hàng',
    'PHẦN 7: TÍNH NĂNG & KHOẢNG TRỐNG SO VỚI ĐỐI THỦ',
    '    7.1 Tính năng bắt buộc (Table Stakes)',
    '    7.2 Khoảng trống ưu tiên P0 — Trước ra mắt',
    '    7.3 Khoảng trống ưu tiên P1',
    '    7.4 Khoảng trống ưu tiên P2-P3',
    'PHẦN 8: PHÁP LÝ & TUÂN THỦ QUY ĐỊNH',
    '    8.1 Tổng quan pháp lý — 35 phát hiện, 5 xung đột liên miền',
    '    8.2 Pháp nhân & Đăng ký thương mại điện tử',
    '    8.3 Thanh toán & Giấy phép trung gian thanh toán',
    '    8.4 Hóa đơn điện tử & Thuế',
    '    8.5 Bảo vệ dữ liệu cá nhân & Lưu trú dữ liệu',
    '    8.6 Viễn thông / SMS / OTP',
    '    8.7 Vận tải — Phân loại nền tảng',
    '    8.8 Bảo vệ người tiêu dùng',
    '    8.9 Điều khoản hợp đồng PSP',
    '    8.10 DPIA Checklist',
    '    8.11 Lao động, AML, SHTT, Bảo hiểm, Tiếp cận',
    '    8.12 Lộ trình tuân thủ 12 tuần trước ra mắt',
    'PHẦN 9: MÔ HÌNH MIỀN (Domain Model)',
    '    9.1 Bản đồ ngữ cảnh giới hạn (12 Bounded Contexts)',
    '    9.2 Ngôn ngữ thống nhất — Bảng thuật ngữ miền',
    '    9.3 Luồng sự kiện chính (6 luồng)',
    '    9.4 Bất biến kinh doanh (14 bất biến)',
    '    9.5 Máy trạng thái (8 máy trạng thái)',
    'PHẦN 10: LỘ TRÌNH CHIẾN LƯỢC & SỔ ĐĂNG KÝ RỦI RO',
    '    10.1 Lộ trình hành động theo giai đoạn',
    '    10.2 Sổ đăng ký rủi ro (10 rủi ro)',
    '    10.3 Thông tin thị trường người dùng (10 insight)',
    'PHẦN 11: TỔNG HỢP QUYẾT ĐỊNH',
]

for item in toc_items:
    p = doc.add_paragraph()
    if item.startswith('    '):
        p.paragraph_format.left_indent = Cm(1.5)
        p.add_run(item.strip()).font.size = Pt(10)
    else:
        r = p.add_run(item)
        r.bold = True
        r.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 0: GIỚI THIỆU & BẢNG THUẬT NGỮ
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 0: GIỚI THIỆU & BẢNG THUẬT NGỮ VIẾT TẮT', level=1)

doc.add_heading('Tổng quan dự án', level=2)
doc.add_paragraph(
    'BB (Bus Booking) là nền tảng đặt vé xe khách trực tuyến được xây dựng riêng cho thị trường Việt Nam. '
    'Định vị là "Shopify cho Nhà Xe" — nhà xe sở hữu kênh bán hàng số của riêng mình thay vì trở thành '
    'nhà cung cấp trên marketplace của bên thứ ba (như VeXeRe). Nền tảng bao gồm: console quản lý cho nhà xe '
    '(quản lý đội xe, tuyến đường, chuyến, nhân viên, doanh thu), giao diện đặt vé cho hành khách, và bảng '
    'điều khiển quản trị nội bộ.'
)
doc.add_paragraph(
    'Báo cáo này tổng hợp toàn bộ 40 tài liệu nghiên cứu kinh doanh trong thư mục documentation/business/, '
    'bao gồm 7 nhóm: chiến lược thị trường, bên liên quan, rủi ro, chân dung người dùng, đối thủ cạnh tranh, '
    'mô hình miền, và pháp lý tuân thủ. Mỗi quyết định chiến lược được trình bày với các phương án, đánh đổi, '
    'lựa chọn và lý do.'
)

doc.add_heading('Bảng thuật ngữ viết tắt', level=2)

acronyms = [
    ('BB', 'Bus Booking', 'Nền tảng đặt vé xe khách — tên dự án'),
    ('AML', 'Anti-Money Laundering', 'Chống rửa tiền'),
    ('AMS', 'Agent Management System', 'Hệ thống quản lý đại lý'),
    ('ARR', 'Annual Recurring Revenue', 'Doanh thu định kỳ hàng năm'),
    ('BMS', 'Bus Management System', 'Hệ thống quản lý nhà xe'),
    ('CAC', 'Customer Acquisition Cost', 'Chi phí thu hút khách hàng'),
    ('CDTIA', 'Cross-border Data Transfer Impact Assessment', 'Đánh giá tác động chuyển dữ liệu xuyên biên giới'),
    ('CIT', 'Corporate Income Tax', 'Thuế thu nhập doanh nghiệp (TNDN)'),
    ('CPL', 'Consumer Protection Law', 'Luật Bảo vệ Quyền lợi Người tiêu dùng'),
    ('DPO', 'Data Protection Officer', 'Nhân viên bảo vệ dữ liệu'),
    ('DPIA', 'Data Protection Impact Assessment', 'Đánh giá tác động bảo vệ dữ liệu cá nhân'),
    ('ERC', 'Enterprise Registration Certificate', 'Giấy chứng nhận đăng ký doanh nghiệp'),
    ('FCT', 'Foreign Contractor Tax', 'Thuế nhà thầu nước ngoài'),
    ('GDT', 'General Department of Taxation', 'Tổng cục Thuế'),
    ('GMV', 'Gross Merchandise Value', 'Tổng giá trị giao dịch hàng hóa'),
    ('IPS', 'Intermediary Payment Services', 'Dịch vụ trung gian thanh toán'),
    ('IRC', 'Investment Registration Certificate', 'Giấy chứng nhận đăng ký đầu tư'),
    ('JWT', 'JSON Web Token', 'Token xác thực dạng JSON'),
    ('KYB', 'Know Your Business', 'Xác minh thông tin doanh nghiệp'),
    ('LTV', 'Lifetime Value', 'Giá trị trọn đời khách hàng'),
    ('MDR', 'Merchant Discount Rate', 'Phí chiết khấu thương mại (phí cổng thanh toán)'),
    ('MISA', 'MISA meInvoice', 'Nhà cung cấp hóa đơn điện tử được GDT chứng nhận'),
    ('MOIT', 'Ministry of Industry and Trade', 'Bộ Công Thương'),
    ('NAPAS', 'National Payment Corporation of VN', 'Công ty Cổ phần Thanh toán Quốc gia Việt Nam'),
    ('NPS', 'Net Promoter Score', 'Chỉ số đo lường sự hài lòng khách hàng'),
    ('OTA', 'Online Travel Agency', 'Đại lý du lịch trực tuyến'),
    ('OTP', 'One-Time Password', 'Mật khẩu dùng một lần'),
    ('PDPL', 'Personal Data Protection Law', 'Luật Bảo vệ Dữ liệu Cá nhân'),
    ('PIT', 'Personal Income Tax', 'Thuế thu nhập cá nhân (TNCN)'),
    ('PSP', 'Payment Service Provider', 'Nhà cung cấp dịch vụ thanh toán'),
    ('RBAC', 'Role-Based Access Control', 'Kiểm soát truy cập dựa trên vai trò'),
    ('SBV', 'State Bank of Vietnam', 'Ngân hàng Nhà nước Việt Nam'),
    ('TOTP', 'Time-based One-Time Password', 'Mật khẩu dùng một lần theo thời gian'),
    ('TTL', 'Time to Live', 'Thời gian tồn tại'),
    ('VAT', 'Value Added Tax', 'Thuế giá trị gia tăng (GTGT)'),
    ('VND', 'Vietnamese Dong', 'Đồng Việt Nam'),
    ('VSIC', 'Vietnam Standard Industrial Classification', 'Hệ thống ngành kinh tế Việt Nam'),
    ('WCAG', 'Web Content Accessibility Guidelines', 'Hướng dẫn truy cập nội dung web'),
    ('ZNS', 'Zalo Notification Service', 'Dịch vụ thông báo Zalo'),
]

add_table(doc, ['Viết tắt', 'Tiếng Anh', 'Giải thích tiếng Việt'],
          [(a, b, c) for a, b, c in acronyms])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 1: THỊ TRƯỜNG VIỆT NAM
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 1: THỊ TRƯỜNG VIỆT NAM & CHIẾN LƯỢC GIA NHẬP', level=1)

doc.add_heading('1.1 Quy mô thị trường', level=2)
add_table(doc, ['Chỉ số', 'Giá trị', 'Nguồn'],
[
    ('Thị trường du lịch trực tuyến VN (2024)', 'USD 2.8-3.0 tỷ', 'IMARC Group'),
    ('Dự kiến 2030-2034', 'USD 6.4-8.0 tỷ', 'IMARC / Google-Temasek'),
    ('CAGR', '~8.5%', 'IMARC'),
    ('Doanh thu vận tải công cộng (2024)', 'USD 0.65 tỷ', 'Statista'),
    ('Thâm nhập trực tuyến (2024)', '63.2%', 'Statista'),
    ('Tỷ lệ đặt vé qua di động', '65-73%', 'IMARC'),
    ('Khách du lịch quốc tế (2024)', '17.5 triệu (+39.5% YoY)', 'Chính phủ VN'),
    ('Ước tính thị trường vé xe khách online', 'USD 140-500 triệu', 'Ước tính nội bộ'),
])

doc.add_paragraph('Động lực tăng trưởng chính:')
add_bullet(doc, 'Tỷ lệ sở hữu smartphone 84.4%, tốc độ download trung bình 50.88 Mbps')
add_bullet(doc, 'Ví điện tử: 57% người trưởng thành sử dụng (từ 14% năm 2018), MoMo 31 triệu user')
add_bullet(doc, 'Du lịch nội địa phục hồi mạnh; 17.5 triệu khách quốc tế 2024')
add_bullet(doc, 'Bắt buộc hóa đơn điện tử từ tháng 7/2022 thúc đẩy số hóa')
add_bullet(doc, 'Nghị định 158/2024 yêu cầu phần mềm đặt vé hiển thị thông tin nhà xe, gửi HĐĐT')

doc.add_heading('1.2 Top 5 hành động ưu tiên trong 90 ngày', level=2)

doc.add_heading('1.2.1 Giải quyết rào cản giấy phép trung gian thanh toán (Tuần 1-2)', level=3)
doc.add_paragraph(
    'Mô hình thu hộ tập trung (central-collection) là bất hợp pháp nếu không có giấy phép SBV. '
    'Đường nhanh nhất: tái cấu trúc sang mô hình split-settlement — mỗi nhà xe mở tài khoản merchant VNPay/MoMo riêng, '
    'thanh toán tách tại nguồn. Phí 6% của nền tảng chảy trực tiếp vào tài khoản nền tảng. '
    'Loại bỏ việc giữ tiền nhà xe, yêu cầu giấy phép SBV, và nghĩa vụ thanh toán T+1.'
)

doc.add_heading('1.2.2 Ký 3 LOI trên một hành lang (Tuần 1-4)', level=3)
doc.add_paragraph(
    'Kill-switch 30 ngày. Chọn MỘT hành lang: Thanh Hóa ↔ TPHCM (giá vé trung bình cao 875K-1,750K VND, '
    'nhu cầu lao động di cư lớn, đỉnh Tết mạnh nhất cả nước). Tiếp cận 3-5 nhà xe với: phí 3% giới thiệu, '
    'onboarding white-glove, thanh toán T+1. Nếu không có 3 LOI trong 30 ngày → product-market fit chưa đạt, pivot.'
)

doc.add_heading('1.2.3 Đóng 5 blocker go-live (Tuần 3-8)', level=3)
add_bullet(doc, 'Issue 094: Khóa thanh toán thực (cần quyết định #1)')
add_bullet(doc, 'Triển khai hoàn tiền MoMo + VNPay API (hiện throw PspRefundNotImplementedError)')
add_bullet(doc, 'Issue 118: Giám sát bên ngoài (phát hiện downtime <2 phút)')
add_bullet(doc, 'Issue 095: Sweeper đối soát thanh toán VietQR')
add_bullet(doc, 'Issue 101: Cổng bảo mật/gian lận (rate-limit fail-closed, hold cap)')
add_bullet(doc, 'Hỗ trợ khách hàng tối thiểu (Liên hệ Hỗ trợ → Zalo OA)')

doc.add_heading('1.2.4 Hoàn thành hồ sơ pháp lý (Tuần 4-12)', level=3)
add_bullet(doc, 'Bổ sung giấy phép kinh doanh VSIC 2025')
add_bullet(doc, 'Đăng ký sàn TMĐT tại online.gov.vn')
add_bullet(doc, 'Nộp hồ sơ DPO + DPIA + CDTIA')
add_bullet(doc, 'Phán quyết thuế về vai trò xuất HĐĐT')
add_bullet(doc, 'Đăng ký SMS Brandname (blocker cứng 2-4 tuần)')

doc.add_heading('1.2.5 Đàm phán MOU với 3 bến xe (Tuần 6-12)', level=3)
doc.add_paragraph(
    'Đảm bảo bến xe chấp nhận mã QR. Cung cấp: cổng manifest chỉ-đọc, logo bến xe trên trang chi tiết chuyến, '
    'phí giới thiệu nhỏ nếu cần.'
)

doc.add_heading('1.3 Điều nên DỪNG và nên BẮT ĐẦU', level=2)

doc.add_heading('Nên DỪNG', level=3)
add_table(doc, ['Dừng', 'Lý do'],
[
    ('Xây dựng cho quy mô multi-operator trước khi chứng minh single-corridor', '1 nhà xe, ~200 bookings/ngày. RBAC admin, dashboard phân tích, bulk API là chưa cần'),
    ('Coi pairedReturn là production-ready', 'S15 đã phê duyệt xóa. Mistake Log Issue 013 ghi nhận xung đột spec'),
    ('Trì hoãn quyết định mô hình thanh toán', 'PAYMENTS_STUB bật từ đầu. Mọi tính năng xây trên central-collection có thể cần tái cấu trúc'),
    ('Mở rộng charter trước khi scheduled booking live', 'Charter = zero revenue (lead-gen). Một hành lang hoạt động với thanh toán thực > state machine charter hoàn hảo'),
])

doc.add_heading('Nên BẮT ĐẦU', level=3)
add_table(doc, ['Bắt đầu', 'Lý do'],
[
    ('Gặp Bộ GTVT trước 10 nhà xe', 'Chủ động thiết lập tính hợp pháp. Đợi đến khi bị phát hiện = đối đầu'),
    ('Đo tỷ lệ gửi SMS theo nhà mạng', 'eSMS duy nhất. Vietnamobile/Gmobile tỷ lệ thấp hơn. Lỗi SMS = khách bị khóa hoặc không có vé'),
    ('Luồng "Liên hệ Hỗ trợ" tối thiểu', 'Luật BVQLNTD 59/2023 yêu cầu kênh khiếu nại. Email + Zalo OA link đáp ứng yêu cầu pháp lý'),
    ('Tình báo cạnh tranh VeXeRe hàng tháng', 'VeXeRe BMS khóa ~300/550 nhà xe. ~250 không khóa + ~1,500 chưa trên VeXeRe = nguồn cung khai thác'),
])

doc.add_heading('1.4 Chuỗi gia nhập thị trường 4 giai đoạn', level=2)

add_table(doc, ['Giai đoạn', 'Thời gian', 'Hành lang', 'Mục tiêu', 'Chỉ số thành công'],
[
    ('Phase 1: Chứng minh single-corridor', 'Tháng 1-3', 'Thanh Hóa ↔ TPHCM', '1-3 nhà xe, 20+ bookings/ngày', '1-3 nhà xe duy trì >10 bookings/ngày trong 30 ngày'),
    ('Phase 2: Mở rộng hành lang', 'Tháng 4-6', '+ Nghệ An, Hà Tĩnh ↔ TPHCM\n+ HCMC-Đà Lạt, Nha Trang', '10-15 nhà xe, 200+ bookings/ngày', '>15% booking từ session tiếng Anh trên tuyến du lịch'),
    ('Phase 3: Hub phía Bắc + Scale', 'Tháng 7-12', '+ Hà Nội-Sa Pa, Hạ Long, Ninh Bình', '30-50 nhà xe, 500+ bookings/ngày', 'Vượt Tết với <1% tỷ lệ lỗi booking'),
    ('Phase 4: Hiệu ứng nền tảng', 'Năm 2', 'Zalo, OTA white-label, B2B shuttle', '100+ nhà xe, Series A metrics', 'Metrics Series A đạt'),
])

add_decision_box(doc,
    'Hành lang beachhead',
    'Thanh Hóa ↔ TPHCM (lao động di cư) vs HCMC-Đà Lạt (du lịch) vs Hà Nội-Sa Pa (du lịch)',
    'Thanh Hóa ↔ TPHCM: giá vé cao (875K-1.75M), nhu cầu ổn định, cạnh tranh OTA thấp hơn — NHƯNG không có phân khúc du lịch, '
    'cần hoãn English UI; tuyến đêm = vận hành phức tạp hơn. Du lịch: AOV thấp hơn nhưng có khách quốc tế.',
    'Thanh Hóa ↔ TPHCM',
    'GMV cao nhất/booking, nhu cầu lao động di cư ổn định quanh năm + đỉnh Tết mạnh nhất cả nước, '
    'cạnh tranh OTA thấp (VeXeRe/redBus tập trung tuyến du lịch), mở rộng tự nhiên sang Nghệ An/Hà Tĩnh ↔ TPHCM.'
)

doc.add_heading('1.5 Điều KHÔNG LÀM trong Năm 1', level=2)
add_bullet(doc, 'Tuyến tỉnh Đồng bằng Sông Cửu Long — cash-centric, smartphone thấp')
add_bullet(doc, 'Xuyên biên giới VN-Campuchia — pháp lý phức tạp, redBus/12Go thống trị')
add_bullet(doc, 'Sơ đồ ghế — mô hình count-based phù hợp thực tế nhà xe')
add_bullet(doc, 'GPS tracking thời gian thực — cần phần cứng nhà xe không có')
add_bullet(doc, 'Cạnh tranh FutaBus trên TPHCM-Cần Thơ — tích hợp dọc, luôn thắng về giá/độ tin cậy trên tuyến riêng')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 2: STAKEHOLDER MAP
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 2: CÁC BÊN LIÊN QUAN (Stakeholder Map)', level=1)

doc.add_heading('2.1 Bên liên quan chính — Người dùng trực tiếp', level=2)

add_table(doc, ['Bên liên quan', 'Quan tâm', 'Ảnh hưởng', 'Nghĩa vụ nền tảng', 'Rủi ro nếu bỏ qua'],
[
    ('Hành khách nội địa\n(ngân sách/SV/gia đình)', 'Giá rẻ nhất, xác nhận chỗ tin cậy, vé QR trên điện thoại, không phí ẩn', 'CAO (tập thể)', 'Sẵn sàng real-time, hold atomic, QR trong 60s, UI tiếng Việt, chạy trên Android giá rẻ/4G', 'Mất so sánh giá cho VeXeRe; double-sell → sụp đổ niềm tin; lỗi Tết = mất vĩnh viễn'),
    ('Doanh nhân', 'HĐĐT MISA cho khai chi phí, xe cao cấp, đặt phút chót', 'TRUNG BÌNH-CAO', 'HĐĐT khi thanh toán (Thông tư 78/2021), phân tách VAT trên biên lai', 'Không HĐĐT = không khai chi phí = mất segment lặp lại doanh nghiệp'),
    ('Nhà xe lớn\n(FUTA-scale)', 'Lượng lớn, phí % thấp, manifest real-time, T+1 payout, kiểm soát thương hiệu', 'RẤT CAO\n(1 nhà xe lớn rời = mất 30-50% supply)', 'Phí đàm phán (FeeConfig effective-dated), T+1, nhập trip bulk, KYB nhanh (48h)', 'Phí quá cao = rút kho, trễ payout = vỡ niềm tin'),
    ('Nhà xe nhỏ/gia đình', 'Tiếp cận khách rộng hơn, setup đơn giản, thông báo Zalo, chi phí cố định thấp', 'THẤP đơn lẻ\nTRUNG BÌNH tập thể', 'Console tiếng Việt, SMS/Zalo thông báo booking, cổng xác minh giấy phép, onboarding hỗ trợ', 'Onboarding phức tạp = bỏ cuộc = thiếu supply tuyến phụ'),
    ('Admin/Ops nền tảng', 'Công cụ duyệt nhà xe, điều tra tranh chấp, xử lý hoàn tiền, quản lý payout batch', 'CAO (trung tâm thần kinh vận hành)', 'Console admin RBAC, audit log bất biến, hoàn tiền 1-click, dashboard payout', 'Lỗi công cụ payout = churn nhà xe hàng loạt'),
    ('Tài xế / Nhân viên vé', 'Scanner QR trên điện thoại cá nhân, manifest rõ ràng, không cần đào tạo', 'THẤP chiến lược\nCAO thực thi', 'Scanner QR hoạt động offline/3G, manifest download PDF, UI đơn nhiệm', 'Tài xế từ chối vé số → khách khó chịu → thiệt hại thương hiệu'),
])

doc.add_heading('2.2 Bên liên quan thứ cấp — Hạ tầng & Quản lý nhà nước', level=2)

add_table(doc, ['Bên liên quan', 'Ảnh hưởng', 'Nghĩa vụ nền tảng', 'Rủi ro nếu bỏ qua'],
[
    ('NAPAS / VietQR', 'HẠ TẦNG QUAN TRỌNG', 'Sweeper đối soát thanh toán, webhook idempotent, fallback khi VietQR unavailable', 'VietQR webhook không đối soát = nhận tiền nhưng không có vé = trải nghiệm tệ nhất'),
    ('MoMo Wallet', 'CAO (31M+ user)', 'FAILURE_RESULT_CODES đúng spec, HMAC-verified IPN, logo MoMo theo co-brand', 'IPN failure codes sai map = thanh toán hợp lệ bị đánh failed = mất doanh thu'),
    ('VNPay Gateway', 'CAO', 'Return URL + IPN dual confirmation, đối soát settlement', 'Return URL sai = VNPay xác nhận nhưng booking không cập nhật = charge đôi khi retry'),
    ('eSMS', 'TRUNG BÌNH', 'SMS delivery status callback, fallback email nếu SMS fail sau 60s', 'eSMS downtime Tết = giao vé fail ở đỉnh = khiếu nại hàng loạt'),
    ('MISA e-Invoice', 'TRUNG BÌNH', 'Invoice khi xác nhận thanh toán (async), VAT breakdown, retry queue', 'HĐĐT không kịp = vi phạm GDT = phạt'),
    ('Bộ GTVT', 'RẤT CAO — có thể yêu cầu đóng nền tảng', 'Hard license gate trong KYB, xác minh lại hàng năm', 'Nhà xe không giấy phép trên nền tảng = cưỡng chế của Bộ GTVT'),
    ('Tổng cục Thuế', 'CAO — có thể đóng băng tài khoản', 'Đăng ký VAT, MISA tích hợp 100% booking thanh toán', 'Thiếu HĐĐT = trigger kiểm toán'),
    ('Bộ Công an / A05 (PDPL)', 'CAO — tăng cường thực thi', 'DPA, lawful basis, cơ chế chuyển dữ liệu xuyên biên giới, redaction PII trong log, runbook breach notification', 'Vi phạm dữ liệu không thông báo 72h = phạt tối đa'),
])

doc.add_heading('2.3 Bên liên quan bậc ba — Hệ sinh thái', level=2)

add_table(doc, ['Bên liên quan', 'Ảnh hưởng', 'Thái độ', 'Rủi ro'],
[
    ('VeXeRe (đối thủ chính)', 'CAO (tham chiếu phí/tính năng cho nhà xe)', 'Thù địch', 'Nhà xe multi-home rồi rời bỏ; FUD về tuân thủ'),
    ('FutaBus/Phương Trang', 'RẤT CAO (chiếm tuyến liên tỉnh lớn)', 'Kháng cự listing bên thứ ba', 'FUTA từ chối = thiếu supply trên tuyến nam trung bộ'),
    ('redBus Vietnam (MakeMyTrip)', 'TRUNG BÌNH-CAO (sâu vốn)', 'Cạnh tranh', 'Khuyến mãi 40%-off liên tục → lôi kéo khách nhạy giá'),
    ('Quản lý bến xe', 'TRUNG BÌNH', 'Dè dặt (phòng vé mất hoa hồng)', 'Bến xe từ chối QR = khách bị quay lại trên tuyến lớn'),
    ('Khách du lịch nước ngoài', 'THẤP đơn lẻ, CAO danh tiếng', 'Trung lập', 'Review tiếng Anh tiêu cực lan ra cộng đồng du lịch quốc tế'),
    ('Nhà đầu tư', 'CAO chiến lược', 'Hỗ trợ', 'Vi phạm pháp lý trong due diligence = term sheet bị rút'),
])

doc.add_heading('2.4 Top 5 bên quyết định sống còn', level=2)
doc.add_paragraph('1. Nhà xe lớn — Không có kho hàng → không có gì để bán. Một nhà xe lớn rời = mất 30-50% supply.')
doc.add_paragraph('2. Bộ GTVT — Có thể đóng nền tảng hoàn toàn. KYB license gate vừa là yêu cầu pháp lý vừa là lợi thế cạnh tranh.')
doc.add_paragraph('3. Hành khách nội địa (tập thể) — Động cơ nhu cầu. Một lỗi Tết = mất vĩnh viễn. Không có equity thương hiệu nào để đệm.')
doc.add_paragraph('4. NAPAS/VietQR + MoMo — Hạ tầng thanh toán. VietQR webhook fail im lặng = mất cả tiền và niềm tin.')
doc.add_paragraph('5. CTO / Tech Lead — Người duy nhất có thể thực thi Mistake Log rules, duy trì ledger immutability, đảm bảo chất lượng code.')

doc.add_heading('2.5 Bên liên quan đặc thù Việt Nam mà đội ngũ phương Tây thường bỏ sót', level=2)
add_bullet(doc, 'Quản lý bến xe (Bến xe): ', 'Thực thể độc lập thu phí bến bãi. Quyền phủ quyết nền tảng nào được hoạt động trong bến. Có thể từ chối mã QR.')
add_bullet(doc, 'Sở GTVT tỉnh (63 tỉnh): ', 'Phê duyệt tuyến là cấp tỉnh, không cấp quốc gia. Tuyến qua 3 tỉnh cần phê duyệt cả 3. Không có registry tập trung.')
add_bullet(doc, 'Zalo: ', 'Nền tảng nhắn tin thống trị VN (100M+ user). Nhà xe nhỏ dùng Zalo, không dùng email. VeXeRe đã tích hợp booking trong Zalo.')
add_bullet(doc, 'Nhà xe không giấy phép: ', '~20-30% chuyến liên tỉnh hoạt động không chính thức. Cho phép = rủi ro pháp lý; từ chối = thu hẹp nguồn cung.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 3: MA TRẬN RỦI RO
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 3: MA TRẬN RỦI RO', level=1)

doc.add_heading('3.1 Top 15 rủi ro xếp hạng theo Khả năng × Tác động', level=2)

add_table(doc, ['#', 'Rủi ro', 'Khả năng', 'Tác động', 'Mức độ', 'Giảm thiểu'],
[
    ('1', '[VN] Vận hành mô hình thanh toán tập trung không có giấy phép SBV (NĐ 52/2024)', 'CHẮC CHẮN', 'NGUY HIỂM', 'NGUY HIỂM', 'Tái cấu trúc sang split-settlement. Quyết định trước mọi khóa thanh toán thực.'),
    ('2', '[VN] Thiếu đăng ký sàn TMĐT (NĐ 52/2013 + 85/2021)', 'CAO', 'CAO', 'CAO', 'Bổ sung VSIC 2025, đăng ký online.gov.vn. Cần rules of operation, ToS, privacy policy.'),
    ('3', 'PSP refund chưa triển khai', 'CHẮC CHẮN', 'CAO', 'CAO', 'Triển khai MoMo + VNPay refund API trước go-live.'),
    ('4', 'Đỉnh Tết (10-20x volume) phá hệ thống', 'CAO', 'CAO', 'CAO', 'Load test 2,000+ concurrent booking; PgBouncer prerequisite; pre-provision Redis.'),
    ('5', '[VN] Vi phạm PDPL 2023 (chuyển dữ liệu xuyên biên giới sang SG)', 'TRUNG BÌNH', 'CAO', 'CAO', 'Hoàn thành DPIA; bổ sung data residency trong Privacy Policy; tạo breach runbook 72h.'),
    ('6', 'Không có kênh hỗ trợ khách hàng', 'CHẮC CHẮN', 'CAO', 'CAO', 'Minimum viable: email trên vé + Contact Support link → Zalo OA.'),
    ('7', 'Con gà-quả trứng: không nhà xe = không chuyến = không khách', 'CAO', 'CAO', 'CAO', 'Kill-switch 30 ngày. Bắt đầu hyper-local. Seed 10 nhà xe đầu tiên với white-glove onboarding.'),
    ('8', 'VietQR đối soát thanh toán fail (memo truncation, user nhập sai)', 'CAO', 'CAO', 'CAO', 'orderRef dưới 25 ký tự; xây dashboard đối soát admin; recon sweeper.'),
    ('9', '[VN] Giấy phép vận tải nhà xe không xác minh trong KYB', 'TRUNG BÌNH', 'CAO', 'CAO', 'Thêm Giấy phép kinh doanh vận tải là KYB bắt buộc; field hạn giấy phép; cron cảnh báo 60 ngày.'),
    ('10', 'Nhà xe rời bỏ khi còn booking tương lai', 'TRUNG BÌNH', 'CAO', 'CAO', 'Giám sát tốc độ booking; outreach khi booking = 0 trong 14 ngày; chặn deactivate nếu còn paid booking.'),
    ('11', '[VN] MISA e-invoice: ai là người bán VAT (nền tảng hay nhà xe)?', 'TRUNG BÌNH', 'CAO', 'CAO', 'Xin phán quyết thuế từ GDT về platform-as-agent vs platform-as-principal.'),
    ('12', 'Thiếu giám sát bên ngoài', 'CHẮC CHẮN', 'TRUNG BÌNH', 'TB-CAO', 'Triển khai external uptime monitoring trước go-live. 2 phút phát hiện.'),
    ('13', 'Redis rate-limit fail open khi UPSTASH_REDIS_REST_URL không set', 'TRUNG BÌNH', 'TRUNG BÌNH', 'TB', 'Rate-limiter fail-closed hoặc Redis URL là yêu cầu khởi động cứng.'),
    ('14', '[VN] Nhà xe phản đối phí (6% cắt margin vs bán tiền mặt ở bến)', 'CAO', 'TRUNG BÌNH', 'TB', 'Phí pilot 3-4%; lượng hóa value prop (manifest số tiết kiệm 1-2h/chuyến).'),
    ('15', '[VN] Quản lý bến xe từ chối QR nền tảng', 'TRUNG BÌNH', 'TRUNG BÌNH', 'TB', 'Đàm phán MOU với top 5 bến xe trước Tết; cung cấp cổng manifest chỉ-đọc.'),
])

doc.add_heading('3.2 Khoảng trống tính năng & Tính năng tương lai', level=2)

add_table(doc, ['#', 'Khoảng trống', 'Khẩn cấp', 'Ai cần', 'Khuyến nghị'],
[
    ('1', 'Kênh hỗ trợ / khiếu nại', 'XÂY NGAY', 'Tất cả khách, Luật BVQLNTD 59/2023', 'Phải ship trước go-live'),
    ('2', 'UI tiếng Anh (i18n)', '6 THÁNG', 'Khách du lịch nước ngoài (17.5M/năm)', 'Bắt khách tuyến du lịch'),
    ('3', 'Thông báo trễ chuyến', '6 THÁNG', 'Tất cả hành khách', 'Chi phí thấp, tác động CX cao'),
    ('4', 'Chính sách hành lý/dịch vụ bổ sung', '6 THÁNG', 'Gia đình, doanh nhân', 'Giảm phí bất ngờ'),
    ('5', 'Đánh giá & nhận xét', 'HOÃN (50+ nhà xe)', 'Tất cả khách', 'Gate trên completed booking'),
    ('6', 'Giá trẻ em/người cao tuổi', 'HOÃN (v2)', 'Gia đình', 'Count-based OK cho v1'),
    ('7', 'Chọn ghế / sơ đồ ghế', 'HOÃN (REMODEL)', 'Khách limousine/sleeper', 'Đã phê duyệt "không trong phiên bản này"'),
    ('8', 'Chuyến nhiều chặng', 'HOÃN (v2)', 'Khách tuyến du lịch', 'Phức tạp cao, khẩn cấp thấp'),
    ('9', 'Giá động/surge pricing', 'HOÃN (v2)', 'Nhà xe (Tết), nền tảng (GMV)', 'v2+'),
    ('10', 'GPS tracking real-time', 'KHÔNG (v1-v3)', 'Khách tech-savvy', 'Không khả thi khi nhà xe không có HW GPS'),
    ('11', 'Bundle khứ hồi', 'HOÃN (v3)', 'Khách thường xuyên', 'pairedReturn đã đánh dấu xóa'),
    ('12', 'Bảo hiểm hủy chuyến', 'KHÔNG', 'Khách tránh rủi ro', 'Cần đối tác bảo hiểm có giấy phép'),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 4: PERSONAS
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 4: CHÂN DUNG NGƯỜI DÙNG (Personas)', level=1)

doc.add_heading('4.1 Nhà xe — 5 phân khúc', level=2)

doc.add_heading('4.1.1 "Bác Tâm" — Nhà xe micro (1-5 xe)', level=3)
add_table(doc, ['Thuộc tính', 'Chi tiết'],
[
    ('Quy mô đội xe', '1-5 xe'),
    ('Loại tuyến', 'Cố định nông thôn/tỉnh'),
    ('Tỷ lệ thị trường', '~60-70% nhà xe theo số lượng'),
    ('Doanh thu năm', 'VND 500M-2 tỷ (~$20K-80K)'),
    ('Khả năng IT', 'Rất thấp — điện thoại/Facebook/bán tại quầy, sổ giấy'),
    ('Người ra quyết định', 'Chủ/trưởng gia đình, 45-65 tuổi, thường tự lái'),
    ('Thanh toán chấp nhận', 'Chỉ tiền mặt, thỉnh thoảng QR ngân hàng'),
    ('Top 3 nhu cầu', '1. Lấp ghế trống (load factor = sống còn)\n2. Thanh toán tin cậy không chậm trễ\n3. Onboarding đơn giản nhất (hỗ trợ qua Zalo)'),
    ('Top 3 phản đối', '1. "Nền tảng lấy phần từ biên lợi nhuận mỏng"\n2. "Khách tôi không đặt online"\n3. Sợ phức tạp / "Tôi không rành công nghệ"'),
])

doc.add_heading('4.1.2 Nhà xe trung bình (6-30 xe)', level=3)
doc.add_paragraph('Doanh thu VND 5-30 tỷ/năm. Có thể có POS cơ bản, MISA AMIS, có thể dùng VeXeRe BMS. '
                   'Nhu cầu chính: quản lý kho liên kênh tránh overbooking, MISA e-invoice tự động, phân tích theo tuyến. '
                   'Phản đối: chia kho nhiều nền tảng tạo rủi ro overbooking, kinh tế hoa hồng vs bán trực tiếp.')

doc.add_heading('4.1.3 Xe limousine/VIP (5-25 xe)', level=3)
doc.add_paragraph('Doanh thu VND 10-50 tỷ/năm. Segment limousine >62% thị phần xe cao cấp. '
                   'Nhu cầu: trưng bày thương hiệu (ảnh, badge tiện ích), giá động peak/off-peak, khách du lịch quốc tế (English UI). '
                   'Phản đối: sợ bị đặt cạnh xe bình dân, mất doanh thu đặt trực tiếp, đánh giá tiêu cực ảnh hưởng thương hiệu cao cấp.')

doc.add_heading('4.1.4 Đội xe lớn — FUTA-scale (50-800+ xe)', level=3)
doc.add_paragraph('Doanh thu $50M-320M+. 5-10 nhà xe quy mô quốc gia (FUTA, Hoàng Long, Thành Bưởi, The Sinh Tourist). '
                   'Có đội IT riêng, ứng dụng riêng, ERP riêng. Nhu cầu: API-first integration, khách mới không tiếp cận organic, báo cáo liên kênh. '
                   'Phản đối: "Tại sao trả hoa hồng khi đã có 20M khách/năm?", rủi ro chia sẻ dữ liệu, effort tích hợp vs doanh thu gia tăng.')

doc.add_heading('4.1.5 HTX Xe khách (10-60 xe)', level=3)
doc.add_paragraph('Mô hình hợp tác xã — thành viên cá nhân sở hữu xe. Tập trung Đồng bằng Sông Cửu Long và Tây Bắc. '
                   'Công nghệ thấp-trung bình. Nhu cầu: hiện diện số, HĐĐT không cần kế toán riêng, payout minh bạch cho từng thành viên. '
                   'Phản đối: quản trị HTX cần bỏ phiếu, giá cố định không linh hoạt, kế toán trợ cấp phức tạp.')

doc.add_heading('4.2 Hành khách — 6 phân khúc', level=2)

add_table(doc, ['Persona', 'Nhân khẩu', 'Tần suất', 'Thanh toán', 'Nhạy giá', 'Nhu cầu nền tảng chính'],
[
    ('"Chị Lan"\nCông nhân tỉnh', '20-45t, lao động di cư,\nthu nhập 5-12M/tháng', '4-8x/năm\nTết đỉnh', 'MoMo wallet,\ntiền mặt, ATM nội', '5/5', 'So giá, mã giảm, MoMo cashback, UI Việt, đảm bảo hoàn tiền'),
    ('"Anh Minh"\nDoanh nhân', '28-50t, chuyên viên,\n20-50M/tháng', '2-4x/tháng\nquanh năm', 'Visa/MC doanh nghiệp,\nVNPay QR', '2/5', 'HĐĐT tự động, chỉ số đúng giờ, filter ghế VIP, PDF biên lai'),
    ('"Trang"\nTrẻ đô thị', '20-32t, đô thị,\n12-30M/tháng', '1-3x/tháng\ncuối tuần', 'MoMo, ZaloPay,\nVNPay QR', '3/5', 'UX mobile-first mượt, sơ đồ ghế visual, badge tiện ích, 1-click rebook'),
    ('"Marco"\nDu khách quốc tế', '20-40t, solo/cặp,\ntrải nghiệm VN 2-4 tuần', '6-12 chặng xe\n/chuyến VN', 'Visa/MC credit,\nWise/Revolut', '3/5', 'UI English đầy đủ, chấp nhận thẻ quốc tế, GPS pickup pin Google Maps'),
    ('"Bà Hoa"\nNgười cao tuổi', '55-75t, hưu trí,\n3-8M/tháng', '2-5x/năm', 'Tiền mặt,\nchuyển khoản có trợ giúp', '4/5', 'Font lớn WCAG AA, vé PDF in được, SMS xác nhận, luồng đặt proxy gia đình'),
    ('"Em Quân"\nSinh viên', '17-24t, sinh viên,\ntiết kiệm data mobile', '4-8x/năm\nnghỉ lễ+Tết', 'MoMo (gia đình nạp),\nATM nội, tiền mặt', '5/5', 'Zalo Mini App (nhẹ data), mã giảm SV, UI tải nhanh, Zalo OA xác nhận'),
])

doc.add_heading('4.3 Quản trị viên — 4 vai trò', level=2)

add_table(doc, ['Vai trò', 'Trách nhiệm chính', 'Điểm đau chính'],
[
    ('Operations Manager', 'Duyệt nhà xe, kích hoạt tuyến, xử lý exception overbooking/hủy, override trip status', 'Không OCR cho tài liệu nhà xe, không bulk-action, không audit trail cho override thủ công'),
    ('Finance/Accounting', 'Chạy payout, đối soát hoa hồng, xuất HĐĐT MISA, khấu trừ thuế, cân đối ledger hoàn tiền', 'MISA API reject vì MST sai, đối soát đa-nhà xe qua file VNPay/MoMo riêng, không tự động tính khấu trừ thuế'),
    ('Compliance Officer', 'Quản lý chương trình PDPL, theo dõi SLA khiếu nại, xử lý DSAR, sàng lọc gian lận', 'Không có PII audit log hợp nhất, theo dõi deadline DSAR thủ công, mơ hồ phân loại nền tảng'),
    ('Customer Support', 'Sửa đổi/hủy booking, xử lý hoàn tiền, escalate khiếu nại, xác minh danh tính khách', 'Không có inbox hợp nhất (Zalo + email riêng), tra cứu booking phải có ref chính xác, SLA nhà xe không enforce'),
])

doc.add_heading('4.4 KPI cho Nhà đầu tư — Khung sẵn sàng Series A', level=2)

add_table(doc, ['Chỉ số', 'Minimum Viable', 'Tín hiệu mạnh'],
[
    ('GMV hàng tháng', '$500K', '$2M+'),
    ('Tăng trưởng GMV MoM', '10%', '20%+ duy trì 6 tháng'),
    ('ARR doanh thu ròng', '$1M', '$2-3M+'),
    ('Take rate', '5%', '8-12%'),
    ('LTV/CAC ratio', '2:1', '4:1+'),
    ('CAC payback period', '<18 tháng', '<9 tháng'),
    ('Tỷ lệ booking lặp lại', '30%', '50%+'),
    ('Nhà xe hoạt động', '100', '500+'),
    ('Phủ sóng tuyến (hành lang)', '30', '100+'),
    ('Gross margin', '30%', '50%+'),
    ('Series A SEA trung vị (2025)', '~$5M', '$7-10M'),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 5: ĐỐI THỦ CẠNH TRANH
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 5: ĐỐI THỦ CẠNH TRANH', level=1)

doc.add_heading('5.1 Bản đồ cạnh tranh', level=2)

doc.add_heading('Tier 1 — Trùng lặp trực tiếp với góc SaaS nhà xe của BB', level=3)
add_table(doc, ['Đối thủ', 'Mô tả', 'Tại sao quan trọng'],
[
    ('VeXeRe', 'OTA marketplace + BMS/AMS SaaS. ~80% thị trường phần mềm xe bus trực tuyến. 700+ nhà xe, 5,000+ đại lý, 2,600+ tuyến. Funding cuối 2019.', 'Đối thủ duy nhất chiếm cả hai: phần mềm nhà xe VÀ marketplace khách hàng. BMS là sản phẩm BB phải vượt trội. Moat = BMS Trojan horse (nhà xe vận hành business trên phần mềm VeXeRe, kho bị khóa).'),
    ('PhanMemNhaXe / Nhanh Travel / XECA', 'Phần mềm quản lý nhà xe standalone', 'Đối thủ SaaS thuần không có marketplace. BB cạnh tranh với họ cho nhà xe muốn phần mềm nhưng không muốn listing marketplace.'),
])

doc.add_heading('Tier 2 — Cạnh tranh phân phối nhà xe (không phần mềm)', level=3)
add_table(doc, ['Đối thủ', 'Mức liên quan'],
[
    ('redBus', 'Vào VN mạnh (cuối 2023/2024) với vốn MakeMyTrip. Cạnh tranh inventory listing, không phần mềm. Có thể là đối tác phân phối cho nhà xe BB.'),
    ('Traveloka', 'App du lịch #1 VN (82% nhận biết, 61% sử dụng 12 tháng). Xe bus là vertical phụ. Cạnh tranh listing.'),
    ('MoMo / ZaloPay / VNPAY', 'Kênh phân phối super-app. Tổng hợp nhu cầu; nhà xe BB nên phân phối qua họ. Đối tác, không phải kẻ thù.'),
])

doc.add_heading('Tier 3 — Segment khác, ít trùng lặp', level=3)
add_bullet(doc, '12Go Asia / Bookaway / Baolau: ', 'OTA khách du lịch quốc tế. Đối tác phân phối tiềm năng cho nhà xe BB.')
add_bullet(doc, 'FUTA Bus (trực tiếp): ', 'Tích hợp dọc với app riêng. Không phải đối thủ nền tảng — benchmark cho nhà xe số hóa tốt. 20-30M khách/năm.')

doc.add_heading('Định vị BB', level=3)
doc.add_paragraph(
    '"Hệ thống đặt vé số thuộc sở hữu nhà xe, không yêu cầu nhượng nhu cầu cho marketplace bên thứ ba."'
)
doc.add_paragraph(
    'VeXeRe BMS cho nhà xe phần mềm nhưng đưa kho hàng vào marketplace VeXeRe — nhà xe thành nhà cung cấp, không phải thương hiệu. '
    'PhanMemNhaXe/Nhanh Travel cho phần mềm nhưng không có kênh booking khách hàng. '
    'BB cho cả hai: console quản lý VÀ trải nghiệm booking khách hàng nơi thương hiệu nhà xe là trung tâm.'
)

add_decision_box(doc,
    'Định vị chiến lược',
    '"Shopify cho Nhà Xe" (sở hữu kênh bán) vs "VeXeRe lite" (marketplace) vs "SaaS thuần" (chỉ phần mềm)',
    '"Shopify cho Nhà Xe": Nhà xe sở hữu kênh, mối quan hệ khách hàng — NHƯNG không có marketplace tổng hợp nhu cầu, nhà xe tự lái traffic. '
    '"VeXeRe lite": có marketplace nhưng cạnh tranh trực tiếp với gã khổng lồ đã có 80% thị phần. '
    '"SaaS thuần": không có kênh booking = giá trị thấp hơn.',
    '"Shopify cho Nhà Xe"',
    'Khác biệt hóa duy nhất so với VeXeRe. Nhà xe lớn (FUTA) đã đầu tư vào stack riêng chính vì muốn kiểm soát thương hiệu. '
    'Giải quyết weakness bằng đối tác phân phối (MoMo/ZaloPay, 12Go, SEO tools, booking link chia sẻ).'
)

doc.add_heading('5.2 Ma trận so sánh tính năng', level=2)

add_table(doc, ['Danh mục / Tính năng', 'BB', 'VeXeRe', 'redBus', 'FUTA'],
[
    ('Tìm kiếm đa nhà xe', '✅', '✅ 2000+ nhà xe', '✅ tuyến chính', '❌ 1 nhà xe'),
    ('Xem sơ đồ ghế', '✅', '✅', '✅ live', '✅'),
    ('Lọc theo loại xe', '✅', '✅', '✅', '🔶 chỉ đội xe riêng'),
    ('Đa phương thức (tàu/bay)', '❌ chỉ xe', '✅ xe+tàu+bay', '🔶 xe+phà', '❌ chỉ xe'),
    ('Chọn ghế', '✅', '✅', '✅', '✅'),
    ('Đặt khứ hồi', '🔶 paired return (op-side)', '✅', '❓', '❓'),
    ('Vé điện tử', '✅', '✅ SMS+email', '✅ M-ticket', '🔶 app'),
    ('Visa/MC/JCB', '✅', '✅', '✅', '✅'),
    ('MoMo', '✅', '✅', '✅', '✅'),
    ('ZaloPay', '🔶 kế hoạch', '✅', '✅', '✅'),
    ('Tự hủy vé', '✅', '✅ one-touch', '✅', '❓'),
    ('Hoàn tiền phương thức gốc', '✅', '✅ 2-7 ngày', '🔶 ưu tiên ví', '❓'),
    ('Bảo đảm 150%', '❌', '✅ 100% cash + 50% voucher', '✅ 1.5x hoàn', '❌'),
    ('Dashboard nhà xe', '✅ RBAC sâu + MISA HĐĐT', '✅ BMS/AMS đầy đủ', '✅ dashboard', '❌ nội bộ'),
    ('Tích hợp API', '🔶 kế hoạch', '✅', '✅', '❌'),
    ('White-label site/app', '❌', '✅ hạ tầng VeXeRe', '❌', '❌'),
    ('SMS xác nhận', '✅', '✅', '✅', '🔶'),
    ('Bus tracking live', '❌', '✅ link chia sẻ', '✅', '❓'),
    ('App native', '❌ PWA/responsive', '✅ iOS 4.8★', '✅ iOS 4.6★', '✅ iOS 3.2★'),
    ('Flash sales', '❌', '✅ đến 50% off', '✅ đến 30%', '🔶'),
    ('Accessibility (WCAG)', '🔶 design-system', '❌', '❌', '❌'),
])

doc.add_paragraph('Điểm khác biệt tiềm năng BB: Accessibility (WCAG), RBAC console sâu, MISA HĐĐT tự động tích hợp, '
                   'nhà xe sở hữu thương hiệu, T+3 settlement minh bạch.')

doc.add_heading('5.3 So sánh hoa hồng & thanh toán', level=2)

add_table(doc, ['Chiều', 'BB', 'VeXeRe', 'redBus', 'FUTA', 'Traveloka'],
[
    ('Hoa hồng', '8-10% (5% intro, sàn 5%, trần 15%)', '~8-12% ước tính', '10-20% xác nhận', '0% (tích hợp dọc)', '~10-20% ước tính'),
    ('Thanh toán', 'T+3', 'Không công bố', 'Không công bố', 'N/A', 'Không công bố'),
    ('Subscription', 'Không', 'BMS subscription', 'Đăng ký miễn phí', 'N/A', 'Đăng ký miễn phí'),
    ('Doanh thu bổ sung', 'SaaS tools nhà xe', 'BMS/AMS, flash sale, bảo hiểm', 'BOSS/SeatSeller SW', 'Bán trực tiếp, FUTA Express', 'Hotel, bay'),
])

add_decision_box(doc,
    'Mức hoa hồng',
    '5% (cạnh tranh tối đa) vs 8-10% (benchmark ngành) vs 12-15% (margin cao)',
    '5% intro: thấp hơn tất cả đối thủ = thu hút nhà xe nhanh nhưng biên lợi nhuận âm sau phí PSP (1.5-2.5%) + chi phí hỗ trợ. '
    '8-10%: dưới VeXeRe (~8-12%) và redBus (10-20%) = cạnh tranh mà vẫn có lãi. '
    '12-15%: margin tốt nhưng nhà xe bypass nền tảng đặt trực tiếp.',
    '8-10% chuẩn, 5% giới thiệu 3 tháng đầu',
    'Cắt dưới market leader VeXeRe về giá; 5% intro tạo incentive thử nghiệm không rủi ro cho nhà xe; '
    'sau 3 tháng chuyển 8-10% vẫn thấp hơn VeXeRe. Net margin ~4% tại 10% = khả thi.'
)

doc.add_heading('5.4 Tâm lý nhà xe & nguyên nhân rời bỏ', level=2)
doc.add_paragraph('Giá trị nhà xe xếp hạng theo hành vi bộc lộ:')
doc.add_paragraph('1. Kiểm soát thương hiệu (CAO NHẤT) — FUTA (30M khách/năm) và Thành Bưởi đầu tư app riêng.')
doc.add_paragraph('2. Tốc độ thanh toán / cash flow — VeXeRe quảng cáo "nạp thẳng tài khoản ngân hàng" vì nhà xe từng bị chậm.')
doc.add_paragraph('3. Hoa hồng — Nhà xe chấp nhận 8-12% vì reach VeXeRe (700+). Ngưỡng ~15% — trên đó đẩy sang kênh trực tiếp.')
doc.add_paragraph('4. Hỗ trợ kỹ thuật / khả năng BMS — VeXeRe BMS (GPS, driver app, cargo, HĐĐT, white-label) là giá trị thực.')

doc.add_paragraph('Nguyên nhân rời bỏ (suy luận):')
add_bullet(doc, 'Tăng hoa hồng → tính toán lại chi phí-lợi ích so với kênh trực tiếp')
add_bullet(doc, 'Flash sale OTA xung đột chiến lược giá nhà xe')
add_bullet(doc, 'Tranh chấp thanh toán hoặc thiếu cash flow')
add_bullet(doc, 'Muốn sở hữu loyalty khách hàng (booking lặp lại qua OTA = nhà xe không bao giờ xây relationship trực tiếp)')
add_bullet(doc, 'OTA onboard đối thủ trên cùng tuyến với promotion chủ động')

doc.add_heading('5.5 Benchmark CAC/LTV', level=2)

add_table(doc, ['Kịch bản', 'Vé TB', 'Take rate', 'Chuyến/năm', 'LTV 3 năm', 'USD'],
[
    ('Bảo thủ', '250K', '8%', '3', '180K VND', '~$7'),
    ('Base case', '350K', '10%', '4', '420K VND', '~$17'),
    ('Lạc quan (loyalty)', '450K', '12%', '6', '972K VND', '~$39'),
])

doc.add_paragraph('LTV:CAC target: 3:1 tối thiểu, 4:1-6:1 mạnh. Tại $5 CAC và $17 LTV: 3.4:1 (đạt ngưỡng). '
                   'CAC payback ~17 tháng — quan ngại chính. Giảm thiểu: (a) giảm CAC qua organic/SEO và Zalo CRM, '
                   '(b) tăng take rate lên 12% khi value prop trưởng thành, (c) thêm doanh thu SaaS/nhà xe.')

doc.add_heading('5.6 Phủ sóng địa lý', level=2)

add_table(doc, ['Khu vực', 'Mô hình nhu cầu', 'Bão hòa OTA', 'Ưu tiên BB'],
[
    ('Trục N-S lao động (TH ↔ TPHCM)', 'Quanh năm cao, đỉnh Tết cực điểm', 'TB (OTA underserved vs tuyến du lịch)', 'Phase 1 Beachhead'),
    ('Trục N-S kề (NA, HT ↔ TPHCM)', 'Cùng mô hình lao động di cư', 'Thấp-TB', 'Phase 2 (mở rộng kề)'),
    ('Tuyến du lịch (HCMC-ĐL, NT)', 'Quanh năm + đỉnh lễ', 'Bão hòa (VeXeRe + redBus mạnh)', 'Phase 2 (bắt du lịch + English UI)'),
    ('Đồng bằng SCL', 'Hạ tầng đang phát triển, OTA ít', 'Thấp', 'Phase 3 (cạnh tranh ít)'),
    ('Tây Nguyên (không ĐL)', 'Du lịch mới nổi, chỉ bến xe', 'Rất thấp', 'Phase 3 (first-mover)'),
    ('Vùng núi phía Bắc', 'Du lịch mùa (Sa Pa bão hòa, xa bắc trống)', 'TB-Thấp', 'Chọn lọc'),
])

doc.add_heading('5.7 Lợi thế cạnh tranh BB', level=2)
doc.add_paragraph('1. MISA e-Invoice tự động (Duy nhất — không đối thủ nào tích hợp native): '
                   'NĐ 158/2024 yêu cầu nhà xe sử dụng phần mềm nền tảng gửi HĐĐT khi hoàn thành chuyến. '
                   'Lead với "Đăng ký và tuân thủ NĐ 158/2024 HĐĐT từ ngày đầu".')
doc.add_paragraph('2. Nhà xe sở hữu thương hiệu (anti-VeXeRe): "Shopify vs Amazon Marketplace".')
doc.add_paragraph('3. T+3 thanh toán nhanh: có ý nghĩa cho nhà xe nhỏ cash flow eo hẹp, đặc biệt Tết.')
doc.add_paragraph('4. Console nhà xe sâu với RBAC: fleet management, route management, trip scheduling, staff accounts phân quyền.')
doc.add_paragraph('5. Compliance Vietnam native: VND integer minor units, +84 phone, Asia/Ho_Chi_Minh timezone, VNPay/MoMo, MISA HĐĐT.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 6: MÔ HÌNH KINH DOANH
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 6: MÔ HÌNH KINH DOANH', level=1)

doc.add_heading('6.1 Cơ cấu hoa hồng', level=2)

add_table(doc, ['Mô hình', 'Tỷ lệ', 'Cơ sở'],
[
    ('Hoa hồng khuyến nghị khởi đầu', '8-10% giá vé', 'Dưới VeXeRe (~12%) và redBus (10-20%). Cắt dưới market leader để thu hút nhà xe.'),
    ('Sàn (Floor)', '5%', 'Dưới mức này: kinh tế đơn vị âm sau PSP (1.5-2.5%) + chi phí hỗ trợ'),
    ('Trần (Ceiling)', '15%', 'Trên mức này: nhà xe bypass nền tảng cho đặt trực tiếp'),
    ('Phí giới thiệu', '5% trong 3 tháng đầu', 'Chiến thuật thu hút nhà xe chuẩn — VeXeRe cung cấp BMS dùng thử miễn phí'),
])
doc.add_paragraph('Hoa hồng do nhà xe chịu, ẩn với khách hàng. Bất kỳ phí đặt vé hiển thị nào sẽ đẩy khách nhạy giá gọi trực tiếp nhà xe.')

doc.add_heading('6.2 Kinh tế đơn vị (10% hoa hồng trên vé trung bình 400,000 VND)', level=2)

add_table(doc, ['Hạng mục', 'Số tiền (VND)'],
[
    ('Doanh thu hoa hồng', '40,000'),
    ('Xử lý thanh toán (2%)', '-8,000'),
    ('Chi phí thông báo (ZNS/SMS)', '-1,000'),
    ('Hỗ trợ khách hàng (ước tính)', '-3,000'),
    ('CAC pha loãng (phân bổ)', '-10,000'),
    ('Hạ tầng (phân bổ)', '-2,000'),
    ('Biên lợi nhuận ròng/booking', '~16,000 (4.0% giá vé)'),
])
doc.add_paragraph('Điểm hòa vốn: ~50,000-100,000 bookings/tháng. Đạt được với 100-200 nhà xe hoạt động trung bình 15-30 bookings/ngày.')

doc.add_heading('6.3 Mô hình giá kép', level=2)

add_decision_box(doc,
    'Cấu trúc giá',
    '(A) Chỉ hoa hồng vs (B) Chỉ SaaS subscription vs (C) Kép: cả hai',
    '(A) Chỉ hoa hồng: đơn giản, nhưng nhà xe volume lớn trả quá nhiều. '
    '(B) Chỉ SaaS: ổn định, nhưng nhà xe nhỏ không chấp nhận chi phí cố định. '
    '(C) Kép: phức tạp hơn nhưng phù hợp mọi quy mô.',
    'Mô hình kép — "Shopify Basic vs Shopify Plus"',
    'Commission (0 VND/tháng, 8-10%/booking): tốt cho nhà xe volume thấp thử nghiệm. '
    'SaaS (1-2M VND/tháng, 3-5%/booking): tốt cho nhà xe volume cao muốn biến phí thấp. '
    'Validated bởi mô hình kép BMS/commission của VeXeRe.'
)

doc.add_heading('6.4 Nguồn doanh thu bổ sung (xếp hạng theo khả thi)', level=2)

add_table(doc, ['#', 'Nguồn doanh thu', 'Khả thi', 'Tiềm năng doanh thu', 'Khi nào'],
[
    ('1', 'SaaS subscription nhà xe (console không marketplace)', 'CAO — sản phẩm đã xây', '500K-2M VND/tháng/nhà xe', 'Ra mắt'),
    ('2', 'Listing nổi bật / promoted (nhà xe trả top placement)', 'CAO — tiêu chuẩn marketplace', '1-5M VND/tháng/nhà xe', 'Tháng 3-6'),
    ('3', 'Bảo hiểm du lịch add-on (Bảo Việt)', 'TRUNG BÌNH — cần đối tác bảo hiểm', '6,000-8,000 VND/conversion', 'Tháng 6-12'),
    ('4', 'Promo/voucher đồng tài trợ', 'TRUNG BÌNH', 'Biến đổi', 'Tháng 3-6'),
    ('5', 'Mạng đại lý/reseller (AMS)', 'THẤP ngắn hạn', 'CPS commission 2-8%', 'Tháng 12+'),
    ('6', 'Sản phẩm dữ liệu/phân tích', 'THẤP ngắn hạn', 'SaaS tier cao cấp', 'Tháng 12+'),
])

doc.add_heading('6.5 Playbook thu hút nhà xe', level=2)

add_table(doc, ['Giai đoạn', 'Mục tiêu', 'Value prop', 'Kênh thu hút', 'Đề xuất'],
[
    ('Phase 1\n(Tháng 0-3)', '10-20 nhà xe\n2-3 hành lang demand cao', '"HĐĐT tuân thủ từ ngày 1 +\ntrang đặt vé thương hiệu riêng +\nthanh toán nhanh hơn VeXeRe"', 'Bán trực tiếp — thăm bến xe,\ngặp quản lý nhà xe', '3 tháng miễn phí (0% hoa hồng),\nsau đó 5% intro'),
    ('Phase 2\n(Tháng 3-6)', '50+ nhà xe\nMở rộng tuyến du lịch', '"Xe của bạn đã trên 12Go/Bookaway\ncho du khách — BB cho khách Việt"', 'Bán trực tiếp + referral\ntừ Phase 1', ''),
    ('Phase 3\n(Tháng 6-12)', '200+ nhà xe\nMở rộng vùng', '"VeXeRe bỏ qua nhà xe <10 xe.\nBB được xây cho bạn."', 'Zalo OA campaigns,\nFB groups nhà xe', ''),
    ('Phase 4\n(Tháng 12+)', 'Đối tác phân phối', 'Inventory nhà xe BB\ntrong super-app booking flows', 'MoMo/ZaloPay,\n12Go/Bookaway API', ''),
])

doc.add_heading('6.6 Kênh thu hút khách hàng (xếp hạng theo ROI)', level=2)

add_table(doc, ['#', 'Kênh', 'CAC ước tính', 'Ưu tiên', 'Ghi chú'],
[
    ('1', 'Kênh riêng nhà xe (FB page, Zalo group, bảng hiệu bến)', 'Gần zero', 'CAO NHẤT', 'Lợi thế cấu trúc BB — nhà xe gửi khách có sẵn đến trang BB thương hiệu riêng'),
    ('2', 'Google SEO — từ khóa long-tail theo tuyến', 'Thấp (đầu tư nội dung, 12-24 tháng)', 'CAO', '"vé xe Sapa Hà Nội", "vé xe Đà Lạt TPHCM". Tuyến tỉnh ít cạnh tranh'),
    ('3', 'Zalo OA / ZNS campaigns', '$3-5 USD CAC first-booking', 'CAO', 'CAC thấp hơn 25-35% so với Facebook cho dịch vụ local'),
    ('4', 'Phân phối MoMo/ZaloPay', 'Gần zero marginal CAC', 'CAO NHẤT (khi có)', 'FUTA tiếp cận hàng chục triệu qua MoMo không cần download app FUTA'),
    ('5', 'Facebook/Meta ads', '$5-8 USD CAC first-booking', 'TRUNG BÌNH', 'Tốt cho awareness; thay đổi thuật toán ảnh hưởng organic reach'),
    ('6', 'Google SEM', '$5-10 USD', 'TRUNG BÌNH', 'Đắt trên từ khóa branded VeXeRe; ROI tốt hơn trên long-tail tuyến'),
    ('7', 'Chương trình referral', '$2-4 USD (chi phí thưởng)', 'TRUNG BÌNH', 'Xây sau khi có user base ban đầu'),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 7: TÍNH NĂNG & KHOẢNG TRỐNG
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 7: TÍNH NĂNG & KHOẢNG TRỐNG SO VỚI ĐỐI THỦ', level=1)

doc.add_heading('7.1 Tính năng bắt buộc (Table Stakes) — Kiểm tra trạng thái', level=2)

add_table(doc, ['Tính năng', 'Trạng thái', 'Tại sao bắt buộc'],
[
    ('Tìm kiếm tuyến (điểm đi + đến + ngày)', '✅ DONE', 'Mọi nền tảng có; khách không đặt vé được nếu thiếu'),
    ('Sơ đồ ghế visual với chọn ghế chính xác', '✅ DONE (hold flow)', 'VeXeRe, FUTA, redBus đều có; khách VN mong đợi chọn ghế/giường'),
    ('Thanh toán online — VNPay', '✅ DONE', 'Thanh toán QR/ngân hàng nội địa thống trị'),
    ('Thanh toán online — MoMo', '✅ DONE', '68% thị phần ví điện tử, 31M user'),
    ('Vé điện tử qua email', '✅ DONE (Resend)', 'Tiêu chuẩn mọi nền tảng'),
    ('Xác nhận đặt vé tức thì', '✅ DONE', 'Tín hiệu tin cậy; xác nhận ngay tách biệt nền tảng uy tín'),
    ('Chọn điểm đón/trả', '✅ DONE', 'Xe khách VN dùng điểm đón, không chỉ bến xe'),
    ('Hiển thị nhiều loại xe', '✅ DONE', 'Sleeper/seat/limousine là cách hành khách VN chọn'),
    ('Mobile-responsive web', '✅ DONE', '65-73% booking là mobile'),
    ('Hủy vé & hoàn tiền khách hàng', '❌ THIẾU', 'P0 LAUNCH BLOCKER — #1 khiếu nại mọi nền tảng. Ra mắt thiếu = phá hủy niềm tin.'),
    ('Đặt vé khứ hồi', '❌ THIẾU', 'P0 LAUNCH BLOCKER — Mọi đối thủ chính có. Thiếu = khách đặt 2 lần riêng, tăng friction.'),
    ('SMS/Zalo xác nhận booking', '⚠️ PARTIAL (chỉ email)', 'P1 — Khách VN mong SMS. Tỷ lệ mở email thấp. Chỉ email là không đủ.'),
    ('Lịch sử đặt vé ("Vé của tôi")', '❌ THIẾU', 'P1 — Mọi nền tảng có account có tính năng này'),
    ('Thanh toán tiền mặt', '⚠️ PARTIAL', 'P1 — Phần đáng kể hành khách không thanh toán online trước'),
])

doc.add_heading('7.2 Khoảng trống P0 — Xây trước ra mắt', level=2)

add_decision_box(doc,
    'Hủy vé & Hoàn tiền tự phục vụ',
    '(A) Ra mắt không có tính năng hủy/hoàn tiền vs (B) Xây trước ra mắt vs (C) Chỉ hủy qua hỗ trợ thủ công',
    '(A): Ra mắt nhanh hơn 2-4 tuần nhưng #1 pain point = tích lũy đánh giá tiêu cực ngay lập tức. '
    '(B): Cần MoMo/VNPay refund API + cấu hình chính sách hủy nhà xe + cơ chế clawback T+3, effort TRUNG BÌNH. '
    '(C): Không scale, chậm xử lý, vẫn gây frustration.',
    'Xây TRƯỚC ra mắt — Launch blocker không thương lượng',
    '#1 pain point user. Mọi đối thủ có. VeXeRe 2.7★ Trustpilot chủ yếu vì "hoàn tiền chậm". '
    'Hoàn tiền tự động trong 24-48h có thể là trust differentiator.'
)

add_decision_box(doc,
    'Đặt vé khứ hồi',
    '(A) Khách đặt 2 lần riêng vs (B) Xây luồng khứ hồi customer-facing',
    '(A): Zero effort nhưng tăng gấp đôi friction, tăng abandonment. '
    '(B): Hạ tầng paired-return đã tồn tại (Issue 013); cần hoàn thiện search/checkout customer-facing.',
    'Xây trước ra mắt',
    'Hạ tầng paired-return phía nhà xe đã có. Cần hoàn thiện search/checkout khách hàng. Effort TRUNG BÌNH.'
)

doc.add_heading('7.3 Khoảng trống P1 — Ra mắt hoặc ngay sau', level=2)
add_bullet(doc, 'SMS/Zalo (ZNS) thông báo booking: ', 'ZNS ưu tiên hơn SMS (open rate cao hơn, rẻ hơn, 70M user Zalo). Effort NHỎ-TB.')
add_bullet(doc, '"Vé của tôi" — lịch sử đặt vé: ', 'Dữ liệu đã có — chỉ cần UI read-only. Effort NHỎ.')
add_bullet(doc, 'Hỗ trợ tiếng Anh: ', 'Khóa segment du khách quốc tế. Quan trọng cho tuyến du lịch (Sa Pa, Đà Lạt, Nha Trang). Effort TB.')
add_bullet(doc, 'Thanh toán tiền mặt: ', 'FUTA hỗ trợ "đặt online, trả tiền mặt cho tài xế/tại quầy". Effort TB.')

doc.add_heading('7.4 Khoảng trống P2-P3 — Tháng 1-12+', level=2)

add_decision_box(doc,
    'App native vs PWA vs Phân phối super-app',
    '(A) Xây app native iOS + Android (effort LỚN, $50-100K+) vs (B) PWA (effort TB) vs (C) Phân phối MoMo/ZaloPay (effort NHỎ-TB)',
    '(A): Home screen, push notification, offline — nhưng thời gian xây lâu, phải maintain 3 codebases. '
    '(B): Nhanh hơn, share codebase — nhưng không có app store presence, push notification hạn chế. '
    '(C): FUTA tiếp cận hàng chục triệu qua MoMo mà không cần download app FUTA — leverage cao nhất.',
    'Hoãn app native. Theo đuổi phân phối MoMo/ZaloPay trước. PWA là bước trung gian. Đánh giá tại tháng 6.',
    'Mobile web viable (Baolau/12Go chứng minh cho tourist). Quyết định phụ thuộc MoMo/ZaloPay partnership có thành không. '
    'Nếu có → hoãn native. Nếu không → đầu tư PWA.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 8: PHÁP LÝ & TUÂN THỦ
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 8: PHÁP LÝ & TUÂN THỦ QUY ĐỊNH', level=1)

doc.add_heading('8.1 Tổng quan pháp lý', level=2)
doc.add_paragraph('Quét pháp lý sâu 16 lĩnh vực quy định, nghiên cứu tháng 6/2026: 35 phát hiện, 26 rủi ro mở, 5 xung đột liên miền.')

doc.add_heading('5 Xung đột liên miền (cần giải quyết với luật sư)', level=3)
add_table(doc, ['#', 'Xung đột', 'Miền', 'Rủi ro'],
[
    ('C1', 'Lưu trú dữ liệu tự động áp dụng cho DN nội địa (NĐ 53/2022) nhưng thực thi chưa mạnh. DB Singapore kỹ thuật vi phạm.', 'Dữ liệu + Privacy', 'BCA yêu cầu bất kỳ lúc nào; buộc di chuyển khẩn cấp'),
    ('C2', 'VNPay thanh toán vào TK nền tảng trước (không nhà xe) trông giống trung gian thanh toán dù contract label marketplace.', 'Thanh toán + HĐĐT', 'SBV phân loại là IPS không giấy phép'),
    ('C3', 'Quyền hủy 3 ngày (Luật BVQLNTD 2023, Đ.29) áp dụng hợp đồng từ xa. Vé xe trước khởi hành có thể đủ điều kiện. Xung đột với thực tiễn ngành no-refund.', 'Consumer + Transport', 'Phải thiết kế chính sách hoàn tiền tuân thủ luật mới'),
    ('C4', 'BCT coi nền tảng là "sàn TMĐT" (NĐ 85/2021). BGTVT chưa phân loại cụ thể cho nền tảng đặt vé bus trực tuyến. Khoảng trống giữa 2 bộ.', 'Transport + E-Commerce', 'Bộ nào cũng có thể chủ trương quyền quản lý'),
    ('C5', 'Không phải thực thể báo cáo AML trực tiếp nếu không có giấy phép thanh toán. Nhưng tracking booking tiền mặt có thể trông giống xử lý thanh toán informal.', 'Thanh toán + AML', 'Vùng xám; tránh tính năng tracking tiền mặt cho đến khi rõ ràng'),
])

doc.add_heading('8.2 Pháp nhân & Đăng ký thương mại điện tử', level=2)

add_decision_box(doc,
    'Loại hình pháp nhân',
    'LLC (TNHH) vs JSC (Cổ phần)',
    'LLC: governance đơn giản hơn, 1-50 thành viên, KHÔNG phát hành cổ phiếu công khai — nhưng chuyển đổi sang JSC sau được. '
    'JSC: cần tối thiểu 3 sáng lập viên, Board of Directors + Ban kiểm soát, CÓ THỂ phát hành cổ phiếu — nhưng phức tạp hơn cho giai đoạn đầu.',
    'LLC (TNHH) cho giai đoạn đầu. Chuyển đổi sang JSC khi chuẩn bị Series A.',
    'Đơn giản hóa governance. ERC 3-5 ngày. Chi phí ~5-15M VND qua dịch vụ pháp lý. '
    'JSC chỉ cần khi huy động vốn từ nhiều nhà đầu tư.'
)

add_decision_box(doc,
    'Mã VSIC',
    'Công nghệ/IT (100% sở hữu nước ngoài) vs Vận tải (giới hạn 49-51% nước ngoài) vs Cả hai',
    'VSIC công nghệ: 100% sở hữu nước ngoài OK (cam kết WTO), không cần giấy phép vận tải. '
    'VSIC vận tải: trigger yêu cầu giấy phép vận tải, giới hạn sở hữu nước ngoài. '
    'Mã sai → giấy phép vận tải bắt buộc, giới hạn sở hữu nước ngoài, bộ quản lý khác.',
    'Mã VSIC công nghệ + TMĐT, KHÔNG vận tải',
    'Tiền lệ Grab Vietnam: đăng ký dưới IT/technology services, 100% sở hữu nước ngoài, '
    'vận hành nền tảng liên quan vận tải không cần giấy phép vận tải. VeXeRe tương tự.'
)

doc.add_heading('Đăng ký sàn TMĐT tại MOIT', level=3)
doc.add_paragraph(
    'BB = "sàn giao dịch TMĐT" (san TMDT) theo NĐ 85/2021. Bắt buộc đăng ký tại Online.gov.vn. '
    'Timeline: 2-4 tuần xét duyệt. Phạt không đăng ký: 40-60 triệu VND. '
    'Yêu cầu: ERC, MST, mô tả nền tảng, cơ chế giải quyết tranh chấp, privacy policy, ToS, quy trình xác minh nhà xe.'
)

doc.add_heading('8.3 Thanh toán & Giấy phép trung gian thanh toán (IPS)', level=2)

add_decision_box(doc,
    'Mô hình luồng tiền',
    'Marketplace (VNPay → TK nhà xe, không qua nền tảng) vs Merchant of Record (nền tảng là chủ thể) vs Hybrid/Pooling (thu qua TK nền tảng rồi chuyển nhà xe)',
    'Marketplace: rủi ro THẤP nhất, không cần giấy phép SBV — nhưng phức tạp (mỗi nhà xe cần TK merchant). '
    'Merchant of Record: rủi ro TRUNG BÌNH, cơ sở pháp lý khác — nhưng nền tảng chịu trách nhiệm thuế bán hàng. '
    'Hybrid/Pooling: RỦI RO CAO — tiền qua TK nền tảng rồi chuyển nhà xe = trông giống "thu hộ chi hộ" = CÓ THỂ CẦN giấy phép SBV (VND 50 tỷ vốn).',
    'Marketplace: VNPay/MoMo thanh toán thẳng vào TK nhà xe',
    'Tiền lệ VeXeRe: nền tảng đặt vé bus lớn nhất VN, hoạt động từ 2013, KHÔNG có giấy phép IPS SBV. '
    'Marketplace model loại bỏ hoàn toàn custody tiền nhà xe. Nền tảng xuất hóa đơn riêng cho nhà xe về phí dịch vụ (B2B).'
)

add_decision_box(doc,
    'Tích hợp PSP',
    'Chỉ VNPay vs Chỉ MoMo vs Cả 3 (VNPay + MoMo + VietQR)',
    'Chỉ VNPay: phủ sóng rộng nhất (thẻ nội + quốc tế + QR) nhưng bỏ 31M user MoMo. '
    'Chỉ MoMo: ví lớn nhất nhưng không có thẻ quốc tế. '
    'Cả 3: effort tích hợp cao nhưng phủ gần 100% payment preferences.',
    'Tích hợp cả 3. VNPay ưu tiên #1 (phủ rộng nhất), MoMo #2 (ví lớn nhất), VietQR #3 (phí thấp nhất)',
    'Adapter architecture đã có (VNPayAdapter, MoMoAdapter). VietQR theo cùng PaymentAdapter interface. '
    'VNPay: MDR 0.5-2% nội + 2.5-4.5% quốc tế, T+1. MoMo: MDR 1.5-2.5%, T+1-T+2. VietQR: MDR <0.5-1%, T+0-T+1.'
)

doc.add_heading('8.4 Hóa đơn điện tử & Thuế', level=2)

add_decision_box(doc,
    'Nhà cung cấp HĐĐT',
    'MISA meInvoice vs VNPT-Invoice vs Viettel S-Invoice vs FPT eInvoice',
    'MISA: phổ biến nhất cho SME, API private per customer, ĐÃ TÍCH HỢP (#74). '
    'VNPT: liên kết chính phủ. Viettel: doanh nghiệp lớn. FPT: giá cạnh tranh.',
    'MISA meInvoice',
    'Đã tích hợp trong codebase (Issue #74). Phổ biến nhất cho SME — phù hợp target nhà xe. '
    'Chi phí ~500-2,000 VND/hóa đơn. Kết nối GDT 24/7, XML/chữ ký số, lưu trữ 10 năm.'
)

add_decision_box(doc,
    'Vai trò xuất HĐĐT',
    'Nền tảng xuất thay nhà xe (ủy quyền) vs Mỗi nhà xe tự xuất vs Nền tảng là bên bán (principal)',
    'Ủy quyền: nền tảng xuất với MST nhà xe, cần thỏa thuận ủy quyền + thông báo cơ quan thuế. '
    'Tự xuất: đơn giản pháp lý nhưng nhà xe nhỏ không có khả năng. '
    'Principal: nền tảng là bên bán = phức tạp thuế, trách nhiệm VAT.',
    'Nền tảng xuất thay nhà xe (mô hình ủy quyền)',
    'Cách tiếp cận tiêu chuẩn ngành (VeXeRe, Ve Xe Nhanh). NĐ 123 Đ.17 + NĐ 70/2025 mở rộng quyền ủy quyền '
    'cho hộ kinh doanh cá thể. HĐĐT hiển thị nhà xe là bên bán (với MST nhà xe). '
    'HĐĐT hoa hồng riêng: nền tảng → nhà xe (B2B, hàng tháng).'
)

doc.add_paragraph('Thuế suất chính:')
add_table(doc, ['Loại thuế', 'Thuế suất', 'Ai chịu'],
[
    ('VAT vé xe', '10%', 'Nhà xe thu và nộp'),
    ('VAT hoa hồng nền tảng', '10%', 'Nền tảng thu từ nhà xe'),
    ('TNDN', '20%', 'Nền tảng trên lợi nhuận ròng'),
    ('Khấu trừ TNCN (nhà xe cá nhân)', '~1.5%', 'Nền tảng khấu trừ (từ tháng 7/2026)'),
    ('Thuế nhà thầu nước ngoài', '5% VAT + 5% CIT', 'Bên VN khấu trừ (nếu áp dụng)'),
])

doc.add_heading('8.5 Bảo vệ dữ liệu cá nhân & Lưu trú dữ liệu', level=2)

add_decision_box(doc,
    'Kiến trúc lưu trú dữ liệu',
    '(A) Giữ toàn bộ trên Vercel Singapore vs (B) Hybrid: Vercel SG cho compute + DB tại VN cho PII vs (C) Chuyển toàn bộ về VN',
    '(A): Kỹ thuật vi phạm NĐ 53/2022, rủi ro BCA buộc di chuyển khẩn cấp. '
    '(B): Thêm 5-15ms latency Vercel SG → VN DB, nhưng tuân thủ hoàn toàn. Prisma directUrl PgBouncer config đã có. '
    '(C): Mất lợi thế CDN/edge functions của Vercel, tăng chi phí hosting.',
    'Hybrid: Vercel Singapore cho compute/CDN + PostgreSQL tại VN (Viettel IDC/VNPT) cho PII',
    'Giữ lợi thế Vercel (auto-scale, edge, CDN) cho serving. PII (tên, SĐT, email, payment tokens, IP) '
    'lưu trên server VN. Latency thêm 5-15ms chấp nhận được cho web app. '
    'NĐ 147/2024 bổ sung: ít nhất 1 server phải ở VN cho điều tra/khiếu nại.'
)

doc.add_paragraph('Yêu cầu PDPL 2025 chính:')
add_bullet(doc, 'Đồng ý tự nguyện, cụ thể, có thông tin, theo mục đích — checkbox mặc định bị cấm')
add_bullet(doc, 'DPIA nộp A05 trong 60 ngày kể từ bắt đầu xử lý — MPS xét 15 ngày')
add_bullet(doc, 'CDTIA nộp A05 trong 60 ngày kể từ chuyển dữ liệu đầu tiên ra nước ngoài')
add_bullet(doc, 'Thông báo vi phạm 72h (24h cho tấn công an ninh mạng ảnh hưởng thông tin người tiêu dùng)')
add_bullet(doc, 'Quyền chủ thể dữ liệu: truy cập 10 ngày, sửa 10 ngày, xóa 20 ngày, rút đồng ý 15 ngày')
add_bullet(doc, 'DPO bắt buộc — startup exemption 5 năm NHƯNG không đủ điều kiện nếu xử lý dữ liệu tài chính (nhạy cảm)')

doc.add_heading('8.6 Viễn thông / SMS / OTP', level=2)

add_decision_box(doc,
    'Kênh OTP & thông báo',
    'Chỉ SMS vs Chỉ ZNS vs Kép: ZNS primary → SMS fallback',
    'Chỉ SMS: phủ 100% nhưng đắt (300-800 VND/msg). '
    'Chỉ ZNS: rẻ hơn (200-500 VND), open rate cao, nội dung phong phú — nhưng chỉ tiếp cận user Zalo (~75M MAU, không 100%). '
    'Kép: tiết kiệm 50-70% chi phí so với chỉ SMS, phủ sóng gần 100%.',
    'ZNS primary → SMS fallback (eSMS làm aggregator)',
    'Mô hình ngành chuẩn (Grab, Be, Tiki, Shopee VN). eSMS đã stub trong codebase. '
    'Đăng ký Brandname SMS: 2-4 tuần/nhà mạng — BLOCKER CỨNG, bắt đầu ngay.'
)

doc.add_heading('8.7 Vận tải — Phân loại nền tảng', level=2)

add_decision_box(doc,
    'Phân loại nền tảng',
    '"Kinh doanh vận tải" (cần giấy phép vận tải, giới hạn sở hữu nước ngoài 49-51%) vs "Nền tảng công nghệ" (đăng ký TMĐT, 100% nước ngoài OK)',
    'Vận tải: cần giấy phép, đội xe, bảo hiểm, vốn đăng ký — nền tảng không sở hữu/vận hành xe. '
    'Công nghệ: chỉ cần đăng ký TMĐT, xác minh nhà xe — phù hợp mô hình marketplace.',
    '"Nền tảng công nghệ" (marketplace kết nối nhà xe có giấy phép với hành khách)',
    'Tiền lệ VeXeRe: hoạt động như công ty công nghệ, không giấy phép vận tải, 10+ năm không bị thách thức pháp lý. '
    'QĐ 24/2018 phân loại app gọi xe (Grab) riêng biệt với kinh doanh vận tải. '
    'Cần ý kiến pháp lý chính thức trước đăng ký: nền tảng không sở hữu/vận hành xe, là marketplace công nghệ.'
)

doc.add_heading('8.8 Bảo vệ người tiêu dùng', level=2)
doc.add_paragraph(
    'Luật BVQLNTD 2023 (có hiệu lực 1/7/2024) — thay đổi lớn: nền tảng phải xác minh danh tính nhà xe, '
    'gỡ listing vi phạm, hợp tác điều tra, ĐỒNG CHỊU TRÁCH NHIỆM với nhà xe trong một số trường hợp.'
)
doc.add_paragraph('Quyền hủy hợp đồng từ xa (Đ.29): người tiêu dùng có quyền hủy trong 3 ngày làm việc — '
                   'TRỪUKS dịch vụ đã thực hiện hoàn toàn. Câu hỏi pháp lý: khởi hành = "đã thực hiện"? '
                   'Trước khởi hành: quyền hủy có thể áp dụng. Cần ý kiến luật sư.')
doc.add_paragraph('Xử lý khiếu nại: tiếp nhận 3 ngày, giải quyết 7-30 ngày. Phải có kênh khiếu nại trên nền tảng.')

doc.add_heading('8.9 Điều khoản hợp đồng PSP', level=2)

add_table(doc, ['Chiều', 'VNPay', 'MoMo', 'VietQR/NAPAS'],
[
    ('Giấy phép', 'IPS SBV (gateway + switching)', 'IPS SBV (ví + trung gian)', 'Không phải PSP — qua ngân hàng'),
    ('MDR nội địa', '0.5-2%', '1.5-2.5%', '<0.5-1%'),
    ('MDR quốc tế', '2.5-4.5%', 'Không có', 'Không có'),
    ('Settlement', 'T+1 chuẩn', 'T+1 đến T+2', 'T+0 đến T+1'),
    ('Chargeback fee', '200K-500K VND/tranh chấp', 'MoMo trung gian', 'Không có (chuyển khoản)'),
    ('Refund API', 'Có (reversal tự động)', 'Có (AIO refund endpoint)', 'Không (chuyển khoản thủ công)'),
    ('Sandbox', 'Có (trước ERC)', 'Có (trước ERC)', 'Qua bank SDK'),
    ('Docs yêu cầu', 'ERC, MST, TK ngân hàng, website T&Cs', 'Tương tự VNPay', 'Qua ngân hàng đối tác'),
])

doc.add_heading('8.10 DPIA Checklist', level=2)
doc.add_paragraph('Kiểm kê dữ liệu: Danh tính (họ tên, ngày sinh), Liên hệ (SĐT, email), Lịch sử booking, '
                   'Payment tokens (nhạy cảm), Vị trí (IP, GPS — nhạy cảm), OTP attempts.')
doc.add_paragraph('Thời hạn lưu trữ: Hồ sơ booking 5 năm (Luật Kế toán), Hồ sơ thanh toán/hóa đơn 10 năm (NĐ 123), '
                   'Log OTP 90 ngày, Session tokens đến khi hết hạn (tối đa 15 phút JWT nhà xe).')
doc.add_paragraph('Thông báo vi phạm: 72h cho A05/BCA, 24h nếu tấn công an ninh mạng, thông báo cá nhân bị ảnh hưởng, '
                   'lưu hồ sơ vi phạm 5 năm. Phạt: đến 5% doanh thu năm VN hoặc 3 tỷ VND.')

doc.add_heading('8.11 Lao động, AML, SHTT, Bảo hiểm, Tiếp cận', level=2)
doc.add_paragraph('Lao động: đóng BHXH 32.5% trên lương gross. Founder nước ngoài góp vốn: miễn work permit (cần giấy xác nhận DOLISA). '
                   'Rủi ro phân loại: nếu nền tảng kiểm soát giá/lịch trình nhà xe quá mức → có thể bị phân loại lại thành quan hệ lao động.')
doc.add_paragraph('AML: nền tảng đặt vé bus KHÔNG phải thực thể báo cáo AML trực tiếp (trừ khi có giấy phép thanh toán). '
                   'KYC cơ bản cho nhà xe: ERC, MST, xác minh TK ngân hàng.')
doc.add_paragraph('SHTT: hệ thống first-to-file — NỘP SỚM. Thời gian: 12-18 tháng. Phí: ~1-3M VND chính thức + ~10-20M VND với luật sư. '
                   'Bảo vệ 10 năm, gia hạn không giới hạn. Đăng ký Nice Class 39 (vận tải), 35 (quản lý kinh doanh), 42 (CNTT).')
doc.add_paragraph('Bảo hiểm: gói SME Bảo Việt/PVI ~15-30M VND/năm bao gồm E&O, trách nhiệm chung, cyber insurance.')
doc.add_paragraph('Tiếp cận: Thông tư 26/2020 KHUYẾN NGHỊ (không bắt buộc) WCAG 2.0 Level AA cho khu vực tư nhân. '
                   'BB đã có ở cấp design-system — không đối thủ nào làm = lợi thế cạnh tranh tiềm năng.')

doc.add_heading('8.12 Lộ trình tuân thủ 12 tuần trước ra mắt', level=2)

add_table(doc, ['Tuần', 'Hành động', 'Ghi chú'],
[
    ('1-2', 'Chọn cơ cấu pháp nhân; thuê luật sư nộp IRC/ERC; bắt đầu legalization tài liệu (apostille/consular)', 'BLOCKER: legalization 15-30 ngày. Bắt đầu Ngày 1.'),
    ('3-4', 'Nộp IRC; bắt đầu đăng ký nhãn hiệu; bắt đầu đàm phán VNPay/MoMo sandbox', 'IRC statutory 15 wd; thực tế 3-5 tuần'),
    ('5-8', 'Nhận IRC → nộp ERC; mở TK ngân hàng; đăng ký thuế; đăng ký HĐĐT GDT + MISA; đăng ký sàn TMĐT MOIT; bắt đầu template SMS Brandname', 'ERC 3-7 wd sau IRC. NĐ 117/2025 áp dụng marketplace.'),
    ('9-12', 'Hoàn thành onboarding VNPay/MoMo production; confirm Brandname SMS; nộp DPIA + CDTIA cho A05; bổ nhiệm DPO; ý kiến pháp lý PSP', 'DPIA trong 60 ngày từ xử lý dữ liệu đầu tiên.'),
    ('Liên tục', 'Cập nhật MOIT, DPIA 6 tháng/lần, template SMS, đối soát PSP hàng ngày, HĐĐT mỗi giao dịch, báo cáo sự cố 72h', ''),
])

doc.add_paragraph('Đường tới hạn: Legalization → IRC → ERC = 43-87 ngày end-to-end. Bắt đầu Ngày 1.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 9: MÔ HÌNH MIỀN
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 9: MÔ HÌNH MIỀN (Domain Model)', level=1)

doc.add_heading('9.1 Bản đồ ngữ cảnh giới hạn (12 Bounded Contexts)', level=2)

add_table(doc, ['#', 'Context', 'Mô tả', 'Models chính'],
[
    ('1', 'Auth', '3 realm riêng biệt: Customer, Operator, Admin. Không bao giờ trộn. Mỗi realm có session, OTP, auth service riêng.', 'Customer, Session, OtpAttempt, OperatorUser, OperatorSession, AdminUser, AdminSession'),
    ('2', 'Fleet/Catalog', 'Phía cung thuộc nhà xe: xe, tuyến, chuyến, pickup areas, lịch định kỳ.', 'Operator, Bus, Route, Trip, RecurringTripTemplate, OperatorPickupArea'),
    ('3', 'Booking', 'Phía cầu khách hàng: holds, bookings, consent, check-in.', 'Hold, Booking, ConsentRecord'),
    ('4', 'Payment', 'Xử lý thanh toán gateway-agnostic và webhook.', 'PaymentEvent + Adapters (MoMo, VNPay, Stub)'),
    ('5', 'Finance/Ledger', 'Kế toán kép, payouts, cấu hình phí, settlements.', 'LedgerEntry, Payout, FeeConfig, PayoutAccount'),
    ('6', 'Notification', 'Đường giao hàng duy nhất cho tất cả domain notifications.', 'NotificationLog (SMS + Email)'),
    ('7', 'Admin/Moderation', 'Quản trị nền tảng, lifecycle nhà xe, kiểm duyệt nội dung.', 'AdminAuditLog, ContentReport'),
    ('8', 'Reporting/Analytics', 'Funnel conversion, báo cáo doanh thu, tracking job.', 'FunnelEvent, JobRunLog'),
    ('9', 'Onboarding/KYB', 'Đăng ký nhà xe, xác minh tài liệu, setup TK payout.', 'KybDocument, PayoutAccount, StoredObject'),
    ('10', 'Charter', 'Yêu cầu thuê xe nhóm/riêng. Chỉ lead-gen, không có payment rail.', 'CharterRequest'),
    ('11', 'E-Invoice', 'Tuân thủ VN per Thông tư 78/2021. HĐĐT GDT.', 'EInvoice'),
    ('12', 'Feature Flags', 'Cổng tính năng runtime.', 'FeatureFlag'),
])

doc.add_heading('9.2 Thuật ngữ miền chính (với tiếng Việt)', level=2)

add_table(doc, ['Thuật ngữ', 'Tiếng Việt', 'Định nghĩa'],
[
    ('Operator', 'Nhà Xe', 'Công ty xe bus đăng ký trên nền tảng. KYB lifecycle: PENDING_REVIEW → APPROVED | REJECTED.'),
    ('Route', 'Tuyến', 'Tuyến xe bus hai chiều. Scoped theo nhà xe. Mang durationMinutes cho tính toán overlap.'),
    ('Trip', 'Chuyến', 'Một khởi hành cụ thể trên Route bằng Bus. Giá cố định VND. Status: scheduled|departed|completed|cancelled.'),
    ('Hold', 'Đặt Chỗ Tạm', 'Giữ chỗ tạm thời TTL 10 phút. Capacity checked atomic qua pg advisory locks.'),
    ('Booking', 'Đặt Vé', 'Mua vé xác nhận. Tạo từ Hold consumed. BookingRef format: BB-YYYY-[0-9a-z]{4}-[0-9a-z]{4}.'),
    ('Payout', 'Giải Ngân', 'Thanh toán cho nhà xe. 2 path: auto-sweep/chuyến (T+1) hoặc on-demand withdrawal.'),
    ('LedgerEntry', '', 'Dòng kế toán kép. Append-only bất biến (DB trigger). Balance luôn derived, không lưu.'),
    ('Platform Fee', '', 'Hoa hồng marketplace. Lưu dạng ratePpm (phần-triệu) trong FeeConfig. Tính bằng BigInt.'),
    ('Charter', 'Thuê Xe', 'Yêu cầu thuê xe nhóm/riêng. Chỉ lead-gen — settlement off-platform.'),
    ('E-Invoice', 'Hóa Đơn Điện Tử', 'HĐĐT tuân thủ GDT. Status: pending → issued → sent. Tích hợp MISA meInvoice.'),
])

doc.add_heading('9.3 Luồng sự kiện chính', level=2)

doc.add_heading('Luồng 1: Search → Hold → Book → Pay → Ticket', level=3)
doc.add_paragraph('Hành trình khách hàng chính. Search: tìm chuyến (diacritics-safe qua unaccent_immutable ILIKE). '
                   'Hold: giữ chỗ 10 phút (advisory locks: phone-level cap → trip-level serialization). '
                   'Book: chuyển hold thành booking awaiting_payment (2 ConsentRecord: no_refund + pii_storage). '
                   'Pay: PSP webhook xác nhận → booking paid (idempotent qua @@unique([adapter, providerTxnId])). '
                   'Ticket: PDF async bởi cron job, lưu object storage.')

doc.add_heading('Luồng 2: Nhà xe tạo chuyến', level=3)
doc.add_paragraph('Xác thực operator scope → $transaction + SELECT FOR UPDATE trên Bus → guards (bus deactivated? maintenance overlap? time overlap?) '
                   '→ INSERT Trip. Giá do nhà xe cung cấp (I7-exempt: nhà xe LÀ price authority).')

doc.add_heading('Luồng 3: Hủy chuyến + Hoàn tiền', level=3)
doc.add_paragraph('Không có hủy do khách — consent no_refund ghi nhận khi booking. Chỉ nhà xe/hệ thống hủy. '
                   '$transaction + FOR UPDATE → idempotent (alreadyCancelled: true → HTTP 200, discriminated result, không throw). '
                   'Cascade: bulk UPDATE Booking → trip_cancelled, Hold → cancelled_trip. Post-commit: refundOut() per paid booking.')

doc.add_heading('Luồng 4: Payout (T+1 Settlement)', level=3)
doc.add_paragraph('completeTripCore → Payout row (scheduledAt = completedAt + 1 ngày). Settlement: revenue eligible khi completedAt + 1 ngày ≤ NOW(). '
                   'Balance derived: available = settledEligible - paidOut. Withdrawal: ≥ 100,000 VND, PayoutAccount verified, idempotent (double-probe).')

doc.add_heading('9.4 Bất biến kinh doanh (14 bất biến)', level=2)

add_table(doc, ['Bất biến', 'Mô tả', 'Hậu quả nếu vi phạm'],
[
    ('I1: Concurrency Control', 'Mọi read-then-write trên shared state chạy trong $transaction với SELECT FOR UPDATE.', 'Oversold trips, double payouts, invalid state transitions'),
    ('I7: No Client-Originated Price', 'Customer-facing endpoints không bao giờ nhận price từ request body. Price derived từ Trip.price × Hold.ticketCount.', 'Khách có thể underpay cho vé'),
    ('I9: No Raw Phone in Payload', 'NotificationLog.payload KHÔNG duplicate SĐT. recipient là sole phone column.', 'Leak PII qua structured log exports'),
    ('Capacity Guard', '3 lớp: Hold creation (advisory locks), Booking-paid webhook (FOR UPDATE recount), Hold cap (per-phone limit).', 'Oversold bus'),
    ('Maintenance Window', 'Bus có maintenance window không thể assigned cho trips overlap window.', 'Xe đi maintenance ngày khởi hành'),
    ('salesClosed Gate', 'Trip với salesClosed=true loại khỏi search và không nhận hold mới.', 'Khách tạo hold trên chuyến đã khởi hành'),
    ('Bus Overlap Guard', 'Cùng xe không thể assigned 2 trips overlap time window (departure + duration + 60min buffer).', 'Xe bị double-booked'),
    ('Ledger Immutability', 'LedgerEntry append-only. Không UPDATE/DELETE — DB trigger enforce.', 'Audit trail không tin cậy, balance retroactive altered'),
    ('Payout Account Verification', 'Withdrawal chỉ khi PayoutAccount.verifiedAt IS NOT NULL. Edit reset verifiedAt về null.', 'Kẻ tấn công đổi TK payout rồi rút tiền'),
    ('BigInt Money Arithmetic', 'Mọi currency math nhân integer minor-unit với fractional rate phải dùng BigInt.', 'IEEE 754 drift gây sai rounding'),
    ('T+1 Settlement Delay', 'Revenue không available cho withdrawal cho đến completedAt + 1 ngày.', 'Nhà xe rút tiền trước khi chargeback/dispute xử lý'),
    ('Operator Bookable Gate', 'Hold creation + booking initiation re-verify operator.status = APPROVED.', 'Khách đặt vé trên chuyến nhà xe đã bị suspend'),
    ('Hold Expiry', 'Active holds hết hạn sau 10 phút. Cron job chuyển expired qua FOR UPDATE SKIP LOCKED batch 500.', 'Hold bị bỏ giảm capacity vĩnh viễn'),
    ('Idempotency Guards', 'Mọi state-changing operation critical là idempotent. PaymentEvent @@unique, Booking ON CONFLICT, LedgerEntry sourceEventId unique...', 'Network retries tạo duplicate bookings/payouts'),
])

doc.add_heading('9.5 Máy trạng thái (8 máy trạng thái)', level=2)

doc.add_paragraph('Trip: scheduled → departed → completed → cancelled (từ bất kỳ state nào). salesClosed orthogonal.')
doc.add_paragraph('Booking: awaiting_payment → paid → completed | trip_cancelled | no_show | refunded. payment_failed_expired từ awaiting_payment.')
doc.add_paragraph('Hold: active → consumed | expired | cancelled_trip.')
doc.add_paragraph('Payout: requested → processing → paid | failed. failed → requested (admin retry).')
doc.add_paragraph('Operator: PENDING_REVIEW → UNDER_REVIEW → APPROVED | REJECTED. REJECTED → PENDING_REVIEW (resubmit). APPROVED ↔ SUSPENDED.')
doc.add_paragraph('OTP: Active → Consumed (correct verify) | Lockout sentinel (3 fail → extend expiresAt 15min) | Expired (5min TTL).')
doc.add_paragraph('EInvoice: pending → issued → sent | failed | cancelled.')
doc.add_paragraph('Charter: SUBMITTED → ADMIN_REVIEW → ASSIGNED_DIRECT | PUBLISHED → ACCEPTED → COMPLETED | CANCELLED. Customer hủy trước ACCEPTED.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 10: LỘ TRÌNH CHIẾN LƯỢC
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 10: LỘ TRÌNH CHIẾN LƯỢC & SỔ ĐĂNG KÝ RỦI RO', level=1)

doc.add_heading('10.1 Lộ trình hành động theo giai đoạn', level=2)

add_table(doc, ['Giai đoạn', '#', 'Hành động', 'Tại sao', 'Effort'],
[
    ('Trước ra mắt', '1', 'Xây luồng hủy & hoàn tiền khách hàng', '#1 khiếu nại user. Mọi đối thủ có. Bản án uy tín nếu thiếu.', 'M'),
    ('', '2', 'Xin ý kiến pháp lý giấy phép IPS', 'T+3 có thể bất hợp pháp. Quyết định go/no-go kiến trúc thanh toán hiện tại.', 'S'),
    ('', '3', 'Giải quyết lưu trú dữ liệu cho PII', 'Vi phạm NĐ 53/2022. Vercel SG không tuân thủ. PostgreSQL tại VN cho PII.', 'M'),
    ('', '4', 'Xây luồng đặt vé khứ hồi', 'Mọi đối thủ chính có. Thiếu = gấp đôi friction. Hạ tầng paired-return đã có.', 'M'),
    ('', '5', 'Thêm Zalo ZNS cho xác nhận booking', 'Chỉ email không đủ cho user VN. Zalo 70M+ user, ZNS open rate cao, rẻ hơn SMS.', 'S-M'),
    ('', '6', 'Xây trang "Vé của tôi"', 'Mọi nền tảng có account có trang này. Dữ liệu có — UI read-only.', 'S'),
    ('Tháng 1-3', '7', 'Hoàn thành hồ sơ pháp lý', 'Đăng ký MOIT, DPO, DPIA, đăng ký standard-form contract, CDTIA.', 'M'),
    ('', '8', 'Xây engine mã giảm giá / voucher', 'Khách VN cực kỳ nhạy giá. Micro-discount 20-50K VND chuyển conversion đáng kể.', 'M'),
    ('', '9', 'Thêm ZaloPay làm phương thức thanh toán', '20M user, embedded trong Zalo (70M+), tăng trưởng nhanh.', 'S-M'),
    ('', '10', 'Xây booking link chia sẻ + Facebook embed widget', 'Lợi thế CAC lớn nhất BB = nhà xe lái traffic đến trang branded. Cho họ công cụ.', 'S'),
    ('', '11', 'Thêm hỗ trợ tiếng Anh', 'Khóa hoàn toàn segment du khách quốc tế. Quan trọng cho tuyến du lịch.', 'M'),
    ('Tháng 3-6', '12', 'Đàm phán đối tác phân phối MoMo/ZaloPay', 'Đường nhanh nhất đến volume khách không cần app native. FUTA tiếp cận hàng chục triệu qua MoMo.', 'M'),
    ('', '13', 'Xây flash sales/promotions nhà xe tự cấu hình', 'VeXeRe BMS cho nhà xe self-serve early-bird, last-minute. Quan trọng lấp xe ngày thấp.', 'M'),
    ('', '14', 'Xây push notifications / nhắc chuyến', 'Giảm no-show, cải thiện CX. Nếu ZNS đã tích hợp, trip reminders là incremental.', 'S'),
    ('', '15', 'Triển khai đánh giá & nhận xét nhà xe', 'Tín hiệu tin cậy. 69% user MoMo thường xuyên đánh giá. Giúp khách chọn trên cùng tuyến.', 'M'),
    ('Tháng 6-12', '16', 'Đàm phán đối tác bảo hiểm du lịch', 'Chỉ VeXeRe có. 20K VND/vé, OTA kiếm 30-40% premium. Cơ hội mở rộng margin rõ nhất.', 'M'),
    ('', '17', 'Xây tích hợp API 12Go/Bookaway', 'Tuyến nhà xe BB hiển thị cho du khách quốc tế không cần BB xây marketing du lịch.', 'M'),
    ('', '18', 'Đánh giá app native vs PWA', 'Quyết định phụ thuộc MoMo/ZaloPay partnership. Nếu có → hoãn native. Nếu không → PWA.', 'M-L'),
    ('', '19', 'Bắt đầu mạng đại lý/reseller', 'Mạng 5,000 đại lý VeXeRe là phân phối offline gần-zero CAC. Cần đủ kho (200+ nhà xe).', 'L'),
    ('', '20', 'Khám phá Zalo mini-app', 'Không nền tảng xe bus liên tỉnh nào có Zalo mini-app. First-mover trong hệ sinh thái 70M user.', 'M-L'),
])

doc.add_heading('10.2 Sổ đăng ký rủi ro thị trường', level=2)

add_table(doc, ['#', 'Rủi ro', 'Khả năng', 'Tác động', 'Giảm thiểu'],
[
    ('1', 'T+3 settlement phân loại IPS không giấy phép', 'TB', 'NGUY HIỂM', 'Xin ý kiến pháp lý. Tái cấu trúc PSP split-settlement.'),
    ('2', 'Thực thi lưu trú dữ liệu (BCA)', 'TB', 'CAO', 'Di chuyển DB PII sang hạ tầng VN.'),
    ('3', 'VeXeRe áp lực độc quyền nhà xe', 'TB', 'CAO', 'Định vị BB bổ sung ("kênh thương hiệu riêng") không cạnh tranh. Không yêu cầu độc quyền.'),
    ('4', 'Đỉnh Tết phá hạ tầng', 'TB', 'CAO', 'Load test 10x baseline. DB connection pooling + read replicas.'),
    ('5', 'Nhà xe rời vì demand online thấp', 'CAO', 'CAO', 'Focus tuyến du lịch demand online đã chứng minh. Cung cấp marketing tools. Đẩy MoMo/ZaloPay.'),
    ('6', 'Tranh chấp hoàn tiền/hủy phá uy tín', 'CAO (nếu thiếu refund flow)', 'CAO', 'Xây refund flow tự động trước ra mắt. Set chính sách hủy cấu hình theo nhà xe.'),
    ('7', 'redBus cắt giá hoa hồng', 'TB', 'TB', 'Khác biệt hóa trên brand ownership + SaaS value. Hoa hồng chỉ là một phần value prop.'),
    ('8', 'SEO traction chậm', 'CAO', 'TB', 'De-prioritize SEO. Dựa vào traffic nhà xe + phân phối super-app.'),
    ('9', 'Hạn chế sở hữu nước ngoài bất ngờ', 'THẤP', 'CAO', 'Đăng ký dưới IT/technology codes. Giữ tách biệt pháp lý giữa nền tảng và thanh toán.'),
    ('10', 'Super-app (MoMo/ZaloPay) tự xây SaaS nhà xe', 'THẤP', 'CAO', 'Họ là công ty thanh toán, không phải SaaS. Xây phần mềm quản lý nhà xe xa core competency.'),
])

doc.add_heading('10.3 Thông tin thị trường người dùng — 10 Insight quan trọng nhất', level=2)

add_table(doc, ['#', 'Insight', 'Ý nghĩa cho sản phẩm'],
[
    ('1', 'Hoàn tiền = #1 pain point mọi nền tảng. VeXeRe 2.7★ Trustpilot.', 'Hoàn tiền tự động 24-48h. Chính sách hoàn tiền nổi bật trên mọi xác nhận booking. Trust differentiator.'),
    ('2', 'Khách VN so sánh giá nhiều tab song song. VeXeRe hold 10 phút vì lý do này.', 'Hold-then-pay flow đúng kiến trúc. Hold 10-15 phút cạnh tranh. Hiển thị countdown "giữ cho bạn X phút".'),
    ('3', 'Cashback 20-50K VND chuyển first-time users. MoMo/ZaloPay chứng minh.', 'Xây engine promo/voucher sớm. Giảm giá 20K VND booking đầu tiên platform-funded có tác động chuyển conversion lớn.'),
    ('4', 'Khan hiếm vé Tết = pain point cảm xúc nhất. Vé hết 1-3 tháng trước.', 'Quản lý kho robust, không oversell. Xem xét waitlist/notification cho chuyến hết. Tết = cửa sổ thu hút nhà xe mạnh nhất.'),
    ('5', 'Xác nhận booking tức thì = tín hiệu tin cậy. 60 giây tách nền tảng uy tín.', 'BB đã gửi email. Thêm Zalo/SMS trong 60 giây. Màn hình "Đã xác nhận" rõ ràng với ref, số ghế, chi tiết đón.'),
    ('6', 'Segment premium (25-35, đô thị) trả 1.5-2x cho xe limousine mini 9-16 chỗ.', 'Console nhà xe hỗ trợ loại xe limousine/VIP với tier giá riêng. Feature premium nổi bật trong search — AOV cao cải thiện unit economics.'),
    ('7', 'Nhà xe nhỏ marketing qua Facebook pages và Zalo groups. Đặt qua Zalo là mặc định.', 'Value prop BB: "số hóa điều bạn đã làm trên Zalo/Facebook" — không "bỏ kênh hiện tại". Feature: tạo booking link chia sẻ trên Facebook/Zalo.'),
    ('8', 'An toàn xe bus rất nhạy cảm. Report tai nạn viral trên MXH gây chuyển nhu cầu.', 'Hiển thị trạng thái xác minh giấy phép nhà xe trên trang booking. Xem xét badge "nhà xe xác minh".'),
    ('9', 'Giao tiếp sau đặt vé gần như không có. Thay đổi gì khách không biết trừ khi gọi hotline.', 'Xây thông báo nhà xe→hành khách (Zalo/SMS). "Xe bạn trễ 30 phút" sẽ là differentiator thực sự.'),
    ('10', 'Email rất ít quan trọng cho khách xe bus điển hình. Dùng cho khiếu nại chính thức, không nhắn giao dịch.', 'De-prioritize email. Zalo (ZNS) chính, SMS phụ, email thứ ba. Chỉ email qua Resend hiện tại là gap.'),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  PHẦN 11: TỔNG HỢP QUYẾT ĐỊNH
# ════════════════════════════════════════════════════════════════════════

doc.add_heading('PHẦN 11: TỔNG HỢP QUYẾT ĐỊNH', level=1)
doc.add_paragraph('Bảng tổng hợp tất cả quyết định chiến lược trong báo cáo, với phương án, đánh đổi, lựa chọn và lý do.')

add_table(doc, ['#', 'Quyết định', 'Phương án', 'Lựa chọn', 'Lý do chính'],
[
    ('1', 'Hành lang beachhead', 'TH↔TPHCM vs HCMC-ĐL vs HN-SP', 'Thanh Hóa ↔ TPHCM', 'GMV/booking cao nhất, nhu cầu lao động ổn định, OTA cạnh tranh thấp'),
    ('2', 'Định vị chiến lược', 'Shopify vs VeXeRe lite vs SaaS thuần', '"Shopify cho Nhà Xe"', 'Khác biệt duy nhất; nhà xe lớn đã đầu tư stack riêng vì brand control'),
    ('3', 'Mức hoa hồng', '5% vs 8-10% vs 12-15%', '8-10% (5% intro 3 tháng)', 'Dưới VeXeRe, vẫn có lãi; 5% intro tạo incentive không rủi ro'),
    ('4', 'Cấu trúc giá', 'Chỉ commission vs Chỉ SaaS vs Kép', 'Kép (commission + SaaS)', 'Phù hợp mọi quy mô nhà xe; validated bởi VeXeRe'),
    ('5', 'Hủy/hoàn tiền', 'Không build vs Build trước vs Thủ công', 'Build TRƯỚC ra mắt', '#1 pain point; mọi đối thủ có; launch blocker'),
    ('6', 'App native', 'Native vs PWA vs Super-app distribution', 'Hoãn native; MoMo/ZaloPay trước; PWA bridge', 'Mobile web viable; FUTA chứng minh super-app reach'),
    ('7', 'Pháp nhân', 'LLC vs JSC', 'LLC (chuyển JSC tại Series A)', 'Đơn giản hơn; chuyển đổi được'),
    ('8', 'Mã VSIC', 'Công nghệ vs Vận tải vs Cả hai', 'Công nghệ + TMĐT, KHÔNG vận tải', '100% sở hữu nước ngoài; tiền lệ Grab/VeXeRe'),
    ('9', 'Luồng tiền thanh toán', 'Marketplace vs Merchant vs Hybrid', 'Marketplace (PSP → TK nhà xe)', 'Không cần giấy phép SBV; tiền lệ VeXeRe 10+ năm'),
    ('10', 'PSP tích hợp', 'Chỉ VNPay vs Chỉ MoMo vs Cả 3', 'VNPay + MoMo + VietQR', 'Phủ gần 100% preferences; adapter architecture có sẵn'),
    ('11', 'Nhà cung cấp HĐĐT', 'MISA vs VNPT vs Viettel vs FPT', 'MISA meInvoice', 'Đã tích hợp (#74); phổ biến nhất cho SME'),
    ('12', 'Vai trò xuất HĐĐT', 'Nền tảng ủy quyền vs Nhà xe tự xuất vs Principal', 'Nền tảng ủy quyền (authorized model)', 'Tiêu chuẩn ngành; NĐ 123 Đ.17 cho phép; nhà xe nhỏ không có khả năng tự xuất'),
    ('13', 'Lưu trú dữ liệu', 'Toàn SG vs Hybrid SG+VN vs Toàn VN', 'Hybrid: Vercel SG compute + DB VN cho PII', 'Tuân thủ NĐ 53/2022; giữ lợi thế Vercel; +5-15ms chấp nhận được'),
    ('14', 'Kênh OTP/thông báo', 'Chỉ SMS vs Chỉ ZNS vs ZNS+SMS', 'ZNS primary → SMS fallback (eSMS)', 'Tiết kiệm 50-70%; mô hình ngành chuẩn (Grab, Tiki)'),
    ('15', 'Phân loại nền tảng', 'Kinh doanh vận tải vs Nền tảng công nghệ', 'Nền tảng công nghệ', 'Tiền lệ VeXeRe 10+ năm; 100% sở hữu nước ngoài'),
])

doc.add_paragraph()
doc.add_paragraph()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— HẾT BÁO CÁO —')
r.bold = True
r.font.name = FONT_NAME
r.font.size = Pt(14)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Tổng hợp từ 40 tài liệu nghiên cứu kinh doanh | documentation/business/\n'
           'Ngày tạo: 17/06/2026 | Phiên bản 1.0 | NỘI BỘ')
r2.font.name = FONT_NAME
r2.font.size = FONT_SIZE

# ── Save ────────────────────────────────────────────────────────────────
output_path = r'D:\Bus-Booking\documentation\business\BAO-CAO-KINH-DOANH-TONG-HOP.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.0f} KB')
