# -*- coding: utf-8 -*-
"""May dinh dang .docx dung chung cho moi bo dung Word cua du an.

VI SAO TON TAI. `build_huong_dan_docx.py` dung tai lieu ngay o CAP MODULE — cac
lenh `H(...)`, `TBL(...)` chay khi import — nen khong bo dung nao khac import
duoc no. Truoc file nay, mot bo dung Word thu hai chi con mot duong: CHEP.

Cai khong duoc phep chep la `lien_ket()`. No mang mot luat mua bang mot lan
hong: khi `relate_to` khong cap duoc rId thi DUNG HAN, khong im lang tra ve chu
tran. Ban .docx la ban cho NGUOI doc soat, va mot URL khong bam duoc trong ban
danh cho nguoi la mot lien ket chet — dung lop loi "mot co che chet nam canh
mot loi chu thich noi rang no dang chay" da ghi trong so loi ngay 30/07. Chep
ham do sang file thu hai la cach hai ban bat dau lech nhau.

Cai KHONG nam o day: `value_run` va `field_table`. Chung gan chat voi ngu nghia
`[CHƯA XÁC MINH]` cua bo huong dan (bo dong khi gia tri chua xac minh, dem
truong theo nhom) — do la CHON LOC doi lot dinh dang, va bo dung khac co chinh
sach khac. Ep chung thanh dung chung se buoc bo dung thu hai nhan mot chinh
sach no khong co.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as _RT

RED = RGBColor(0xC0, 0x1C, 0x1C)        # chua xac minh — cam noi ra
AMBER = RGBColor(0xB0, 0x6A, 0x00)      # suy dien — phai rao
GREY = RGBColor(0x5A, 0x5A, 0x5A)       # da xac minh — nguon va ngay
INK = RGBColor(0x1A, 0x1A, 0x1A)
LINK = RGBColor(0x0B, 0x4F, 0xA0)


class Khoi:
    """Mot tai lieu Word dang dung + may dinh dang cua no."""

    def __init__(self, font="Calibri", co=10):
        self.doc = Document()
        st = self.doc.styles["Normal"]
        st.font.name = font
        st.font.size = Pt(co)

    # ── van ban ────────────────────────────────────────────────────────────
    def H(self, t, lvl=1):
        h = self.doc.add_heading(t, level=lvl)
        # Khong de tieu de mac ket mot minh o cuoi trang, cach doi noi dung
        # sang trang sau. `keep_with_next` la API san cua python-docx.
        h.paragraph_format.keep_with_next = True
        return h

    def P(self, t="", bold=False, italic=False, size=10, color=None):
        p = self.doc.add_paragraph()
        r = p.add_run(t)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)
        r.font.color.rgb = color or INK
        return p

    def B(self, t):
        self.doc.add_paragraph(t, style="List Bullet")

    # ── lien ket ───────────────────────────────────────────────────────────
    def lien_ket(self, par, url, chu=None):
        """Chen mot HYPERLINK THAT vao doan van (python-docx khong co san).

        `relate_to` cap mot rId trong phan quan he cua tai lieu. Neu no khong
        cap duoc thi DUNG HAN — xem docstring dau file.
        """
        rid = par.part.relate_to(url, _RT.HYPERLINK, is_external=True)
        if not rid:
            raise SystemExit(f"DUNG — khong cap duoc rId cho lien ket: {url}")
        h = OxmlElement("w:hyperlink")
        h.set(qn("r:id"), rid)
        run = OxmlElement("w:r")
        pr = OxmlElement("w:rPr")
        for ten, khoa, gia in (("w:color", "w:val", "0B4FA0"),
                               ("w:u", "w:val", "single"),
                               ("w:sz", "w:val", "19")):
            el = OxmlElement(ten)
            el.set(qn(khoa), gia)
            pr.append(el)
        run.append(pr)
        t = OxmlElement("w:t")
        t.text = chu or url
        run.append(t)
        h.append(run)
        par._p.append(h)

    # ── bang ───────────────────────────────────────────────────────────────
    def shade(self, cell, hexcolor):
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), hexcolor)
        cell._tc.get_or_add_tcPr().append(el)

    def lap_dong_dau(self, t):
        """Lap dong tieu de tren moi trang khi bang tran sang trang sau.

        Khong ap cho bang co NHIEU dong tieu de nhom rai rac (vd `field_table`
        cua bo huong dan): dat co nay se dun moi tieu de nhom len dau moi trang
        tiep — sai han.
        """
        if not t.rows:
            return
        trPr = t.rows[0]._tr.get_or_add_trPr()
        el = OxmlElement("w:tblHeader")
        el.set(qn("w:val"), "true")
        trPr.append(el)

    def TBL(self, headers, data, widths=None, size=8.5):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        # `autofit` mac dinh True, nghia la Word TINH LAI chieu rong khi mo file
        # va bo qua moi `cell.width` dat o duoi. Tat no thi cac so kia moi co
        # tac dung.
        t.autofit = False
        self.lap_dong_dau(t)
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ""
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(size)
        for row in data:
            cs = t.add_row().cells
            for i, v in enumerate(row):
                cs[i].text = ""
                r = cs[i].paragraphs[0].add_run(str(v))
                r.font.size = Pt(size)
        if widths:
            for rr in t.rows:
                for i, wd in enumerate(widths):
                    rr.cells[i].width = Cm(wd)
        self.doc.add_paragraph()
        return t
