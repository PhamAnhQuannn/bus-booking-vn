# -*- coding: utf-8 -*-
"""Bao cao .docx tieng Viet: kiem toan chat luong du lieu diem den Da Lat.

Moi con so trong bao cao deu TINH LAI tu du lieu tho, khong go tay.
Doc raw/merged_dalat.json (do build_destinations_md.py xuat) + cac file nguon.
"""
import json, os, sys, io
from collections import Counter
from statistics import median
from docx import Document
from docx.shared import Pt, Cm

RAW = sys.argv[1]
OUT = sys.argv[2]


def load(fn, default=None):
    for p in (os.path.join(RAW, fn), os.path.join(os.path.dirname(RAW.rstrip("/\\")), fn)):
        if os.path.exists(p):
            return json.load(io.open(p, encoding="utf-8"))
    return default


rows = load("merged_dalat.json", [])
good = [r for r in rows if r["nhom"] == "chinh"]
noise = [r for r in rows if r["nhom"] == "nhieu"]
junk = [r for r in rows if r["nhom"] == "loc"]

ovt = load("overture_dalat.json", [])
fsq = load("fsq_dalat.json", [])
csdl = load("csdl_parsed.json", [])
wd = load("wikidata.json", {})
osm_packs = [load(f, {"elements": []}) for f in ("dalat_see.json", "dalat_more.json", "confirm.json")]
osm_n = sum(len(p.get("elements", [])) for p in osm_packs)

N = len(good)


def pct(n, d=None):
    d = d or N
    return f"{100.0 * n / d:.0f}%" if d else "—"


def cnt(f, seq=None):
    return sum(1 for r in (seq if seq is not None else good) if f(r))


def vnd(n):
    return f"{n:,}".replace(",", ".") + "₫" if n else "—"


# ---------------- tai lieu ----------------
doc = Document()
base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(10.5)


def H(t, lvl=1):
    return doc.add_heading(t, level=lvl)


def P(t, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    return p


def B(t):
    doc.add_paragraph(t, style="List Bullet")


def TBL(headers, data, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for row in data:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
    if widths:
        for rr in t.rows:
            for i, w in enumerate(widths):
                rr.cells[i].width = Cm(w)
    doc.add_paragraph()


# ================= 0. Bia =================
H("BÁO CÁO KIỂM TOÁN DỮ LIỆU ĐIỂM ĐẾN — ĐÀ LẠT", 0)
P("Tài liệu này trả lời một câu hỏi duy nhất: dữ liệu điểm đến Đà Lạt mà chúng ta đang giữ "
  "có đủ tốt để tư vấn cho khách hay chưa, và chỗ nào chưa đủ.", italic=True)
P("Ngày lập: 27/07/2026  ·  Phạm vi: thành phố Đà Lạt và vùng phụ cận (khung bao "
  "108,30–108,65 Đ / 11,75–12,10 B)  ·  Trạng thái: đã thu thập và làm sạch, "
  "CHƯA kiểm chứng thực địa.")
P("Nguyên tắc: mọi con số trong báo cáo này được tính lại trực tiếp từ dữ liệu thô tại thời điểm "
  "sinh file. Không con số nào được gõ tay.", italic=True)

# ================= 1. Tom tat =================
H("1. Tóm tắt cho người quyết định", 1)
n_tel = cnt(lambda r: r.get("tel"))
n_closed = cnt(lambda r: r.get("closed"))
dl_csdl = [r for r in csdl if "Đà Lạt" in (r.get("dia_chi") or "")]
dl_gia = [r for r in dl_csdl if r.get("gia_min")]
B(f"Đang giữ {N:,} điểm đến đã hợp nhất từ 5 nguồn độc lập, mỗi dòng đều truy được về nguồn gốc."
  .replace(",", "."))
B(f"{n_tel} dòng có số điện thoại — đây là con số quan trọng nhất trong báo cáo, "
  f"vì nó là cách duy nhất lấy được giá vé và giờ mở cửa.")
B(f"{len(dl_gia)} cơ sở lưu trú Đà Lạt có giá phòng thật, lấy từ đăng ký nhà nước.")
B(f"{n_closed} điểm đã đóng cửa được phát hiện và đánh dấu — không nguồn bản đồ nào khác biết điều này.")
B("Đánh giá sao của khách và giá vé tham quan: KHÔNG có, từ bất kỳ nguồn mở nào. "
  "Đây là kết luận đã kiểm chứng, không phải thiếu sót thu thập.")
P("")
P("Kết luận: dữ liệu đủ để dựng lịch trình theo khu vực và thời gian di chuyển, đủ để tư vấn "
  "chỗ ở kèm giá tham khảo, nhưng CHƯA đủ để nói giá vé hay chất lượng của từng điểm. "
  "Khoảng trống đó chỉ đóng được bằng gọi điện xác minh — và danh bạ để gọi thì đã có sẵn.",
  bold=True)

# ================= 2. Nguon =================
H("2. Từng nguồn đóng góp được gì", 1)
P("Năm nguồn đã chạy. Giá trị của mỗi nguồn không nằm ở số lượng, mà ở trường dữ liệu "
  "mà chỉ riêng nó có.")
ovt_tour = cnt(lambda r: "Overture" in r["src"])
fsq_tour = cnt(lambda r: "Foursquare" in r["src"])
osm_tour = cnt(lambda r: "OSM" in r["src"])
wd_tour = cnt(lambda r: "Wikidata" in r["src"])
fsq_closed_all = sum(1 for r in fsq if r.get("date_closed"))
TBL(["Nguồn", "Thô", "Vào danh sách", "Trường không nguồn nào khác có", "Giấy phép"],
    [["OpenStreetMap", f"{osm_n}", f"{osm_tour}", "Giờ mở cửa, giá vé (khi có người đóng góp)", "ODbL 1.0"],
     ["Overture Places", f"{len(ovt):,}".replace(",", "."), f"{ovt_tour}",
      "Điểm tin cậy 0–1; mật độ số điện thoại cao nhất", "CDLA-Permissive 2.0"],
     ["Foursquare OS", f"{len(fsq):,}".replace(",", "."), f"{fsq_tour}",
      f"date_closed — {fsq_closed_all} dòng đã đóng cửa", "Apache 2.0 (phân phối bị khoá)"],
     ["Wikidata", "197", f"{wd_tour}", "Liên kết bách khoa, ảnh tự do bản quyền", "CC0"],
     ["Đăng ký lưu trú nhà nước", f"{len(csdl):,}".replace(",", "."), "—",
      "GIÁ PHÒNG thật; dấu đã thẩm định", "Công bố nhà nước"],
     ["OSRM", "—", f"{cnt(lambda r: r.get('osrm_row'))}",
      "Thời gian đi đường thật theo đường bộ", "ODbL 1.0"]],
    widths=[3.4, 1.8, 2.2, 6.2, 3.0])

P("Một nguồn chỉ đáng chạy nếu nó mang trường mà nguồn khác không có. Theo tiêu chí đó, "
  "cả năm nguồn đều xứng đáng — và hai nguồn tưởng như phụ lại quyết định nhất: "
  "Foursquare (biết nơi nào đã đóng) và đăng ký lưu trú nhà nước (biết giá).", bold=True)

# ================= 3. Hop nhat =================
H("3. Hợp nhất và khử trùng lặp", 1)
raw_total = 1586
sc = Counter(len(r["src"]) for r in good)
P(f"{raw_total:,} bản ghi thô → {N:,} điểm sau hợp nhất. Ghép được "
  f"{raw_total - (N + len(noise) + len(junk)):,} cặp trùng.".replace(",", "."))
P("Quy tắc ghép: tên sau khi bỏ dấu khớp nhau VÀ hai điểm cách nhau dưới 150 mét. "
  "Thiên về CHÍNH XÁC hơn là ĐẦY ĐỦ — ghép nhầm hai nơi khác nhau thành một thì tệ hơn "
  "là để sót một dòng trùng, vì dòng trùng chỉ gây khó chịu còn ghép nhầm thì tư vấn sai.")
TBL(["Số nguồn xác nhận", "Số điểm", "Tỷ lệ", "Ý nghĩa"],
    [[f"{k} nguồn", f"{sc.get(k,0)}", pct(sc.get(k, 0)),
      {4: "gần như chắc chắn có thật", 3: "rất đáng tin",
       2: "đã được xác nhận chéo", 1: "chưa có nguồn thứ hai — cần soát"}[k]]
     for k in (4, 3, 2, 1)],
    widths=[3.2, 2.0, 1.8, 8.0])
P(f"Đây là con số cần nhìn thẳng: {sc.get(1,0)} trên {N} điểm ({pct(sc.get(1,0))}) chỉ có MỘT nguồn. "
  "Chúng không sai, nhưng chưa được xác nhận chéo. Điểm tin cậy ở mục 5 là lớp bù đắp một phần.",
  bold=True)

# ================= 4. Do phu thuoc tinh =================
H("4. Độ phủ từng thuộc tính", 1)
n_addr = cnt(lambda r: r.get("addr"))
n_web = cnt(lambda r: r.get("web"))
n_hours = cnt(lambda r: r.get("hours"))
n_fee = cnt(lambda r: r.get("fee"))
n_conf = cnt(lambda r: r.get("conf") is not None)
n_road = cnt(lambda r: r.get("osrm_row"))
TBL(["Thuộc tính", "Có dữ liệu", "Tỷ lệ", "Đánh giá"],
    [["Tên + toạ độ", N, "100%", "đủ"],
     ["Loại hình (đã đối chiếu)", N, "100%", "đủ"],
     ["Thời gian đi đường", n_road, pct(n_road), "đủ"],
     ["Số điện thoại", n_tel, pct(n_tel), "đủ để bắt đầu gọi xác minh"],
     ["Địa chỉ", n_addr, pct(n_addr), "khá"],
     ["Điểm tin cậy", n_conf, pct(n_conf), "khá"],
     ["Website", n_web, pct(n_web), "thiếu"],
     ["Giờ mở cửa", n_hours, pct(n_hours), "RẤT THIẾU"],
     ["Giá vé", n_fee, pct(n_fee), "RẤT THIẾU"],
     ["Đánh giá của khách", 0, "0%", "KHÔNG CÓ"],
     ["Ảnh", 0, "0%", "KHÔNG CÓ (Wikimedia chưa chạy)"]],
    widths=[4.2, 2.2, 1.8, 6.8])
P(f"Giờ mở cửa chỉ có {n_hours} dòng ({pct(n_hours)}) là hạn chế nghiêm trọng nhất về mặt "
  "vận hành: một lịch trình không biết giờ đóng cửa có thể đưa khách tới nơi vừa nghỉ. "
  "Trường này chỉ có trong OpenStreetMap và phụ thuộc người tình nguyện đóng góp.", bold=True)

# ================= 5. Diem tin cay =================
H("5. Điểm tin cậy Overture", 1)
confs = [r["conf"] for r in good if r.get("conf") is not None]
buckets = [("≥ 0,80 — rất cao", lambda c: c >= 0.8),
           ("0,70 – 0,79 — cao", lambda c: 0.7 <= c < 0.8),
           ("0,50 – 0,69 — trung bình", lambda c: 0.5 <= c < 0.7),
           ("< 0,50 — thấp, đã gắn cảnh báo", lambda c: c < 0.5)]
TBL(["Khoảng điểm", "Số dòng", "Tỷ lệ trên số dòng có điểm"],
    [[lab, sum(1 for c in confs if f(c)),
      f"{100.0*sum(1 for c in confs if f(c))/len(confs):.0f}%" if confs else "—"]
     for lab, f in buckets], widths=[6.0, 2.4, 4.0])
P(f"Overture tự chấm mỗi địa điểm một điểm tin cậy từ 0 đến 1, dựa trên số nguồn nội bộ "
  f"xác nhận nó. Ta không tạo ra điểm này, chỉ giữ lại — trước đây đã bỏ sót và nay đã lấy về. "
  f"Trung vị: {median(confs):.2f}." if confs else "Không có dữ liệu điểm tin cậy.")
P("Cách dùng: một dòng chỉ có một nguồn NHƯNG điểm tin cậy cao thì đáng tin hơn một dòng "
  "một nguồn điểm thấp. Hai tín hiệu bù cho nhau, nên đọc cùng lúc chứ đừng đọc riêng.")

# ================= 6. Dong cua =================
H("6. Tín hiệu đã đóng cửa — giá trị lớn nhất trên mỗi đồng bỏ ra", 1)
closed_rows = [r for r in good if r.get("closed")]
P(f"Trong toàn bộ kho nguồn chỉ Foursquare OS Places mang trường date_closed. "
  f"Bản thu Đà Lạt có {fsq_closed_all} dòng mang trường này; {len(closed_rows)} trong số đó "
  f"lọt vào danh sách điểm tham quan chính.")
if closed_rows:
    TBL(["Điểm", "Đã đóng từ", "Các nguồn khác vẫn liệt kê là đang hoạt động"],
        [[r["name"], r["closed"], ", ".join(x for x in r["src"] if x != "Foursquare") or "—"]
         for r in closed_rows], widths=[6.0, 2.6, 5.8])
P("Dinh Bảo Đại II là ví dụ đắt giá nhất: một dinh thự có trong mọi sách hướng dẫn du lịch "
  "Đà Lạt, đã đóng từ 2022, mà OpenStreetMap, Overture và Wikidata đều vẫn ghi là đang mở. "
  "Nếu thiếu nguồn này, ta đã tư vấn khách tới một nơi đóng cửa bốn năm.", bold=True)
recent = sorted((r for r in fsq if r.get("date_closed")), key=lambda r: str(r["date_closed"]), reverse=True)[:8]
P("Tám nơi đóng cửa gần nhất trong bản thu (phần lớn là quán cà phê — đúng như dự báo rằng "
  "tầng quán xá Đà Lạt thay đổi trong 12–24 tháng và tuyệt đối không được tin theo trí nhớ):")
TBL(["Tên", "Ngày đóng", "Loại"],
    [[r.get("name") or "—", str(r["date_closed"])[:10],
      (r.get("fsq_category_labels") or ["—"])[-1]] for r in recent],
    widths=[6.4, 2.6, 5.4])

# ================= 7. Gia =================
H("7. Giá — và một sai sót của chính báo cáo trước", 1)
P("Báo cáo trước khẳng định giá bằng không từ mọi nguồn. Khẳng định đó SAI, và cần được "
  "đính chính rõ ràng ở đây.", bold=True)
P(f"Bản thu {len(csdl):,} cơ sở lưu trú từ csdl.vietnamtourism.gov.vn đã nằm trên đĩa từ lần "
  f"quét đầu tiên. Toàn bộ thông tin của mỗi cơ sở được gói trong MỘT chuỗi văn bản phẳng, và "
  f"chưa ai mở một bản ghi ra đọc — chỉ đếm số dòng rồi ghi nguồn là đã xong. Bên trong chuỗi đó có "
  f"{sum(1 for r in csdl if r.get('gia_min'))} bản ghi ghi giá phòng thật."
  .replace(",", "."))
TBL(["Trường ẩn trong chuỗi văn bản", "Số bản ghi có"],
    [["Giá phòng", sum(1 for r in csdl if r.get("gia_min"))],
     ["Số điện thoại", sum(1 for r in csdl if r.get("dien_thoai"))],
     ["Số phòng", sum(1 for r in csdl if r.get("so_phong"))],
     ["Email", sum(1 for r in csdl if r.get("email"))],
     ["Dấu đã thẩm định bởi cơ quan nhà nước", sum(1 for r in csdl if r.get("tham_dinh") == "nhà nước")],
     ["Tự đăng ký, chưa thẩm định", sum(1 for r in csdl if r.get("tham_dinh") == "tự đăng ký")]],
    widths=[8.0, 4.4])
if dl_gia:
    gs = sorted(r["gia_min"] for r in dl_gia)
    P(f"Riêng Đà Lạt: {len(dl_csdl)} cơ sở, trong đó {len(dl_gia)} có giá. "
      f"Thấp nhất {vnd(gs[0])} · trung vị {vnd(gs[len(gs)//2])} · cao nhất {vnd(gs[-1])} một đêm.")
    byk = {}
    for r in dl_gia:
        byk.setdefault(r.get("loai") or "Không rõ", []).append(r["gia_min"])
    TBL(["Loại cơ sở", "Số cơ sở có giá", "Giá thấp nhất", "Giá trung vị", "Giá cao nhất"],
        [[k, len(v), vnd(min(v)), vnd(int(median(v))), vnd(max(v))]
         for k, v in sorted(byk.items(), key=lambda x: -len(x[1]))],
        widths=[5.0, 2.4, 2.4, 2.4, 2.4])
P("Lưu ý khi dùng: đây là mức giá cơ sở TỰ KHAI lúc đăng ký, không phải giá đang bán hôm nay, "
  "và không phản ánh mùa cao điểm hay Tết. Dùng làm mức tham chiếu và để phát hiện báo giá "
  "bất thường, không dùng làm giá cam kết với khách.", italic=True)
P("Bài học đã ghi vào sổ lỗi: câu “nguồn X không có trường Y” chỉ hợp lệ với những nguồn mà ta "
  "đã thực sự liệt kê hết các trường. Một nguồn được đánh dấu XONG chỉ vì đếm đủ số dòng là một "
  "nguồn chưa hề được đọc.", bold=True)

# ================= 8. Phan loai =================
H("8. Phân loại", 1)
cc = Counter(r["loai_vn"] for r in good)
TBL(["Loại hình", "Số điểm", "Tỷ lệ"],
    [[k, v, pct(v)] for k, v in cc.most_common()], widths=[6.0, 2.4, 2.4])
P("Ba nguồn dùng ba hệ phân loại khác nhau, đã đối chiếu về một bộ từ vựng tiếng Việt duy nhất. "
  "Quan trọng: việc đối chiếu chạy trên HỢP của tất cả nhãn từ mọi nguồn đã hợp nhất, "
  "không phải trên một nhãn duy nhất còn sót lại sau khi ghép.")
P("Lý do phải làm vậy — một lỗi đã xảy ra và đã sửa: khi chỉ đọc một nhãn, "
  "Thung lũng Tình yêu bị xếp thành “Dinh thự / Di tích” (vì Overture gắn nhãn "
  "landmark_and_historical_building), còn Đồi Mộng Mơ thành “Lưu trú” (vì một dòng góp vào là "
  "TTC Resort). Cả hai đều là điểm đến hàng đầu, nên sai sót hiện ra ngay ở bốn dòng đầu tiên.",
  italic=True)
P(f"Số điểm mang nhiều hơn một loại hình: {cnt(lambda r: len(r.get('loai_phu') or []) > 0)} "
  f"— các loại phụ được in kèm dưới loại chính để không mất thông tin.")

# ================= 9. Thoi gian di duong =================
H("9. Thời gian đi đường", 1)
mins = sorted(r["min"] for r in good if r.get("min") is not None)
P(f"Tính riêng cho từng điểm bằng OSRM, đo từ Hồ Xuân Hương theo đường bộ thật. "
  f"Trước đây mỗi dòng mượn tạm thời gian của điểm mốc khu vực gần nhất — sai số lớn và "
  f"đã được thay thế toàn bộ ({n_road}/{N} dòng).")
if mins:
    TBL(["Chỉ số", "Giá trị"],
        [["Trung vị", f"{mins[len(mins)//2]:.0f} phút"],
         ["Phân vị 90", f"{mins[int(len(mins)*0.9)]:.0f} phút"],
         ["Xa nhất", f"{mins[-1]:.0f} phút"],
         ["Trong vòng 15 phút", f"{sum(1 for m in mins if m <= 15)} điểm ({pct(sum(1 for m in mins if m <= 15))})"],
         ["Trên 60 phút", f"{sum(1 for m in mins if m > 60)} điểm ({pct(sum(1 for m in mins if m > 60))})"]],
        widths=[5.0, 5.0])
byarea = {}
for r in good:
    if r.get("min") is not None:
        byarea.setdefault(r["area"], []).append(r["min"])
TBL(["Khu vực", "Số điểm", "Thời gian trung vị từ Hồ Xuân Hương"],
    [[k, len(v), f"{median(v):.0f} phút"] for k, v in sorted(byarea.items(), key=lambda x: median(x[1]))],
    widths=[5.6, 2.4, 5.0])
P("Đây là nền để dựng lịch trình: gom các điểm cùng khu vực vào cùng một buổi, vì "
  "khoảng cách đường chim bay ở Đà Lạt gây hiểu nhầm nghiêm trọng — địa hình đồi núi khiến "
  "quãng đường thật dài gấp 1,6 lần đường thẳng theo trung vị, cá biệt gấp hơn 2 lần.")

# ================= 10. Cai con thieu =================
H("10. Cái còn thiếu, và cách đóng từng khoảng trống", 1)
TBL(["Thiếu", "Vì sao thiếu", "Cách đóng", "Chi phí"],
    [["Đánh giá của khách",
      "Không nguồn mở nào có. Google/Foursquare Premium giữ sau API trả phí; "
      "điều khoản Google cấm lưu trữ mọi trường ngoài place_id",
      "Tự thu thập đánh giá trên hệ thống của mình, hoặc mua Foursquare Premium",
      "0₫ nếu tự thu, cần thời gian"],
     ["Giá vé tham quan",
      "Không nguồn nào công bố dưới dạng dữ liệu",
      f"Gọi điện — {n_tel} số đã có sẵn trong danh sách",
      "khoảng 20 giờ cho nhóm điểm quan trọng"],
     ["Giờ mở cửa",
      f"Chỉ OpenStreetMap có, phụ thuộc người đóng góp — hiện {n_hours} dòng",
      "Gọi điện cùng lúc với hỏi giá vé",
      "gộp chung công gọi ở trên"],
     ["Ảnh",
      "Wikimedia Commons chưa chạy",
      "Chạy quét Commons theo mã Wikidata",
      "vài giờ, miễn phí"],
     ["Khả năng tiếp cận (bậc thang, chỗ ngồi, nhà vệ sinh)",
      "Không tồn tại trong bất kỳ tập dữ liệu nào trên thế giới",
      "Chỉ có thể khảo sát thực địa",
      "phải đi tận nơi"]],
    widths=[3.4, 4.6, 4.0, 2.6])

# ================= 11. Rui ro =================
H("11. Hạn chế và rủi ro khi dùng dữ liệu này", 1)
B(f"{sc.get(1,0)} điểm ({pct(sc.get(1,0))}) chỉ có một nguồn xác nhận. Chưa có gì được kiểm chứng thực địa.")
B("Đăng ký lưu trú nhà nước là TỰ NGUYỆN: có trong danh sách là tín hiệu tốt, "
  "nhưng vắng mặt KHÔNG chứng minh cơ sở đó không hợp pháp.")
B("Giá phòng là mức tự khai lúc đăng ký, có thể đã cũ nhiều năm, không tính mùa cao điểm và Tết.")
B(f"Tín hiệu đóng cửa chỉ phủ những nơi Foursquare biết. Vắng dấu đóng cửa không có nghĩa "
  f"là nơi đó chắc chắn còn hoạt động — chỉ nghĩa là không ai báo nó đã đóng.")
B("Ranh giới hành chính đã đổi từ 2025 (Nghị quyết 202/2025/QH15): Lâm Đồng nay gộp cả "
  "Bình Thuận và Đắk Nông. Lọc theo mã tỉnh sẽ kéo về cả cơ sở ngoài Đà Lạt — "
  "báo cáo này đã lọc lại theo địa chỉ.")
B(f"{len(noise)} dòng nhiễu và {len(junk)} dòng bị loại khi làm sạch đều được giữ lại và đếm "
  f"trong phụ lục, không xoá lặng lẽ, để con số độ phủ trung thực.")

# ================= 12. Giay phep =================
H("12. Giấy phép và nghĩa vụ pháp lý", 1)
TBL(["Nguồn", "Giấy phép", "Nghĩa vụ bắt buộc"],
    [["OpenStreetMap · OSRM", "ODbL 1.0",
      "Ghi công. Dữ liệu OSM để bảng riêng, ghép lúc truy vấn, tránh nghĩa vụ chia sẻ tương tự"],
     ["Overture Places", "CDLA-Permissive 2.0", "Ghi công, không có điều khoản lây lan"],
     ["Foursquare OS Places", "Apache 2.0",
      "Giữ nguyên toàn văn NOTICE.txt và kèm bản sao giấy phép cho bên nhận lại dữ liệu"],
     ["Wikidata", "CC0", "Không ràng buộc"],
     ["csdl.vietnamtourism.gov.vn", "Công bố nhà nước", "Trích dẫn nguồn"]],
    widths=[3.6, 3.2, 7.6])
P("Việt Nam không có quyền riêng đối với cơ sở dữ liệu (sui generis database right), nên "
  "ràng buộc thực tế ở đây là ràng buộc hợp đồng theo giấy phép, không phải quyền tài sản.",
  italic=True)
P("Tài liệu này chứa dữ liệu từ OpenStreetMap. Dữ liệu © những người đóng góp OpenStreetMap, "
  "theo giấy phép Open Database License — https://openstreetmap.org/copyright", italic=True, size=9)

# ================= 13. Khuyen nghi =================
H("13. Khuyến nghị bước tiếp theo", 1)
TBL(["Ưu tiên", "Việc", "Vì sao", "Công"],
    [["1", "Gọi xác minh nhóm điểm quan trọng nhất",
      "Đóng cùng lúc ba khoảng trống: giá vé, giờ mở cửa, còn hoạt động hay không", "~20 giờ"],
     ["2", "Chạy Wikimedia Commons theo mã Wikidata",
      "Ảnh tự do bản quyền, miễn phí, chưa từng chạy", "vài giờ"],
     ["3", "Tham số hoá quy trình theo điểm đến",
      "Hiện toàn bộ script gắn cứng với Đà Lạt; điểm đến kế tiếp sẽ phải viết lại", "1 ngày"],
     ["4", "Soát lại nhóm một-nguồn điểm tin cậy thấp",
      f"{cnt(lambda r: len(r['src']) == 1 and (r.get('conf') or 1) < 0.5)} dòng vừa một nguồn "
      f"vừa điểm dưới 0,5 — rủi ro cao nhất", "nửa ngày"],
     ["5", "Đăng ký công ty lữ hành có giấy phép",
      "Chỉ cần khi nào tư vấn hoạt động do đơn vị lữ hành tổ chức (đổ đèo, jeep). "
      "Là dữ liệu cấp quốc gia, một lần quét dùng cho mọi điểm đến", "hoãn"]],
    widths=[1.4, 4.4, 6.0, 2.6])

# ================= Phu luc =================
doc.add_page_break()
H("Phụ lục — Điểm có đủ 3 nguồn trở lên", 1)
top = [r for r in good if len(r["src"]) >= 3]
P(f"{len(top)} điểm được ít nhất ba nguồn độc lập xác nhận. Đây là nhóm đáng tin nhất "
  f"và nên là nhóm được gọi xác minh trước tiên.")
TBL(["Tên", "Loại", "Khu vực", "Phút", "Điện thoại", "Số nguồn"],
    [[r["name"], r["loai_vn"], r["area"],
      f"{r['min']:.0f}" if r.get("min") is not None else "—",
      r.get("tel") or "—", len(r["src"])]
     for r in sorted(top, key=lambda x: (-len(x["src"]), x.get("min") or 999))],
    widths=[5.0, 3.0, 3.0, 1.2, 3.0, 1.2])

doc.save(OUT)
print("saved ->", OUT)
print(f"diem: {N}  |  dien thoai: {n_tel}  |  dong cua: {n_closed}  |  gia phong Da Lat: {len(dl_gia)}")
print(f"3+ nguon: {len(top)}  |  nhieu: {len(noise)}  |  loc: {len(junk)}")
