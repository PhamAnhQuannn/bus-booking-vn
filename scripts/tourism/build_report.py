# -*- coding: utf-8 -*-
"""Sinh bao cao .docx tieng Viet: tham dinh chien luoc thu thap du lieu + du lieu Da Lat."""
import json, os, sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from places_dalat import PLACES, AREA_NAMES

RAW = sys.argv[1]
OUT = sys.argv[2]
osrm = json.load(open(os.path.join(RAW, "osrm_matrix.json"), encoding="utf-8"))
D, T = osrm["distances"], osrm["durations"]
idx = {p[0]: i for i, p in enumerate(PLACES)}
nm = {p[0]: p[1] for p in PLACES}

doc = Document()
base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(10.5)


def H(t, lvl=1):
    return doc.add_heading(t, level=lvl)


def P(t, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def B(t):
    doc.add_paragraph(t, style="List Bullet")


def TBL(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9.5)
    if widths:
        for rr in t.rows:
            for i, w in enumerate(widths):
                rr.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# ---------- TRANG BIA ----------
ttl = doc.add_heading("BÁO CÁO THẨM ĐỊNH CHIẾN LƯỢC THU THẬP DỮ LIỆU", 0)
ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Điểm đến thí điểm: ĐÀ LẠT (tỉnh Lâm Đồng)")
r.bold = True
r.font.size = Pt(14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Ngày lập: 27/07/2026  ·  Lần chạy thứ nhất  ·  Tài liệu nội bộ").italic = True
doc.add_paragraph()
P("Câu hỏi mà báo cáo này trả lời:", bold=True)
P("“Chiến lược thu thập dữ liệu mà chúng ta đã thiết kế có thực sự lấy được toàn bộ thông tin cần "
  "thiết để tư vấn một lịch trình chi tiết cho khách hàng hay không?”", italic=True)
doc.add_paragraph()
P("Giá trị của báo cáo nằm ở cả những gì các nguồn KHÔNG cung cấp được, không chỉ ở những gì thu về.",
  italic=True)
doc.add_page_break()

# ---------- PHAN 1 ----------
H("PHẦN 1 — THẨM ĐỊNH CHIẾN LƯỢC", 1)

H("1.1 Kết luận", 2)
P("Chiến lược hoạt động — nhưng mới phủ được khoảng một nửa số trường thông tin cần thiết.", bold=True)
P("Đã chạy thật 5 trên 10 nguồn. Kết quả chính:")
B("ĐỘT PHÁ LỚN NHẤT — OSRM cho ma trận thời gian đi đường thật, đủ 1.225/1.225 ô, miễn phí. "
  "Đây là trường dữ liệu quan trọng nhất mà trước đó hoàn toàn trống.")
B("Đăng ký lưu trú CHÍNH THỨC lấy được: 1.230 cơ sở tại Lâm Đồng, kèm phân hạng sao nhà nước.")
B("Wikidata (CC0) bổ sung hai trường không nguồn nào khác có: ảnh tự do bản quyền và website chính thức.")
B("VẪN KHÔNG CÓ, và không nguồn nào cung cấp được: khả năng tiếp cận cho người cao tuổi "
  "(bậc thang, chỗ ngồi, nhà vệ sinh), giá phòng và giá món thực tế, cấu hình phòng cho nhóm 7 người.")
B("Chưa chạy: Overture, Foursquare OS, quanlyluhanh.vn, Wikimedia Commons, thống kê chi tiết.")

H("1.2 Bảng đối chiếu thông tin", 2)
P("Từng trường thông tin đã xác định là cần thiết, đối chiếu với thực tế thu được.", italic=True)
TBL(["Trường thông tin", "Có?", "Nguồn", "Độ phủ đo được"], [
    ["Địa điểm tồn tại, toạ độ", "CÓ", "OpenStreetMap", "1.467 thực thể"],
    ["Khoảng cách & thời gian đi đường", "CÓ (mới)", "OSRM", "1.225/1.225 ô = 100%"],
    ["Phân hạng sao chính thức", "CÓ (mới)", "csdl.vietnamtourism.gov.vn", "1.230 cơ sở"],
    ["Ảnh tự do bản quyền", "CÓ (mới)", "Wikidata P18", "45/197 = 23%"],
    ["Website chính thức", "CÓ (mới)", "Wikidata P856", "9/197"],
    ["Giờ mở cửa (nhà hàng)", "MỘT PHẦN", "OpenStreetMap", "13,3%"],
    ["Giờ mở cửa (điểm tham quan)", "YẾU", "OpenStreetMap", "8/37 điểm chính"],
    ["Giá vé vào cửa", "RẤT YẾU", "OpenStreetMap", "3/114 = 1,8%"],
    ["Món ăn đặc sản theo vùng", "CÓ", "Hội đồng tác nhân AI", "11 món"],
    ["Cảnh báo “bẫy du lịch”", "CÓ", "Hội đồng tác nhân AI", "6 địa điểm"],
    ["Giá phòng / giá món", "KHÔNG", "—", "0%"],
    ["Đánh giá (rating)", "KHÔNG", "—", "0%"],
    ["Bậc thang, chỗ ngồi, nhà vệ sinh", "KHÔNG", "—", "0% — chỉ khảo sát thực địa"],
    ["Phòng nhóm 7 người, thang máy", "KHÔNG", "—", "0%"],
], [6.2, 2.0, 4.6, 3.4])

H("1.3 Hai lỗi tự phát hiện trong quá trình chạy", 2)
P("Cả hai đã được sửa trước khi đưa vào báo cáo. Ghi lại vì đây là bài học phương pháp.", italic=True)
P("Lỗi 1 — Ánh xạ mã lọc bị ĐẢO NGƯỢC.", bold=True)
P("Trên csdl.vietnamtourism.gov.vn, ô lọc rate[] có value=1 ứng với nhãn “5 sao”, còn value=5 ứng với "
  "“1 sao” — thứ tự giá trị NGƯỢC với thứ tự nhãn. Lần chạy đầu tôi đọc xuôi và ra kết quả "
  "“294 khách sạn 5 sao ở Lâm Đồng”. Con số đó vô lý, nên tôi kiểm tra lại và phát hiện lỗi.")
P("Bài học: luôn đối chiếu kết quả của một bộ lọc với một phép thử hợp lý trước khi tin nó. "
  "Thứ tự value trong form có thể ngược với thứ tự nhãn hiển thị.", italic=True)
P("Lỗi 2 — Khoảng cách đường chim bay đánh giá thấp quãng đường thật.", bold=True)
P("Trước đây tôi dùng khoảng cách đường chim bay tính từ toạ độ. So với đường thật của OSRM, tỷ lệ "
  "trung vị là 1,59 lần, phân vị 90 lên tới 2,15 lần. Một số cặp sai lệch nghiêm trọng — xem mục 1.5.")
doc.add_page_break()

H("1.4 Kết quả từng nguồn", 2)
TBL(["#", "Nguồn", "Trạng thái", "Thu được"], [
    ["1", "OSRM (định tuyến đường bộ)", "ĐÃ CHẠY", "Ma trận 35×35, đủ 1.225 ô"],
    ["2", "csdl.vietnamtourism.gov.vn", "ĐÃ CHẠY", "1.230 cơ sở lưu trú, 83 trang"],
    ["3", "Wikidata SPARQL (CC0)", "ĐÃ CHẠY", "197 thực thể bán kính 25 km"],
    ["4", "OpenStreetMap / Overpass", "ĐÃ CHẠY", "1.467 thực thể, 4 truy vấn"],
    ["5", "Hội đồng tác nhân AI", "ĐÃ CHẠY", "5 tác nhân độc lập"],
    ["6", "quanlyluhanh.vn", "CHƯA CHẠY", "Đã kiểm tra kết nối HTTP 200"],
    ["7", "Wikimedia Commons", "CHƯA CHẠY", "Đã kiểm tra kết nối HTTP 200"],
    ["8", "thongke.tourism.vn (chi tiết)", "CHƯA CHẠY", "Mới có số liệu cấp tỉnh"],
    ["9", "Overture Places", "CHƯA CHẠY", "Cần cài pyarrow"],
    ["10", "Foursquare OS Places", "CHƯA CHẠY", "Cần cài pyarrow"],
], [1.2, 5.2, 3.0, 6.8])

H("Phân hạng sao CHÍNH THỨC — Lâm Đồng (sau sáp nhập, mã 68+60+67)", 3)
TBL(["Hạng", "Số cơ sở"], [
    ["5 sao", "9"], ["4 sao", "50"], ["3 sao", "48"], ["2 sao", "146"], ["1 sao", "294"],
    ["Đạt chuẩn phục vụ khách du lịch", "27"], ["Khác / chưa xếp hạng", "526"], ["TỔNG", "1.230"],
], [7.0, 4.0])
P("Lưu ý quan trọng: đăng ký xếp hạng là TỰ NGUYỆN theo Luật Du lịch 2017. Danh sách này là một tập "
  "con, không phải điều tra toàn bộ. Có sao là dấu hiệu tích cực; KHÔNG có sao không chứng minh điều "
  "gì. Có 130/1.230 cơ sở (11%) ở dạng “tự đăng ký”, chưa qua thẩm định nhà nước.", italic=True)
P("Giá trị phụ đáng chú ý: mã tỉnh của Lâm Đồng trên cổng này là “68,60,67” — tức ba mã tỉnh cũ "
  "(Lâm Đồng, Bình Thuận, Đắk Nông) gộp làm một. Chính hệ thống nhà nước xác nhận việc sáp nhập 2025.",
  italic=True)
doc.add_page_break()

H("1.5 Định tuyến đường bộ đã làm thay đổi kế hoạch như thế nào", 2)
P("Ba cặp địa điểm từng được ghép chung một buổi dựa trên đường chim bay, nhưng đường thật bác bỏ:",
  bold=True)
TBL(["Cặp địa điểm", "Chim bay", "Đường thật", "Kết luận"], [
    ["Hồ Tuyền Lâm → Đường hầm điêu khắc", "1,7 km", "14,2 km / 21,5 phút", "KHÔNG ghép được — lệch 8,4 lần"],
    ["Núi Lang Biang → Chợ Lạc Dương", "4,7 km", "19,4 km / 46,9 phút", "KHÔNG ghép được"],
    ["Đồi trà Cầu Đất → Chùa Linh Phước", "9,4 km", "13 phút", "NGƯỢC LẠI: ghép rất tốt"],
], [6.4, 2.4, 3.6, 4.0])
P("Phát hiện quan trọng nhất — Núi Lang Biang.", bold=True)
P("Đường chim bay 11,4 km khiến nó có vẻ gần. Đường thật: 23,8 km và 46,4 phút mỗi chiều. Riêng thời "
  "gian ngồi xe đã gần 1 giờ 40 phút. Đây là hoạt động CHIẾM TRỌN MỘT NGÀY, không phải một buổi sáng "
  "như kế hoạch cũ. Với đoàn có người cao tuổi, điều này thay đổi hoàn toàn bài toán.")
P("Và một đảo chiều tích cực: Đồi trà Cầu Đất từng bị loại vì “biệt lập”. Thực tế chỉ cách Chùa Linh "
  "Phước 13 phút đi đường — hai điểm này ghép thành một buổi hợp lý.", italic=True)

H("1.6 Những gì không nguồn nào cung cấp được", 2)
TBL(["Thông tin", "Vì sao không có", "Cách duy nhất để có", "Chi phí"], [
    ["Bậc thang, chỗ ngồi cho người già, nhà vệ sinh ngồi hay xổm",
     "Không tồn tại trong bất kỳ tập dữ liệu nào", "Khảo sát thực địa", "~10 điểm, vài ngày"],
    ["Khoảng cách từ chỗ đỗ xe đến cổng", "Không có trong dữ liệu bản đồ", "Khảo sát thực địa", "cùng chuyến"],
    ["Xe 45 chỗ có đỗ được không", "Không có", "Khảo sát / hỏi nhà xe", "thấp"],
    ["Giá phòng, giá món thực tế",
     "OSM, Overture, Foursquare OS đều không có. Google có nhưng cấm lưu trữ",
     "Gọi điện, chụp thực đơn, hoặc mua Foursquare Premium", "~10 phút/cơ sở"],
    ["Phòng chứa 7 người, có thang máy", "Không có", "Gọi điện cho khách sạn", "~5 phút/cơ sở"],
    ["Giá vé vào cửa (111/114 điểm)", "OSM chỉ có 1,8%", "Gọi điện hoặc đến nơi", "~10 phút/điểm"],
], [4.6, 5.0, 4.4, 2.4])
doc.add_page_break()

# ---------- PHAN 2 ----------
H("PHẦN 2 — DỮ LIỆU ĐÀ LẠT", 1)
P("Toàn bộ số liệu dưới đây đều có nguồn và ngày lấy. Không có giá hay đánh giá gắn với cơ sở cụ thể.",
  italic=True)

H("2.1 Khoảng cách và thời gian đi đường thật, tính từ Hồ Xuân Hương", 2)
P("Nguồn: OSRM trên dữ liệu OpenStreetMap, lấy ngày 27/07/2026.", italic=True)
ref = idx["V-01"]
rows = []
for pid, name, lat, lon, area in PLACES:
    j = idx[pid]
    if j == ref:
        continue
    rows.append((D[ref][j], [pid, name, AREA_NAMES[area],
                             "%.1f km" % (D[ref][j] / 1000.0),
                             "%.0f phút" % (T[ref][j] / 60.0)]))
rows.sort()
TBL(["Mã", "Địa điểm", "Khu vực", "Đường bộ", "Thời gian"], [r[1] for r in rows],
    [1.6, 5.4, 4.2, 2.2, 2.2])

H("2.2 Các cặp địa điểm ghép được trong một buổi", 2)
P("Tính từ ma trận OSRM. Đây là căn cứ để xếp lịch, thay cho phỏng đoán.", italic=True)
pairs = [("V-05", "V-06"), ("V-10", "V-09"), ("V-07", "V-15"), ("V-16", "V-15"),
         ("V-26", "V-27"), ("V-26", "V-25"), ("V-20", "V-22"), ("V-19", "V-20"),
         ("V-24", "V-22"), ("V-32", "V-31"), ("V-13", "V-14"), ("V-31", "V-15")]
TBL(["Địa điểm A", "Địa điểm B", "Đường bộ", "Thời gian"],
    [[nm[a], nm[b], "%.1f km" % (D[idx[a]][idx[b]] / 1000.0),
      "%.0f phút" % (T[idx[a]][idx[b]] / 60.0)] for a, b in pairs],
    [5.0, 5.0, 2.6, 2.6])

H("2.3 Cảnh báo mùa vụ", 2)
B("Tháng 11–4: khô ráo, trời trong, đẹp nhất — NHƯNG các thác nước ít nước.")
B("Tháng 5–10: mưa chiều hằng ngày, thác nhiều nước, đường đèo dễ sạt lở.")
B("Tháng 12–2: lạnh nhất, ban đêm có thể xuống một chữ số. Người cao tuổi cảm nhận rõ.")
B("30/4 và Tết: giá tăng 2–3 lần, QL20 kẹt xe, phải đặt trước nhiều tháng. Tết 2027: thứ Bảy 06/02/2027.")
B("Giờ đóng cửa dễ làm hỏng lịch: cáp treo 17:00 · Thiền Viện Trúc Lâm nghỉ trưa 11:30–13:30 · "
  "Dinh III nghỉ 11:00–13:00 · Bảo tàng Lâm Đồng nghỉ 11:30–13:30 · Thung lũng Tình Yêu đóng 17:15.")

H("2.4 Địa điểm gây tranh cãi — nên nói thẳng với khách", 2)
P("Nguồn: hội đồng tác nhân AI, chưa kiểm chứng. Hãy trình bày sự bất đồng, đừng kết luận thay khách.",
  italic=True)
TBL(["Địa điểm", "Sách hướng dẫn nói", "Người địa phương nói"], [
    ["Thung lũng Tình Yêu", "Đáng đi, có xe điện", "Đắt, trang trí tim nhựa — nên bỏ qua"],
    ["Tàu hoả cổ → Chùa Linh Phước", "Combo đặc trưng Đà Lạt", "Chậm, đắt, chùa trang trí loè loẹt"],
    ["Đồi trà Cầu Đất", "Điểm chụp ảnh kinh điển", "Nay rào lại thu phí chỗ trước đây miễn phí"],
    ["Thác Cam Ly", "Trung tâm, có lịch sử", "Nước ô nhiễm nhiều năm — dân địa phương né"],
    ["Núi Lang Biang", "Có xe jeep lên đài quan sát", "Chỉ là bãi đỗ xe có view, trừ khi leo sớm"],
    ["Đồi Mộng Mơ", "Vườn chủ đề", "Nghi đã đóng cửa hoặc xuống cấp — phải kiểm tra"],
], [4.6, 5.4, 6.4])

H("2.5 Đề xuất bước tiếp theo", 2)
TBL(["#", "Việc cần làm", "Chi phí", "Đóng được khoảng trống nào"], [
    ["1", "Cài pyarrow, chạy Overture + Foursquare OS", "1 lệnh cài",
     "Tín hiệu date_closed — biết nơi nào đã đóng cửa"],
    ["2", "Chạy quanlyluhanh.vn", "vài giờ", "Danh sách công ty lữ hành có giấy phép"],
    ["3", "Wikimedia Commons theo QID", "vài giờ", "Ảnh tự do bản quyền cho từng điểm"],
    ["4", "KHẢO SÁT THỰC ĐỊA khoảng 10 điểm", "vài ngày", "Khả năng tiếp cận — không có cách nào khác"],
    ["5", "Gọi điện lấy giá vé", "~10 phút/điểm", "111/114 điểm còn thiếu giá vé"],
], [1.2, 6.0, 2.6, 6.6])

H("2.6 Nguồn và giấy phép", 2)
TBL(["Nguồn", "Giấy phép", "Nghĩa vụ"], [
    ["OpenStreetMap", "ODbL 1.0", "Bắt buộc ghi công; dữ liệu OSM để bảng riêng, ghép khi truy vấn"],
    ["OSRM (chạy trên dữ liệu OSM)", "ODbL 1.0", "Như trên"],
    ["Wikidata", "CC0", "Không ràng buộc"],
    ["csdl.vietnamtourism.gov.vn", "Công bố nhà nước", "Trích dẫn nguồn"],
    ["Overture Places", "CDLA-Permissive 2.0", "Không copyleft"],
    ["Foursquare OS Places", "Apache 2.0", "Giữ file NOTICE"],
], [5.0, 3.6, 7.8])
doc.add_paragraph()
P("Tài liệu này chứa dữ liệu từ OpenStreetMap. Dữ liệu © những người đóng góp OpenStreetMap, theo "
  "giấy phép Open Database License — https://openstreetmap.org/copyright", italic=True, size=9)

doc.save(OUT)
print("Da luu:", OUT)
print("So doan:", len(doc.paragraphs), "| So bang:", len(doc.tables))
