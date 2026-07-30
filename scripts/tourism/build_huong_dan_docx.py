# -*- coding: utf-8 -*-
"""Ban .docx cua huong dan diem den Da Lat.

Doc raw/guide_data.json do build_huong_dan.py xuat ra — MOT lan chon, MOT lan
hop nhat, hai dinh dang dau ra. Khong lam lai logic chon/hop nhat o day.

═══════════════════════════════════════════════════════════════════════════════
DOCSTRING CU CUA FILE NAY LA SAI, VA DAY LA BAN SUA.

No ghi: "Word lam duoc mot viec Markdown khong lam duoc: TO MAU. Ca gia tri cua
tai lieu nay nam o cho [CHƯA XÁC MINH] khong the bo qua duoc — mau do lam duoc
dieu do, chu in dam thi khong."

Co che do DA CHET. `field_table()` bo MOI dong co gia tri bat dau bang UNV
TRUOC khi render, va `value_run()` chi duoc goi tu trong `field_table()` — nen
nhanh to do trong `value_run` khong bao gio chay duoc, va vong quet the cua no
lap tren mot tuple rong. Chinh sach da doi tu "danh dau bang mau" sang "bo han
dong", va cau tren khong duoc sua theo.

VAY BAN .docx CON DE LAM GI: no la ban cho NGUOI doc soat, khong phai cho tac
nhan AI (ban .md moi la ban cua tac nhan — muc 0 noi ro the). Mau va bong to
khong co nghia gi voi mot bo trich van ban, nhung co nghia voi mot nguoi dang
quet tim cho thieu truoc khi goi dien xac minh o muc 9.

Tin hieu cho thieu gio nam o CAP NHOM, khong o cap dong: moi nhom in mot dong
xam "— N/M truong da xac minh —". Nguoi doc thay do day cua khoang trong ma
khong phai doc 1.336 dong [CHƯA XÁC MINH].
═══════════════════════════════════════════════════════════════════════════════
"""
import json, os, sys, io
from collections import Counter, defaultdict
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import hoat_dong_data as _hoat_dong   # CUNG module chon loc voi ban .md
import an_ngu_data as _an_ngu

# Duong ra co MAC DINH, va do la de chan mot loi da xay ra: ten file tung la
# lua chon cua tung lan goi, nen hai phien lam viec song song da sinh ra
# `-v4.docx` va `-v5.docx` ben canh `-v3.docx`. Ket qua la so phien ban NGUOC
# voi do moi — v5 dung 14:59 truoc khi cac ban va vao, v3 dung 16:55 sau do —
# va nguoi doc mo file so cao nhat lai thay ban cu nhat, thieu han lop mon.
# Mot duong mac dinh nghia la goi khong tham so thi GHI DE ban chinh thay vi
# them mot so moi.
OUT_MAC_DINH = "docs/Huong-Dan-Da-Lat.docx"
RAW = sys.argv[1]
OUT = sys.argv[2] if len(sys.argv) > 2 else OUT_MAC_DINH
_LAN_CAN = _an_ngu.tai_lan_can(RAW)
_LCKV = _an_ngu.tai_lan_can_khu_vuc(RAW)   # khoi theo KHU VUC, in mot lan
# ── THU TU CHAY LA LOAD-BEARING, va day la cho no de vo im lang ────────────
# `guide_data.json` do build_huong_dan.py ghi ra; file nay chi DOC. Nen chay
# rieng bo dung .docx se dung lai bo diem cua lan chay .md gan nhat — khong loi,
# khong canh bao, chi la du lieu cu. Viec tach logic chon ra `diem_den_data.py`
# se xoa han lop loi nay va da duoc hoan sang commit rieng (do la thay doi rui
# ro nhat: DL-xx la id THEO VI TRI nen mot phep sort lech se doi id cua ca 36
# diem, va 11 script khac gan du lieu theo id do).
# Trong khi cho: kiem tuoi file. Re, va bien truong hop im lang thanh on ao.
_gp = os.path.join(RAW, "guide_data.json")
if not os.path.exists(_gp):
    raise SystemExit(f"KHONG co {_gp}.\nChay build_huong_dan.py truoc — no sinh ra"
                     " file nay, ban .docx chi doc lai.")
for _phu in ("enrichment.json", "lan_can_khu_vuc.json", "lan_can.json"):
    _pp = os.path.join(RAW, _phu)
    if os.path.exists(_pp) and os.path.getmtime(_pp) > os.path.getmtime(_gp) + 1:
        raise SystemExit(
            f"DUNG — {_phu} moi hon guide_data.json.\n"
            "Nghia la du lieu nguon da doi nhung bo diem chua duoc chon lai, nen"
            " ban .docx se\ndung bo diem cu. Chay lai theo dung thu tu:\n"
            "   python scripts/tourism/build_huong_dan.py .tourism-data/raw\n"
            "   python scripts/tourism/build_huong_dan_docx.py .tourism-data/raw")
G = json.load(io.open(_gp, encoding="utf-8"))
picked, NEAR, mat = G["picked"], G["near"], G["matrix"]
BUILD_DATE = G["build_date"]
byid = {r["id"]: r for r in picked}
UNV = "[CHƯA XÁC MINH]"

# ── SO MUC: mot cho duy nhat, giong ban .md ────────────────────────────────
# Lan renumber thu ba trong mot phien; hai lan truoc deu sot tham chieu, mot
# lan sot ngay trong file nay ("muc 2" trong khi ban .md da la "muc 3").
S_QUYTAC, S_TONGQUAN, S_KHUVUC, S_DIAHINH, S_DIEMDEN = 0, 1, 2, 3, 4
S_HOATDONG, S_ANNGU, S_SOSANH = 5, 6, 7
S_TUYEN, S_MATRAN, S_KIEMCHUNG = 8, 9, 10


# lop lam giau — doc cung mot file enrichment.json nhu ban .md
_ep = os.path.join(RAW, "enrichment.json")
ENR = {}
if os.path.exists(_ep):
    for _e in json.load(io.open(_ep, encoding="utf-8")):
        ENR.setdefault(_e["id"], {}).setdefault(_e["field"], _e)


def ev(pid, field, tail=""):
    e = ENR.get(pid, {}).get(field)
    if not e:
        return UNV + tail
    return str(e["value"])


def has(pid, field):
    return field in ENR.get(pid, {})

RED = RGBColor(0xC0, 0x1C, 0x1C)        # chua xac minh — cam noi ra
AMBER = RGBColor(0xB0, 0x6A, 0x00)      # suy dien — phai rao
GREY = RGBColor(0x5A, 0x5A, 0x5A)       # da xac minh — nguon va ngay
INK = RGBColor(0x1A, 0x1A, 0x1A)


# Tai lieu chi danh cho nguoi doc tieng Viet. Ten dia diem trong nguon co the
# kem chu Han/Hangul/Kana/Kirin (vd "Thiên Vương Cổ Sát Chùa Tàu - 大叻市 天王古剎").
# Do la du lieu that, nhung khong danh cho doc gia nay.
_FOREIGN = __import__("re").compile(r"[가-힯一-鿿぀-ヿЀ-ӿ]+")


def vn_only(s):
    """Bo chu ngoai (Han, Hangul, Kana, Kirin). Giu tieng Viet va chu Latin."""
    s = _FOREIGN.sub("", str(s or ""))
    s = __import__("re").sub(r"\s{2,}", " ", s)
    return s.strip(" ,-·/")

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10)


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def H(t, lvl=1):
    h = doc.add_heading(t, level=lvl)
    # Khong de tieu de mac ket mot minh o cuoi trang, cach doi noi dung sang
    # trang sau. `keep_with_next` la API san cua python-docx, khong can OXML.
    # Day la cach nhe hon nhieu so voi "moi the mot trang" — the dai ngan khac
    # nhau nen moi trang mot the vua thua giay vua khong giu duoc loi hua do,
    # con nguoi soat thi can thay nhieu the cung luc.
    h.paragraph_format.keep_with_next = True
    return h


def P(t="", bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    r.font.color.rgb = color or INK
    return p


def B(t):
    doc.add_paragraph(t, style="List Bullet")


def value_run(par, text):
    """To mau theo muc tin cay. Day la ly do ban .docx ton tai."""
    if text.startswith(UNV):
        r = par.add_run(UNV)
        r.bold = True
        r.font.color.rgb = RED
        r.font.size = Pt(9.5)
        rest = text[len(UNV):]
        if rest:
            r2 = par.add_run(rest)
            r2.italic = True
            r2.font.size = Pt(8.5)
            r2.font.color.rgb = RED
        return
    head, tag = text, ""
    for mark in ():
        if mark in text:
            i = text.index(mark)
            head, tag = text[:i], text[i:]
            break
    if head:
        r = par.add_run(head)
        r.font.size = Pt(9.5)
        r.font.color.rgb = INK
    if tag:
        r = par.add_run(tag)
        r.font.size = Pt(8)
        r.italic = True
        r.font.color.rgb = AMBER if tag.startswith("[SUY RA") else GREY


def lap_dong_dau(t):
    """Lap dong tieu de tren moi trang khi bang tran sang trang sau.

    Khong ap cho `field_table`: bang do co NHIEU dong tieu de nhom rai rac chu
    khong phai mot dong dau, nen dat co nay se dun moi tieu de nhom len dau moi
    trang tiep — sai han.
    """
    if not t.rows:
        return
    trPr = t.rows[0]._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def field_table(rows_spec):
    """rows_spec: ('group', title) | ('f', nhan, gia tri) | ('note', text)

    Chi giu truong DA XAC MINH: moi dong co gia tri [CHƯA XÁC MINH] bi bo, va nhom
    nao khong con dong nao thi bo luon tieu de nhom.
    """
    # ── DEM truoc khi loc, de con biet da bo bao nhieu ────────────────────
    # Ly do ton tai cua ban .docx, theo docstring cua chinh no, la TO MAU dong
    # [CHƯA XÁC MINH] cho khong the bo qua. Nhung dong loc ngay duoi xoa han
    # nhung dong do TRUOC khi render, va `value_run` chi duoc goi tu ham nay —
    # nen nhanh to do trong `value_run` khong bao gio chay duoc. Co che ma tai
    # lieu vien dan de ton tai da chet tu luc chinh sach doi tu "danh dau" sang
    # "bo han".
    # Khong dua dong do tro lai (do la quyet dinh co y). Thay vao do dua tin
    # hieu len MOT CAP: moi nhom ghi ro con bao nhieu tren tong bao nhieu. Doc
    # gia thay duoc do day cua khoang trong ma khong phai doc 1.336 dong do.
    tong_nhom, con_nhom, nhom_ht = {}, {}, None
    for x in rows_spec:
        if x[0] == "group":
            nhom_ht = x[1]
            tong_nhom.setdefault(nhom_ht, 0)
            con_nhom.setdefault(nhom_ht, 0)
        elif x[0] == "f" and nhom_ht:
            tong_nhom[nhom_ht] += 1
            if not str(x[2]).startswith(UNV):
                con_nhom[nhom_ht] += 1

    kept = [x for x in rows_spec if not (x[0] == "f" and str(x[2]).startswith(UNV))]
    out = []
    for i, x in enumerate(kept):
        if x[0] == "group":
            nxt = kept[i + 1] if i + 1 < len(kept) else None
            if nxt is None or nxt[0] == "group":
                continue
        out.append(x)
    # Chen dong "con N/M truong" vao cuoi moi nhom con song.
    with_note, nhom_ht = [], None
    for i, x in enumerate(out):
        if x[0] == "group":
            if nhom_ht and tong_nhom.get(nhom_ht, 0) > con_nhom.get(nhom_ht, 0):
                with_note.append(("done", f"{con_nhom[nhom_ht]}/{tong_nhom[nhom_ht]}"
                                          " trường đã xác minh"))
            nhom_ht = x[1]
        with_note.append(x)
    if nhom_ht and tong_nhom.get(nhom_ht, 0) > con_nhom.get(nhom_ht, 0):
        with_note.append(("done", f"{con_nhom[nhom_ht]}/{tong_nhom[nhom_ht]}"
                                  " trường đã xác minh"))
    rows_spec = with_note
    # Danh so o DAY, sau khi da loc. Dem luc dung spec thi van de lai lo: nhom
    # "Luu y quan trong" nhan so 5 roi bi xoa vi ca 7 truong con deu 0/36, va moi
    # ho so hien A.1-A.4 roi nhay sang A.6 — dung cai day so nhay ma viec danh so
    # lai sinh ra de sua, chi dich mot buoc.
    _muc[0] = 0
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for spec in rows_spec:
        cells = t.add_row().cells
        if spec[0] == "group":
            cells[0].merge(cells[1])
            c = t.rows[-1].cells[0]
            c.text = ""
            _muc[0] += 1
            r = c.paragraphs[0].add_run(f"A.{_muc[0]} — {spec[1]}")
            r.bold = True
            r.font.size = Pt(9.5)
            shade(c, "E8EDF3")
            continue
        if spec[0] == "done":
            # Xam nhat, khong phai do: day khong phai canh bao, chi la thuoc do
            # do day. Do danh cho `note` — thu nguoi doc phai hanh dong.
            cells[0].merge(cells[1])
            c = t.rows[-1].cells[0]
            c.text = ""
            r = c.paragraphs[0].add_run("— " + spec[1] + " —")
            r.italic = True
            r.font.size = Pt(7.5)
            r.font.color.rgb = GREY
            continue
        if spec[0] == "note":
            cells[0].merge(cells[1])
            c = t.rows[-1].cells[0]
            c.text = ""
            r = c.paragraphs[0].add_run(spec[1])
            r.italic = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RED
            shade(c, "FDF2F2")
            continue
        cells[0].text = ""
        r = cells[0].paragraphs[0].add_run(spec[1])
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        cells[0].width = Cm(4.6)
        cells[1].text = ""
        value_run(cells[1].paragraphs[1] if False else cells[1].paragraphs[0], spec[2])
        cells[1].width = Cm(11.4)
    doc.add_paragraph()
    return t


def TBL(headers, data, widths=None, size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    # `autofit` mac dinh True, nghia la Word TINH LAI chieu rong khi mo file va
    # bo qua moi `cell.width` dat o duoi. Tat no thi cac so kia moi co tac dung.
    t.autofit = False
    lap_dong_dau(t)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(size)
    for row in data:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = ""
            r = cs[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(size)
    if widths:
        for rr in t.rows:
            for i, wd in enumerate(widths):
                rr.cells[i].width = Cm(wd)
    doc.add_paragraph()
    return t


# ==================================================== bia + quy tac doc
H("HƯỚNG DẪN ĐIỂM ĐẾN ĐÀ LẠT", 0)
P(f"Hồ sơ chi tiết {len(picked)} điểm đến · sinh tự động ngày {BUILD_DATE}", italic=True, size=11)
P("Nguồn: OpenStreetMap · Overture Maps · Foursquare OS · Wikidata · OSRM · "
  "đăng ký lưu trú Cục Du lịch Quốc gia · "
  # Wikipedia PHAI co o day: cac doan mo ta trich nguyen van tu Wikipedia, va
  # CC BY-SA 4.0 doi ghi cong kem lien ket.
  "Wikipedia tiếng Việt (CC BY-SA 4.0, https://vi.wikipedia.org)",
  italic=True, size=9, color=GREY)

H(f"{S_QUYTAC}. QUY TẮC ĐỌC — BẮT BUỘC", 1)
P("QUAN TRỌNG — tài liệu này CHỈ liệt kê những trường ĐÃ XÁC MINH. "
  "Một trường KHÔNG xuất hiện trong hồ sơ nghĩa là CHƯA BIẾT, không phải là “không có”. "
  "Không được suy đoán giá trị cho trường vắng mặt.", bold=True, color=RED)
P("Ba điều tuyệt đối không được làm:", bold=True)
B("KHÔNG thay [CHƯA XÁC MINH] bằng một giá trị thường gặp. “Giờ mở cửa 08:00–17:00” là hình "
  "dạng của một giờ mở cửa, không phải giờ mở cửa của nơi này.")
B("KHÔNG suy ra giá vé, giờ mở cửa, thời lượng thăm hay mức độ dễ đi lại từ loại hình. "
  "Chỉ ba suy diễn được duyệt: trong nhà/ngoài trời, link bản đồ, điểm lân cận.")
B("KHÔNG viết mô tả, “lý do nên đến” hay “điểm nhấn” — tài liệu cố ý KHÔNG sinh "
  "những mục đó vì mọi chữ trong đó sẽ là bịa. Mục 3 liệt kê hoạt động kèm nơi và "
  "đơn vị cụ thể, đó là dữ kiện; “Đà Lạt lãng mạn” thì không.")
P("Nhịp độ mặc định (chuyến thư giãn): tối đa 4 điểm/ngày · tối đa 2 giờ di chuyển/ngày · "
  "mỗi ngày chừa một khoảng trống.", bold=True)
P("Bay flycam: mặc định COI NHƯ BỊ CẤM trừ khi có xác nhận ngược lại. Sai theo hướng an toàn "
  "thì mất một tấm ảnh; sai theo hướng kia thì khách bị phạt.", bold=True, color=RED)

# ==================================================== 1. tong quan
H(f"{S_TONGQUAN}. Tổng quan điểm đến", 1)
TBL(["Mục", "Giá trị"],
    [["Thành phố", "Đà Lạt, tỉnh Lâm Đồng"],
     ["Số điểm trong hồ sơ", str(len(picked))],
     ["Kho dữ liệu đầy đủ", "diem-tham-quan.md — 1.361 điểm"],
     *([["Hướng mặt trời mọc", (_an_ngu.tai_dia_hinh(RAW) or {}).get("binh_minh")
         + " — chung cho cả thành phố, không khác nhau giữa các điểm"]]
       if (_an_ngu.tai_dia_hinh(RAW) or {}).get("binh_minh") else []),
     ["Thời tiết theo tháng", UNV],
     ["Lịch lễ hội", UNV],
     ["Ảnh hưởng Tết", UNV + " — nhiều nơi đóng cửa, giá tăng mạnh"],
     ["Đi lại tới Đà Lạt", UNV + " — chưa thu thập tuyến xe / máy bay"],
     ["Phương tiện tại chỗ", UNV + " — chưa thu thập giá thuê xe / taxi"]],
    widths=[5.0, 11.0], size=9)
P("⚠ Năm hàng cuối là khoảng trống có thật, không phải lỗi hiển thị. Một lịch trình không biết "
  "khách tới bằng gì và đi lại bằng gì thì chưa phải một lịch trình.", bold=True, color=RED)

# ==================================================== 2. ho so tung diem
# ── muc MOI. Cau hoi thuong gap nhat — "di 3-5 ngay thi chia the nao" —
# truoc day khong muc nao tra loi o cap KHU VUC, la cap nguoi ta thuc su chia
# ngay. Moi con so rut tu guide_data.json + lan_can_khu_vuc.json.
doc.add_page_break()
H(f"{S_KHUVUC}. Tổng quan theo khu vực", 1)
P("Chia ngày theo khu vực, không theo từng điểm: các điểm trong một khu vực đủ gần "
  "để đi liền trong cùng buổi.", italic=True, size=9, color=GREY)
_rows_kv = []
for _a in sorted({r["area"] for r in picked}):
    _ps = [r for r in picked if r["area"] == _a]
    _kv = _LCKV.get(_a) or {}
    _n_ks = sum(len(b["khach_san"]) for b in _kv.get("bac_khach_san") or [])
    if _kv.get("khong_co_khach_san"):
        _luu = f"không có — xem mục {S_ANNGU}"
    elif _n_ks:
        _luu = f"{_kv.get('tong_ks', 0)} cơ sở có giá"
    else:
        _luu = "—"
    _rong = _kv.get("ban_kinh_khu_vuc", "—")
    if _kv.get("canh_bao_khoang_cach"):
        _rong += "  ⚠"
    _rows_kv.append([_a, str(len(_ps)),
                     f"{min(r['km'] for r in _ps):.1f} – {max(r['km'] for r in _ps):.1f}",
                     _rong, _luu])
TBL(["Khu vực", "Điểm", "Km từ trung tâm", "Rộng", "Lưu trú trong khu vực"],
    _rows_kv, widths=[4.6, 1.2, 3.2, 2.2, 4.8], size=9)
P("⚠ Khu vực có dấu ⚠ ở cột Rộng thì các điểm trong đó cách nhau xa hơn bán kính 5 km "
  "dùng để tìm cơ sở gần — chúng KHÔNG dùng chung một thị trường lưu trú, nên đừng "
  "gộp vào một đêm nghỉ.", size=8.5, color=AMBER)

# Lop nay do Phase L thu thap (SRTM 30 m) va CHUA TUNG duoc in: do_cao 36/36,
# do_nho 36/36, huong_mo 26/36 nam trong enrichment.json tu 28/07. Lan thu nam
# trong du an nay mot lop du lieu duoc thu roi khong ai doc.
# KHONG co cot phuong vi mat troi moc: gia tri giong nhau o ca 36 dong nen no
# khong phan biet duoc diem nao voi diem nao. No o muc 1, mot dong.
_DH = _an_ngu.tai_dia_hinh(RAW)
if _DH and _DH["hang"]:
    doc.add_page_break()
    H(f"{S_DIAHINH}. Ngắm cảnh · săn mây · chụp ảnh", 1)
    P("Suy ra từ mô hình độ cao SRTM 30 m — KHÔNG phải quan sát thực địa. Độ nhô là "
      "độ cao của điểm trừ trung vị vùng xung quanh: số dương nghĩa là cao hơn cảnh "
      "quan quanh nó nên tầm nhìn thoáng, số âm nghĩa là bị che.",
      italic=True, size=9, color=AMBER)
    TBL(["Điểm", "Độ cao", "Độ nhô", "Hướng mở"],
        [[f"{x['id']} · {x['ten']}"[:38] + ("  ⚠" if x.get("canh_bao") else ""),
          x["do_cao"], x["do_nho"] or "", x["huong_mo"] or "—"]
         for x in _DH["hang"]], widths=[7.0, 2.4, 2.4, 4.2], size=8.5)
    for x in _DH["hang"]:
        if x.get("canh_bao"):
            P(f"⚠ {x['id']} · {x['ten']} — {x['canh_bao']}. Mọi con số ở dòng này nói "
              "về khu cổng, không nói về đỉnh.", bold=True, size=9, color=RED)
    P("Toạ độ Đỉnh Langbiang chưa sửa. Vì vậy dòng cao nhất bảng này là "
      f"{_DH['hang'][0]['ten']} ({_DH['hang'][0]['do_cao']}), không phải Langbiang — "
      "đó là hệ quả của toạ độ sai, không phải sự thật về địa hình Đà Lạt.",
      size=9, color=AMBER)

doc.add_page_break()
H(f"{S_DIEMDEN}. Danh sách điểm đến", 1)
P("Thứ tự các mục trong mỗi hồ sơ là cổng lọc trước, mô tả sau: nhận dạng → khả năng tiếp cận "
  "→ kế hoạch thăm → giờ giấc. Một ràng buộc về đi lại loại bỏ địa điểm trước khi chi tiết "
  "chụp ảnh có ý nghĩa gì.", italic=True)

INDOOR_NOTE = {"trong nhà": "không gian trong nhà", "có mái": "phần lớn có mái che",
               "hỗn hợp": "vừa có mái vừa ngoài trời", "ngoài trời": "địa hình ngoài trời"}
_muc = [0]


def sec(nhan):
    """Nhan tieu muc, DEM theo tung ho so — giong `sec()` cua ban .md.

    Chuoi cung `A.1/A.11/A.9/A.3/A.4/A.10/A.13` de lai lo `A.2`, `A.5`-`A.8`,
    `A.12` khong ton tai o BAT KY ho so nao, va nguoi doc ket luan tai lieu bi
    thieu muc. Dem theo the thi luon lien tuc, ke ca khi mot nhom bi bo loc
    rong xoa han.
    """
    _muc[0] += 1
    return f"A.{_muc[0]} — {nhan}"


def khoi_khu_vuc(kv):
    """Khach san + quan an cho CA KHU VUC, in mot lan o dau khu vuc."""
    v = _LCKV.get(kv)
    if not v:
        return
    P(f"Lưu trú & ăn uống trong khu vực — {v['so_diem']} điểm, cách nhau tối đa "
      f"{v['ban_kinh_khu_vuc']}", bold=True, size=9.5)
    if v.get("canh_bao_khoang_cach"):
        P("⚠ " + v["canh_bao_khoang_cach"], size=8.5, color=AMBER)
    if v.get("khong_co_khach_san"):
        P(v["khong_co_khach_san"] + f" Xem mục {S_ANNGU} để chọn theo bậc giá.",
          size=9, color=AMBER)
    for b in v["bac_khach_san"]:
        P(f"{b['ten']} — {b['tong']} cơ sở trong khu vực, {b['tong_thanh_pho']} "
          "trên toàn Đà Lạt", italic=True, size=9, color=GREY)
        TBL(["Khách sạn", "Giá/đêm", "Cách", "Gần", "Phòng", "Điện thoại", "Thẩm định"],
            [[h["ten"][:30], h["gia"] or "", h["khoang_cach"], h["gan_diem"],
              str(h["so_phong"] or ""), h["dien_thoai"] or "", h["tham_dinh"] or ""]
             for h in b["khach_san"]],
            widths=[4.4, 2.6, 1.5, 1.4, 1.2, 2.6, 2.3], size=8)
    if v["loai_quan"]:
        P(f"Quán ăn — {v['tong_quan']} quán còn mở trong khu vực",
          italic=True, size=9, color=GREY)
        TBL(["Loại", "Quán", "Cách", "Gần", "Điện thoại"],
            [[l["ten"], q["ten"][:32], q["khoang_cach"], q["gan_diem"],
              q["dien_thoai"] or ""]
             for l in v["loai_quan"] for q in l["quan"]],
            widths=[3.0, 5.4, 1.5, 1.4, 2.7], size=8)


cur_area = None
for r in picked:
    if r["area"] != cur_area:
        cur_area = r["area"]
        doc.add_page_break()
        H(f"Khu vực: {cur_area}", 2)
        khoi_khu_vuc(cur_area)
    H(f"{r['id']} · {r['name']}", 3)
    srcs = "+".join(r["src"])
    # Mo ta trich nguyen van tu Wikipedia — cung nguon voi ban .md.
    # Khong dien dat lai: sua van ban CC BY-SA 4.0 la tao "Adapted Material" va
    # lam phat sinh nghia vu chia se tuong tu cho chinh doan da sua.
    if has(r["id"], "mo_ta_wikipedia"):
        _e = ENR[r["id"]]["mo_ta_wikipedia"]
        P(_e["value"], italic=True, size=9.5)
        P(f"— {_e.get('source') or 'Wikipedia tiếng Việt'} · trích nguyên văn · "
          f"{_e.get('date') or ''}" + (f" · {_e['url']}" if _e.get("url") else ""),
          italic=True, size=8, color=GREY)
    spec = [("group", "Nhận dạng"),
            ("f", "Tên", r["name"]),
            ("f", "Tên khác", ", ".join(vn_only(x) for x in (r.get("alt") or []) if vn_only(x)) or UNV),
            ("f", "Loại hình", r["loai_vn"]
             + (f" (phụ: {', '.join(r['loai_phu'])})" if r.get("loai_phu") else "")
             ),
            ("f", "Địa chỉ", ev(r["id"], "dia_chi_day_du")),
            ("f", "Điện thoại", r.get("tel") or UNV),
            ("f", "Website", r.get("web") or UNV),
            ("f", "Tình trạng hoạt động", "đang hoạt động")]
    # Muc do pho bien de o day, KHONG phai o "Danh gia cua khach" — ti le de
    # xuat cua Facebook la thang do khac, khong quy doi sang sao duoc.
    for _lab, _k in (("Email", "email"),
                     ("Ảnh (tự do bản quyền)", "anh"),
                     ("Website chính thức", "website_chinh_thuc"),
                     ("⚠ Website đã lưu", "canh_bao_website"),
                     ("Kiểm tra trang web", "kiem_tra_website")):
        # Nam truong xuat xu (Trang Facebook, Email FB, Luot check-in, Nguoi
        # theo doi, Ti le de xuat) da chuyen sang PHU LUC NGHIEN CUU — chung la
        # xuat xu, khong phai thong tin di choi.
        # `canh_bao_website` o tren thi KHONG chuyen: no noi rang URL in ngay
        # duoi no khong phai trang cua dia diem nay. Mot canh bao phai nam canh
        # thu no canh bao.
        if has(r["id"], _k):
            spec.append(("f", _lab, ev(r["id"], _k)))
    spec += [
            ("group", "Tiện nghi tại chỗ  ·  cổng lọc đầu tiên")]
    _fac = [("Nhà vệ sinh", "nha_ve_sinh"), ("Bãi đỗ xe", "bai_do_xe"),
            ("Chỗ ngồi nghỉ", "cho_ngoi"), ("Hàng ăn", None), ("Hàng nước", "nuoc_uong"),
            ("Quà lưu niệm", "qua_luu_niem"), ("Hướng dẫn viên", None),
            ("Quầy thông tin", "quay_thong_tin"), ("Lối cho xe lăn", "loi_xe_lan"),
            ("Sơ cứu y tế", "so_cuu")]
    _nf = 0
    for _lab, _k in _fac:
        if _k and has(r["id"], _k):
            _nf += 1
            spec.append(("f", _lab, ev(r["id"], _k)))
        else:
            spec.append(("f", _lab, UNV))
    if has(r["id"], "wifi"):
        spec.append(("f", "Wifi", ev(r["id"], "wifi")))
    ind = None
    for lab, note in INDOOR_NOTE.items():
        pass
    from_cat = {"Bảo tàng": "trong nhà", "Nghệ thuật / Triển lãm": "trong nhà",
                "Chợ / Mua sắm": "có mái", "Nhà thờ": "có mái",
                "Chùa / Thiền viện": "hỗn hợp", "Dinh thự / Di tích": "hỗn hợp",
                "Khu vui chơi": "hỗn hợp", "Thác nước": "ngoài trời", "Hồ / Đập": "ngoài trời",
                "Công viên / Vườn hoa": "ngoài trời", "Điểm ngắm cảnh": "ngoài trời",
                "Núi / Đèo / Đường mòn": "ngoài trời", "Nông trại / Vườn": "ngoài trời",
                "Cáp treo": "ngoài trời"}
    ind = from_cat.get(r["loai_vn"])
    spec += [("group", "Kế hoạch thăm"),
             ("f", "Thời lượng thăm", UNV + " ← không có mục này thì không xếp được lịch một ngày"),
             ("f", "Thời điểm tốt trong ngày", UNV),
             ("f", "Mùa tốt nhất", UNV),
             ("f", "Trong nhà / ngoài trời",
              ind if ind else UNV),
             ("f", "Quãng đi bộ", UNV),
             ("f", "Độ khó", UNV),
             ("f", "Phù hợp trẻ nhỏ", UNV),
             ("f", "Phù hợp người cao tuổi", UNV),
             ("group", "Giờ giấc và chi phí"),
             ("f", "Ngày mở cửa", UNV),
             ("f", "Giờ mở cửa", r["hours"] if r.get("hours")
              else ev(r["id"], "gio_mo_cua")),
             ("f", "Giờ nhận khách cuối", UNV),
             ("f", "Giá vé", (str(r["fee"])) if r.get("fee")
              else UNV + " ← KHÔNG nêu số tiền; nói giá có thể thay đổi"),
             ("f", "Phí gửi xe", UNV),
             ("f", "Cần đặt trước", ev(r["id"], "can_dat_truoc"))]
    if has(r["id"], "gia_ve_tham_khao"):
        spec.append(("f", "Giá vé THAM KHẢO", ev(r["id"], "gia_ve_tham_khao")))
    if has(r["id"], "khoang_gia_facebook"):
        spec.append(("f", "Khoảng giá (Facebook)",
                     ev(r["id"], "khoang_gia_facebook") + "  (thang 4 bậc, chủ cơ sở tự khai)"))
    if has(r["id"], "gia_ve_ghi_chu"):
        spec.append(("f", "Ghi chú giá vé", ev(r["id"], "gia_ve_ghi_chu")))
    if r.get("csdl_gia_min"):
        g = f"{r['csdl_gia_min']:,}".replace(",", ".")
        if r.get("csdl_gia_max") and r["csdl_gia_max"] != r["csdl_gia_min"]:
            g += "–" + f"{r['csdl_gia_max']:,}".replace(",", ".")
        spec.append(("f", "Giá phòng (lưu trú)", f"{g}₫/đêm"))
    spec.append(("group", "Lưu ý quan trọng"))
    for f_ in ("Trang phục", "Giày dép", "Lưu ý thời tiết", "An toàn", "Giờ đông khách",
               "Nên mang theo", "Điều cấm"):
        spec.append(("f", f_, UNV))
    spec += [("group", "Vị trí và di chuyển"),
             ("f", "Từ hồ Xuân Hương",
              f"{r['km']:.1f} km · {r['min']:.0f} phút"
              if r.get("min") is not None else UNV),
             ("f", "Từ khách sạn", "→ thuộc hồ sơ chuyến đi, tính khi biết khách ở đâu"),
             ("f", "Đường chính gần nhất", ev(r["id"], "duong_gan_nhat")),
             ("f", "Tình trạng đường", UNV),
             ("f", "Phương tiện tới được", UNV),
             ("f", "Bãi đỗ xe", UNV),
             ("group", "Chụp ảnh")]
    for f_ in ("Điểm chụp đẹp", "Giờ chụp đẹp", "Ngắm bình minh / hoàng hôn",
               "Phí chụp ảnh", "Lưu ý chụp ảnh"):
        spec.append(("f", f_, UNV))
    spec.append(("f", "Bay flycam", "COI NHƯ BỊ CẤM cho tới khi có xác nhận ngược lại"))
    for _lab, _k in (("Năm khánh thành", "nam_khanh_thanh"), ("Năm xây dựng", "nam_xay_dung"),
                     ("Kiến trúc", "kien_truc"), ("Xếp hạng di tích", "xep_hang_di_tich"),
                     ("Tôn giáo", "ton_giao"), ("Diện tích", "dien_tich")):
        if has(r["id"], _k):
            spec.append(("f", _lab, ev(r["id"], _k)))
    field_table(spec)

    if NEAR.get(r["id"]):
        P(sec("Điểm lân cận") + " (thời gian đường bộ thật, không phải đường chim bay)", bold=True, size=9.5)
        TBL(["#", "Điểm", "Loại", "Km", "Phút"],
            [[i, f"{oid} · {byid[oid]['name']}", byid[oid]["loai_vn"], f"{dkm:.1f}", f"{tmin:.0f}"]
             for i, (oid, dkm, tmin) in enumerate(NEAR[r["id"]], 1)],
            widths=[1.0, 7.0, 4.0, 2.0, 2.0])

    # ── Khach san & quan an GAN diem nay ────────────────────────────────────
    # So luong BIEN, khong phai 3 co dinh: chi 24/36 diem lap du 3 bac va 9 diem
    # khong co co so luu tru dang ky nao trong 5 km. Quan an nhom theo LOAI MON
    # chu khong theo bac gia — 0/5.559 quan co gia, va mot tieu de bac gia tu no
    # la mot khang dinh du kien. Khong in `Đánh giá`/`Tiện nghi`: 0 du lieu, va
    # rao can la giay phep luu tru; da noi mot lan o muc 0.
    # Khoi khach san + quan an KHONG con o day — in mot lan o dau khu vuc.
    # Truoc day no chiem 1.420 dong tren ca tai lieu, trung vi 43% moi ho so, de
    # in lai gan nhu cung mot danh sach 36 lan; chi co 52 khach san khac nhau.
    _lc = _LAN_CAN.get(r["id"], {})
    _pt = P(sec("Lưu trú & ăn uống") + f" → xem khối đầu khu vực {r['area']}",
            bold=True, size=9.5)
    if _lc.get("tong_ks_trong_bk") is not None:
        P(f"Quanh riêng điểm này: {_lc['tong_ks_trong_bk']} cơ sở lưu trú trong 5 km"
          f" · {_lc['tong_quan_trong_bk']} quán trong 2 km", size=8.5, color=GREY)
    if _lc.get("khong_co_khach_san"):
        P(_lc["khong_co_khach_san"], size=9, color=AMBER)

    pv = doc.add_paragraph()
    rr = pv.add_run(f"Kiểm chứng: CHƯA GỌI · gọi {r.get('tel') or 'CHƯA CÓ SỐ'} để đóng "
                    "giờ mở cửa, giá vé, thời lượng thăm và điều kiện đi lại.")
    rr.bold = True
    rr.font.size = Pt(8.5)
    rr.font.color.rgb = RED

# ==================================================== 3. HOAT DONG
# Cung mot module chon loc voi ban .md — `hoat_dong_data.tai()`. Neu viet lai
# logic cat gon o day thi hai ban se lech nhau va khong ai biet cho toi khi doc
# canh nhau; du an da dinh dung lop loi do (hai bo trich cung payload VNPay).
_HD, _HDTK = _hoat_dong.tai(RAW)
_MON = _hoat_dong.tai_mon_an(RAW)
_PC = _hoat_dong.tai_phong_cach(RAW)

doc.add_page_break()
H(f"{S_HOATDONG}. Hoạt động — làm gì ở Đà Lạt", 1)
P(f"{_HDTK['so_hoat_dong']} hoạt động, {_HDTK['so_nhom']} nhóm. "
  f"Mã DL-xx dẫn về mục chi tiết ở mục {S_DIEMDEN}.", italic=True, size=9, color=GREY)

_nhom_hien = None
for _a in _HD:
    if _a["nhom"] != _nhom_hien:
        _nhom_hien = _a["nhom"]
        H(_nhom_hien.upper(), 2)
    H(_a["ten"], 3)
    if _a["noi"]:
        P("Làm ở đâu — %d nơi — %d nơi tiêu biểu:" % (_a["tong_noi"], len(_a["noi"]))
          if _a["tong_noi"] > len(_a["noi"]) else "Làm ở đâu:", bold=True, size=9)
        for _n in _a["noi"]:
            B((f"[{_n['ma']}] " if _n["ma"] else "") + _n["ten"]
              + (f" — {_n['khu_vuc']}" if _n.get("khu_vuc") else ""))
    if _a["don_vi"]:
        P("Đơn vị tổ chức — %d đơn vị — %d đơn vị tiêu biểu:"
          % (_a["tong_don_vi"], len(_a["don_vi"]))
          if _a["tong_don_vi"] > len(_a["don_vi"]) else "Đơn vị tổ chức:",
          bold=True, size=9)
        for _d in _a["don_vi"]:
            B(_d["ten"] + (f" — {_d['dien_thoai']}" if _d.get("dien_thoai")
                           else " — chưa có số"))
    if _a.get("tour_web"):
        P("Trang tour:", bold=True, size=9)
        for _t in _a["tour_web"]:
            B(f"{_t['ten']} — {_t['url']}")
            if _t["khoang_gia_don_vi"]:
                _pp = doc.add_paragraph(style="List Bullet 2")
                _rr = _pp.add_run(f"Khoảng giá cả gói: {_t['khoang_gia_don_vi']}")
                _rr.font.size = Pt(8.5)
                _rr.font.color.rgb = AMBER
            for _n in _t["ten_tour"]:
                doc.add_paragraph(_n, style="List Bullet 2")
    if not _a["don_vi"]:
        P("Không cần đơn vị tổ chức — tự đi được.", italic=True, size=9, color=GREY)

# BA KHOI, khong phai mot bang phang — xep tren toan bo theo so quan dao nguoc
# cau tra loi cho "an gi o Da Lat". `nhom` la phan doan bien tap, khong phai so do.
if _MON:
    _tong_q = sum(n for _, rows in _MON for _, n, _ in rows)
    _tong_m = sum(len(rows) for _, rows in _MON)
    H(f"ẨM THỰC — {_tong_q} quán trên {_tong_m} món", 2)
    P("Chia nhóm là phán đoán biên tập, không phải số đo: “đặc sản” nghĩa là món gắn "
      "với Đà Lạt, “phổ thông” nghĩa là món có ở mọi thành phố và cũng có ở đây. "
      "Trong từng nhóm xếp theo số cơ sở bán.", italic=True, size=9, color=GREY)
    # Cot "vlog" chi hien khi co du lieu — han muc YouTube theo NGAY nen file
    # co the rong hoac mot phan; in cot toan trong trong nhu du lieu bi mat.
    _co_vlog = any(q.get("vlog") for _, rows in _MON for _, _, qs in rows for q in qs)
    for _nhom, _rows in _MON:
        H(f"{_nhom.upper()} — {len(_rows)} món", 3)
        TBL(["Món", "Số quán", "Gợi ý (ưu tiên quán có số gọi)"],
            [[_m, str(_sl),
              " · ".join(q["ten"][:26] + (f" ({q['vlog']} kênh)" if q.get("vlog") else "")
                         for q in _q[:3])]
             for _m, _sl, _q in _rows], widths=[3.6, 1.6, 10.6])
    if _co_vlog:
        P("(N kênh) — số kênh du lịch tiếng Việt KHÁC NHAU nhắc tên quán, ngưỡng ≥2 "
          "kênh. Đếm theo kênh chứ không theo video, vì một kênh đăng nhiều video "
          "không phải nhiều lời khuyên độc lập. Quán không có ghi chú nghĩa là chưa "
          "quét đến, không phải không được nhắc.", italic=True, size=8.5, color=GREY)

if _PC:
    H("QUÁN CÓ PHONG CÁCH ĐẶC BIỆT", 2)
    P(f"{len(_PC)} quán được vlog nhắc kèm một cách làm riêng. Thẻ lấy từ bộ từ vựng "
      "cố định, không phải mô tả tự do.", italic=True, size=9, color=GREY)
    TBL(["Quán", "Phong cách", "Kênh nhắc", "Điện thoại"],
        [[_q["ten"][:34], ", ".join(_q["the_phong_cach"]),
          str(_q.get("so_kenh_nhac", _q.get("so_video_nhac", 0))),
          _q.get("dien_thoai") or ""] for _q in _PC], widths=[5.0, 5.4, 1.4, 3.0])

# Ba truong mua/gio/thoi luong trong tren ca 28 hoat dong. Cho thieu + viec can
# lam thuoc muc 12, khong thuoc giua chuong tra cuu.

# ==================================================== 4. LUU TRU & AN UONG
_LT = _an_ngu.tai_luu_tru(RAW)
_AU = _an_ngu.tai_an_uong(RAW)

if _LT or _AU:
    doc.add_page_break()
    H(f"{S_ANNGU}. Lưu trú & ăn uống", 1)

if _LT:
    H(f"{S_ANNGU}.1 Lưu trú", 2)
    # Nhom DUY NHAT trong ca tai lieu co GIA THAT — tu dang ky luu tru cua Cuc
    # Du lich Quoc gia, khong phai tu blog. Nen gia ghi thang.
    P(f"{_LT['tong']} cơ sở trong đăng ký lưu trú nhà nước · {_LT['nha_nuoc']} đã thẩm "
      f"định, {_LT['tu_dang_ky']} tự đăng ký · {_LT['co_gia']} cơ sở có giá công bố · "
      f"{_LT['co_dien_thoai']} có số gọi.", italic=True, size=9, color=GREY)
    P("Giá là giá phòng/đêm do cơ sở công bố với cơ quan quản lý — đổi theo mùa, gọi "
      "xác nhận trước khi báo khách.", bold=True, size=9)
    for _b in _LT["bac"]:
        if not _b["co_so"]:
            continue
        H(f"{_b['ten']} — {_b['tong']} cơ sở", 3)
        TBL(["Cơ sở", "Giá/đêm", "Phòng", "Điện thoại", "Địa chỉ"],
            [[_c["ten"][:38], _c["gia"], str(_c["so_phong"] or ""),
              _c["dien_thoai"] or "", (_c["dia_chi"] or "").split(",")[0][:26]]
             for _c in _b["co_so"]], widths=[5.4, 3.4, 1.4, 2.8, 3.6])
    if _LT["dong_cua"]:
        _pd = doc.add_paragraph()
        _rd = _pd.add_run("Đã đóng cửa — không giới thiệu: "
                          + " · ".join(f"{r['ten']} ({r['ngay']})" for r in _LT["dong_cua"]))
        _rd.bold = True
        _rd.font.size = Pt(9)
        _rd.font.color.rgb = RED

if _AU:
    H(f"{S_ANNGU}.2 Ăn uống", 2)
    P(f"{_AU['tong_mo']:,} quán còn hoạt động · {_AU['co_dien_thoai']:,} có số gọi. "
      f"Món đặc trưng xem mục {S_HOATDONG}.".replace(",", "."), italic=True, size=9, color=GREY)
    for _n in _AU["nhom"]:
        H(f"{_n['ten']} — {_n['tong']} quán", 3)
        for _q in _n["quan"]:
            B(_q["ten"] + (f" — {_q['dien_thoai']}" if _q["dien_thoai"] else "")
              + (f" · {', '.join(_q['mon'])}" if _q["mon"] else ""))
    # Muc quan trong nhat cua chuong nay. Thac khong dong cua; quan an thi co.
    if _AU["dong_cua"]:
        H("⚠ Đã đóng cửa — KHÔNG giới thiệu", 3)
        P(f"{_AU['tong_dong']} quán đã đóng, dưới đây {len(_AU['dong_cua'])} gần nhất. "
          "Nhiều quán trong số này vẫn còn trong các hướng dẫn cũ.",
          bold=True, size=9, color=RED)
        TBL(["Ngày đóng", "Quán"],
            [[_r["ngay"], _r["ten"][:52]] for _r in _AU["dong_cua"]], widths=[3.0, 12.0])

# ==================================================== 5-13 chi muc
doc.add_page_break()
H(f"{S_SOSANH}. Bảng so sánh", 1)
P(f"Sinh tự động từ mục {S_DIEMDEN} — không sửa tay.", italic=True, size=9, color=GREY)
TBL(["ID", "Điểm", "Loại", "Khu vực", "Km", "Phút", "Vé", "Nguồn"],
    [[r["id"], r["name"][:34], r["loai_vn"], r["area"][:20], f"{r['km']:.1f}",
      f"{r['min']:.0f}", r.get("fee") or UNV, len(r["src"])] for r in picked],
    widths=[1.4, 5.0, 3.0, 3.2, 1.2, 1.2, 1.6, 1.2], size=8)

# DA CAT muc 6 "Theo loai hinh" · 7 "Theo khu vuc" · 8 "Theo khoang cach"
# · 9 "Theo thoi diem tham" — ca bon la cach SAP XEP LAI cung 36 diem. `loai_vn`
# va `km` da la COT cua bang so sanh; muc 2 gio da nhom theo khu vuc; va truong
# "thoi diem tot trong ngay" la 0/36. Phan duy nhat co that trong muc 9 — nhom
# "di duoc khi troi mua" — thanh mot cot cua bang so sanh.
H(f"{S_TUYEN}. Tuyến gợi ý theo khu vực", 1)
P("Thứ tự dựng bằng thuật toán láng giềng gần nhất trên ma trận OSRM, xuất phát từ điểm gần "
  "trung tâm nhất.", italic=True, size=9, color=GREY)
if mat:
    idx = {pid: i for i, pid in enumerate(mat["ids"])}
    for a in sorted({r["area"] for r in picked}):
        lst = [r for r in picked if r["area"] == a]
        if len(lst) < 2:
            continue
        cur = min(lst, key=lambda x: x["min"])
        route, left, tot = [cur], [x for x in lst if x is not cur], 0.0
        while left:
            nxt = min(left, key=lambda x: mat["durations"][idx[cur["id"]]][idx[x["id"]]] or 9e9)
            tot += (mat["durations"][idx[cur["id"]]][idx[nxt["id"]]] or 0) / 60.0
            route.append(nxt)
            left.remove(nxt)
            cur = nxt
        P(f"{a} — {len(route)} điểm · di chuyển giữa các điểm ~{tot:.0f} phút", bold=True, size=9.5)
        P("   " + " → ".join(f"{x['id']} {x['name']}" for x in route), size=9)

H(f"{S_MATRAN}. Ma trận thời gian trong từng khu vực (phút)", 1)
# Truoc day la mot bang 36x36 = 1.296 o o co chu 6pt — khong vua mot trang doc
# va khong ai tra thoi gian giua hai diem o hai dau thanh pho. Chia theo khu vuc:
# chin bang nho, moi bang vua mot trang. Cap xa nhau van con trong osrm_selected.
if mat:
    _idx8 = {pid: i for i, pid in enumerate(mat["ids"])}
    P("Chia theo khu vực: thời gian giữa các điểm trong cùng một buổi. Cặp ở hai khu "
      f"vực khác nhau thì tra mục {S_TUYEN} (tuyến gợi ý) hoặc cột Phút ở mục {S_SOSANH}.",
      italic=True, size=9, color=GREY)
    for _a in sorted({r["area"] for r in picked}):
        _lst = [r for r in picked if r["area"] == _a]
        if len(_lst) < 2:
            continue
        P(f"{_a} ({len(_lst)} điểm)", bold=True, size=9.5)
        TBL([""] + [r["id"].replace("DL-", "") for r in _lst],
            [[r["id"].replace("DL-", "")] +
             ["—" if r is o or mat["durations"][_idx8[r["id"]]][_idx8[o["id"]]] is None
              else f"{mat['durations'][_idx8[r['id']]][_idx8[o['id']]]/60:.0f}"
              for o in _lst]
             for r in _lst], size=8)

# DA CHUYEN muc 12 "Danh sach rut gon" -> phu luc: thu hang do do MUC DO HIEN
# DIEN TREN BAN DO, tuc chat luong DU LIEU, khong phai chat luong trai nghiem.
# Muc 9 ngay duoi dung dung thu hang do lam THU TU GOI DIEN, la cong dung dung.
#
# Thu tu nay den TU build_huong_dan.py qua khoa `xep_hang`, khong tu tinh lai o
# day. Truoc day dong nay la `sorted(picked, key=lambda r: (-len(r["src"]),
# -conf))` trong khi ban .md xep theo `-_score` — va `_score` bi bo loc
# `startswith("_")` chan lai nen khong bao gio den duoc file nay. Ket qua: hai
# ban tai lieu bat dong ve THU TU GOI DIEN xac minh, tu hang thu 9 tro di.
_xh = {pid: i for i, pid in enumerate(G.get("xep_hang") or [])}
if not _xh:
    raise SystemExit("guide_data.json thieu khoa 'xep_hang' — chay"
                     " build_huong_dan.py truoc de sinh lai.")
rank = sorted(picked, key=lambda r: _xh.get(r["id"], 10 ** 6))

H(f"{S_KIEMCHUNG}. Sổ kiểm chứng — việc cần làm", 1)
n_tel = sum(1 for r in picked if r.get("tel"))
TBL(["Chỉ số", "Giá trị"],
    [["Điểm trong hồ sơ", len(picked)],
     ["Có số điện thoại để gọi", f"{n_tel} / {len(picked)}"],
     ["Chưa có số — cần tìm", len(picked) - n_tel],
     ["Có giờ mở cửa", sum(1 for r in picked if r.get("hours"))],
     ["Có giá vé", sum(1 for r in picked if r.get("fee"))],
     ["Có đánh giá sao", "0 — không nguồn mở nào có"],
     ["Trường [CHƯA XÁC MINH] ước tính", f"~{41*len(picked)}"]],
    widths=[7.0, 9.0], size=9)
# Cho thieu cua muc 3 thuoc VE DAY, khong thuoc giua chuong tra cuu.
if _HDTK["thieu"]:
    _pm = doc.add_paragraph()
    _rm = _pm.add_run(f"Mục 3 — cả {_HDTK['so_hoat_dong']} hoạt động đều chưa có mùa, "
                      "giờ trong ngày và thời lượng. Trống nghĩa là chưa xác minh, "
                      "không nghĩa là quanh năm. Đừng khẳng định với khách cỏ hồng "
                      "tháng nào hay tour đi mấy giờ khi chưa gọi.")
    _rm.bold = True
    _rm.font.size = Pt(9)
    _rm.font.color.rgb = RED
    P(f"Tra đơn vị tour bằng Facebook, đừng tra tên miền — {_HDTK['ten_mien_chet']}/"
      f"{_HDTK['tong_website_thu']} tên miền đã kiểm không còn hoạt động "
      "(canyoningdalat.com, dalatjeep.com, toursanmaydalat.com).", size=9, color=AMBER)

P("Gọi một cuộc đóng được khoảng 9 trường. Danh sách cần gọi, theo thứ tự ưu tiên:", bold=True)
TBL(["#", "Điểm", "Điện thoại"],
    [[i, f"{r['id']} · {r['name']}", r["tel"]]
     for i, r in enumerate([x for x in rank if x.get("tel")][:20], 1)],
    widths=[1.2, 9.0, 5.0], size=9)

# Chon loc nam o `an_ngu_data.tai_nghien_cuu` — CUNG ham ma ban .md goi. Ban
# truoc khoi nay chi co o ban .md nen ban .docx mat trang: 32 link Facebook, 21
# luot check-in, 4 ti le de xuat va ca bang thu hang.
_NC = _an_ngu.tai_nghien_cuu(RAW)
if _NC and _NC["hang"]:
    doc.add_page_break()
    H("Phụ lục nghiên cứu — xuất xứ, không phải thông tin đi chơi", 1)
    P("Các trường dưới đây đã được đưa ra khỏi hồ sơ điểm đến: người lập kế hoạch "
      "không dùng chúng, người kiểm chứng nguồn thì cần. Tra theo mã DL-xx.",
      italic=True, size=9, color=GREY)
    TBL(["ID", "Điểm", "Trường", "Giá trị"],
        [[i, t[:26], nhan, gt[:44]] for i, t, nhan, gt in _NC["hang"]],
        widths=[1.6, 4.6, 3.4, 6.4], size=8)
    P(f"{_NC['so_diem']}/{_NC['tong']} điểm có ít nhất một trường xuất xứ.",
      italic=True, size=8.5, color=GREY)
    H("Thứ hạng theo mức độ hiện diện trên bản đồ", 2)
    P("⚠ Đây là chất lượng DỮ LIỆU, không phải chất lượng trải nghiệm. Mục "
      f"{S_KIEMCHUNG} dùng đúng thứ hạng này làm thứ tự gọi điện xác minh — đó là "
      "công dụng đúng của nó. Không dùng làm lời khuyên “nơi này hay hơn nơi kia”.",
      bold=True, size=9, color=RED)
    for lab, seg in _NC["xep_hang"]:
        P(f"{lab} ({len(seg)}): " + " · ".join(f"{i} {t}" for i, t in seg), size=9)

doc.add_page_break()
H("Phụ lục — Tóm tắt độ phủ (đọc lại trước khi trả lời khách)", 1)
P(f"Tài liệu này theo dõi khoảng {41*len(picked)} trường trên {len(picked)} điểm đến. "
  "Phần lớn mang dấu [CHƯA XÁC MINH].")
P("Nếu khách hỏi về một trường mang dấu [CHƯA XÁC MINH] ở bất kỳ đâu phía trên: nói "
  "“chỗ này em chưa có số liệu đã xác minh, để em gọi hỏi rồi báo lại mình”, và KHÔNG đưa ra "
  "con số, giờ giấc hay đánh giá nào thay thế.", bold=True)
P("Ba trường tuyệt đối không được suy đoán, vì sai là khách mất tiền hoặc mất cả ngày:",
  bold=True, color=RED)
B("giờ mở cửa")
B("giá vé")
B("mức độ dễ đi lại cho người cao tuổi")
P("Tài liệu này chứa dữ liệu từ OpenStreetMap. Dữ liệu © những người đóng góp OpenStreetMap, "
  "theo giấy phép Open Database License — https://openstreetmap.org/copyright",
  italic=True, size=8, color=GREY)

doc.save(OUT)
print("saved ->", OUT)
print(f"diem: {len(picked)}  |  co dien thoai: {n_tel}  |  bang: {len(doc.tables)}")
