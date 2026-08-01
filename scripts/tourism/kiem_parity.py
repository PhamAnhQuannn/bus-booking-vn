# -*- coding: utf-8 -*-
"""Kiem HAI BAN DAU RA co khop nhau khong. Exit 1 neu lech.

VI SAO CAN. Du an nay co quy tac "mot nguon chon loc, hai nguon dinh dang" —
nhung khong co gi kiem tra rang quy tac do dang duoc tuan thu. Va no da bi pha
hai lan, ca hai lan im lang:

  1. Khoi phu luc xuat xu duoc viet TRUC TIEP trong build_huong_dan.py, nen no
     khong bao gio den duoc ban .docx. Do duoc: 32 link Facebook, 21 luot
     check-in, 4 ti le de xuat va ca bang thu hang co trong .md, BANG KHONG
     trong .docx — trong khi nam truong do da bi bo khoi the cua ca hai ban.
     Ban .docx mat trang du lieu.
  2. Muc 12/13 xep hang bang hai cong thuc khac nhau: `.md` theo `_score`,
     `.docx` theo `(-len(src), -conf)`, vi `guide_data.json` loc bo moi khoa bat
     dau bang `_`. Hai ban bat dong ve THU TU GOI DIEN xac minh tu hang thu 9.

Ca hai chi lo ra khi so hai file BANG TAY. Bo kiem nay lam viec do tu dong.

No doc DUY NHAT hai file dau ra — khong import bo dung, khong doc du lieu goc —
nen no khong the lech theo cung mot loi voi thu no dang kiem.

Chay:  python scripts/tourism/kiem_parity.py [duong-md] [duong-docx]
"""
import io
import os
import re
import sys
import time

MD = sys.argv[1] if len(sys.argv) > 1 else \
    "documentation/tourism/destinations/da-lat/huong-dan-diem-den.md"
# Ban GOP, khong nam trong `docs/` — xem chu thich `_CAU_HINH` trong
# `build_huong_dan_docx.py`. `docs/` chi giu ba ban giao, va ba ban do danh so
# muc lai theo tung tai lieu nen khong so truc tiep voi ban .md duoc.
DOCX = sys.argv[2] if len(sys.argv) > 2 else ".tourism-data/build/Huong-Dan-Da-Lat.docx"

# ── 0. DIEU KIEN TIEN QUYET: hai file phai tu CUNG MOT lan build ───────────
# 31/07: `build_huong_dan.py` ghi xong `.md`, roi `build_huong_dan_docx.py`
# chet vi Word dang giu file (`PermissionError`), nen `.docx` tren dia van la
# ban HOM TRUOC. Bo kiem nay doc cap file lech mot ngay va bao KHOP, exit 0.
# Do duoc luc do: `Nhà Chung` 3 (.md) / 1 (.docx), `Bánh Căn Lệ` 3/1,
# `Lẩu Gà Lá É Tao Ngộ` 6/2.
#
# No khong bi danh lua bang gi tinh vi — no khong co khai niem "do tuoi" de ma
# hoi. Khi hai dau vao khong tu cung mot lan build thi MOI ket luan sau do vo
# nghia, ke ca ket luan "khop"; nen day la DIEU KIEN TIEN QUYET — dung truoc,
# exit 1, khong so noi dung — chu khong phai mot muc lech thu bay.
LECH_TOI_DA = 300     # giay. Du cho hai bo dung chay noi nhau, khong du cho
                      # hai lan build khac nhau.


def _khac_gio(t):
    return time.strftime("%d/%m %H:%M:%S", time.localtime(t))


_t_md, _t_dx = os.path.getmtime(MD), os.path.getmtime(DOCX)
if abs(_t_md - _t_dx) > LECH_TOI_DA:
    print(f"{MD}\n{DOCX}\n")
    print("LECH BUILD — hai file khong tu cung mot lan build, cach nhau"
          f" {abs(_t_md - _t_dx) / 60:.0f} phut:")
    print(f"      .md   {_khac_gio(_t_md)}")
    print(f"      .docx {_khac_gio(_t_dx)}")
    print("\nKhong so noi dung: mot bo dung da that bai hoac chua chay lai."
          " Chay lai CA HAI roi kiem lai.")
    sys.exit(1)

# Noi ro bo kiem nay so gi va KHONG so gi. "KHOP" tran trui doc thanh "hai tai
# lieu dong y", trong khi nghia that la "may thu toi nhin thi dong y" — va lop
# am thuc (quan, mon, cot vlog) hien nam NGOAI pham vi, nen mot thay doi nam
# tron trong do di qua bo kiem nay ma khong ai biet.
PHAM_VI = ("tieu de muc cap 1 · day so cap 1 · tap tham chieu `mục N` ·"
           " danh sach goi dien · day so A.x moi the · vang mat truong danh gia")
KHONG_SO = "lop AM THUC (quan an, mon, cot vlog) — CHUA co trong bo kiem nay"

loi = []


def sai(muc, a, b, ghi=""):
    loi.append((muc, a, b, ghi))


md = io.open(MD, encoding="utf-8").read()
try:
    from docx import Document
except ImportError:
    print("khong co python-docx — khong kiem duoc")
    sys.exit(0)
d = Document(DOCX)
dx_para = [p.text for p in d.paragraphs]
dx_cell = [c.text for t in d.tables for r in t.rows for c in r.cells]
dx = "\n".join(dx_para + dx_cell)


def fold(s):
    """Bo dau va ha chu de so tieu de: hai bo dung viet hoa khac nhau
    ("DANH SÁCH ĐIỂM ĐẾN" vs "Danh sách điểm đến") — do la dinh dang, khong
    phai noi dung, nen khong tinh la lech."""
    import unicodedata
    s = unicodedata.normalize("NFD", (s or "").lower())
    return "".join(c for c in s if unicodedata.category(c) != "Mn").strip()


# ── 1. tieu de muc cap 1, dung thu tu ──────────────────────────────────────
md_h = [l[3:].strip() for l in md.split("\n") if l.startswith("## ")]
dx_h = [p.text.strip() for p in d.paragraphs if p.style.name == "Heading 1"]
if len(md_h) != len(dx_h):
    sai("so muc cap 1", len(md_h), len(dx_h))
for i in range(min(len(md_h), len(dx_h))):
    a, b = fold(md_h[i]), fold(dx_h[i])
    # So 18 ky tu dau: du de bat lech so muc va lech ten, bo qua phan duoi
    # ngoac dai khac nhau giua hai dinh dang.
    if a[:18] != b[:18]:
        sai(f"tieu de muc #{i}", md_h[i][:40], dx_h[i][:40])

# ── 2. tap tham chieu "muc N" ──────────────────────────────────────────────
# Neu mot ban doi so muc ma ban kia khong doi, tap nay lech ngay.
md_ref = sorted({int(x) for x in re.findall(r"mục (\d+)", md)})
dx_ref = sorted({int(x) for x in re.findall(r"mục (\d+)", dx)})
if md_ref != dx_ref:
    sai("tap tham chieu 'muc N'", md_ref, dx_ref)

# Moi tham chieu phai tro den mot muc CO THAT.
so_muc = {int(m.group(1)) for h in md_h
          for m in [re.match(r"(\d+)\.", h)] if m}
for n in md_ref:
    if n not in so_muc:
        sai("tham chieu tro vao muc khong ton tai", f"mục {n}", sorted(so_muc))

# ── 3. cac khoi phai co o CA HAI ban ──────────────────────────────────────
KHOI = [("phu luc nghien cuu", "Trang Facebook", 1),
        ("luot check-in", "Lượt check-in", 1),
        ("ti le de xuat", "Tỉ lệ đề xuất", 1),
        ("thu hang du lieu", "Ưu tiên xác minh trước", 1),
        ("khoi khu vuc", "Lưu trú & ăn uống trong khu vực", 9),
        ("tro ve khoi khu vuc", "xem khối đầu khu vực", 36),
        ("canh bao website", "Website đã lưu", 8),
        ("lop dia hinh", "Độ nhô", 1),
        # Them sau khi bo kiem nay BO SOT dung mot lan: 28 doan mo ta Wikipedia
        # vao ban .md ma khong vao ban .docx, va bo kiem van bao KHOP — vi danh
        # sach nay khong co dong nao ve mo ta. Mot bo kiem chi bat duoc thu no
        # duoc bao la phai bat; moi khoi noi dung moi phai them mot dong o day,
        # neu khong no nam ngoai vung phu.
        ("mo ta Wikipedia", "trích nguyên văn", 20),
        ("ghi cong CC BY-SA", "CC BY-SA 4.0", 1),
        # Link ban do — mot trong ba suy dien duoc muc 0 cho phep, va la cai duy
        # nhat trong ba tung khong duoc sinh ra. 36/36 diem co toa do nen con so
        # nay phai dung bang so diem; tut xuong la mot the mat link.
        ("link ban do", "google.com/maps", 36),
        # Khoi 207 co so khong cong bo gia. Cot cua no da duoc muc 3a phu (tieu
        # de bat dau bang "Cơ sở"), dong nay phu truong hop khac: ca khoi bien
        # mat khoi mot ban.
        ("khoi khong cong bo gia", "Không công bố giá", 1)]
for ten, chuoi, toi_thieu in KHOI:
    a, b = md.count(chuoi), dx.count(chuoi)
    if a == 0 or b == 0:
        sai(f"khoi '{ten}' chi co o mot ban", f".md={a}", f".docx={b}")
    elif a < toi_thieu or b < toi_thieu:
        sai(f"khoi '{ten}' it hon mong doi ({toi_thieu})", f".md={a}", f".docx={b}")

# ── 3a. TIEU DE COT cua cac bang du lieu phai GIONG HET ───────────────────
# Them mot cot vao mot bo dung ma quen bo dung kia la kieu lech de xay ra nhat:
# danh sach tieu de va tuple dong duoc viet RIENG trong tung bo dung (bang khu
# vuc: build_huong_dan.py:716 va build_huong_dan_docx.py:547), du du lieu den tu
# cung mot ham chon loc. Muc 3 khong bat duoc — no dem chuoi con, ma "Địa chỉ"
# da co san o bang toan thanh pho lan trong the tung diem, nen mot cot moi bi bo
# quen o mot ban van chim trong con so tong.
#
# Chi doi chieu cac bang co o dau tien nam trong CHOT_BANG: cac bang the diem
# (ban .docx dung bang cho moi the, ban .md dung khoi ```) va ma tran thoi gian
# (nhan rut gon o ban .docx) khac nhau MOT CACH CO CHU DICH, khong phai lech.
CHOT_BANG = {"Khách sạn", "Loại", "Cơ sở"}


def _tieu_de_md(txt):
    ra, dong = [], txt.split("\n")
    for i, l in enumerate(dong[:-1]):
        if l.startswith("|") and re.match(r"^\|[\s\-:|]+\|$", dong[i + 1] or ""):
            ra.append(tuple(c.strip() for c in l.strip("|").split("|")))
    return ra


_md_td = [t for t in _tieu_de_md(md) if t and t[0] in CHOT_BANG]
_dx_td = [tuple(c.text.strip() for c in t.rows[0].cells)
          for t in d.tables if t.rows]
_dx_td = [t for t in _dx_td if t and t[0] in CHOT_BANG]
from collections import Counter as _Counter
_a, _b = _Counter(_md_td), _Counter(_dx_td)
for _t in sorted(set(_a) | set(_b)):
    if _a[_t] != _b[_t]:
        sai(f"tieu de cot lech: {' | '.join(_t)}", f".md={_a[_t]}", f".docx={_b[_t]}")

# ── 3b. CANH BAO: so luong phai BANG NHAU, khong chi "cung co mat" ────────
# Muc 3 chi hoi "co o ca hai ban khong" va "co du toi thieu khong", nen mot ban
# in 3 canh bao con ban kia in 2 van duoc bao KHOP. Do khong phai gia dinh — no
# vua xay ra: canh bao mau thuan so nha cua XQ Su Quan co trong .md va KHONG co
# trong .docx, vi `field_table()` bo moi dong co gia tri bat dau bang UNV, ma
# dia chi cua diem do dung la UNV sau khi ban ghi ban do bi loai. Bo kiem cu
# doc "3 va 2, ca hai deu >= 1" la dat.
#
# Canh bao la lop noi dung KHONG duoc phep lech mot don vi: mot canh bao thieu
# la mot khang dinh sai duoc doc nhu su that o dung mot ban.
CANH_BAO = ["MÂU THUẪN SỐ NHÀ",
            "lớp bản đồ ghi khác",
            "HAI TRANG CỦA CÙNG BẢO TÀNG",
            "ĐÃ BỊ LOẠI vì khớp nhầm",
            "CÓ thu phí",
            "số trần trên thẻ OSM",
            "KHÔNG cơ sở nào công bố giá",
            "Website đã lưu"]
for chuoi in CANH_BAO:
    a, b = md.count(chuoi), dx.count(chuoi)
    if a != b:
        sai(f"so canh bao '{chuoi[:34]}' KHAC NHAU", f".md={a}", f".docx={b}")

# ── 4. thu tu goi dien — dung thu hai ban tung bat dong ───────────────────
md_goi = re.findall(r"\| \d+ \| (DL-\d\d) ·", md.split("theo thứ tự ưu tiên")[-1])
dx_goi = []
for t in d.tables:
    if [c.text.strip() for c in t.rows[0].cells][:3] == ["#", "Điểm", "Điện thoại"]:
        dx_goi = [m.group(1) for r in t.rows[1:]
                  for m in [re.match(r"(DL-\d\d)", r.cells[1].text.strip())] if m]
if not md_goi or not dx_goi:
    sai("khong doc duoc danh sach goi dien", bool(md_goi), bool(dx_goi))
elif md_goi != dx_goi:
    k = next((i for i, (x, y) in enumerate(zip(md_goi, dx_goi)) if x != y), 0)
    sai(f"thu tu goi dien lech tu hang {k+1}", md_goi[:6], dx_goi[:6])

# ── 5. day so A.x lien tuc tu A.1, o CA HAI ban ───────────────────────────
def day_so_md(txt):
    ra, cur = [], []
    for ln in txt.split("\n"):
        if ln.startswith("#### DL-"):
            if cur:
                ra.append(cur)
            cur = []
        m = re.match(r"\*\*A\.(\d+) — ", ln.lstrip())
        if m:
            cur.append(int(m.group(1)))
    if cur:
        ra.append(cur)
    return ra


for nhan, days in (("md", day_so_md(md)),
                   ("docx", [[int(m.group(1)) for r in t.rows
                              for m in [re.match(r"A\.(\d+) — ", r.cells[0].text.strip())]
                              if m] for t in d.tables])):
    xau = [x for x in days if x and x != list(range(1, len(x) + 1))]
    if xau:
        sai(f"day so A.x khong lien tuc ({nhan})", len(xau), xau[:2])

# ── 6. rating phai VANG o ca hai ban ──────────────────────────────────────
for nhan, txt in (("md", md), ("docx", dx)):
    n = txt.count("Đánh giá:")
    if n:
        sai(f"co truong danh gia ({nhan}) — bi chan boi giay phep luu tru", n, 0)

# ── ket qua ───────────────────────────────────────────────────────────────
print(f"{MD}\n{DOCX}\n")
print(f"muc cap 1: {len(md_h)} / {len(dx_h)}   ·   tham chieu: {md_ref}")
if not loi:
    print(f"\nda so   : {PHAM_VI}")
    print(f"KHONG so: {KHONG_SO}")
    print("\nKHOP — hai ban dau ra khong lech o CAC DIEM TREN.")
    sys.exit(0)
print(f"\nLECH {len(loi)} diem:\n")
for muc, a, b, _ in loi:
    print(f"  {muc}")
    print(f"      .md   {a}")
    print(f"      .docx {b}")
sys.exit(1)
