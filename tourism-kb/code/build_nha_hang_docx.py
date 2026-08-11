# -*- coding: utf-8 -*-
"""Ban .docx RUT GON danh sach nha hang & quan an — chuan (1 bang 5 cot).

Doc export/<slug>/nha-hang.json (curated ~250). GIU nguyen THU TU cua export = MUC DO
ANH HUONG noi bo (VQS) — quan dong khach/anh huong cao len truoc (QUYET DINH chu san pham
2026-08-05; xem anh_huong.py). KHONG in diem/sao/luot — chi THU TU. Cot: Nha hang · Loai mon ·
Gia TB/nguoi · Dia chi · Diem Google. Gia + Diem Google = "Chua xac minh" (chi luu place_id).

Chay TU GOC REPO:  python tourism-kb/code/build_nha_hang_docx.py tourism-kb/raw[/<slug>]
"""
import json, os, sys, io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import duong_dan_ra as _dr
import anh_huong as _ah
from dia_diem_config import cfg as _cfg, slug_of as _slug_of

RAW = sys.argv[1]
SLUG = _slug_of(RAW)
C = _cfg(RAW)
CITY = C["city"]
CITY_DIR = os.path.dirname(RAW.rstrip("/\\"))       # raw/<slug> — chứa noi-bo/ (rank ảnh hưởng)
_CITY_FILE = "-".join(p.capitalize() for p in SLUG.split("-"))
EXPORT = os.path.join("tourism-kb", "export", SLUG, "nha-hang.json")
OUT = "tourism-kb/output/%s/Nha-Hang-%s.docx" % (SLUG, _CITY_FILE)
CHUA = "Chưa xác minh"

HANG_MUC_VN = {
    "coffee_shop": "Cà phê", "cafe": "Cà phê", "hong_kong_style_cafe": "Cà phê",
    "restaurant": "Nhà hàng", "eat_and_drink": "Quán ăn", "diner": "Quán ăn",
    "food": "Quán ăn", "food_stand": "Quán ăn", "bistro": "Quán ăn",
    "vietnamese_restaurant": "Nhà hàng Việt", "southern_restaurant": "Nhà hàng Việt",
    "chinese_restaurant": "Nhà hàng Trung Hoa", "taiwanese_restaurant": "Nhà hàng Đài Loan",
    "asian_restaurant": "Nhà hàng châu Á",
    "korean_restaurant": "Nhà hàng Hàn", "japanese_restaurant": "Nhà hàng Nhật",
    "sushi_restaurant": "Sushi", "thai_restaurant": "Nhà hàng Thái",
    "indian_restaurant": "Nhà hàng Ấn", "italian_restaurant": "Nhà hàng Ý",
    "pizza_restaurant": "Pizza", "french_restaurant": "Nhà hàng Pháp",
    "american_restaurant": "Nhà hàng Mỹ", "european_restaurant": "Nhà hàng Âu",
    "mediterranean_restaurant": "Nhà hàng Địa Trung Hải", "greek_restaurant": "Nhà hàng Hy Lạp",
    "malaysian_restaurant": "Nhà hàng Malaysia", "afghan_restaurant": "Nhà hàng Afghanistan",
    "australian_restaurant": "Nhà hàng Úc", "polish_restaurant": "Nhà hàng Ba Lan",
    "bakery": "Tiệm bánh", "desserts": "Tráng miệng", "ice_cream_shop": "Kem",
    "delicatessen": "Đồ nguội",
    "bubble_tea": "Trà sữa", "smoothie_juice_bar": "Nước ép · sinh tố", "tea_room": "Quán trà",
    "bar": "Quán bar", "pub": "Pub", "cocktail_bar": "Cocktail bar", "beer_bar": "Beer bar",
    "beer_garden": "Beer garden", "wine_bar": "Wine bar", "sports_bar": "Sports bar",
    "sake_bar": "Sake bar", "gastropub": "Gastropub", "brewery": "Brewery",
    "barbecue_restaurant": "Nướng · BBQ", "bar_and_grill_restaurant": "Nướng · bar",
    "steakhouse": "Steak", "seafood_restaurant": "Hải sản",
    "chicken_restaurant": "Gà", "burger_restaurant": "Burger", "taco_restaurant": "Taco",
    "noodles_restaurant": "Mì · phở", "soup_restaurant": "Súp",
    "fast_food_restaurant": "Ăn nhanh", "food_truck": "Xe đồ ăn",
    "breakfast_and_brunch_restaurant": "Điểm tâm · brunch",
    "vegetarian_restaurant": "Chay", "health_food_restaurant": "Đồ ăn healthy",
    "buffet_restaurant": "Buffet", "theme_restaurant": "Nhà hàng chủ đề",
    "molecular_gastronomy_restaurant": "Ẩm thực phân tử",
}


def shade(cell, hexc):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(el)


d = json.load(io.open(EXPORT, encoding="utf-8"))
# GIU nguyen thu tu export = MUC DO ANH HUONG noi bo (VQS). KHONG sort lai o day.
# (Doctrine cu sap theo khoang cach; QUYET DINH 2026-08-05 doi sang thu tu anh huong.)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)
_t = doc.add_paragraph().add_run("DANH SÁCH %d NHÀ HÀNG & QUÁN ĂN %s - RÚT GỌN" % (len(d), CITY.upper()))
_t.bold = True
_t.font.size = Pt(15)
# Subtitle TRUNG THỰC theo thứ tự thực tế: có noi-bo rank -> ảnh hưởng; không -> danh bạ chưa xếp hạng.
if _ah.thu_tu(CITY_DIR, "nha_hang"):
    _sub = ("Giá trung bình và điểm Google không có trong tài liệu nguồn nên được ghi "
            "“Chưa xác minh”. Danh sách sắp theo MỨC ĐỘ ẢNH HƯỞNG (quán đông khách / ảnh hưởng cao "
            "lên trước) — tài liệu KHÔNG in điểm, sao hay lượt đánh giá; thứ tự là tín hiệu.")
else:
    _sub = ("Giá trung bình và điểm Google không có trong nguồn nên ghi “Chưa xác minh”. "
            "DANH BẠ CHƯA XẾP HẠNG — sắp theo mức độ đầy đủ thông tin; THỨ TỰ KHÔNG phản ánh "
            "độ phổ biến hay ảnh hưởng (xếp hạng ảnh hưởng cần dữ liệu Google, chưa áp dụng). "
            "Tài liệu KHÔNG in điểm, sao hay lượt.")
_s = doc.add_paragraph().add_run(_sub)
_s.italic = True
_s.font.size = Pt(10)
_s.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

COLS = ["Nhà hàng / quán ăn", "Loại món", "Giá trung bình/người", "Địa chỉ", "Điểm Google"]
tbl = doc.add_table(rows=1, cols=len(COLS))
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(COLS):
    c = tbl.rows[0].cells[i]
    c.text = ""
    run = c.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(9.5)
    shade(c, "FCE4D6")

for r in d:
    row = tbl.add_row().cells
    vals = [r["name"], HANG_MUC_VN.get(r["category"].get("primary"), "Nhà hàng"),
            CHUA, r["address"].get("full_address") or CHUA, CHUA]
    for j, v in enumerate(vals):
        row[j].text = ""
        run = row[j].paragraphs[0].add_run(v)
        run.font.size = Pt(9)

for row in tbl.rows:
    for w, cell in zip([4.6, 3.0, 3.0, 4.4, 2.6], row.cells):
        cell.width = Cm(w)

p = _dr.kiem_loi_ra(OUT)
os.makedirs(os.path.dirname(p), exist_ok=True)
tmp = p + ".tmp"
doc.save(tmp)
os.replace(tmp, p)
print("saved -> %s  (%d quán)" % (OUT, len(d)))
