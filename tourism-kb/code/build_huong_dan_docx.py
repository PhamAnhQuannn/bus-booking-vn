# -*- coding: utf-8 -*-
"""Ban .docx cua huong dan diem den Da Lat.

Doc raw/guide_data.json do build_huong_dan.py xuat ra — MOT lan chon, MOT lan
hop nhat, hai dinh dang dau ra. Khong lam lai logic chon/hop nhat o day.

═══════════════════════════════════════════════════════════════════════════════
DOCSTRING CU CUA FILE NAY LA SAI, VA DAY LA BAN SUA.

No ghi: "Word lam duoc mot viec Markdown khong lam duoc: TO MAU. Ca gia tri cua
tai lieu nay nam o cho [CHƯA XÁC MINH] khong the bo qua duoc — mau do lam duoc
dieu do, chu in dam thi khong."

Co che do DA CHET. `field_table()` bo MOI dong co gia tri bat dau bang UNV
TRUOC khi render, va `value_run()` chi duoc goi tu trong `field_table()` — nen
nhanh to do trong `value_run` khong bao gio chay duoc, va vong quet the cua no
lap tren mot tuple rong. Chinh sach da doi tu "danh dau bang mau" sang "bo han
dong", va cau tren khong duoc sua theo.

VAY BAN .docx CON DE LAM GI: no la ban cho NGUOI doc soat, khong phai cho tac
nhan AI (ban .md moi la ban cua tac nhan — muc 0 noi ro the). Mau va bong to
khong co nghia gi voi mot bo trich van ban, nhung co nghia voi mot nguoi dang
quet tim cho thieu truoc khi goi dien xac minh o muc 9.

Tin hieu cho thieu gio nam o CAP NHOM, khong o cap dong: moi nhom in mot dong
xam "— N/M truong da xac minh —". Nguoi doc thay do day cua khoang trong ma
khong phai doc 1.336 dong [CHƯA XÁC MINH].
═══════════════════════════════════════════════════════════════════════════════
"""
import json, os, sys, io
import re as _re
from collections import Counter, defaultdict
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as _RT

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import hoat_dong_data as _hoat_dong   # CUNG module chon loc voi ban .md
import an_ngu_data as _an_ngu
import docx_chung as _dx             # may dinh dang Word dung chung
import xep_hang as _xep_hang         # nguong xep hang — MOT nguon, khong go tay
import duong_dan_ra as _duong_dan_ra # chan ghi ra ngoai vung da gitignore

# Duong ra co MAC DINH, va do la de chan mot loi da xay ra: ten file tung la
# lua chon cua tung lan goi, nen hai phien lam viec song song da sinh ra
# `-v4.docx` va `-v5.docx` ben canh `-v3.docx`. Ket qua la so phien ban NGUOC
# voi do moi — v5 dung 14:59 truoc khi cac ban va vao, v3 dung 16:55 sau do —
# va nguoi doc mo file so cao nhat lai thay ban cu nhat, thieu han lop mon.
# Mot duong mac dinh nghia la goi khong tham so thi GHI DE ban chinh thay vi
# them mot so moi.
OUT_MAC_DINH = "tourism-kb/raw/build/Huong-Dan-Da-Lat.docx"
RAW = sys.argv[1]

# ── BA TAI LIEU, MOT BO DUNG ───────────────────────────────────────────────
# Tach lam ba file: diem den · nha hang · khach san. Nhung KHONG tach lam ba
# script — file nay co ~450 dong may dinh dang dung chung (H, P, TBL,
# field_table, value_run, ev, has…) va docstring o tren da ghi ro thu tu chay
# la load-bearing. Chep may do sang ba noi la tai lap dung loi 30/07: mot quy
# tac dung chung boi hai bo dung ma nam trong tung bo dung thi mot ben se
# lech, va lan do ban .docx mat trang 32 link Facebook + 21 luot check-in.
#
# Nen: MOT script, chay ba lan voi mot co. Cac khoi phat noi dung duoc CONG
# bang `if TAI_LIEU in (...)`; phan nap du lieu o cap module giu nguyen vo
# dieu kien vi no re va dung thu tu san.
#
# Muc 0 (QUY TAC DOC) va khoi ghi cong CO Y khong bi cong — chung vao ca ba
# file. Nguoi chi cam file Nha hang ma khong co muc 0 se doc o trong thanh
# "quan nay khong co gio mo cua" thay vi "chua ai xac minh", va ~1.476 truong
# trong bo du lieu nay mang dau [CHƯA XÁC MINH].
TAI_LIEU = "tatca"
for _i, _a in enumerate(sys.argv):
    if _a == "--tai-lieu" and _i + 1 < len(sys.argv):
        TAI_LIEU = sys.argv[_i + 1]
_HOP_LE = {"tatca", "diemden", "nhahang", "khachsan"}
if TAI_LIEU not in _HOP_LE:
    print(f"--tai-lieu phải là một trong {sorted(_HOP_LE)}; nhận được {TAI_LIEU!r}")
    sys.exit(1)

# Ten file + tieu de + phu de theo tung tai lieu. `tatca` giu nguyen hanh vi cu
# de mot lenh khong tham so van ra dung ban gop nhu truoc.
# Ban GOP khong con la mot ban giao: `docs/` chi giu BA file (diem den · nha
# hang · khach san), va ban gop chua lai dung noi dung cua ca ba.
#
# Nhung no KHONG bi xoa, vi no la MOC NEO cua `kiem_parity.py`: bo chan do so
# ban .md (van la mot file, danh cho tac nhan AI) voi ban .docx tuong ung, va
# chi ban gop moi danh so giong ban .md. Ba file tach danh so lai theo tung tai
# lieu (3.1/3.2 thay vi 6.1/6.2), nen so ban .md voi chung se bao lech o moi
# tieu de — dung theo thiet ke, sai theo phep kiem.
#
# Nen ban gop chuyen sang thu muc build noi bo (tourism-kb/raw/build/), khong
# nam trong tourism-kb/output/ (noi chi giu ba ban phat hanh).
_CAU_HINH = {
    "tatca":    ("tourism-kb/raw/build/Huong-Dan-Da-Lat.docx",
                 "HƯỚNG DẪN ĐIỂM ĐẾN ĐÀ LẠT", None),
    "diemden":  ("tourism-kb/output/Diem-Den-Da-Lat.docx", "ĐIỂM THAM QUAN ĐÀ LẠT",
                 "Điểm tham quan · công viên · hoạt động ngoài trời"),
    "nhahang":  ("tourism-kb/output/Nha-Hang-Da-Lat.docx", "NHÀ HÀNG & QUÁN ĂN ĐÀ LẠT",
                 "Quán ăn · cà phê · đặc sản theo món"),
    "khachsan": ("tourism-kb/output/Khach-San-Da-Lat.docx", "KHÁCH SẠN & LƯU TRÚ ĐÀ LẠT",
                 "Cơ sở lưu trú đã đăng ký, theo phân khúc giá"),
}
_OUT_TL, TIEU_DE, PHU_DE = _CAU_HINH[TAI_LIEU]
# Chan truoc khi ghi. Bon dich mac dinh o tren deu dat (ba ban trong
# `tourism-kb/output/`, ban gop trong `tourism-kb/raw/build/` — deu da ignore);
# chi co ban ghi de tren dong lenh moi co the sai. Xem duong_dan_ra.py.
OUT = _duong_dan_ra.kiem_loi_ra(
    sys.argv[2] if len(sys.argv) > 2 and not sys.argv[2].startswith("--") else _OUT_TL)


def phat(*tai_lieu):
    """Khoi nay co thuoc tai lieu dang dung khong?

    `tatca` luon dung — ban gop giu nguyen moi khoi nhu truoc khi tach.
    """
    return TAI_LIEU == "tatca" or TAI_LIEU in tai_lieu
_LAN_CAN = _an_ngu.tai_lan_can(RAW)
_LCKV = _an_ngu.tai_lan_can_khu_vuc(RAW)   # khoi theo KHU VUC, in mot lan
# ── THU TU CHAY LA LOAD-BEARING, va day la cho no de vo im lang ────────────
# `guide_data.json` do build_huong_dan.py ghi ra; file nay chi DOC. Nen chay
# rieng bo dung .docx se dung lai bo diem cua lan chay .md gan nhat — khong loi,
# khong canh bao, chi la du lieu cu. Viec tach logic chon ra `diem_den_data.py`
# se xoa han lop loi nay va da duoc hoan sang commit rieng (do la thay doi rui
# ro nhat: DL-xx la id THEO VI TRI nen mot phep sort lech se doi id cua ca 36
# diem, va 11 script khac gan du lieu theo id do).
# Trong khi cho: kiem tuoi file. Re, va bien truong hop im lang thanh on ao.
_gp = os.path.join(RAW, "guide_data.json")
if not os.path.exists(_gp):
    raise SystemExit(f"KHONG co {_gp}.\nChay build_huong_dan.py truoc — no sinh ra"
                     " file nay, ban .docx chi doc lai.")
for _phu in ("enrichment.json", "lan_can_khu_vuc.json", "lan_can.json"):
    _pp = os.path.join(RAW, _phu)
    if os.path.exists(_pp) and os.path.getmtime(_pp) > os.path.getmtime(_gp) + 1:
        raise SystemExit(
            f"DUNG — {_phu} moi hon guide_data.json.\n"
            "Nghia la du lieu nguon da doi nhung bo diem chua duoc chon lai, nen"
            " ban .docx se\ndung bo diem cu. Chay lai theo dung thu tu:\n"
            "   python tourism-kb/code/build_huong_dan.py tourism-kb/raw\n"
            "   python tourism-kb/code/build_huong_dan_docx.py tourism-kb/raw")
G = json.load(io.open(_gp, encoding="utf-8"))
picked, NEAR, mat = G["picked"], G["near"], G["matrix"]
BUILD_DATE = G["build_date"]
byid = {r["id"]: r for r in picked}
UNV = "[CHƯA XÁC MINH]"

# Trong nha / ngoai troi suy tu loai hinh — mot trong ba suy dien duoc duyet.
# Truoc day bang nay nam TRONG vong lap the (bien `from_cat`), nen hai bang so
# sanh khong dung duoc, va bang so sanh cua ban .docx vi vay THIEU cot Mua ma
# ban .md co. Dua ra day de mot bang phuc vu ca ba cho.
INDOOR_CAT = {
    "Bảo tàng": "trong nhà", "Nghệ thuật / Triển lãm": "trong nhà",
    "Chợ / Mua sắm": "có mái", "Nhà thờ": "có mái",
    "Chùa / Thiền viện": "hỗn hợp", "Dinh thự / Di tích": "hỗn hợp",
    "Khu vui chơi": "hỗn hợp", "Thác nước": "ngoài trời", "Hồ / Đập": "ngoài trời",
    "Công viên / Vườn hoa": "ngoài trời", "Điểm ngắm cảnh": "ngoài trời",
    "Núi / Đèo / Đường mòn": "ngoài trời", "Nông trại / Vườn": "ngoài trời",
    "Cáp treo": "ngoài trời",
}

# Thu hang den TU build_huong_dan.py qua khoa `xep_hang` — khong tu tinh lai.
# Tinh o day vi CA bang chon nhanh (dau tai lieu) va thu tu goi dien (cuoi tai
# lieu) deu dung no.
_xh = {pid: i for i, pid in enumerate(G.get("xep_hang") or [])}
if not _xh:
    raise SystemExit("guide_data.json thieu khoa 'xep_hang' — chay"
                     " build_huong_dan.py truoc de sinh lai.")
rank = sorted(picked, key=lambda r: _xh.get(r["id"], 10 ** 6))

# ── SO MUC: mot cho duy nhat, giong ban .md ────────────────────────────────
# Lan renumber thu ba trong mot phien; hai lan truoc deu sot tham chieu, mot
# lan sot ngay trong file nay ("muc 2" trong khi ban .md da la "muc 3").
# Sau khi tach lam ba tai lieu, so muc phai danh lai THEO TUNG FILE — de nguyen
# thi file Nha hang mo dau bang "muc 6" va nguoi doc di tim muc 1-5 khong co.
# Cap phat theo dung thu tu khoi se phat trong tai lieu dang dung; khoi khong
# thuoc tai lieu do nhan None, va bat ky cho nao lo dung no se vo NGAY (TypeError
# khi format) chu khong in ra mot so sai trong im lang.
_DEM_CON = {}


def muc_con(cha):
    """Cap so muc con 1,2,3… theo THU TU PHAT THAT trong tai lieu nay.

    Song song voi `_cap_so` o cap 1, va can vi cung ly do. Truoc khi co ham nay,
    so muc con la hang so viet tay: ban khach san bo muc "An uong" nen file nha
    hang in ra "3.3" va "3.4" ma khong co 3.1 hay 3.2 — nguoi doc thay lo hong
    so va tuong tai lieu bi thieu noi dung.

    Goi NGAY TAI cho phat tieu de, khong goi truoc: bo dem chi tang khi mot muc
    that su duoc in, nen muc bi `phat()` chan se khong an mat mot so.
    """
    _DEM_CON[cha] = _DEM_CON.get(cha, 0) + 1
    return f"{cha}.{_DEM_CON[cha]}"


def _cap_so(*ten_khoi):
    """Danh so 1,2,3… cho cac khoi CO trong tai lieu nay; con lai None."""
    ra, n = {}, 0
    for t, thuoc in ten_khoi:
        if thuoc:
            n += 1
            ra[t] = n
        else:
            ra[t] = None
    return ra


_D = TAI_LIEU == "tatca"          # ban gop: giu du moi muc
# XEPHANG dat NGAY SAU muc 0, truoc moi danh sach. Quy uoc nay ap cho MOI tai
# lieu sinh ra sau nay, ke ca cac tinh khac: nguoi doc mo file ra la thay ngay
# cach doc thu tu, thay vi doc het 353 dong roi moi gap mot chu thich o cuoi noi
# rang bang khong sap theo chat luong.
_S = _cap_so(
    ("XEPHANG",  _D or TAI_LIEU in ("nhahang", "khachsan")),
    # Dieu kien o day PHAI khop cong `phat()` tai cho phat tieu de, neu khong
    # thi mot so bi cap ma khong ai chiem cho -> tai lieu thung so. Do dung la
    # loi da co: TONGQUAN truoc day cap so cho ca ba tai lieu nhung tieu de chi
    # phat duoi `if phat("diemden")`, nen ban nha hang va khach san thieu han
    # muc 1 tu lau ma khong ai de y.
    ("TONGQUAN", _D or TAI_LIEU == "diemden"),
    ("KHUVUC",   _D or TAI_LIEU == "diemden"),
    ("DIAHINH",  _D or TAI_LIEU == "diemden"),
    ("DIEMDEN",  _D or TAI_LIEU == "diemden"),
    ("HOATDONG", _D or TAI_LIEU in ("diemden", "nhahang")),
    ("ANNGU",    _D or TAI_LIEU in ("nhahang", "khachsan")),
    ("SOSANH",   _D or TAI_LIEU == "diemden"),
    ("TUYEN",    _D or TAI_LIEU == "diemden"),
    ("MATRAN",   _D or TAI_LIEU == "diemden"),
    ("KIEMCHUNG", True),
)
S_QUYTAC = 0                      # muc 0 co trong CA BA tai lieu
S_XEPHANG = _S["XEPHANG"]
S_TONGQUAN, S_KHUVUC, S_DIAHINH = _S["TONGQUAN"], _S["KHUVUC"], _S["DIAHINH"]
S_DIEMDEN, S_HOATDONG, S_ANNGU = _S["DIEMDEN"], _S["HOATDONG"], _S["ANNGU"]
S_SOSANH, S_TUYEN, S_MATRAN = _S["SOSANH"], _S["TUYEN"], _S["MATRAN"]
S_KIEMCHUNG = _S["KIEMCHUNG"]


# lop lam giau — doc cung mot file enrichment.json nhu ban .md, va CUNG mot phep
# loc element-khop-qua-xa (`an_ngu_data.loc_khop_xa`). Neu chi mot ban loc thi
# ban kia in gio mo cua cua hang xom, va khong co gi bao.
_ep = os.path.join(RAW, "enrichment.json")
ENR = {}
_GIO_LOAI, _ENR_BO = set(), []
if os.path.exists(_ep):
    _rows, _ENR_BO = _an_ngu.loc_khop_xa(json.load(io.open(_ep, encoding="utf-8")))
    if _ENR_BO:
        print(f"  loai {len(_ENR_BO)} dong enrichment tu element OSM khop >= "
              f"{_an_ngu.BAN_KINH_VAN_HANH} m (khong phai dia diem nay)")
    _GIO_LOAI = _an_ngu.gio_bi_loai(_ENR_BO)
    for _e in _rows:
        ENR.setdefault(_e["id"], {}).setdefault(_e["field"], _e)


def ev(pid, field, tail=""):
    e = ENR.get(pid, {}).get(field)
    if not e:
        return UNV + tail
    return str(e["value"])


def has(pid, field):
    return field in ENR.get(pid, {})


def _canh_bao_dia_chi(pid):
    """CUNG cach dien dat voi ban .md — xem chu thich o `an_ngu_data.dia_chi_mau_thuan`."""
    mt = _an_ngu.dia_chi_mau_thuan(pid, ENR, _ENR_BO)
    if not mt:
        return ""
    md, bg, pho, loai = mt
    return (f"  ⚠ MÂU THUẪN SỐ NHÀ trên đường {pho}: bản mô tả ghi “{md}”, bản ghi "
            f"bản đồ ghi “{bg}”"
            + (" — bản ghi bản đồ ĐÃ BỊ LOẠI vì khớp nhầm sang một cơ sở khác, nên "
               "không còn số nhà nào đã xác minh." if loai else ".")
            + " CHƯA hoà giải — hỏi lại cổng vào hiện tại trước khi chỉ đường cho khách.")

RED, AMBER, GREY, INK = _dx.RED, _dx.AMBER, _dx.GREY, _dx.INK


# Tai lieu chi danh cho nguoi doc tieng Viet. Ten dia diem trong nguon co the
# kem chu Han/Hangul/Kana/Kirin (vd "Thiên Vương Cổ Sát Chùa Tàu - 大叻市 天王古剎").
# Do la du lieu that, nhung khong danh cho doc gia nay.
_FOREIGN = __import__("re").compile(r"[가-힯一-鿿぀-ヿЀ-ӿ]+")


def vn_only(s):
    """Bo chu ngoai (Han, Hangul, Kana, Kirin). Giu tieng Viet va chu Latin."""
    s = _FOREIGN.sub("", str(s or ""))
    s = __import__("re").sub(r"\s{2,}", " ", s)
    return s.strip(" ,-·/")

# May dinh dang o `docx_chung`, KHONG dinh nghia lai o day — xem docstring cua
# module do. `value_run`/`field_table` o lai ben duoi vi chung gan voi ngu nghia
# [CHƯA XÁC MINH] cua rieng bo huong dan.
_KHOI = _dx.Khoi()
doc = _KHOI.doc
H, P, B = _KHOI.H, _KHOI.P, _KHOI.B
TBL, shade, lap_dong_dau = _KHOI.TBL, _KHOI.shade, _KHOI.lap_dong_dau
them_lien_ket = _KHOI.lien_ket

LINK = _dx.LINK
_URL = _re.compile(r"https?://\S+")


def value_run(par, text):
    """To mau theo muc tin cay. Day la ly do ban .docx ton tai."""
    # URL -> lien ket that. Dat TRUOC nhanh UNV vi mot o chi chua URL thi khong
    # bao gio bat dau bang UNV, con o vua co chu vua co URL thi van di duong duoi.
    m = _URL.search(text)
    if m and not text.startswith(UNV):
        truoc, sau = text[:m.start()], text[m.end():]
        if truoc:
            r = par.add_run(truoc)
            r.font.size = Pt(9.5)
            r.font.color.rgb = INK
        them_lien_ket(par, m.group(0))
        if sau:
            r = par.add_run(sau)
            r.font.size = Pt(9.5)
            r.font.color.rgb = INK
        return
    if text.startswith(UNV):
        r = par.add_run(UNV)
        r.bold = True
        r.font.color.rgb = RED
        r.font.size = Pt(9.5)
        rest = text[len(UNV):]
        if rest:
            r2 = par.add_run(rest)
            r2.italic = True
            r2.font.size = Pt(8.5)
            r2.font.color.rgb = RED
        return
    head, tag = text, ""
    for mark in ():
        if mark in text:
            i = text.index(mark)
            head, tag = text[:i], text[i:]
            break
    if head:
        r = par.add_run(head)
        r.font.size = Pt(9.5)
        r.font.color.rgb = INK
    if tag:
        r = par.add_run(tag)
        r.font.size = Pt(8)
        r.italic = True
        r.font.color.rgb = AMBER if tag.startswith("[SUY RA") else GREY


def field_table(rows_spec):
    """rows_spec: ('group', title) | ('f', nhan, gia tri) | ('note', text)

    Chi giu truong DA XAC MINH: moi dong co gia tri [CHƯA XÁC MINH] bi bo, va nhom
    nao khong con dong nao thi bo luon tieu de nhom.
    """
    # ── DEM truoc khi loc, de con biet da bo bao nhieu ────────────────────
    # Ly do ton tai cua ban .docx, theo docstring cua chinh no, la TO MAU dong
    # [CHƯA XÁC MINH] cho khong the bo qua. Nhung dong loc ngay duoi xoa han
    # nhung dong do TRUOC khi render, va `value_run` chi duoc goi tu ham nay —
    # nen nhanh to do trong `value_run` khong bao gio chay duoc. Co che ma tai
    # lieu vien dan de ton tai da chet tu luc chinh sach doi tu "danh dau" sang
    # "bo han".
    # Khong dua dong do tro lai (do la quyet dinh co y). Thay vao do dua tin
    # hieu len MOT CAP: moi nhom ghi ro con bao nhieu tren tong bao nhieu. Doc
    # gia thay duoc do day cua khoang trong ma khong phai doc 1.336 dong do.
    tong_nhom, con_nhom, nhom_ht = {}, {}, None
    for x in rows_spec:
        if x[0] == "group":
            nhom_ht = x[1]
            tong_nhom.setdefault(nhom_ht, 0)
            con_nhom.setdefault(nhom_ht, 0)
        elif x[0] == "f" and nhom_ht:
            tong_nhom[nhom_ht] += 1
            if not str(x[2]).startswith(UNV):
                con_nhom[nhom_ht] += 1

    kept = [x for x in rows_spec if not (x[0] == "f" and str(x[2]).startswith(UNV))]
    out = []
    for i, x in enumerate(kept):
        if x[0] == "group":
            nxt = kept[i + 1] if i + 1 < len(kept) else None
            if nxt is None or nxt[0] == "group":
                continue
        out.append(x)
    # Chen dong "con N/M truong" vao cuoi moi nhom con song.
    with_note, nhom_ht = [], None
    for i, x in enumerate(out):
        if x[0] == "group":
            if nhom_ht and tong_nhom.get(nhom_ht, 0) > con_nhom.get(nhom_ht, 0):
                with_note.append(("done", f"{con_nhom[nhom_ht]}/{tong_nhom[nhom_ht]}"
                                          " trường đã xác minh"))
            nhom_ht = x[1]
        with_note.append(x)
    if nhom_ht and tong_nhom.get(nhom_ht, 0) > con_nhom.get(nhom_ht, 0):
        with_note.append(("done", f"{con_nhom[nhom_ht]}/{tong_nhom[nhom_ht]}"
                                  " trường đã xác minh"))
    rows_spec = with_note
    # Danh so o DAY, sau khi da loc. Dem luc dung spec thi van de lai lo: nhom
    # "Luu y quan trong" nhan so 5 roi bi xoa vi ca 7 truong con deu 0/36, va moi
    # ho so hien A.1-A.4 roi nhay sang A.6 — dung cai day so nhay ma viec danh so
    # lai sinh ra de sua, chi dich mot buoc.
    _muc[0] = 0
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for spec in rows_spec:
        cells = t.add_row().cells
        if spec[0] == "group":
            cells[0].merge(cells[1])
            c = t.rows[-1].cells[0]
            c.text = ""
            _muc[0] += 1
            r = c.paragraphs[0].add_run(f"A.{_muc[0]} — {spec[1]}")
            r.bold = True
            r.font.size = Pt(9.5)
            shade(c, "E8EDF3")
            continue
        if spec[0] == "done":
            # Xam nhat, khong phai do: day khong phai canh bao, chi la thuoc do
            # do day. Do danh cho `note` — thu nguoi doc phai hanh dong.
            cells[0].merge(cells[1])
            c = t.rows[-1].cells[0]
            c.text = ""
            r = c.paragraphs[0].add_run("— " + spec[1] + " —")
            r.italic = True
            r.font.size = Pt(7.5)
            r.font.color.rgb = GREY
            continue
        if spec[0] == "note":
            cells[0].merge(cells[1])
            c = t.rows[-1].cells[0]
            c.text = ""
            r = c.paragraphs[0].add_run(spec[1])
            r.italic = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RED
            shade(c, "FDF2F2")
            continue
        cells[0].text = ""
        r = cells[0].paragraphs[0].add_run(spec[1])
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        cells[0].width = Cm(4.6)
        cells[1].text = ""
        value_run(cells[1].paragraphs[1] if False else cells[1].paragraphs[0], spec[2])
        cells[1].width = Cm(11.4)
    doc.add_paragraph()
    return t


# ==================================================== bia + quy tac doc
H(TIEU_DE, 0)
if PHU_DE:
    P(PHU_DE, italic=True, size=11)
    # Ba tai lieu la BA phan cua mot bo. Noi ra, de nguoi cam mot file biet hai
    # file kia ton tai thay vi tuong tai lieu nay thieu noi dung.
    P("Một trong ba tài liệu: Điểm tham quan · Nhà hàng & quán ăn · Khách sạn"
      " & lưu trú. Ba file dùng chung quy tắc đọc ở mục 0.",
      italic=True, size=9, color=GREY)
else:
    P(f"Hồ sơ chi tiết {len(picked)} điểm đến · sinh tự động ngày {BUILD_DATE}",
      italic=True, size=11)
if PHU_DE:
    P(f"Sinh tự động ngày {BUILD_DATE}", italic=True, size=9, color=GREY)
P("Nguồn: OpenStreetMap · Overture Maps · Foursquare OS · Wikidata · OSRM · "
  "đăng ký lưu trú Cục Du lịch Quốc gia · "
  # Wikipedia PHAI co o day: cac doan mo ta trich nguyen van tu Wikipedia, va
  # CC BY-SA 4.0 doi ghi cong kem lien ket.
  "Wikipedia tiếng Việt (CC BY-SA 4.0, https://vi.wikipedia.org)",
  italic=True, size=9, color=GREY)

H(f"{S_QUYTAC}. QUY TẮC ĐỌC — BẮT BUỘC", 1)
P("QUAN TRỌNG — tài liệu này CHỈ liệt kê những trường ĐÃ XÁC MINH. "
  "Một trường KHÔNG xuất hiện trong hồ sơ nghĩa là CHƯA BIẾT, không phải là “không có”. "
  "Không được suy đoán giá trị cho trường vắng mặt.", bold=True, color=RED)
P("Ba điều tuyệt đối không được làm:", bold=True)
B("KHÔNG thay [CHƯA XÁC MINH] bằng một giá trị thường gặp. “Giờ mở cửa 08:00–17:00” là hình "
  "dạng của một giờ mở cửa, không phải giờ mở cửa của nơi này.")
B("KHÔNG suy ra giá vé, giờ mở cửa, thời lượng thăm hay mức độ dễ đi lại từ loại hình. "
  "Chỉ ba suy diễn được duyệt: trong nhà/ngoài trời, link bản đồ, điểm lân cận.")
B("KHÔNG TỰ VIẾT mô tả, “lý do nên đến” hay “điểm nhấn”. Đoạn mô tả trong hồ sơ (nếu có) là TRÍCH NGUYÊN VĂN từ Wikipedia tiếng Việt, kèm nguồn và ngày — dẫn nguồn khi đọc cho khách, và KHÔNG sửa lời. Tài liệu cố ý KHÔNG sinh "
  "những mục đó vì mọi chữ trong đó sẽ là bịa. Mục 3 liệt kê hoạt động kèm nơi và "
  "đơn vị cụ thể, đó là dữ kiện; “Đà Lạt lãng mạn” thì không.")
P("Nhịp độ mặc định (chuyến thư giãn): tối đa 4 điểm/ngày · tối đa 2 giờ di chuyển/ngày · "
  "mỗi ngày chừa một khoảng trống.", bold=True)
# ── Ba quy tac ap cho CA 36 ho so, noi mot lan ──────────────────────────
P("Ba điều dưới đây áp cho cả 36 hồ sơ, không nhắc lại ở từng điểm:", bold=True, size=9.5)
B("Mỗi hồ sơ chỉ nêu những trường KHÔNG mang dấu [CHƯA XÁC MINH]. Trường vắng mặt "
  "nghĩa là chưa xác minh được — nói với khách đúng như vậy, đừng suy ra.")
B("Khoảng cách từ khách sạn của khách không có trong tài liệu này: nó thuộc hồ sơ "
  "chuyến đi và chỉ tính được khi biết khách ở đâu.")
B("Khoảng cách và thời gian đường bộ là ước lượng trong điều kiện bình thường, chưa "
  f"tính chỗ đậu xe hay tắc đường. Cột Rộng ở mục {S_KHUVUC} là ĐƯỜNG CHIM BAY, còn Km từ trung "
  "tâm và bảng điểm lân cận là ĐƯỜNG BỘ THẬT theo OSRM — hai thước đo khác nhau.")
P("Bay flycam: mặc định COI NHƯ BỊ CẤM trừ khi có xác nhận ngược lại. Sai theo hướng an toàn "
  "thì mất một tấm ảnh; sai theo hướng kia thì khách bị phạt.", bold=True, color=RED)

# ==================================================== XEP HANG — dat DAU tien
# Quy uoc cho MOI tai lieu sinh ra sau nay, ke ca cac tinh khac: muc nay dung
# ngay sau muc 0, TRUOC moi danh sach. Truoc day no nam o cuoi muc 3.x, nen
# nguoi doc phai luot het 353 dong moi gap dong noi rang bang khong sap theo
# chat luong — mot canh bao dat sau thu no canh bao thi khong con la canh bao.
#
# Muc nay chua QUY TAC, khong chua con so. Nguong doc tu `mo_ta_nguong()`, KHONG
# go tay o day — go tay la cach tai lieu va code bat dau noi hai con so khac
# nhau, va khong co gi bao khi chung lech.
if phat("nhahang", "khachsan"):
    H(f"{S_XEPHANG}. XẾP HẠNG THEO ĐÁNH GIÁ — ĐỌC MỤC NÀY TRƯỚC", 1)
    P(_xep_hang.mo_ta_nguong(), bold=True, size=10)
    P("Cận dưới Wilson trả lời “với bằng chứng đang có, chất lượng thật thấp "
      "nhất có thể là bao nhiêu”. Nên một nơi đạt điểm tuyệt đối với vài chục "
      "lượt KHÔNG chắc hơn một nơi kém hơn chút ít với vài trăm lượt — xếp theo "
      "điểm thô sẽ cho hạng cao nhất cho cơ sở ÍT khách nhất, tức đảo ngược "
      "đúng thứ người ta muốn biết.", size=9.5)
    P("Hai trạng thái âm KHÔNG được gộp:", bold=True, size=9.5)
    B("“— dưới chuẩn” là một KẾT LUẬN: đủ đánh giá để nói, và kết quả không đạt.")
    B(f"“chưa đủ đánh giá” (dưới {_xep_hang.SAN_TOI_THIEU} lượt) là một KHOẢNG "
      "TRỐNG: chưa nói được gì. Một quán mới mở tốt thật nằm ở nhóm này, không "
      "phải nhóm trên. Đừng đọc nó thành “dở”.")
    P("Tài liệu này KHÔNG in điểm, số lượt hay hạng của bất kỳ cơ sở nào.",
      bold=True, color=RED)
    P("Điều khoản Google Places cho lưu place_id vĩnh viễn nhưng cấm lưu mọi nội "
      "dung khác, và không có ngoại lệ cho dữ liệu suy ra từ chúng. Lý do thứ hai "
      "độc lập với giấy phép: một thứ hạng in ra giấy sẽ cũ đi mà KHÔNG có gì "
      "trong tờ giấy báo rằng nó đã sai. Bảng được tính LẠI mỗi lần chạy:",
      size=9, color=GREY)
    for _lenh, _mo in (("quan_hxh", "quán ăn"), ("luu_tru_hxh", "cơ sở lưu trú")):
        _p = P("")
        _r = _p.add_run(f"python tourism-kb/code/xep_hang_song.py tourism-kb/raw {_lenh}")
        _r.font.name, _r.font.size = "Consolas", Pt(9)
        _r2 = _p.add_run(f"   → {_mo}")
        _r2.font.size, _r2.italic = Pt(9), True
        _r2.font.color.rgb = GREY
    P("Hoặc mở liên kết ở cột Bản đồ của từng bảng — trang Google hiện ra là "
      "chính cơ sở đó, kèm điểm và số lượt của hôm nay.", size=9, color=GREY)
    P("Các bảng bên dưới sắp theo KHOẢNG CÁCH, không theo chất lượng.",
      bold=True, size=9.5)

# ==================================================== 1. tong quan
if phat("diemden"):
    H(f"{S_TONGQUAN}. Tổng quan điểm đến", 1)
    TBL(["Mục", "Giá trị"],
        [["Thành phố", "Đà Lạt, tỉnh Lâm Đồng"],
         ["Số điểm trong hồ sơ", str(len(picked))],
         ["Kho dữ liệu đầy đủ", "diem-tham-quan.md — 1.361 điểm"],
         *([["Hướng mặt trời mọc", (_an_ngu.tai_dia_hinh(RAW) or {}).get("binh_minh")
             + " — chung cho cả thành phố, không khác nhau giữa các điểm"]]
           if (_an_ngu.tai_dia_hinh(RAW) or {}).get("binh_minh") else []),
         ["Thời tiết theo tháng", UNV],
         ["Lịch lễ hội", UNV],
         ["Ảnh hưởng Tết", UNV + " — nhiều nơi đóng cửa, giá tăng mạnh"],
         ["Đi lại tới Đà Lạt", UNV + " — chưa thu thập tuyến xe / máy bay"],
         ["Phương tiện tại chỗ", UNV + " — chưa thu thập giá thuê xe / taxi"]],
        widths=[5.0, 11.0], size=9)
    P("⚠ Năm hàng cuối là khoảng trống có thật, không phải lỗi hiển thị. Một lịch trình không biết "
      "khách tới bằng gì và đi lại bằng gì thì chưa phải một lịch trình.", bold=True, color=RED)

# ==================================================== 2. ho so tung diem
# ── muc MOI. Cau hoi thuong gap nhat — "di 3-5 ngay thi chia the nao" —
# truoc day khong muc nao tra loi o cap KHU VUC, la cap nguoi ta thuc su chia
# ngay. Moi con so rut tu guide_data.json + lan_can_khu_vuc.json.
if phat("diemden"):
    doc.add_page_break()
    H(f"{S_KHUVUC}. Tổng quan theo khu vực", 1)
    P("Chia ngày theo khu vực, không theo từng điểm: các điểm trong một khu vực đủ gần "
      "để đi liền trong cùng buổi.", italic=True, size=9, color=GREY)
    _rows_kv = []
    for _a in sorted({r["area"] for r in picked}):
        _ps = [r for r in picked if r["area"] == _a]
        _kv = _LCKV.get(_a) or {}
        _n_ks = sum(len(b["khach_san"]) for b in _kv.get("bac_khach_san") or [])
        if _kv.get("khong_co_khach_san"):
            _luu = f"không có — xem mục {S_ANNGU}"
        elif _n_ks:
            _luu = f"{_kv.get('tong_ks', 0)} cơ sở có giá"
        else:
            _luu = "—"
        _rong = _kv.get("ban_kinh_khu_vuc", "—")
        if _kv.get("canh_bao_khoang_cach"):
            _rong += "  ⚠"
        _rows_kv.append([_a, str(len(_ps)),
                         f"{min(r['km'] for r in _ps):.1f} – {max(r['km'] for r in _ps):.1f}",
                         _rong, _luu])
    TBL(["Khu vực", "Điểm", "Km từ trung tâm", "Rộng", "Lưu trú trong khu vực"],
        _rows_kv, widths=[4.6, 1.2, 3.2, 2.2, 4.8], size=9)
    P("⚠ Khu vực có dấu ⚠ ở cột Rộng thì các điểm trong đó cách nhau xa hơn bán kính 5 km "
      "dùng để tìm cơ sở gần — chúng KHÔNG dùng chung một thị trường lưu trú, nên đừng "
      "gộp vào một đêm nghỉ.", size=8.5, color=AMBER)

    # ── Bang chon nhanh: 16 dong dau cua bang so sanh, dat NGAY DAU ───────────
    # Bang day du van o muc 7 nhung nam o ~92% chieu dai tai lieu.
    _TOP = 16
if phat("diemden"):
    doc.add_page_break()
    H(f"{S_KHUVUC}b. Chọn nhanh — 16 điểm có dữ liệu đầy đủ nhất", 1)
    P("⚠ Xếp theo độ đầy đủ của DỮ LIỆU, không phải chất lượng trải nghiệm. Điểm đứng đầu "
      "là điểm ta biết rõ nhất, không phải điểm đáng đi nhất — thứ đó chưa có dữ liệu. "
      f"Bảng đầy đủ 36 điểm ở mục {S_SOSANH}.", bold=True, size=9, color=RED)
    TBL(["ID", "Điểm", "Loại", "Khu vực", "Km", "Phút", "Vé", "Mưa", "Mô tả"],
        [[r["id"], r["name"][:26], r["loai_vn"], r["area"][:16], f"{r['km']:.1f}",
          f"{r['min']:.0f}", _an_ngu.the_fee_ngan(r.get("fee")),
          INDOOR_CAT.get(r["loai_vn"], "ngoài trời"),
          "có" if has(r["id"], "mo_ta_wikipedia") else "—"]
         for r in rank[:_TOP]],
        widths=[1.4, 4.0, 2.6, 2.4, 1.1, 1.1, 1.6, 1.4, 1.2], size=8)

    # Lop nay do Phase L thu thap (SRTM 30 m) va CHUA TUNG duoc in: do_cao 36/36,
    # do_nho 36/36, huong_mo 26/36 nam trong enrichment.json tu 28/07. Lan thu nam
    # trong du an nay mot lop du lieu duoc thu roi khong ai doc.
# KHONG co cot phuong vi mat troi moc: gia tri giong nhau o ca 36 dong nen no
# khong phan biet duoc diem nao voi diem nao. No o muc 1, mot dong.
_DH = _an_ngu.tai_dia_hinh(RAW)
if phat("diemden"):
    if _DH and _DH["hang"]:
        doc.add_page_break()
        H(f"{S_DIAHINH}. Ngắm cảnh · săn mây · chụp ảnh", 1)
        P("Suy ra từ mô hình độ cao SRTM 30 m — KHÔNG phải quan sát thực địa. Độ nhô là "
          "độ cao của điểm trừ trung vị vùng xung quanh: số dương nghĩa là cao hơn cảnh "
          "quan quanh nó nên tầm nhìn thoáng, số âm nghĩa là bị che.",
          italic=True, size=9, color=AMBER)
        TBL(["Điểm", "Độ cao", "Độ nhô", "Hướng mở"],
            [[f"{x['id']} · {x['ten']}"[:38] + ("  ⚠" if x.get("canh_bao") else ""),
              x["do_cao"], x["do_nho"] or "", x["huong_mo"] or "—"]
             for x in _DH["hang"]], widths=[7.0, 2.4, 2.4, 4.2], size=8.5)
        for x in _DH["hang"]:
            if x.get("canh_bao"):
                P(f"⚠ {x['id']} · {x['ten']} — {x['canh_bao']}. Mọi con số ở dòng này nói "
                  "về khu cổng, không nói về đỉnh.", bold=True, size=9, color=RED)
        P("Toạ độ Đỉnh Langbiang chưa sửa. Vì vậy dòng cao nhất bảng này là "
          f"{_DH['hang'][0]['ten']} ({_DH['hang'][0]['do_cao']}), không phải Langbiang — "
          "đó là hệ quả của toạ độ sai, không phải sự thật về địa hình Đà Lạt.",
          size=9, color=AMBER)

if phat("diemden"):
    doc.add_page_break()
    H(f"{S_DIEMDEN}. Danh sách điểm đến", 1)
    P("Thứ tự các mục trong mỗi hồ sơ là cổng lọc trước, mô tả sau: nhận dạng → khả năng tiếp cận "
      "→ kế hoạch thăm → giờ giấc. Một ràng buộc về đi lại loại bỏ địa điểm trước khi chi tiết "
      "chụp ảnh có ý nghĩa gì.", italic=True)

    INDOOR_NOTE = {"trong nhà": "không gian trong nhà", "có mái": "phần lớn có mái che",
                   "hỗn hợp": "vừa có mái vừa ngoài trời", "ngoài trời": "địa hình ngoài trời"}
    _muc = [0]


    def sec(nhan):
        """Nhan tieu muc, DEM theo tung ho so — giong `sec()` cua ban .md.

        Chuoi cung `A.1/A.11/A.9/A.3/A.4/A.10/A.13` de lai lo `A.2`, `A.5`-`A.8`,
        `A.12` khong ton tai o BAT KY ho so nao, va nguoi doc ket luan tai lieu bi
        thieu muc. Dem theo the thi luon lien tuc, ke ca khi mot nhom bi bo loc
        rong xoa han.
        """
        _muc[0] += 1
        return f"A.{_muc[0]} — {nhan}"


    def khoi_khu_vuc(kv):
        """Khach san + quan an cho CA KHU VUC, in mot lan o dau khu vuc."""
        v = _LCKV.get(kv)
        if not v:
            return
        P(f"Lưu trú & ăn uống trong khu vực — {v['so_diem']} điểm, cách nhau tối đa "
          f"{v['ban_kinh_khu_vuc']}", bold=True, size=9.5)
        if v.get("canh_bao_khoang_cach"):
            P("⚠ " + v["canh_bao_khoang_cach"], size=8.5, color=AMBER)
        if v.get("khong_co_khach_san"):
            P(v["khong_co_khach_san"] + f" Xem mục {S_ANNGU} để chọn theo bậc giá.",
              size=9, color=AMBER)
        for b in v["bac_khach_san"]:
            P(f"{b['ten']} — {b['tong']} cơ sở trong khu vực, {b['tong_thanh_pho']} "
              "trên toàn Đà Lạt", italic=True, size=9, color=GREY)
            # Cot `Địa chỉ` — cung mot thay doi voi ban .md, cung phep cat o dau phay.
            # Tong be rong giu nguyen 16.0 cm, thu hep cac cot cu de nhuong cho.
            TBL(["Khách sạn", "Giá/đêm", "Cách", "Gần", "Phòng", "Điện thoại",
                 "Địa chỉ", "Thẩm định"],
                [[h["ten"][:30], h["gia"] or "", h["khoang_cach"], h["gan_diem"],
                  str(h["so_phong"] or ""), h["dien_thoai"] or "",
                  (h.get("dia_chi") or "").split(",")[0][:24], h["tham_dinh"] or ""]
                 for h in b["khach_san"]],
                widths=[3.6, 2.3, 1.2, 1.2, 1.0, 2.4, 2.5, 1.8], size=8)
        if v["loai_quan"]:
            P(f"Quán ăn — {v['tong_quan']} quán còn mở trong khu vực",
              italic=True, size=9, color=GREY)
            TBL(["Loại", "Quán", "Cách", "Gần", "Điện thoại", "Địa chỉ"],
                [[l["ten"], q["ten"][:32], q["khoang_cach"], q["gan_diem"],
                  q["dien_thoai"] or "", (q.get("dia_chi") or "").split(",")[0][:26]]
                 for l in v["loai_quan"] for q in l["quan"]],
                widths=[2.6, 4.2, 1.3, 1.2, 2.4, 4.3], size=8)


    cur_area = None
    for r in picked:
        if r["area"] != cur_area:
            cur_area = r["area"]
            doc.add_page_break()
            H(f"Khu vực: {cur_area}", 2)
            khoi_khu_vuc(cur_area)
        H(f"{r['id']} · {r['name']}", 3)
        srcs = "+".join(r["src"])
        # Mo ta trich nguyen van tu Wikipedia — cung nguon voi ban .md.
        # Khong dien dat lai: sua van ban CC BY-SA 4.0 la tao "Adapted Material" va
        # lam phat sinh nghia vu chia se tuong tu cho chinh doan da sua.
        if has(r["id"], "mo_ta_wikipedia"):
            _e = ENR[r["id"]]["mo_ta_wikipedia"]
            P(_e["value"], italic=True, size=9.5)
            P(f"— {_e.get('source') or 'Wikipedia tiếng Việt'} · trích nguyên văn · "
              f"{_e.get('date') or ''}" + (f" · {_e['url']}" if _e.get("url") else ""),
              italic=True, size=8, color=GREY)
        spec = [("group", "Nhận dạng"),
                ("f", "Tên", r["name"]),
                ("f", "Tên khác", ", ".join(vn_only(x) for x in (r.get("alt") or []) if vn_only(x)) or UNV),
                ("f", "Loại hình", r["loai_vn"]
                 + (f" (phụ: {', '.join(r['loai_phu'])})" if r.get("loai_phu") else "")
                 ),
                # Canh bao dat TRUOC gia tri, khong sau: `field_table()` bo moi dong
                # co gia tri BAT DAU bang UNV (`:280`), nen voi XQ Su Quan — noi ban
                # ghi dia chi da bi loai va gia tri con lai dung la UNV — canh bao
                # mau thuan se bi xoa cung ca dong, va ban .docx im lang trong khi
                # ban .md canh bao. Do la dung lop lech-hai-bo-dung ma muc nay ton
                # tai de chan.
                ("f", "Địa chỉ", (_canh_bao_dia_chi(r["id"]).strip() + "  "
                                  if _canh_bao_dia_chi(r["id"]) else "")
                 + ev(r["id"], "dia_chi_day_du")),
                ("f", "Điện thoại", r.get("tel") or UNV),
                ("f", "Website", r.get("web") or UNV),
                ("f", "Tình trạng hoạt động", "đang hoạt động")]
        # Muc do pho bien de o day, KHONG phai o "Danh gia cua khach" — ti le de
        # xuat cua Facebook la thang do khac, khong quy doi sang sao duoc.
        for _lab, _k in (("Email", "email"),
                         ("Ảnh (tự do bản quyền)", "anh"),
                         ("Website chính thức", "website_chinh_thuc"),
                         ("⚠ Website đã lưu", "canh_bao_website"),
                         ("Kiểm tra trang web", "kiem_tra_website")):
            # Nam truong xuat xu (Trang Facebook, Email FB, Luot check-in, Nguoi
            # theo doi, Ti le de xuat) da chuyen sang PHU LUC NGHIEN CUU — chung la
            # xuat xu, khong phai thong tin di choi.
            # `canh_bao_website` o tren thi KHONG chuyen: no noi rang URL in ngay
            # duoi no khong phai trang cua dia diem nay. Mot canh bao phai nam canh
            # thu no canh bao.
            if has(r["id"], _k):
                spec.append(("f", _lab, ev(r["id"], _k)))
        spec += [
                ("group", "Tiện nghi tại chỗ  ·  cổng lọc đầu tiên")]
        _fac = [("Nhà vệ sinh", "nha_ve_sinh"), ("Bãi đỗ xe", "bai_do_xe"),
                ("Chỗ ngồi nghỉ", "cho_ngoi"), ("Hàng ăn", None), ("Hàng nước", "nuoc_uong"),
                ("Quà lưu niệm", "qua_luu_niem"), ("Hướng dẫn viên", None),
                ("Quầy thông tin", "quay_thong_tin"), ("Lối cho xe lăn", "loi_xe_lan"),
                ("Sơ cứu y tế", "so_cuu")]
        _nf = 0
        for _lab, _k in _fac:
            if _k and has(r["id"], _k):
                _nf += 1
                spec.append(("f", _lab, ev(r["id"], _k)))
            else:
                spec.append(("f", _lab, UNV))
        if has(r["id"], "wifi"):
            spec.append(("f", "Wifi", ev(r["id"], "wifi")))
        ind = None
        from_cat = INDOOR_CAT if False else {
                    "Bảo tàng": "trong nhà", "Nghệ thuật / Triển lãm": "trong nhà",
                    "Chợ / Mua sắm": "có mái", "Nhà thờ": "có mái",
                    "Chùa / Thiền viện": "hỗn hợp", "Dinh thự / Di tích": "hỗn hợp",
                    "Khu vui chơi": "hỗn hợp", "Thác nước": "ngoài trời", "Hồ / Đập": "ngoài trời",
                    "Công viên / Vườn hoa": "ngoài trời", "Điểm ngắm cảnh": "ngoài trời",
                    "Núi / Đèo / Đường mòn": "ngoài trời", "Nông trại / Vườn": "ngoài trời",
                    "Cáp treo": "ngoài trời"}
        ind = from_cat.get(r["loai_vn"])
        # CUNG thu tu uu tien voi ban .md: dong enrichment (co nguon + ngay) thang
        # `r["hours"]` cua lop hop nhat (khong nguon). Xem chu thich o build_huong_dan.py.
        _h = r.get("hours")
        if (r["id"], (_h or "").strip()) in _GIO_LOAI:
            _h = None
        _he = ev(r["id"], "gio_mo_cua") if has(r["id"], "gio_mo_cua") else None
        if _he:
            _gio = _he
            if _h and _h.strip() != _he.strip():
                _gio += (f"  ⚠ lớp bản đồ ghi khác: “{_h}” — hai nguồn CHƯA hoà giải, "
                         "gọi xác nhận trước khi báo khách")
        else:
            _gio = _h if _h else UNV
        spec += [("group", "Kế hoạch thăm"),
                 ("f", "Thời lượng thăm", UNV + " ← không có mục này thì không xếp được lịch một ngày"),
                 ("f", "Thời điểm tốt trong ngày", UNV),
                 ("f", "Mùa tốt nhất", UNV),
                 ("f", "Trong nhà / ngoài trời",
                  ind if ind else UNV),
                 ("f", "Quãng đi bộ", UNV),
                 ("f", "Độ khó", UNV),
                 ("f", "Phù hợp trẻ nhỏ", UNV),
                 ("f", "Phù hợp người cao tuổi", UNV),
                 ("group", "Giờ giấc và chi phí"),
                 ("f", "Ngày mở cửa", UNV),
                 ("f", "Giờ mở cửa", _gio),
                 ("f", "Giờ nhận khách cuối", UNV),
                 # `fee` cua OSM khong phai so tien — xem `an_ngu_data.doc_the_fee`.
                 # Chu thich dat TRUOC UNV vi ly do o o "Địa chỉ" ngay tren.
                 ("f", "Giá vé", (f"({_an_ngu.doc_the_fee(r.get('fee'))})  "
                                  if _an_ngu.doc_the_fee(r.get("fee")) else "")
                  + UNV + " ← KHÔNG nêu số tiền; nói giá có thể thay đổi"),
                 ("f", "Phí gửi xe", UNV),
                 ("f", "Cần đặt trước", ev(r["id"], "can_dat_truoc"))]
        if has(r["id"], "gia_ve_tham_khao"):
            spec.append(("f", "Giá vé THAM KHẢO", ev(r["id"], "gia_ve_tham_khao")))
        if has(r["id"], "khoang_gia_facebook"):
            spec.append(("f", "Khoảng giá (Facebook)",
                         ev(r["id"], "khoang_gia_facebook") + "  (thang 4 bậc, chủ cơ sở tự khai)"))
        if has(r["id"], "gia_ve_ghi_chu"):
            spec.append(("f", "Ghi chú giá vé", ev(r["id"], "gia_ve_ghi_chu")))
        if r.get("csdl_gia_min"):
            g = f"{r['csdl_gia_min']:,}".replace(",", ".")
            if r.get("csdl_gia_max") and r["csdl_gia_max"] != r["csdl_gia_min"]:
                g += "–" + f"{r['csdl_gia_max']:,}".replace(",", ".")
            spec.append(("f", "Giá phòng (lưu trú)", f"{g}₫/đêm"))
        spec.append(("group", "Lưu ý quan trọng"))
        for f_ in ("Trang phục", "Giày dép", "Lưu ý thời tiết", "An toàn", "Giờ đông khách",
                   "Nên mang theo", "Điều cấm"):
            spec.append(("f", f_, UNV))
        spec += [("group", "Vị trí và di chuyển"),
                 ("f", "Từ hồ Xuân Hương",
                  f"{r['km']:.1f} km · {r['min']:.0f} phút"
                  if r.get("min") is not None else UNV),
                 # "Tu khach san" da chuyen len muc 0 — cung mot cau cho ca 36 diem.
                 ("f", "Đường chính gần nhất", ev(r["id"], "duong_gan_nhat")),
                 # Cung ham sinh URL voi ban .md; `value_run` bien no thanh lien ket
                 # THAT o day, vi day la ban cho nguoi doc soat bam vao.
                 ("f", "Bản đồ", _an_ngu.lien_ket_ban_do(r["lat"], r["lon"])),
                 ("f", "Tình trạng đường", UNV),
                 ("f", "Phương tiện tới được", UNV),
                 ("f", "Bãi đỗ xe", UNV),
                 ("group", "Chụp ảnh")]
        for f_ in ("Điểm chụp đẹp", "Giờ chụp đẹp", "Ngắm bình minh / hoàng hôn",
                   "Phí chụp ảnh", "Lưu ý chụp ảnh"):
            spec.append(("f", f_, UNV))
        # Dong flycam da chuyen len muc 0: quy tac giong nhau cho ca 36 diem.
        for _lab, _k in (("Năm khánh thành", "nam_khanh_thanh"), ("Năm xây dựng", "nam_xay_dung"),
                         ("Kiến trúc", "kien_truc"), ("Xếp hạng di tích", "xep_hang_di_tich"),
                         ("Tôn giáo", "ton_giao"), ("Diện tích", "dien_tich")):
            if has(r["id"], _k):
                spec.append(("f", _lab, ev(r["id"], _k)))
        field_table(spec)

        if NEAR.get(r["id"]):
            P(sec("Điểm lân cận") + " (thời gian đường bộ thật, không phải đường chim bay)", bold=True, size=9.5)
            TBL(["#", "Điểm", "Loại", "Km", "Phút"],
                [[i, f"{oid} · {byid[oid]['name']}", byid[oid]["loai_vn"], f"{dkm:.1f}", f"{tmin:.0f}"]
                 for i, (oid, dkm, tmin) in enumerate(NEAR[r["id"]], 1)],
                widths=[1.0, 7.0, 4.0, 2.0, 2.0])

        # ── Khach san & quan an GAN diem nay ────────────────────────────────────
        # So luong BIEN, khong phai 3 co dinh: chi 24/36 diem lap du 3 bac va 9 diem
        # khong co co so luu tru dang ky nao trong 5 km. Quan an nhom theo LOAI MON
        # chu khong theo bac gia — 0/5.559 quan co gia, va mot tieu de bac gia tu no
        # la mot khang dinh du kien. Khong in `Đánh giá`/`Tiện nghi`: 0 du lieu, va
        # rao can la giay phep luu tru; da noi mot lan o muc 0.
        # Khoi khach san + quan an KHONG con o day — in mot lan o dau khu vuc.
        # Truoc day no chiem 1.420 dong tren ca tai lieu, trung vi 43% moi ho so, de
        # in lai gan nhu cung mot danh sach 36 lan; chi co 52 khach san khac nhau.
        _lc = _LAN_CAN.get(r["id"], {})
        _pt = P(sec("Lưu trú & ăn uống") + f" → xem khối đầu khu vực {r['area']}",
                bold=True, size=9.5)
        if _lc.get("tong_ks_trong_bk") is not None:
            P(f"Quanh riêng điểm này: {_lc['tong_ks_trong_bk']} cơ sở lưu trú trong 5 km"
              f" · {_lc['tong_quan_trong_bk']} quán trong 2 km", size=8.5, color=GREY)
        if _lc.get("khong_co_khach_san"):
            P(_lc["khong_co_khach_san"], size=9, color=AMBER)

        pv = doc.add_paragraph()
        rr = pv.add_run(f"Kiểm chứng: CHƯA GỌI · gọi {r.get('tel') or 'CHƯA CÓ SỐ'} để đóng "
                        "giờ mở cửa, giá vé, thời lượng thăm và điều kiện đi lại.")
        rr.bold = True
        rr.font.size = Pt(8.5)
        rr.font.color.rgb = RED

# ==================================================== 3. HOAT DONG
# Cung mot module chon loc voi ban .md — `hoat_dong_data.tai()`. Neu viet lai
# logic cat gon o day thi hai ban se lech nhau va khong ai biet cho toi khi doc
# canh nhau; du an da dinh dung lop loi do (hai bo trich cung payload VNPay).
_HD, _HDTK = _hoat_dong.tai(RAW)
_MON = _hoat_dong.tai_mon_an(RAW)
_PC = _hoat_dong.tai_phong_cach(RAW)

if phat("diemden", "nhahang"):
    # Nhom "ẩm thực" (ăn vặt chợ đêm · tour ẩm thực) thuoc tai lieu NHA HANG;
    # sau nhom con lai thuoc tai lieu DIEM DEN. Ban gop `tatca` giu ca bay.
    _HD_TL = [a for a in _HD
              if TAI_LIEU == "tatca"
              or (a["nhom"] == "ẩm thực") == (TAI_LIEU == "nhahang")]
    _nhom_tl = len({a["nhom"] for a in _HD_TL})

    doc.add_page_break()
    H(f"{S_HOATDONG}. Hoạt động — làm gì ở Đà Lạt", 1)
    _dan = (f" Mã DL-xx dẫn về mục chi tiết ở mục {S_DIEMDEN}."
            if S_DIEMDEN else
            " Mã DL-xx dẫn về tài liệu Điểm tham quan Đà Lạt.")
    P(f"{len(_HD_TL)} hoạt động, {_nhom_tl} nhóm." + _dan,
      italic=True, size=9, color=GREY)

    _nhom_hien = None
    for _a in _HD_TL:
        if _a["nhom"] != _nhom_hien:
            _nhom_hien = _a["nhom"]
            H(_nhom_hien.upper(), 2)
        H(_a["ten"], 3)
        if _a["noi"]:
            P("Làm ở đâu — %d nơi — %d nơi tiêu biểu:" % (_a["tong_noi"], len(_a["noi"]))
              if _a["tong_noi"] > len(_a["noi"]) else "Làm ở đâu:", bold=True, size=9)
            for _n in _a["noi"]:
                B((f"[{_n['ma']}] " if _n["ma"] else "") + _n["ten"]
                  + (f" — {_n['khu_vuc']}" if _n.get("khu_vuc") else ""))
        if _a["don_vi"]:
            P("Đơn vị tổ chức — %d đơn vị — %d đơn vị tiêu biểu:"
              % (_a["tong_don_vi"], len(_a["don_vi"]))
              if _a["tong_don_vi"] > len(_a["don_vi"]) else "Đơn vị tổ chức:",
              bold=True, size=9)
            for _d in _a["don_vi"]:
                B(_d["ten"] + (f" — {_d['dien_thoai']}" if _d.get("dien_thoai")
                               else " — chưa có số"))
        if _a.get("tour_web"):
            P("Trang tour:", bold=True, size=9)
            for _t in _a["tour_web"]:
                B(f"{_t['ten']} — {_t['url']}")
                if _t["khoang_gia_don_vi"]:
                    _pp = doc.add_paragraph(style="List Bullet 2")
                    _rr = _pp.add_run(f"Khoảng giá cả gói: {_t['khoang_gia_don_vi']}")
                    _rr.font.size = Pt(8.5)
                    _rr.font.color.rgb = AMBER
                for _n in _t["ten_tour"]:
                    doc.add_paragraph(_n, style="List Bullet 2")
        if not _a["don_vi"]:
            P("Không cần đơn vị tổ chức — tự đi được.", italic=True, size=9, color=GREY)

# BA KHOI, khong phai mot bang phang — xep tren toan bo theo so quan dao nguoc
# cau tra loi cho "an gi o Da Lat". `nhom` la phan doan bien tap, khong phai so do.
if phat("nhahang"):
    if _MON:
        _tong_q = sum(n for _, rows in _MON for _, n, _ in rows)
        _tong_m = sum(len(rows) for _, rows in _MON)
        H(f"ẨM THỰC — {_tong_q} quán trên {_tong_m} món", 2)
        P("Chia nhóm là phán đoán biên tập, không phải số đo: “đặc sản” nghĩa là món gắn "
          "với Đà Lạt, “phổ thông” nghĩa là món có ở mọi thành phố và cũng có ở đây. "
          "Trong từng nhóm xếp theo số cơ sở bán.", italic=True, size=9, color=GREY)
        # Cot "vlog" chi hien khi co du lieu — han muc YouTube theo NGAY nen file
        # co the rong hoac mot phan; in cot toan trong trong nhu du lieu bi mat.
        _co_vlog = any(q.get("vlog") for _, rows in _MON for _, _, qs in rows for q in qs)
        for _nhom, _rows in _MON:
            H(f"{_nhom.upper()} — {len(_rows)} món", 3)
            TBL(["Món", "Số quán", "Gợi ý (ưu tiên quán có số gọi)"],
                [[_m, str(_sl),
                  " · ".join(q["ten"][:26] + (f" ({q['vlog']} kênh)" if q.get("vlog") else "")
                             for q in _q[:3])]
                 for _m, _sl, _q in _rows], widths=[3.6, 1.6, 10.6])
        if _co_vlog:
            P("(N kênh) — số kênh du lịch tiếng Việt KHÁC NHAU nhắc tên quán, ngưỡng ≥2 "
              "kênh. Đếm theo kênh chứ không theo video, vì một kênh đăng nhiều video "
              "không phải nhiều lời khuyên độc lập. Quán không có ghi chú nghĩa là chưa "
              "quét đến, không phải không được nhắc.", italic=True, size=8.5, color=GREY)

if phat("nhahang"):
    if _PC:
        H("QUÁN CÓ PHONG CÁCH ĐẶC BIỆT", 2)
        P(f"{len(_PC)} quán được vlog nhắc kèm một cách làm riêng. Thẻ lấy từ bộ từ vựng "
          "cố định, không phải mô tả tự do.", italic=True, size=9, color=GREY)
        TBL(["Quán", "Phong cách", "Kênh nhắc", "Điện thoại"],
            [[_q["ten"][:34], ", ".join(_q["the_phong_cach"]),
              str(_q.get("so_kenh_nhac", _q.get("so_video_nhac", 0))),
              _q.get("dien_thoai") or ""] for _q in _PC], widths=[5.0, 5.4, 1.4, 3.0])

    # Ba truong mua/gio/thoi luong trong tren ca 28 hoat dong. Cho thieu + viec can
    # lam thuoc muc 12, khong thuoc giua chuong tra cuu.

    # ==================================================== 4. LUU TRU & AN UONG
_LT = _an_ngu.tai_luu_tru(RAW)
_AU = _an_ngu.tai_an_uong(RAW)

if phat("nhahang", "khachsan"):
    if _LT or _AU:
        doc.add_page_break()
        H(f"{S_ANNGU}. Lưu trú & ăn uống", 1)

if phat("khachsan"):
    if _LT:
        H(f"{muc_con(S_ANNGU)} Lưu trú", 2)
        # Nhom DUY NHAT trong ca tai lieu co GIA THAT — tu dang ky luu tru cua Cuc
        # Du lich Quoc gia, khong phai tu blog. Nen gia ghi thang.
        P(f"{_LT['tong']} cơ sở trong đăng ký lưu trú nhà nước · {_LT['nha_nuoc']} đã thẩm "
          f"định, {_LT['tu_dang_ky']} tự đăng ký · {_LT['co_gia']} cơ sở có giá công bố · "
          f"{_LT['co_dien_thoai']} có số gọi.", italic=True, size=9, color=GREY)
        P("Giá là giá phòng/đêm do cơ sở công bố với cơ quan quản lý — đổi theo mùa, gọi "
          "xác nhận trước khi báo khách.", bold=True, size=9)
        for _b in _LT["bac"]:
            if not _b["co_so"]:
                continue
            H(f"{_b['ten']} — {_b['tong']} cơ sở", 3)
            TBL(["Cơ sở", "Giá/đêm", "Phòng", "Điện thoại", "Địa chỉ"],
                [[_c["ten"][:38], _c["gia"], str(_c["so_phong"] or ""),
                  _c["dien_thoai"] or "", (_c["dia_chi"] or "").split(",")[0][:26]]
                 for _c in _b["co_so"]], widths=[5.4, 3.4, 1.4, 2.8, 3.6])
        # Khoi "khong cong bo gia" — CUNG du lieu, cung cach dien dat voi ban .md.
        if _LT.get("khong_gia"):
            H(f"Không công bố giá — {_LT['khong_gia_tong']} cơ sở đã đăng ký", 3)
            P(f"{_LT['khong_gia_co_dt']} cơ sở có số gọi. Không có giá KHÔNG phải là "
              "không có chỗ: hỏi giá khi gọi.", size=9, color=GREY)
            TBL(["Cơ sở", "Địa chỉ", "Điện thoại", "Thẩm định"],
                [[_c["ten"][:38] + (f" ({_c['loai']})" if _c["loai"] else ""),
                  (_c["dia_chi"] or "").split(",")[0][:30],
                  _c["dien_thoai"] or "", _c["tham_dinh"] or ""]
                 for _c in _LT["khong_gia"]], widths=[6.0, 5.0, 2.8, 2.2])
        if _LT["dong_cua"]:
            _pd = doc.add_paragraph()
            _rd = _pd.add_run("Đã đóng cửa — không giới thiệu: "
                              + " · ".join(f"{r['ten']} ({r['ngay']})" for r in _LT["dong_cua"]))
            _rd.bold = True
            _rd.font.size = Pt(9)
            _rd.font.color.rgb = RED

# ── 3.2 luu tru quanh Ho Xuan Huong ────────────────────────────────────────
# Truoc day khoi nay la MOT FILE RIENG (`Khach-San-Ho-Xuan-Huong.docx`, do
# `build_khachsan_hxh.py` dung). Gop vao day vi no la mot LAT CAT cua cung lop
# luu tru, khong phai mot tai lieu khac — va vi mot bo dung thu ba cho cung mot
# lop du lieu la mot cho nua de hai ban lech nhau.
_LH = _an_ngu.tai_luu_tru_quanh_ho(RAW)
if _LH:
    _an_ngu.gan_place_id(_LH["trong"], RAW, "place_id_hxh.json")
if phat("khachsan") and _LH:
    _lo, _hi = _LH["gia"]
    H(f"{muc_con(S_ANNGU)} Lưu trú quanh Hồ Xuân Hương — bảng giá "
      f"{_lo:,}–{_hi:,}₫".replace(",", "."), 2)
    P("“3 sao” ở đây là QUY ƯỚC GIÁ của tài liệu, KHÔNG phải hạng sao nhà nước.",
      bold=True, color=RED)
    B(f"Quy ước: giá phòng thấp nhất cơ sở công bố nằm trong {_lo:,}–{_hi:,}₫/đêm."
      .replace(",", "."))
    B("Hạng sao chính thức KHÔNG có trong dữ liệu: 0/420 cơ sở mang hạng sao, và "
      "nguồn công bố nó (đăng ký Cục Du lịch Quốc gia) trả lỗi 403 từ 31/07/2026. "
      "Đừng nói với khách đây là khách sạn 3 sao được xếp hạng.")
    B(f"Nhãn giá chỉ ràng buộc PHÒNG RẺ NHẤT: {_LH['vuot_tran']}/{_LH['tong_bang_gia']} "
      f"cơ sở có phòng đắt hơn {_hi:,}₫. Cột giá in cả hai đầu — HỎI GIÁ PHÒNG CỤ "
      f"THỂ khi gọi.".replace(",", "."))
    _nlh = sum(1 for r in _LH["trong"] if r["place_id"])
    _xlh = sum(1 for r in _LH["trong"] if r["so_nha_lech"])
    P(f"{len(_LH['trong'])} cơ sở trong {_LH['ban_kinh']:,} m quanh hồ, sắp theo "
      f"khoảng cách (đường chim bay tới mép hồ). {_nlh} có liên kết bản đồ."
      .replace(",", "."), size=9, color=GREY)
    if _xlh:
        P(f"⚠ {_xlh} dòng có dấu ⚠ ở cột Địa chỉ: số nhà trong đăng ký chưa khớp "
          "số nhà trên bản đồ, dù tên và vị trí đều khớp. Hỏi lại khi gọi.",
          bold=True, size=9, color=RED)
    _tlh = TBL(["Khách sạn", "Địa chỉ", "Cách hồ", "Giá/đêm", "Phòng",
                "Điện thoại", "Bản đồ"],
               [[r["ten"], ("⚠ " if r["so_nha_lech"] else "") + (r["dia_chi"] or "—"),
                 f"{r['m']:.0f}", r["gia"] or "—", r["so_phong"] or "—",
                 r["dien_thoai"] or "—", ""] for r in _LH["trong"]],
               widths=[3.3, 3.9, 1.3, 2.8, 1.0, 2.2, 1.5], size=8)
    for _i, _r in enumerate(_LH["trong"], start=1):
        _o = _tlh.rows[_i].cells[6]
        _o.text = ""
        _u = _an_ngu.lien_ket_place_id(_r["place_id"])
        if _u:
            them_lien_ket(_o.paragraphs[0], _u, "Mở")
        else:
            _rr = _o.paragraphs[0].add_run("chưa phân giải")
            _rr.font.size, _rr.italic = Pt(7), True
            _rr.font.color.rgb = GREY
    # Nhom KHONG do duoc khoang cach — co mat de goi dien duoc, khong gia vo da do.
    _dolh = _an_ngu.do_pho_khong_dinh_vi(RAW, _LH["khong_toa_do"])
    H(f"{len(_LH['khong_toa_do'])} cơ sở có địa chỉ nhưng CHƯA đo được khoảng cách", 3)
    P("Cùng bảng giá, có địa chỉ đầy đủ, nhưng dữ liệu chưa có toạ độ nên không "
      "nói được chúng cách hồ bao xa. Có mặt ở đây để gọi điện được — KHÔNG phải "
      "vì đã xác minh là gần hồ.", bold=True, size=9)
    P(f"Đã thử suy vị trí theo TÊN PHỐ và phép đo bác bỏ: đối chiếu với "
      f"{_dolh['tham_chieu']:,} bản ghi có cả địa chỉ lẫn toạ độ trên "
      f"{_dolh['so_pho']:,} phố, chỉ {_dolh['trong']}/{len(_LH['khong_toa_do'])} "
      f"địa chỉ nằm trên phố mà mọi bản ghi đều trong bán kính; {_dolh['vat']} phố "
      f"vắt qua ranh giới. Phố Đà Lạt dài hàng km nên tên phố không định vị được "
      f"một toà nhà.".replace(",", "."), size=9, color=GREY)
    TBL(["Khách sạn", "Địa chỉ", "Giá/đêm", "Phòng", "Điện thoại"],
        [[r["ten"], r["dia_chi"] or "—", r["gia"] or "—",
          r["so_phong"] or "—", r["dien_thoai"] or "—"]
         for r in _LH["khong_toa_do"]],
        widths=[3.6, 5.4, 2.9, 1.1, 3.0], size=8)
    # Quy tac xep hang da chuyen len muc S_XEPHANG (ngay sau muc 0). Chi de lai
    # mot dong tro ve — de tai lieu khong noi cung mot quy tac o hai cho roi
    # mot cho troi di.
    P(f"Bảng trên sắp theo khoảng cách, KHÔNG theo chất lượng. Cách xếp theo "
      f"đánh giá và câu lệnh chạy bảng: xem mục {S_XEPHANG}.",
      bold=True, size=9.5)
    # Pham vi: danh sach nay KHONG vet het, va con so phai noi ra dieu do.
    H("Phạm vi — danh sách quanh hồ KHÔNG vét hết", 3)
    TBL(["Mục", "Số"],
        [["Cơ sở lưu trú còn hoạt động (toàn thành phố)", f"{_LH['tong_hoat_dong']:,}"],
         ["…trong đó có công bố giá", f"{_LH['tong_co_gia']:,}"],
         [f"…giá phòng thấp nhất trong {_lo:,}–{_hi:,}₫", f"{_LH['tong_bang_gia']:,}"],
         ["…có toạ độ để đo khoảng cách", f"{_LH['co_toa_do']:,}"],
         [f"…trong {_LH['ban_kinh']:,} m quanh hồ → bảng trên", f"{len(_LH['trong']):,}"],
         ["Có địa chỉ nhưng không có toạ độ → bảng dưới", f"{len(_LH['khong_toa_do']):,}"],
         ["Có liên kết bản đồ", f"{_nlh}/{len(_LH['trong'])}"]],
        widths=[11.0, 5.0], size=9)
    P(f"Hai khoảng trống người đọc phải biết: "
      f"{_LH['tong_hoat_dong'] - _LH['tong_co_gia']:,} cơ sở đang hoạt động KHÔNG "
      f"công bố giá nên không lọt vào bất kỳ bảng giá nào; và "
      f"{len(_LH['khong_toa_do']):,} cơ sở ở bảng dưới có thể gần hồ hơn nhiều cơ "
      f"sở ở bảng trên — chưa ai đo.".replace(",", "."), size=9, color=GREY)

if phat("nhahang"):
    if _AU:
        H(f"{muc_con(S_ANNGU)} Ăn uống", 2)
        P(f"{_AU['tong_mo']:,} quán còn hoạt động · {_AU['co_dien_thoai']:,} có số gọi. "
          f"Món đặc trưng xem mục {S_HOATDONG}.".replace(",", "."), italic=True, size=9, color=GREY)
        for _n in _AU["nhom"]:
            H(f"{_n['ten']} — {_n['tong']} quán", 3)
            for _q in _n["quan"]:
                B(_q["ten"] + (f" — {_q['dien_thoai']}" if _q["dien_thoai"] else "")
                  + (f" · {', '.join(_q['mon'])}" if _q["mon"] else ""))
        # Muc quan trong nhat cua chuong nay. Thac khong dong cua; quan an thi co.
        if _AU["dong_cua"]:
            H("⚠ Đã đóng cửa — KHÔNG giới thiệu", 3)
            P(f"{_AU['tong_dong']} quán đã đóng, dưới đây {len(_AU['dong_cua'])} gần nhất. "
              "Nhiều quán trong số này vẫn còn trong các hướng dẫn cũ.",
              bold=True, size=9, color=RED)
            TBL(["Ngày đóng", "Quán"],
                [[_r["ngay"], _r["ten"][:52]] for _r in _AU["dong_cua"]], widths=[3.0, 12.0])

# ── 3.3 quan an quanh Ho Xuan Huong ────────────────────────────────────────
# Muc 3.2 nhom theo hang muc va lay MAU 8 quan moi nhom — de biet thanh pho co
# gi. Muc nay khac han muc dich: de TIM MOT CHO CU THE quanh ho, nen no mang
# dia chi, khoang cach do duoc va lien ket tro dung co so, va KHONG cat bot.
#
# KHONG chua diem danh gia. Dieu khoan Places chi cho luu `place_id`; va ngay
# ca khi duoc phep thi mot thu hang in ra giay cung se cu di ma khong co gi bao
# rang no da sai. Nguoi doc bam lien ket la thay so cua hom nay.
_QH = _an_ngu.tai_quan_an_quanh_ho(RAW)
if _QH:
    _an_ngu.gan_place_id(_QH["trong"], RAW, "place_id_quan_hxh.json")
if phat("nhahang") and _QH:
    H(f"{muc_con(S_ANNGU)} Quán ăn quanh Hồ Xuân Hương", 2)
    _nl = sum(1 for r in _QH["trong"] if r["place_id"])
    _nx = sum(1 for r in _QH["trong"] if r["so_nha_lech"])
    P(f"Đủ {len(_QH['trong'])} quán còn hoạt động trong {_QH['ban_kinh']} m quanh hồ — "
      f"danh sách này KHÔNG cắt bớt, khác các danh sách gợi ý ở mục {S_ANNGU}.2. "
      f"{_QH['co_dia_chi']} có địa chỉ · {_QH['co_dien_thoai']} có số gọi · "
      f"{_nl} có liên kết bản đồ.", italic=True, size=9, color=GREY)
    P("Sắp theo khoảng cách tới hồ, KHÔNG theo đánh giá.", bold=True, size=9.5)
    P("Tài liệu không chứa điểm đánh giá hay số lượt đánh giá của Google: điều khoản "
      "chỉ cho lưu place_id, và một thứ hạng in ra giấy sẽ cũ đi mà không có gì trong "
      "tờ giấy báo rằng nó đã sai. Bấm cột Bản đồ để thấy điểm và số lượt HIỆN TẠI.",
      size=9, color=GREY)
    if _nx:
        # ⚠ dat TRUOC gia tri — hau to bi bo loc tien to xoa mat (bai hoc 30/07).
        P(f"⚠ {_nx} dòng có dấu ⚠ ở cột Địa chỉ: số nhà trong dữ liệu của ta chưa khớp "
          f"số nhà trên bản đồ, dù tên và vị trí đều khớp. Hỏi lại khi gọi.",
          bold=True, size=9, color=RED)
    _t = TBL(["Quán", "Địa chỉ", "Cách hồ", "Loại", "Điện thoại", "Bản đồ"],
             [[r["ten"], ("⚠ " if r["so_nha_lech"] else "") + (r["dia_chi"] or "—"),
               f"{r['m']:.0f}", (r["hang_muc"] or "—").replace("_", " "),
               r["dien_thoai"] or "—", ""] for r in _QH["trong"]],
             widths=[3.9, 4.3, 1.2, 2.6, 2.5, 1.5], size=8)
    for _i, _r in enumerate(_QH["trong"], start=1):
        _o = _t.rows[_i].cells[5]
        _o.text = ""
        _u = _an_ngu.lien_ket_place_id(_r["place_id"])
        if _u:
            them_lien_ket(_o.paragraphs[0], _u, "Mở")
        else:
            _rr = _o.paragraphs[0].add_run("chưa phân giải")
            _rr.font.size = Pt(7)
            _rr.italic = True
            _rr.font.color.rgb = GREY
    P(f"{len(_QH['trong'])-_nl} quán chưa có liên kết: luật nhận đòi tên khớp theo biên từ "
      "VÀ vị trí khớp trong 100 m. Chỉ một trục đúng thì không nhận — một liên kết dẫn "
      "khách tới nhầm chỗ tệ hơn hẳn một ô trống.", size=9, color=GREY)
    # Quy tac xep hang da chuyen len muc S_XEPHANG — xem chu thich o khoi luu tru.
    P(f"Bảng trên sắp theo khoảng cách, KHÔNG theo chất lượng. Cách xếp theo "
      f"đánh giá và câu lệnh chạy bảng: xem mục {S_XEPHANG}.",
      bold=True, size=9.5)

    _tr = _an_ngu.nhom_trung_co_so(_QH["trong"])
    if _tr:
        P(f"⚠ {sum(len(g) for g in _tr)} dòng ở trên là {len(_tr)} cơ sở, không phải "
          f"{sum(len(g) for g in _tr)} cơ sở — các dòng trong mỗi nhóm dưới đây trỏ về "
          "CÙNG một địa điểm trên Google (cùng place_id). Chúng không bị xoá vì bản ghi "
          "gốc của chúng khác nhau và chưa biết bản nào đúng, nhưng đừng gọi hai lần:",
          bold=True, size=9, color=RED)
        for _g in _tr:
            B(" = ".join(_g))

# ==================================================== 5-13 chi muc
if phat("diemden"):
    doc.add_page_break()
    H(f"{S_SOSANH}. Bảng so sánh", 1)
    P(f"Sinh tự động từ mục {S_DIEMDEN} — không sửa tay.", italic=True, size=9, color=GREY)
    # Cot "Mua" khop ban .md — truoc day ban .docx thieu han cot nay.
    TBL(["ID", "Điểm", "Loại", "Khu vực", "Km", "Phút", "Vé", "Mưa", "Nguồn"],
        [[r["id"], r["name"][:30], r["loai_vn"], r["area"][:18], f"{r['km']:.1f}",
          f"{r['min']:.0f}", _an_ngu.the_fee_ngan(r.get("fee")),
          INDOOR_CAT.get(r["loai_vn"], "ngoài trời"), len(r["src"])] for r in picked],
        widths=[1.4, 4.4, 2.8, 2.8, 1.1, 1.1, 1.5, 1.4, 1.1], size=8)

# DA CAT muc 6 "Theo loai hinh" · 7 "Theo khu vuc" · 8 "Theo khoang cach"
# · 9 "Theo thoi diem tham" — ca bon la cach SAP XEP LAI cung 36 diem. `loai_vn`
# va `km` da la COT cua bang so sanh; muc 2 gio da nhom theo khu vuc; va truong
# "thoi diem tot trong ngay" la 0/36. Phan duy nhat co that trong muc 9 — nhom
# "di duoc khi troi mua" — thanh mot cot cua bang so sanh.
if phat("diemden"):
    H(f"{S_TUYEN}. Tuyến gợi ý theo khu vực", 1)
    P("Thứ tự dựng bằng thuật toán láng giềng gần nhất trên ma trận OSRM, xuất phát từ điểm gần "
      "trung tâm nhất.", italic=True, size=9, color=GREY)
    if mat:
        idx = {pid: i for i, pid in enumerate(mat["ids"])}
        for a in sorted({r["area"] for r in picked}):
            lst = [r for r in picked if r["area"] == a]
            if len(lst) < 2:
                continue
            cur = min(lst, key=lambda x: x["min"])
            route, left, tot = [cur], [x for x in lst if x is not cur], 0.0
            while left:
                nxt = min(left, key=lambda x: mat["durations"][idx[cur["id"]]][idx[x["id"]]] or 9e9)
                tot += (mat["durations"][idx[cur["id"]]][idx[nxt["id"]]] or 0) / 60.0
                route.append(nxt)
                left.remove(nxt)
                cur = nxt
            P(f"{a} — {len(route)} điểm · di chuyển giữa các điểm ~{tot:.0f} phút", bold=True, size=9.5)
            P("   " + " → ".join(f"{x['id']} {x['name']}" for x in route), size=9)

if phat("diemden"):
    H(f"{S_MATRAN}. Ma trận thời gian trong từng khu vực (phút)", 1)
    # Truoc day la mot bang 36x36 = 1.296 o o co chu 6pt — khong vua mot trang doc
    # va khong ai tra thoi gian giua hai diem o hai dau thanh pho. Chia theo khu vuc:
    # chin bang nho, moi bang vua mot trang. Cap xa nhau van con trong osrm_selected.
    if mat:
        _idx8 = {pid: i for i, pid in enumerate(mat["ids"])}
        P("Chia theo khu vực: thời gian giữa các điểm trong cùng một buổi. Cặp ở hai khu "
          f"vực khác nhau thì tra mục {S_TUYEN} (tuyến gợi ý) hoặc cột Phút ở mục {S_SOSANH}.",
          italic=True, size=9, color=GREY)
        for _a in sorted({r["area"] for r in picked}):
            _lst = [r for r in picked if r["area"] == _a]
            if len(_lst) < 2:
                continue
            P(f"{_a} ({len(_lst)} điểm)", bold=True, size=9.5)
            TBL([""] + [r["id"].replace("DL-", "") for r in _lst],
                [[r["id"].replace("DL-", "")] +
                 ["—" if r is o or mat["durations"][_idx8[r["id"]]][_idx8[o["id"]]] is None
                  else f"{mat['durations'][_idx8[r['id']]][_idx8[o['id']]]/60:.0f}"
                  for o in _lst]
                 for r in _lst], size=8)

# DA CHUYEN muc 12 "Danh sach rut gon" -> phu luc: thu hang do do MUC DO HIEN
# DIEN TREN BAN DO, tuc chat luong DU LIEU, khong phai chat luong trai nghiem.
# Muc 9 ngay duoi dung dung thu hang do lam THU TU GOI DIEN, la cong dung dung.
#
# Thu tu nay den TU build_huong_dan.py qua khoa `xep_hang`, khong tu tinh lai o
# day. Truoc day dong nay la `sorted(picked, key=lambda r: (-len(r["src"]),
# -conf))` trong khi ban .md xep theo `-_score` — va `_score` bi bo loc
# `startswith("_")` chan lai nen khong bao gio den duoc file nay. Ket qua: hai
# ban tai lieu bat dong ve THU TU GOI DIEN xac minh, tu hang thu 9 tro di.
if phat("diemden"):
    H(f"{S_KIEMCHUNG}. Sổ kiểm chứng — việc cần làm", 1)
    n_tel = sum(1 for r in picked if r.get("tel"))
    TBL(["Chỉ số", "Giá trị"],
        [["Điểm trong hồ sơ", len(picked)],
         ["Có số điện thoại để gọi", f"{n_tel} / {len(picked)}"],
         ["Chưa có số — cần tìm", len(picked) - n_tel],
         # Dem CAI DA IN — CUNG ham voi ban .md.
         ["Có giờ mở cửa", sum(1 for r in picked
                               if _an_ngu.co_gio_mo_cua(r, ENR, _GIO_LOAI))],
         # Cung phep dem voi ban .md: hai the `fee` deu khong phai so tien.
         ["Có giá vé đã xác minh", "0 — hai thẻ `fee` của OSM đều không phải số tiền"],
         ["Có giá vé THAM KHẢO (nguồn thương mại, có thể lệch nhau)",
          sum(1 for r in picked if has(r["id"], "gia_ve_tham_khao"))],
         ["Có đánh giá sao", "0 — không nguồn mở nào có"],
         ["Trường [CHƯA XÁC MINH] ước tính", f"~{41*len(picked)}"]],
        widths=[7.0, 9.0], size=9)
    # Cho thieu cua muc 3 thuoc VE DAY, khong thuoc giua chuong tra cuu.
    if _HDTK["thieu"]:
        _pm = doc.add_paragraph()
        _rm = _pm.add_run(f"Mục 3 — cả {_HDTK['so_hoat_dong']} hoạt động đều chưa có mùa, "
                          "giờ trong ngày và thời lượng. Trống nghĩa là chưa xác minh, "
                          "không nghĩa là quanh năm. Đừng khẳng định với khách cỏ hồng "
                          "tháng nào hay tour đi mấy giờ khi chưa gọi.")
        _rm.bold = True
        _rm.font.size = Pt(9)
        _rm.font.color.rgb = RED
        P(f"Tra đơn vị tour bằng Facebook, đừng tra tên miền — {_HDTK['ten_mien_chet']}/"
          f"{_HDTK['tong_website_thu']} tên miền đã kiểm không còn hoạt động "
          "(canyoningdalat.com, dalatjeep.com, toursanmaydalat.com).", size=9, color=AMBER)

    P("Gọi một cuộc đóng được khoảng 9 trường. Danh sách cần gọi, theo thứ tự ưu tiên:", bold=True)
    TBL(["#", "Điểm", "Điện thoại"],
        [[i, f"{r['id']} · {r['name']}", r["tel"]]
         for i, r in enumerate([x for x in rank if x.get("tel")][:20], 1)],
        widths=[1.2, 9.0, 5.0], size=9)

# Chon loc nam o `an_ngu_data.tai_nghien_cuu` — CUNG ham ma ban .md goi. Ban
# truoc khoi nay chi co o ban .md nen ban .docx mat trang: 32 link Facebook, 21
# luot check-in, 4 ti le de xuat va ca bang thu hang.
_NC = _an_ngu.tai_nghien_cuu(RAW)
if phat("diemden"):
    if _NC and _NC["hang"]:
        doc.add_page_break()
        H("Phụ lục nghiên cứu — xuất xứ, không phải thông tin đi chơi", 1)
        P("Các trường dưới đây đã được đưa ra khỏi hồ sơ điểm đến: người lập kế hoạch "
          "không dùng chúng, người kiểm chứng nguồn thì cần. Tra theo mã DL-xx.",
          italic=True, size=9, color=GREY)
        TBL(["ID", "Điểm", "Trường", "Giá trị"],
            [[i, t[:26], nhan, gt[:44]] for i, t, nhan, gt in _NC["hang"]],
            widths=[1.6, 4.6, 3.4, 6.4], size=8)
        P(f"{_NC['so_diem']}/{_NC['tong']} điểm có ít nhất một trường xuất xứ.",
          italic=True, size=8.5, color=GREY)
        H("Thứ hạng theo mức độ hiện diện trên bản đồ", 2)
        P("⚠ Đây là chất lượng DỮ LIỆU, không phải chất lượng trải nghiệm. Mục "
          f"{S_KIEMCHUNG} dùng đúng thứ hạng này làm thứ tự gọi điện xác minh — đó là "
          "công dụng đúng của nó. Không dùng làm lời khuyên “nơi này hay hơn nơi kia”.",
          bold=True, size=9, color=RED)
        for lab, seg in _NC["xep_hang"]:
            P(f"{lab} ({len(seg)}): " + " · ".join(f"{i} {t}" for i, t in seg), size=9)

doc.add_page_break()
H("Phụ lục — Tóm tắt độ phủ (đọc lại trước khi trả lời khách)", 1)
P(f"Tài liệu này theo dõi khoảng {41*len(picked)} trường trên {len(picked)} điểm đến. "
  "Phần lớn mang dấu [CHƯA XÁC MINH].")
P("Nếu khách hỏi về một trường mang dấu [CHƯA XÁC MINH] ở bất kỳ đâu phía trên: nói "
  "“chỗ này em chưa có số liệu đã xác minh, để em gọi hỏi rồi báo lại mình”, và KHÔNG đưa ra "
  "con số, giờ giấc hay đánh giá nào thay thế.", bold=True)
P("Ba trường tuyệt đối không được suy đoán, vì sai là khách mất tiền hoặc mất cả ngày:",
  bold=True, color=RED)
B("giờ mở cửa")
B("giá vé")
B("mức độ dễ đi lại cho người cao tuổi")
P("Tài liệu này chứa dữ liệu từ OpenStreetMap. Dữ liệu © những người đóng góp OpenStreetMap, "
  "theo giấy phép Open Database License — https://openstreetmap.org/copyright",
  italic=True, size=8, color=GREY)

os.makedirs(os.path.dirname(OUT) or ".", exist_ok=True)
doc.save(OUT)
print("saved ->", OUT)
# `n_tel` chi duoc gan trong khoi so kiem chung (tai lieu diem den), nen dong
# tong ket phai tinh lai o day thay vi doc bien do — mot bien gan CO DIEU KIEN
# la loi ma bo kiem "ten co duoc gan o dau do khong" khong the thay.
_n_tel = sum(1 for r in picked if r.get("tel"))
print(f"[{TAI_LIEU}] diem: {len(picked)}  |  co dien thoai: {_n_tel}"
      f"  |  bang: {len(doc.tables)}  |  doan: {len(doc.paragraphs)}")
