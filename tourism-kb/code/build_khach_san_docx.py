# -*- coding: utf-8 -*-
"""Ban .docx RUT GON danh sach khach san / luu tru — chuan (1 bang).

Doc export/<slug>/khach-san.json. GIU nguyen THU TU export = MUC DO ANH HUONG noi bo (VQS)
— co so noi bat/dong khach len truoc (QUYET DINH 2026-08-05; xem anh_huong.py). KHONG in
diem/sao/luot — chi THU TU. Cot: Khach san · Loai hinh · Dia chi · Dien thoai · Ban do(link).

KHONG co cot GIA / HANG SAO: city nguon Overture (NT/DN) khong co gia; bo cot thay vi pad
"Chua xac minh" (mot cot 100% trong = claim field khong ton tai). Dia Lat co gia nhung docx
nay dung export chung -> giu doctrine "khong in so"; muon xem gia dung ban verbose legacy.

Chay TU GOC REPO:  python tourism-kb/code/build_khach_san_docx.py tourism-kb/raw/<slug>/scrape
"""
import json, os, sys, io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import duong_dan_ra as _dr
import anh_huong as _ah
from dia_diem_config import cfg as _cfg, slug_of as _slug_of

RAW = sys.argv[1]
SLUG = _slug_of(RAW)
CITY = _cfg(RAW)["city"]
CITY_DIR = os.path.dirname(RAW.rstrip("/\\"))       # raw/<slug> — chứa noi-bo/ (rank ảnh hưởng)
_CITY_FILE = "-".join(p.capitalize() for p in SLUG.split("-"))
EXPORT = os.path.join("tourism-kb", "export", SLUG, "khach-san.json")
EXPORT_DC = os.path.join("tourism-kb", "export", SLUG, "khach-san-chi-dia-chi.json")
OUT = "tourism-kb/output/%s/Khach-San-%s.docx" % (SLUG, _CITY_FILE)
CHUA = "Chưa xác minh"

# Overture LODGE category -> nhan VN (DL registry loai da la tieng Viet -> giu nguyen qua .get fallback)
LOAI_VN = {
    "hotel": "Khách sạn", "accommodation": "Lưu trú", "resort": "Khu nghỉ dưỡng",
    "beach_resort": "Resort biển", "hostel": "Hostel", "motel": "Nhà nghỉ",
    "bed_and_breakfast": "B&B", "guest_house": "Nhà khách", "inn": "Nhà trọ",
    "holiday_rental_home": "Nhà cho thuê",
}


def shade(cell, hexc):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(el)


def add_link(cell, url, text):
    part = cell.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rpr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18"); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    hl.append(run)
    cell.paragraphs[0]._p.append(hl)


def maps_url(r):
    pid = (r.get("identity") or {}).get("place_id") or (r.get("external_ids") or {}).get("google_place_id")
    if pid:
        return "https://www.google.com/maps/place/?q=place_id:" + pid
    c = r.get("coordinates") or {}
    if c.get("latitude") is not None:
        return "https://www.google.com/maps/search/?api=1&query=%s,%s" % (c["latitude"], c["longitude"])
    return None


d = json.load(io.open(EXPORT, encoding="utf-8"))
dc = json.load(io.open(EXPORT_DC, encoding="utf-8")) if os.path.exists(EXPORT_DC) else []

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)
_t = doc.add_paragraph().add_run("DANH SÁCH %d KHÁCH SẠN & LƯU TRÚ %s - RÚT GỌN" % (len(d), CITY.upper()))
_t.bold = True
_t.font.size = Pt(15)
# Subtitle TRUNG THỰC: có noi-bo rank -> ảnh hưởng; không -> danh bạ chưa xếp hạng.
if _ah.thu_tu(CITY_DIR, "khach_san"):
    _sub = ("Danh sách sắp theo MỨC ĐỘ ẢNH HƯỞNG (cơ sở nổi bật / đông khách lên trước); phần chưa đủ dữ liệu "
            "ảnh hưởng giữ thứ tự thu thập. Tài liệu KHÔNG in điểm, sao hay lượt đánh giá — thứ tự là tín hiệu. "
            "Nguồn dữ liệu không có mức giá / hạng sao nên không hiển thị cột giá; bấm “bản đồ” để xem trên Google.")
else:
    _sub = ("DANH BẠ CHƯA XẾP HẠNG — sắp theo mức độ đầy đủ thông tin; THỨ TỰ KHÔNG phản ánh độ phổ biến hay "
            "ảnh hưởng (xếp hạng ảnh hưởng cần dữ liệu Google, chưa áp dụng). Tài liệu KHÔNG in điểm, sao hay lượt. "
            "Nguồn không có mức giá / hạng sao nên không hiển thị cột giá; bấm “bản đồ” để xem trên Google.")
_s = doc.add_paragraph().add_run(_sub)
_s.italic = True
_s.font.size = Pt(10)
_s.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

COLS = ["Khách sạn / lưu trú", "Loại hình", "Địa chỉ", "Điện thoại", "Bản đồ"]
tbl = doc.add_table(rows=1, cols=len(COLS))
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(COLS):
    c = tbl.rows[0].cells[i]
    c.text = ""
    run = c.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(9.5)
    shade(c, "DDEBF7")

for r in d:
    row = tbl.add_row().cells
    loai = (r.get("category") or {}).get("primary")
    vals = [r.get("name") or "", LOAI_VN.get(loai, loai or ""),
            (r.get("address") or {}).get("full_address") or CHUA,
            (r.get("contact") or {}).get("phone") or CHUA]
    for j, v in enumerate(vals):
        row[j].text = ""
        run = row[j].paragraphs[0].add_run(v)
        run.font.size = Pt(9)
    row[4].text = ""
    u = maps_url(r)
    if u:
        add_link(row[4], u, "mở")
    else:
        row[4].paragraphs[0].add_run("—").font.size = Pt(9)

for row in tbl.rows:
    for w, cell in zip([5.0, 2.6, 4.6, 2.6, 1.2], row.cells):
        cell.width = Cm(w)

# ghi chu: co so chi co dia chi (khong toa do) — in "0" neu rong, khong bien mat
doc.add_paragraph()
_c = doc.add_paragraph().add_run("Cơ sở chỉ có địa chỉ (không toạ độ để ghim bản đồ): %d" % len(dc))
_c.bold = True
_c.font.size = Pt(9.5)
for r in dc:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("%s — %s" % (r.get("name") or "", (r.get("address") or {}).get("full_address") or CHUA))
    run.font.size = Pt(8.5)

p = _dr.kiem_loi_ra(OUT)
os.makedirs(os.path.dirname(p), exist_ok=True)
tmp = p + ".tmp"
doc.save(tmp)
os.replace(tmp, p)
print("saved -> %s  (%d cơ sở + %d chỉ-địa-chỉ)" % (OUT, len(d), len(dc)))
