# -*- coding: utf-8 -*-
"""Ban .docx RUT GON danh sach diem den — chuan moi (1 bang 6 cot).

Thay ban verbose (build_huong_dan_docx.py) lam DELIVERABLE diem den. Doc
export/<slug>/diem-den.json (da assemble). Cot: STT · Ten · Loai hinh/trai nghiem ·
Dia chi · Gio mo-dong · Gia ve. O trong -> "Chua xac minh".

Chay TU GOC REPO:  python tourism-kb/code/build_diem_den_docx.py tourism-kb/raw[/<slug>]
"""
import json, os, sys, io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import duong_dan_ra as _dr
import trai_nghiem as _tn
from dia_diem_config import cfg as _cfg, slug_of as _slug_of, GHI_CHU_DIEM_DEN as _GHI_CHU

RAW = sys.argv[1]
SLUG = _slug_of(RAW)
C = _cfg(RAW)
CITY = C["city"]
_CITY_FILE = "-".join(p.capitalize() for p in SLUG.split("-"))     # nha-trang -> Nha-Trang
EXPORT = os.path.join("tourism-kb", "export", SLUG, "diem-den.json")
OUT = "tourism-kb/output/%s/Diem-Den-%s.docx" % (SLUG, _CITY_FILE)
CHUA = "Chưa xác minh"

def shade(cell, hexc):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(el)


def gia_ve(ext):
    vals = []
    for t in ext.get("ticketing") or []:
        v = (t.get("value") or "").strip()
        if not v:
            continue
        vals.append("Có thu phí (chưa rõ mức)" if v.lower() == "yes" else v)
    return " · ".join(dict.fromkeys(vals)) or CHUA


def dia_chi(r):
    a = (r["address"].get("full_address") or "").strip()
    warns = [w for w in r["data_quality"].get("warnings", [])
             if any(k in w.lower() for k in ("địa chỉ", "số nhà", "đường"))]
    conf = r["data_quality"].get("conflicts") or {}
    pre = ""
    if warns or any(k in ("address", "dia_chi", "dia_chi_day_du") for k in conf):
        pre = "⚠ " + (warns[0] + " — " if warns else "")
    return (pre + a) if a else (pre + CHUA if pre else CHUA)


def _maps_link(r):
    # store-vs-link: giờ mở DRIFT — thay vì in "Chưa xác minh", link Google LIVE khi có place_id (ToS: chỉ link)
    pid = (r.get("external_ids") or {}).get("google_place_id") or (r.get("identity") or {}).get("place_id")
    return ("https://www.google.com/maps/place/?q=place_id:" + pid) if pid else None


def gio(r):
    oh = r["ext"]["destination"].get("opening_hours")
    if oh and oh.get("raw"):
        return oh["raw"]
    link = _maps_link(r)
    return ("Xem giờ live trên Google ↗ " + link) if link else CHUA


d = json.load(io.open(EXPORT, encoding="utf-8"))
doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

_t = doc.add_paragraph()
_r = _t.add_run("DANH SÁCH ĐIỂM ĐẾN %s - BẢN RÚT GỌN" % CITY.upper())
_r.bold = True
_r.font.size = Pt(15)
_s = doc.add_paragraph()
_sr = _s.add_run("Chỉ gồm: tên điểm đến, loại hình, địa chỉ, giờ mở cửa - đóng cửa và giá vé")
_sr.italic = True
_sr.font.size = Pt(10)
_sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

COLS = ["STT", "Tên điểm đến", "Loại hình / trải nghiệm", "Địa chỉ",
        "Giờ mở cửa - đóng cửa", "Giá vé"]
tbl = doc.add_table(rows=1, cols=len(COLS))
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(COLS):
    c = tbl.rows[0].cells[i]
    c.text = ""
    run = c.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(9.5)
    shade(c, "D9E2F3")

for i, r in enumerate(d, 1):
    ext = r["ext"]["destination"]
    row = tbl.add_row().cells
    _tnh = ext.get("trai_nghiem") or _tn.nhan_trai_nghiem(r["category"]["primary"])  # field luu; fallback export cu
    _gc = _GHI_CHU.get(r["name"])                    # note curator (cáp treo…) cho attraction cha
    if _gc:
        _tnh += " · " + _gc
    vals = [str(i), r["name"], _tnh,
            dia_chi(r), gio(r), gia_ve(ext)]
    for j, v in enumerate(vals):
        row[j].text = ""
        run = row[j].paragraphs[0].add_run(v)
        run.font.size = Pt(9)

# do rong cot (cm)
for row in tbl.rows:
    for w, cell in zip([1.0, 4.0, 3.6, 4.4, 3.2, 3.8], row.cells):
        cell.width = Cm(w)

p = _dr.kiem_loi_ra(OUT)
os.makedirs(os.path.dirname(p), exist_ok=True)
tmp = p + ".tmp"
doc.save(tmp)
os.replace(tmp, p)
print("saved -> %s  (%d điểm)" % (OUT, len(d)))
