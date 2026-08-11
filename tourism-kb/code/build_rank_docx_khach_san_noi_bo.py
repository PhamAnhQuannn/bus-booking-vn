# -*- coding: utf-8 -*-
"""⚠ NỘI BỘ — docx sinh ra CẤM đưa vào sản phẩm/website/khách. ⚠

Render raw/<slug>/noi-bo/rank_noi_bo_khach_san.json (Google + VQS) thành docx xếp hạng
khách sạn để TỰ chọn lịch trình. KHÔNG gọi API. Ghi vào raw/<slug>/docx/ (gitignored) — KHÔNG output/.
★ là Google user-rating, KHÔNG phải hạng sao nhà nước.

Chạy:  PYTHONIOENCODING=utf-8 python tourism-kb/code/build_rank_docx_khach_san_noi_bo.py tourism-kb/raw/<slug>/scrape
"""
import json, os, sys, io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import duong_dan_ra as _dr
from dia_diem_config import cfg as _cfg, slug_of as _slug_of

RAW = sys.argv[1]                                    # raw/<slug>/scrape
SLUG = _slug_of(RAW)
CITY = _cfg(RAW)["city"]
CITY_DIR = os.path.dirname(RAW.rstrip("/\\"))        # raw/<slug>
_CITY_FILE = "-".join(p.capitalize() for p in SLUG.split("-"))
SRC = os.path.join(CITY_DIR, "noi-bo", "rank_noi_bo_khach_san.json")
OUT = os.path.join(CITY_DIR, "docx", "Rank-Noi-Bo-Khach-San-%s.docx" % _CITY_FILE)


def shade(cell, hexc):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(el)


def add_link(cell, url, text):
    part = cell.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rpr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16"); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    hl.append(run)
    cell.paragraphs[0]._p.append(hl)


d = json.load(io.open(SRC, encoding="utf-8"))
xh = d["xep_hang"]
chua = d.get("chua_du_danh_gia", [])

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

_t = doc.add_paragraph().add_run("⚠ NỘI BỘ — XẾP HẠNG KHÁCH SẠN %s (Google, %s)" % (CITY.upper(), d.get("lay_luc", "")))
_t.bold = True
_t.font.size = Pt(14)
_t.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
_s = doc.add_paragraph().add_run(
    "CẤM đưa vào sản phẩm / website / cho khách. Rating + lượt là Google content, "
    "cũ dần theo thời gian — chỉ để TỰ chọn lịch trình; chạy lại builder để refresh. "
    "★ là Google user-rating, KHÔNG phải hạng sao nhà nước (0/420 cơ sở có hạng chính thức). "
    "Xếp theo VQS = √lượt × chất-lượng³ (ưu tiên SỐ LƯỢT/độ đông khách, dìm rating tệ); dưới 5 lượt không xếp.")
_s.italic = True
_s.font.size = Pt(9)
_s.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

COLS = ["STT", "Khách sạn", "★", "Lượt", "VQS", "Địa chỉ", "Google"]
tbl = doc.add_table(rows=1, cols=len(COLS))
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(COLS):
    c = tbl.rows[0].cells[i]
    c.text = ""
    run = c.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    shade(c, "FFF2CC")

for r in xh:
    row = tbl.add_row().cells
    vals = [str(r["rank"]), r["ten"], "%.1f" % r["R"], "{:,}".format(r["n"]),
            "%.2f" % r["vqs"], r.get("dia_chi") or ""]
    for j, v in enumerate(vals):
        row[j].text = ""
        run = row[j].paragraphs[0].add_run(v)
        run.font.size = Pt(8.5)
    row[6].text = ""
    add_link(row[6], r["google_maps_url"], "mở")

for row in tbl.rows:
    for w, cell in zip([0.9, 5.2, 1.0, 1.5, 1.5, 4.4, 1.0], row.cells):
        cell.width = Cm(w)

# ghi chu: chua du danh gia
doc.add_paragraph()
_c = doc.add_paragraph().add_run("Chưa đủ đánh giá (< 5 lượt) — mẫu quá nhỏ (%d cơ sở):" % len(chua))
_c.bold = True
_c.font.size = Pt(9.5)
for r in sorted(chua, key=lambda x: -(x.get("n") or 0)):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run("%s — %.1f★ / %s lượt" % (r["ten"], r["R"], r["n"]))
    run.font.size = Pt(8.5)

p = _dr.kiem_loi_ra(OUT)
os.makedirs(os.path.dirname(p), exist_ok=True)
tmp = p + ".tmp"
doc.save(tmp)
os.replace(tmp, p)
print("saved -> %s  (%d xếp hạng, %d chưa đủ)  [NỘI BỘ, gitignored]" % (OUT, len(xh), len(chua)))
