# -*- coding: utf-8 -*-
"""Build guide .docx GỘP 10 thành phố theo ĐÚNG format file Đà Lạt (giữ style + shading).

Load file Đà Lạt làm BASE (thừa hưởng styles CardTitle/TagLine/Small + table shading), APPEND 9 tp
từ guide-data.json (engine+KB, không LLM, không bịa giờ-clock/duration/tip). Ghi file MỚI ở Downloads.

Chạy:  python trip-planner/scripts/build_guide_docx.py
"""
import copy
import io
import json
import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.table import Table

HERE = os.path.dirname(os.path.abspath(__file__))
BASE = os.path.join(os.path.expanduser("~"), "Downloads", "Lich_trinh_Da_Lat_Travel_Guide_2026.docx")
OUT = os.path.join(os.path.expanduser("~"), "Downloads", "Lich_trinh_10_ThanhPho_Travel_Guide_2026.docx")
DATA = os.path.join(HERE, "guide-data.json")

doc = docx.Document(BASE)
data = json.load(io.open(DATA, encoding="utf-8"))

# Template tables (giữ shading/format): 5-cột lịch = tables[0]; card body 1x1 = tables[8].
TMPL_ITIN = doc.tables[0]._tbl
TMPL_CARD = doc.tables[8]._tbl
STYLES = {s.name for s in doc.styles}


def _last_p():
    return doc.paragraphs[-1]._p


def append_table(template_tbl):
    """Deepcopy template (giữ nguyên shading/width/border) rồi chèn sau paragraph cuối."""
    new = copy.deepcopy(template_tbl)
    _last_p().addnext(new)
    return Table(new, doc)


def clear_cell(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)


def label_para(cell, label, text=None, style=None):
    p = cell.add_paragraph()
    if style and style in STYLES:
        p.style = style
    if label:
        r = p.add_run(label)
        r.bold = True
    if text:
        p.add_run((" " if label else "") + text)
    return p


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(26)


def itinerary_table(rows):
    t = append_table(TMPL_ITIN)
    # giữ header row[0] (đã có text + shading); xoá data rows còn lại, giữ 1 row làm mẫu
    trs = t.rows
    row_tmpl = copy.deepcopy(trs[1]._tr) if len(trs) > 1 else copy.deepcopy(trs[0]._tr)
    for r in list(trs[1:]):
        r._tr.getparent().remove(r._tr)
    for row in rows:
        tr = copy.deepcopy(row_tmpl)
        t._tbl.append(tr)
        cells = Table(t._tbl, doc).rows[-1].cells
        vals = [row["buoi"], row["ten"], row["loai"], "—", row["ghi_chu"]]
        for c, v in zip(cells, vals):
            c.text = v or ""


def card(c):
    doc.add_paragraph(c["ten"], style="CardTitle" if "CardTitle" in STYLES else None)
    if c.get("tagline"):
        doc.add_paragraph(c["tagline"], style="TagLine" if "TagLine" in STYLES else None)
    t = append_table(TMPL_CARD)
    cell = t.rows[0].cells[0]
    clear_cell(cell)
    if c.get("gioi_thieu"):
        label_para(cell, "Giới thiệu nhanh.", c["gioi_thieu"])
    if c.get("hoat_dong"):
        label_para(cell, "Có gì ở đây?")
        for h in c["hoat_dong"]:
            label_para(cell, None, "• " + h)
    if c.get("loi_vao_dac_trung"):
        label_para(cell, "Lối vào đặc trưng.", "🚡 " + c["loi_vao_dac_trung"])
    if c.get("gio_mo"):
        label_para(cell, "Giờ mở.", c["gio_mo"])
    if c.get("gia_ve"):
        label_para(cell, "Giá vé (tham khảo).", c["gia_ve"])
    if c.get("phu_hop"):
        label_para(cell, "Phù hợp.", c["phu_hop"])
    label_para(cell, None, "Nguồn: %d nguồn dữ liệu · thông tin có thể thay đổi, kiểm tra lại trước ngày đi."
               % c.get("nguon", 0), style="Small")


for city in data:
    doc.add_page_break()
    title("KẾ HOẠCH DU LỊCH " + city["ten"])
    p = doc.add_paragraph()
    p.add_run(city["cauHoi"]).italic = True
    doc.add_paragraph(city["meta"] + f"  ·  (dữ liệu: {city['generated_from']})")

    doc.add_heading("1. Phương án lịch trình", level=1)
    for d in city["days"]:
        doc.add_heading(f"Ngày {d['day']}" + (f" · khu {d['region']}" if d.get("region") else ""), level=2)
        if d["rows"]:
            itinerary_table(d["rows"])
        else:
            doc.add_paragraph("(chưa đủ điểm cho ngày này)")
    if city.get("hotel"):
        doc.add_paragraph("🏨 Gợi ý khách sạn: " + city["hotel"])
    if city.get("notes"):
        drop = [n for n in city["notes"] if "ngoài vùng thuận tiện" in n]
        other = [n for n in city["notes"] if "ngoài vùng thuận tiện" not in n]
        for n in other:
            doc.add_paragraph("• " + n, style="Small" if "Small" in STYLES else None)
        if drop:
            doc.add_paragraph(f"• (+{len(drop)} cụm ngoài vùng thuận tiện — chưa đưa vào lịch)",
                              style="Small" if "Small" in STYLES else None)

    doc.add_heading("2. Destination Guide — chi tiết từng điểm", level=1)
    for c in city["cards"]:
        card(c)

doc.save(OUT)
print("Da ghi:", OUT)
print("So tp append:", len(data), "| tong bang:", len(doc.tables), "| tong para:", len(doc.paragraphs))
